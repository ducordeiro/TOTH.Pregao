import re
import shutil
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TEMPLATE = Path(r"C:\Users\ducor\Downloads\Document 2.docx")
SOURCE_TR = Path(
    r"C:\Users\ducor\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm"
    r"\LocalState\sessions\6A86223F22A57B5C2750F0831822A5157DFFE0E7"
    r"\transfers\2026-28\19TR+-+Mobiliarios++e+eletrodomesticos.pdf"
)
OUT_DOCX = Path("Proposta_Final_Tabela_Mobiliarios_Eletrodomesticos.docx")


def compact(value):
    if value is None:
        return ""
    value = str(value).replace("\r", "\n")
    lines = [" ".join(line.split()) for line in value.splitlines()]
    return " ".join(line for line in lines if line).strip()


def extract_items():
    items = []
    current = None
    with pdfplumber.open(str(SOURCE_TR)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for raw_row in table:
                    cells = [compact(cell) for cell in raw_row]
                    while len(cells) < 3:
                        cells.append("")
                    item, spec, quant = cells[0], cells[1], cells[2]

                    if item.startswith("ANEXO") or item.lower() == "item":
                        current = None
                        continue

                    if re.fullmatch(r"\d+", item) and spec and quant:
                        current = {
                            "item": item,
                            "quantidade": quant,
                            "unidade": "",
                            "descricao": spec,
                            "marca": "",
                            "valor_unitario": "",
                            "valor_total": "",
                        }
                        items.append(current)
                        continue

                    if not item and spec and current:
                        current["descricao"] = compact(current["descricao"] + " " + spec)
    return items


def set_font(run, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=100, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def write_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run)
    run.bold = bold
    run.font.size = Pt(size)


def build_docx(items):
    if not items:
        raise RuntimeError("informações não encontradas")

    shutil.copyfile(TEMPLATE, OUT_DOCX)
    doc = Document(str(OUT_DOCX))

    doc.add_paragraph()
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    total_width = 15800
    widths = [700, 700, 620, 8200, 1400, 1900, 2280]
    set_table_width(table, total_width)

    headers = [
        "ITEM",
        "QTD",
        "UN",
        "DESCRIÇÃO.",
        "MARCA",
        "VALOR\nUNITÁRIO",
        "VALOR TOTAL",
    ]
    header_row = table.rows[0]
    repeat_header(header_row)
    for cell, header, width in zip(header_row.cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        write_cell(cell, header, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    for item in items:
        values = [
            item["item"],
            item["quantidade"],
            item["unidade"],
            item["descricao"],
            item["marca"],
            item["valor_unitario"],
            item["valor_total"],
        ]
        row = table.add_row()
        for idx, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 3 else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cell, value, bold=False, size=7.2, align=align)

    doc.save(str(OUT_DOCX))


def main():
    items = extract_items()
    build_docx(items)
    print(OUT_DOCX.resolve())
    print(f"items={len(items)}")


if __name__ == "__main__":
    main()
