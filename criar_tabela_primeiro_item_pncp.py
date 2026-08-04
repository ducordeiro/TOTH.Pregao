from decimal import Decimal
import json
from pathlib import Path
import urllib.request

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


API_URL = (
    "https://pncp.gov.br/pncp-api/v1/orgaos/"
    "45370707000128/compras/2026/94/itens?pagina=1&tamanhoPagina=10"
)
OUT = Path("Tabela_Primeiro_Item_PNCP_45370707000128_2026_94.docx")


def fix_mojibake(value):
    if not isinstance(value, str):
        return value
    if "Ã" in value or "Â" in value:
        try:
            return value.encode("latin1").decode("utf-8")
        except UnicodeError:
            return value
    return value


def money(value):
    if value is None:
        return ""
    amount = Decimal(str(value))
    formatted = f"{amount:,.2f}"
    return "R$ " + formatted.replace(",", "X").replace(".", ",").replace("X", ".")


def number(value):
    if value is None:
        return ""
    amount = Decimal(str(value))
    if amount == amount.to_integral_value():
        return str(int(amount))
    return str(amount).replace(".", ",")


def set_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor.from_string("2E74B5")
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)


def write_cell(cell, text, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_font(run)
    run.bold = bold
    run.font.size = Pt(9)


def main():
    with urllib.request.urlopen(API_URL, timeout=30) as response:
        data = json.loads(response.read().decode("utf-8"))

    items = data["value"] if isinstance(data, dict) and "value" in data else data
    item = items[0]
    item = {key: fix_mojibake(value) for key, value in item.items()}

    doc = Document()
    setup_doc(doc)

    title = doc.add_heading("Tabela de Itens - PNCP", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    r = note.add_run("Fonte: PNCP, edital 45370707000128/2026/94. Linha extraida da aba Itens.")
    set_font(r)
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(89, 89, 89)

    headers = [
        "Numero",
        "Descricao",
        "Quantidade",
        "Unidade",
        "Valor unitario estimado",
        "Valor total",
    ]
    values = [
        item.get("numeroItem", ""),
        item.get("descricao", ""),
        number(item.get("quantidade")),
        item.get("unidadeMedida", ""),
        money(item.get("valorUnitarioEstimado")),
        money(item.get("valorTotal")),
    ]
    widths = [720, 4680, 900, 1260, 900, 900]

    table = doc.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table)

    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        write_cell(cell, header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for cell, value, width, idx in zip(table.rows[1].cells, values, widths, range(len(values))):
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        align = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
        write_cell(cell, value, align=align)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
