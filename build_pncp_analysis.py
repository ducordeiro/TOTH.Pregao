from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


BASE = Path(__file__).resolve().parent
PDF = BASE / "ED_292026_PE0232026_MOBILIARIOS.pdf"
OUT = BASE / "Analise_PNCP_46583654000196_2026_115.docx"

SOURCE_URL = "https://pncp.gov.br/app/editais/46583654000196/2026/115"

HEADER_LINES = {
    "PREFEITURA MUNICIPAL DE MIRACATU",
    "DEPARTAMENTO DE COMPRAS E PROJETOS",
    "www.miracatu.sp.gov.br",
}


def clean_line(line: str) -> str:
    s = " ".join(line.replace("\u00a0", " ").split())
    s = re.sub(r"\s+([,.;:])", r"\1", s)
    s = re.sub(r"(\d{2})/\s+(\d{2})", r"\1/\2", s)
    s = s.replace(" - ", "-")
    return s.strip()


def load_pages():
    reader = PdfReader(str(PDF))
    pages = []
    for idx, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        lines = []
        for line in raw.splitlines():
            s = clean_line(line)
            if not s:
                continue
            if s in HEADER_LINES:
                continue
            if s.startswith("Av. Dona Evarista") or s.startswith("CEP:"):
                continue
            if s.startswith("Pregão Eletrônico n") or s.startswith("Processo Digital 1DOC"):
                continue
            lines.append(s)
        pages.append({"page": idx, "lines": lines, "text": "\n".join(lines)})
    return pages


def normalize_for_match(text: str) -> str:
    return re.sub(r"\s+", " ", text.upper()).strip()


def find_line(pages, pattern, start_page=1):
    rx = re.compile(pattern, re.I)
    for page in pages:
        if page["page"] < start_page:
            continue
        for i, line in enumerate(page["lines"]):
            if rx.search(line):
                return page["page"], i
    raise ValueError(f"Pattern not found: {pattern}")


def slice_between(pages, start_pattern, end_pattern=None, start_page=1, end_page=None):
    sp, si = find_line(pages, start_pattern, start_page=start_page)
    ep = None
    ei = None
    if end_pattern:
        ep, ei = find_line(pages, end_pattern, start_page=sp)
    elif end_page:
        ep = end_page
        ei = len(pages[ep - 1]["lines"])
    else:
        raise ValueError("Provide end_pattern or end_page")

    out = []
    for page in pages:
        pn = page["page"]
        if pn < sp or pn > ep:
            continue
        start_idx = si if pn == sp else 0
        stop_idx = ei if pn == ep else len(page["lines"])
        chunk = page["lines"][start_idx:stop_idx]
        if chunk:
            out.append((pn, chunk))
    return out


def section_pages(chunks):
    nums = [pn for pn, lines in chunks if lines]
    if not nums:
        return ""
    if min(nums) == max(nums):
        return f"p. {min(nums)}"
    return f"p. {min(nums)}-{max(nums)}"


def extract_key_value(pages, pattern, default=""):
    rx = re.compile(pattern, re.I)
    for page in pages:
        for line in page["lines"]:
            m = rx.search(line)
            if m:
                return m.group(1).strip()
    return default


def extract_item_index(pages):
    lines = []
    capture = False
    for page in pages:
        if page["page"] < 22 or page["page"] > 47:
            continue
        for line in page["lines"]:
            if "1.6. ESPECIFICA" in normalize_for_match(line):
                capture = True
            if capture:
                lines.append((page["page"], line))

    items = []
    current = None
    rx = re.compile(r"^(\d{2})\s+(\d+)\s+([A-Z]{2,4})\s*(.*)$")
    for pn, line in lines:
        m = rx.match(line)
        if m:
            if current:
                items.append(current)
            current = {
                "item": m.group(1),
                "quant": m.group(2),
                "und": m.group(3),
                "desc": [m.group(4).strip()] if m.group(4).strip() else [],
                "tipo": "",
                "page": pn,
            }
            continue
        if not current:
            continue
        up = normalize_for_match(line)
        if up in {"EXCLUSIVO", "AMPLA"}:
            current["tipo"] = up.title()
            continue
        if len(" ".join(current["desc"])) < 170:
            current["desc"].append(line)
    if current:
        items.append(current)

    # Some long table rows are split before the item number; keep the index useful
    # without pretending it is a perfect catalog extraction.
    compact = []
    seen = set()
    for item in items:
        if item["item"] in seen:
            continue
        seen.add(item["item"])
        desc = " ".join(item["desc"])
        desc = re.sub(r"\s+", " ", desc).strip()
        compact.append({
            "item": item["item"],
            "quant": item["quant"],
            "und": item["und"],
            "desc": desc[:160] + ("..." if len(desc) > 160 else ""),
            "tipo": item["tipo"],
            "page": item["page"],
        })
    return compact


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [2700, 6660]
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Informação"
    for cell, width in zip(hdr, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for cell, width in zip(cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(89, 89, 89)


def add_block(doc, title, chunks, note=None):
    doc.add_heading(title, level=2)
    add_source_note(doc, f"Origem: {section_pages(chunks)}." + (f" {note}" if note else ""))
    for pn, lines in chunks:
        if len(chunks) > 1:
            p = doc.add_paragraph()
            run = p.add_run(f"Página {pn}")
            set_font(run)
            run.bold = True
            run.font.color.rgb = RGBColor(31, 77, 120)
        para = []
        for line in lines:
            if re.match(r"^\d+(\.\d+)*\.", line) or re.match(r"^(ANEXO|OBJETO:|ITEM\s+)", line, re.I):
                if para:
                    doc.add_paragraph(" ".join(para))
                    para = []
                doc.add_paragraph(line)
            else:
                para.append(line)
                if len(" ".join(para)) > 600:
                    doc.add_paragraph(" ".join(para))
                    para = []
        if para:
            doc.add_paragraph(" ".join(para))


def add_item_index(doc, items):
    doc.add_heading("Índice de itens do Anexo I", level=2)
    add_source_note(
        doc,
        "Origem: p. 22-47. Índice automatizado para triagem; descrições completas permanecem no PDF oficial.",
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["Item", "Qtde", "Und.", "Descrição inicial", "Cota", "Página"]
    widths = [700, 800, 750, 5000, 1200, 910]
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        cell.text = text
        set_cell_width(cell, width)
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for item in items[:60]:
        cells = table.add_row().cells
        values = [
            item["item"],
            item["quant"],
            item["und"],
            item["desc"],
            item["tipo"],
            str(item["page"]),
        ]
        for cell, value, width in zip(cells, values, widths):
            cell.text = value
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def main():
    pages = load_pages()
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Análise de Edital PNCP - Blocos Extraídos")
    set_font(run)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Edital nº 029/2026 | Pregão Eletrônico nº 023/2026 | Município de Miracatu/SP")
    set_font(r)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(89, 89, 89)

    add_metadata_table(
        doc,
        [
            ("Fonte PNCP", SOURCE_URL),
            ("Órgão", "Prefeitura Municipal de Miracatu"),
            ("CNPJ", "46.583.654/0001-96"),
            ("Objeto", "Registro de preços para aquisição de mobiliários"),
            ("Data da sessão", extract_key_value(pages, r"DATA DA REALIZAÇÃO:\s*(.*)", "14/07/2026")),
            ("Recebimento das propostas", extract_key_value(pages, r"RECEBIMENTO DAS PROPOSTAS:\s*(.*)", "até 8h50min de 14/07/2026")),
            ("Critério de julgamento", "Menor preço por item"),
            ("Arquivo PNCP", PDF.name),
            ("Observação técnica", "PDF com camada de texto; OCR não foi necessário para este arquivo."),
        ],
    )

    doc.add_heading("Blocos selecionados", level=1)
    blocks = [
        (
            "Preâmbulo e dados da sessão",
            slice_between(pages, r"PREÂMBULO", r"EDITAL N", start_page=1),
            "Inclui objeto resumido, datas, plataforma, critério e legislação.",
        ),
        (
            "Objeto",
            slice_between(pages, r"^1\.\s+DO OBJETO$", r"^2\.\s+DA DESPESA", start_page=2),
            None,
        ),
        (
            "Participação na licitação",
            slice_between(pages, r"^5\.\s+DA PARTICIPAÇÃO", r"^6\.\s+DA PARTICIPAÇÃO DE EMPRESAS", start_page=5),
            None,
        ),
        (
            "Proposta e classificação da proposta",
            slice_between(pages, r"^7\.\s+DA PROPOSTA", r"^8\.\s+DA ABERTURA", start_page=7),
            None,
        ),
        (
            "Catálogos",
            slice_between(pages, r"^14\.\s+DOS CATÁLOGOS", r"^15\.\s+DA HABILITAÇÃO", start_page=12),
            None,
        ),
        (
            "Habilitação",
            slice_between(pages, r"^15\.\s+DA HABILITAÇÃO", r"^16\.\s+DO RECURSO", start_page=13),
            None,
        ),
        (
            "Recebimento provisório e definitivo",
            slice_between(pages, r"^20\.\s+DO RECEBIMENTO", r"^21\.\s+DAS SANÇÕES", start_page=19),
            None,
        ),
        (
            "Sanções",
            slice_between(pages, r"^21\.\s+DAS SANÇÕES", r"^22\.\s+DISPOSIÇÕES", start_page=19),
            None,
        ),
        (
            "Anexos do edital",
            slice_between(pages, r"^23\.\s+DOS ANEXOS", r"VINICIUS", start_page=21),
            None,
        ),
        (
            "Termo de Referência - justificativa",
            slice_between(pages, r"^1\.\s+JUSTIFICATIVA", r"^1\.6\.\s+ESPECIFICA", start_page=22),
            None,
        ),
        (
            "Termo de Referência - pagamento, entrega e obrigações",
            slice_between(pages, r"^2\.\s+FONTE DE RECURSO", r"JULIE MORAES", start_page=47),
            None,
        ),
    ]

    for title_text, chunks, note in blocks:
        add_block(doc, title_text, chunks, note=note)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_item_index(doc, extract_item_index(pages))

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("Documento gerado a partir do edital oficial disponível no PNCP.")
    set_font(rf)
    rf.font.size = Pt(9)
    rf.font.color.rgb = RGBColor(89, 89, 89)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
