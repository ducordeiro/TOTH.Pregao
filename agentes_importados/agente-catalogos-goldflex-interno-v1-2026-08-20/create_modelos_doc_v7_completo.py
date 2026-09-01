from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\modelos-de-cadeiras-analisados-v7-completo.docx")

GROUPS = {
"Cadeiras fixas": {
"Cadeiras em polipropileno injetado": [
{"arms":"Braço/prancheta fixa","model":"Cadeira Universitária Fixa em Polipropileno com Prancheta","seat":"aprox. 42 cm profundidade x 46 cm largura","back":"aprox. 28,0 cm altura x 46 cm largura","tech":"Assento e encosto em PP de alta resistência. Suportes do encosto: 2 tubos oblongos 16 x 30 mm, espessura 1,5 mm. Suportes do assento: 2 tubos redondos 3/4, espessura 1,2 mm. Pés: 2 tubos oblongos 16 x 30 mm, espessura 1,2 mm. Solda MIG. Prancheta MDP/MDF entre 18 mm, laminado melamínico nas duas faces e fita PVC, aprox. 49 x 23 x 1,8 cm.","extra":"Tratamento antiferrugem, desengraxe, estabilização, fosforização, pintura a pó eletrostática e estufa a 250 °C. Porta-livros mencionado. Capacidade, norma, cor, garantia e dimensões globais não evidenciadas.","source":"UFRGS"},
{"arms":"Não evidenciado","model":"Item 1 — Cadeira Fixa Empilhável","seat":"aprox. 43 cm profundidade x 45 cm largura","back":"aprox. 30 cm altura x 45 cm largura","tech":"Assento e encosto em PP de alta resistência, injetados termoplasticamente, com curvatura anatômica. Encosto com furos de aeração, encaixes e pino-tampão de PP. Assento travado por parafusos. Estrutura de aço com pés/suportes 16 x 30 mm, espessura 1,5 mm; travessas 22,22 mm, espessura 1,5 mm; solda MIG.","extra":"Modelo empilhável, cor preta, pintura a pó eletrostática, sapatas fixas e ponteiras plásticas. Peso até 120 kg e garantia de 12 meses declarados. Concha dupla, ausência de braços, base e sapatas em poliamida não evidenciadas.","source":"Assistência Jurídica"},
],
"Cadeiras estofadas": [
{"arms":"Braço/prancheta fixa","model":"Cadeira Universitária Fixa Estofada com Prancheta","seat":"aprox. 39 x 42 x 5 cm","back":"aprox. 29 x 36 x 4 cm","tech":"Compensado multilaminado mínimo 12 mm, espuma PU moldada/injetada, densidade mínima 50 kg/m³, perfil PVC e tecido 100% poliéster. Estrutura em tubos redondos 7/8, espessura 1,2 mm, quatro pés, solda MIG e deslizantes. Prancheta MDP/MDF entre 18 mm, laminado nas duas faces, fita PVC, aprox. 49 x 23 x 1,8 cm.","extra":"Carga declarada 130 kg, atendimento à NBR 16671/2018 e opção destra/canhota. Cores, código, garantia e embalagem não evidenciados.","source":"UFRGS"},
{"arms":"Sem braços","model":"Modelo Fixa Estofada","seat":"aprox. 42 cm profundidade x 43 cm largura","back":"aprox. 27 cm altura x 39 cm largura","tech":"Compensado multilaminado, moldagem anatômica a quente, bordas arredondadas, porcas-garras, espuma PU moldada/injetada e perfil PVC. Estrutura tubular 7/8, quatro pés, solda MIG, ponteiras/sapatas de borracha e reforço declarado a 25 cm do chão.","extra":"Revestimento em couro ecológico preto. Tratamento antiferrugem, pintura a pó eletrostática e estufa a 250 °C. Capacidade, norma, garantia, densidade/espessura da espuma e laudos não evidenciados.","source":"Prefeitura BH"},
{"arms":"Com braços","model":"Cadeira Fixa Universitária Executiva Com Braço Escamoteável","seat":"aprox. 43,5 cm profundidade x 46 cm largura","back":"aprox. 37 cm altura x 42 cm largura","tech":"Compensado multilaminado mínimo 12 mm, moldagem anatômica, espuma PU d55, perfil PVC e couro sintético preto. Estrutura aço carbono 7/8, solda MIG, suporte para livros e ponteiras PP/PVC. Braço escamoteável em aço/PU integral skin com suporte para prancheta.","extra":"Confirmações registradas: espuma de nylon injetada, PVC conforme tipo/cor, estrutura monobloco, quatro parafusos sextavados e pintura epóxi preto-fosco. Laudos/ensaios informados como aprovados pelo fabricante.","source":"Prefeitura BH universitária"},
{"arms":"Não evidenciado","model":"Item 6 — Cadeira Secretaria Fixa","seat":"aprox. 39 cm profundidade x 42 cm largura","back":"aprox. 29 cm altura x 37 cm largura","tech":"Compensado multilaminado 12 mm, moldagem anatômica, espuma PU moldada/injetada d50, perfil PVC. Estrutura de aço carbono 7/8, espessura 1,2 mm, quatro pés, solda MIG, ponteiras PP/PVC e deslizantes.","extra":"Cor a definir; revestimento tecido ou sintético; peso 130 kg. Tratamento antiferrugem e pintura eletrostática. Catálogos apresentam divergência de norma: NBR 13692/2018 na identificação e NBR 13962/2018 nas observações.","source":"Patrocínio Paulista"},
{"arms":"Não evidenciado","model":"Item 10 — Cadeira Fixa","seat":"aprox. 39 cm profundidade x 42 cm largura","back":"aprox. 29 cm altura x 37 cm largura","tech":"Construção igual à do Item 6: compensado multilaminado 12 mm, espuma PU d50, perfil PVC, tubos 7/8 de aço carbono 1,2 mm, quatro pés, solda MIG, ponteiras e deslizantes.","extra":"Cor preta, revestimento couro ecológico com menção a malha 50% poliéster e peso 130 kg. Divergência de referência normativa entre as páginas do catálogo; composição do revestimento deve ser esclarecida.","source":"Patrocínio Paulista"},
],
},
"Cadeiras giratórias": {
"Cadeiras com encosto em tela e assento estofado": [
{"arms":"Com braços","model":"Cadeira Ergonômica Tela Com Braços","seat":"aprox. 48 x 49 x 8 cm","back":"aprox. 54 x 45 x 8 cm","tech":"Encosto em Tela Mesh sobre tubo de aço. Assento em compensado multilaminado mínimo 12 mm, PU moldada/injetada d50, espessura mínima 55 mm e perfil PVC. Base aço carbono 25 x 25 mm, espessura 1,2 mm, capa PP e solda MIG. Cinco rodízios de nylon com esferas de aço, pistão classe 4 e Back System.","extra":"Braço digitador com seis regulagens. Altura do assento aprox. 455–565 mm; reclinação aprox. -5° a 20°. Carga 140 kg, NR 17 e NBR 13962/2018 declaradas.","source":"São José dos Pinhais"},
{"arms":"Com braços","model":"Modelo Tela Ergonômica Relax com Braço","seat":"aprox. 48 cm profundidade x 48 cm largura","back":"aprox. 56 cm altura x 48 cm largura","tech":"Assento em compensado multilaminado 12 mm, PU d50 moldada/injetada, perfil PVC e Tela Mesh no encosto. Base estrela em aço carbono 25 x 25 mm, capa PP e solda MIG. Cinco rodízios com esferas de aço, pistão a gás, rotação 360° e capa telescópica PP.","extra":"Mecanismo Relax monobloco, travamento ou livre flutuação, ajuste de tensão, inclinação sincronizada 2:1 e braço digitador com sete regulagens. Carga 120 kg, NR 17 e NBR 13962/2018 declaradas; courvin preto.","source":"Material Bélico Tela"},
],
"Cadeiras estofadas": [
{"arms":"Com braços","model":"Cadeira Ergonômica Com Braços","seat":"aprox. 46 cm profundidade x 48 cm largura","back":"aprox. 47 cm altura x 44 cm largura","tech":"Assento em compensado multilaminado mínimo 15 mm, PU d55 e 50 mm, carenagem PP texturizada. Encosto regulável com apoio lombar, PU d55 e 45 mm, concha e carenagem PP. Base aço carbono 25 x 25 mm, capa PP, solda MIG; cinco rodízios duplo giro de 60 mm, PU/antirrisco; pistão a gás 360°.","extra":"Back System com sete posições e ajuste de 75 mm; reclinação aprox. -5° a 20°. Braço digitador tipo T. Carga 135 kg, NR 17 e NBR 13962/2018 declaradas; courvin azul escuro.","source":"CRECI"},
{"arms":"Com braços","model":"Item 8 — Cadeira Executiva Lâmina Ergonômica","seat":"aprox. 42 cm profundidade x 46 cm largura","back":"aprox. 37 cm altura x 40 cm largura","tech":"Compensado multilaminado com espessura extraída como “1 3 mm”, a confirmar; espuma PU d50, perfil PVC e moldagem anatômica. Base estrela aço 25 x 25 mm, espessura 1,2 mm, capa PP, cinco rodízios de nylon, pistão a gás 360° e capa telescópica. Braço digitador com seis regulagens.","extra":"Mecanismo Relax monobloco, travamento ou livre flutuação e ajuste de tensão. Altura do assento aprox. 450–560 mm; peso 120 kg. Cor e revestimento a definir; confirmar a norma específica.","source":"Patrocínio Paulista"},
],
"Cadeiras em poliuretano integral skin": [
{"arms":"Não evidenciado","model":"Item 3 — Cadeira Injetada Industrial","seat":"aprox. 41,5 cm profundidade x 43 cm largura x 4 cm espessura","back":"aprox. 25 cm altura x 40,5 cm largura x 4 cm espessura","tech":"Assento e encosto em PU integral skin preto texturizado. Base aço carbono 25 x 25 mm, capa PP e solda MIG. Cinco rodízios PU de 50 mm, pistão classe 3 indicado como 125 mm, rotação 360°, capa telescópica PP, suporte do encosto oval 16 x 30 x 1,9 mm e flange de aço 3 mm.","extra":"Encosto regulável em 100 mm e flange com inclinação indicada de 3°. Altura catalogada 510–620 mm; peso 120 kg e NR 17 declarados. O pedido associado exige 350–480 mm e curso 130 mm, gerando pontos de conflito.","source":"Material Bélico Injetado"},
],
"Bancos/cadeiras semi-sentados em PU": [
{"arms":"Não evidenciado","model":"Item 2 — Banco Semi Sentado com Rodízios","seat":"largura 35 cm x profundidade 28,5 cm","back":"altura 13 cm x largura 22 cm","tech":"Assento PU integral skin preto texturizado, inclinação para repouso semi-sentado, base aço giratória, ajuste pneumático, mecanismo de ângulo/inclinação e base estrela 25 x 25 mm com capa PP e solda MIG.","extra":"Giro 360° e peso máximo 120 kg. O catálogo menciona rodízios no nome, mas descreve cinco hastes com sapatas fixas. Altura aparece como “máxima 385 mm | mínima 510 mm”, ordem inconsistente; NR 17 declarada.","source":"Material Bélico Injetado"},
],
},
"Cadeiras rebatíveis de auditório": {
"Cadeiras de auditório estofadas": [
{"arms":"Com braços / prancheta escamoteável","model":"Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável","seat":"aprox. 48 cm profundidade x 49 cm largura","back":"aprox. 48 cm altura x 46 cm largura","tech":"Assento e encosto em compensado multilaminado aprox. 15 mm, moldados a quente, espuma PU d50 e perfil PVC. Estrutura aço carbono SAE 1010/1020, seção 50 x 30 mm, ponteiras PP/PVC, tratamento antiferrugem, pintura eletrostática e estufa a 250 °C.","extra":"Braços em PU integral skin preto e prancheta escamoteável em MDF dupla face texturizado, fórmica e bordas PVC. Tecido de polipropileno azul-marinho e NBR 13962/2018 declarada. Catálogo individual; longarina não evidenciada.","source":"Almirante Alexandrino"},
],
},
}

doc = Document(); s = doc.sections[0]; s.top_margin = Pt(54); s.bottom_margin = Pt(54); s.left_margin = Pt(60); s.right_margin = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Modelos de cadeiras analisados"); r.bold = True; r.font.size = Pt(16)
doc.add_paragraph("Classificação por tipo de movimento e categoria construtiva oficial.")
doc.add_paragraph("Registro técnico consolidado. As informações abaixo foram extraídas dos catálogos e laudos já avaliados. Quando um dado não foi comprovado, ele permanece indicado como não evidenciado ou a confirmar.")
for movement, categories in GROUPS.items():
    h = doc.add_paragraph(); rr = h.add_run(movement); rr.bold = True; rr.font.size = Pt(13)
    for category, models in categories.items():
        h = doc.add_paragraph(); rr = h.add_run(category); rr.bold = True; rr.font.size = Pt(11)
        for m in models:
            doc.add_paragraph(f"{m['arms']} — {m['model']}", style="List Bullet")
            doc.add_paragraph(f"    Assento: {m['seat']}; Encosto: {m['back']}")
            doc.add_paragraph(f"    Construção e componentes: {m['tech']}")
            doc.add_paragraph(f"    Capacidade, mecanismos, acabamento e evidência: {m['extra']}")
            doc.add_paragraph(f"    Fonte/registro: {m['source']}")
doc.save(OUT); print(OUT)
