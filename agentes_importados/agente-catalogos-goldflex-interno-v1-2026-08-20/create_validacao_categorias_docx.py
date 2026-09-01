from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos")
SOURCE = ROOT / "catalogos de pedidos" / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v4.docx"
OUTDIR = ROOT / "validação de dados"
OUT = OUTDIR / "validacao-modelos-cadeiras-por-categoria-v3.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-validacao-categorias.docx"
OUTDIR.mkdir(parents=True, exist_ok=True)
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def heading(text, size=14):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True
    r.font.size = Pt(size); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00); return p

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VALIDAÇÃO DA BIBLIOTECA DE MODELOS"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Modelos de cadeiras organizados por categoria"); r.bold = True; r.font.size = Pt(13)

heading("Critério de organização")
doc.add_paragraph("Os 36 modelos distintos foram distribuídos em uma categoria principal, para evitar dupla contagem. Quando um produto poderia se relacionar a mais de um grupo, foi usada a característica predominante do modelo declarado no catálogo: mocho, universitária, auditório ou presença de rodízios.")

groups = {
    "Cadeiras fixas": [
        ("Cadeira Fixa Empilhável", "Assistência Jurídica"),
        ("Cadeira Injetada Industrial", "Material Bélico Injetada"),
        ("Cadeira Fixa", "Patrocínio Paulista — Item 10"),
        ("Cadeira Secretaria Fixa", "Patrocínio Paulista — Item 6"),
        ("Cadeira Executiva Lâmina Ergonômica", "Patrocínio Paulista — Item 8"),
        ("Cadeira Fixa Executiva", "Bombeiros"),
        ("Cadeira Chile Fixa Hotel TH", "Lagoa Santa 33"),
        ("Cadeira Fixa Estofada", "Prefeitura BH"),
        ("Cadeira Diretor Fixa", "Catálogos geral — Item 1"),
        ("Cadeira/Poltrona Chile Diretor Fixa com Braço", "Catálogos geral — Item 1"),
        ("Cadeira Fixa Executiva Chile Contra Capa", "Catálogos geral — Item 2"),
        ("Cadeira Fixa Executiva Contra Capa", "Catálogos geral — Item 4"),
        ("Longarina 2 Lugares sem Braços", "Catálogos geral — Item 8"),
    ],
    "Cadeiras universitárias": [
        ("Cadeira Universitária Fixa Estofada com Prancheta", "UFRGS"),
        ("Cadeira Universitária Fixa em Polipropileno com Prancheta", "UFRGS"),
        ("Cadeira Universitária Escamoteável", "Pró-Reitoria"),
        ("Cadeira Universitária ISO", "Pró-Reitoria"),
        ("Cadeira Universitária Prefeitura BH", "Prefeitura BH Universitária"),
    ],
    "Cadeiras com rodízios": [
        ("Cadeira Ergonômica Tela com Braços", "São José dos Pinhais"),
        ("Cadeira Chile Executiva Giratória Ergonômica", "Del Rei"),
        ("Cadeira Gerente Chile", "Lagoa Santa 12"),
        ("Cadeira Diretor Ergonômica com Braço", "Material Bélico Tela / Pró-Reitoria"),
        ("Cadeira Chile Diretor Gomo com Braço Cromada", "Piumhi"),
        ("Cadeira Ergonômica com Braços", "CRECI"),
        ("Cadeira Chile Giratória Plus Size", "Catálogos geral — Item 6"),
    ],
    "Mochos": [
        ("Mocho Sela Bipartido", "Piçarras / Estância Velha / Tijucas"),
        ("Banco Semi-Sentado com Rodízios", "Material Bélico Injetada"),
        ("Mocho Odontológico Sela Bipartido com Aro", "Catálogos geral — Item 14"),
    ],
    "Poltronas de auditório": [
        ("Poltrona de Auditório Itália Rebatível com Braço e Prancheta Escamoteável", "Paranaguá"),
        ("Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável", "Almirante Alexandrino / Artilharia"),
        ("Poltrona Auditório Roma com Prancheta Escamoteável", "Pró-Reitoria"),
        ("Poltrona Auditório Obeso Roma com Prancheta Escamoteável", "Pró-Reitoria"),
        ("Poltrona Concha Única Anatômica", "São Roque 9"),
        ("Poltrona Auditório", "Secretaria de Segurança Pública"),
        ("Cadeira Auditório", "UNITAU"),
        ("Poltrona Auditório Roma", "Londrina Auditório"),
    ],
}

heading("Resumo por categoria")
summary = doc.add_table(rows=1, cols=2); summary.style = "Table Grid"
for i, h in enumerate(["Categoria", "Quantidade"]): summary.rows[0].cells[i].text = h; shade(summary.rows[0].cells[i])
for name, items in groups.items():
    cells = summary.add_row().cells; cells[0].text = name; cells[1].text = str(len(items))
cells = summary.add_row().cells; cells[0].text = "Total de modelos distintos"; cells[1].text = str(sum(len(x) for x in groups.values()))

for name, items in groups.items():
    heading(name)
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
    for i, h in enumerate(["Nº", "Modelo", "Fonte/pasta"]): table.rows[0].cells[i].text = h; shade(table.rows[0].cells[i])
    for i, (model, source) in enumerate(items, 1):
        cells = table.add_row().cells; cells[0].text = str(i); cells[1].text = model; cells[2].text = source

heading("Observação")
doc.add_paragraph("A classificação é uma organização da biblioteca documental. Ela não substitui a classificação técnica individual do laudo e não significa que modelos de uma mesma categoria sejam equivalentes entre si. As versões anteriores do documento de validação foram preservadas.")
doc.save(OUT)
print(OUT)
