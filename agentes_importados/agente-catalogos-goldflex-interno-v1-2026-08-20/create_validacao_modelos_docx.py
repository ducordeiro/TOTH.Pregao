from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos\catalogos de pedidos")
SOURCE = ROOT / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v4.docx"
OUTDIR = ROOT / "validação de dados"
OUT = OUTDIR / "validacao-modelos-cadeiras-sem-repeticao-v1.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-validacao.docx"
OUTDIR.mkdir(parents=True, exist_ok=True)
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def heading(text, size=14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
    return p

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("VALIDAÇÃO DA BIBLIOTECA DE MODELOS")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Modelos de cadeiras avaliados sem repetição"); r.bold = True; r.font.size = Pt(13)

heading("Resultado da contagem")
p = doc.add_paragraph()
r = p.add_run("29 modelos distintos")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
doc.add_paragraph("Foram identificados 33 PDFs de produtos nas pastas de catálogos. A contagem final agrupou somente repetições comprovadas pelo mesmo nome/modelo declarado, sem fundir modelos apenas semelhantes.")

heading("Critério de consolidação")
for text in [
    "3 registros de Mocho Sela Bipartido foram consolidados em 1 modelo.",
    "2 registros de Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável foram consolidados em 1 modelo.",
    "2 registros de Cadeira Diretor Ergonômica com Braço foram consolidados em 1 modelo.",
    "Modelos com nomes diferentes, configurações diferentes ou evidência insuficiente de identidade foram mantidos separados.",
]:
    doc.add_paragraph(text, style="List Bullet")

heading("Modelos distintos registrados")
models = [
    ("Cadeira Fixa Empilhável", "Assistência Jurídica"),
    ("Poltrona de Auditório Itália Rebatível com Braço e Prancheta Escamoteável", "Paranaguá"),
    ("Mocho Sela Bipartido", "Piçarras / Estância Velha / Tijucas"),
    ("Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável", "Almirante Alexandrino / Artilharia"),
    ("Banco Semi-Sentado com Rodízios", "Material Bélico Injetada"),
    ("Cadeira Injetada Industrial", "Material Bélico Injetada"),
    ("Cadeira Fixa", "Patrocínio Paulista — Item 10"),
    ("Cadeira Secretaria Fixa", "Patrocínio Paulista — Item 6"),
    ("Cadeira Executiva Lâmina Ergonômica", "Patrocínio Paulista — Item 8"),
    ("Cadeira Ergonômica Tela com Braços", "São José dos Pinhais"),
    ("Cadeira Universitária Fixa Estofada com Prancheta", "UFRGS"),
    ("Cadeira Universitária Fixa em Polipropileno com Prancheta", "UFRGS"),
    ("Cadeira Fixa Executiva", "Bombeiros"),
    ("Cadeira Chile Executiva Giratória Ergonômica", "Del Rei"),
    ("Cadeira Gerente Chile", "Lagoa Santa 12"),
    ("Cadeira Chile Fixa Hotel TH", "Lagoa Santa 33"),
    ("Cadeira Diretor Ergonômica com Braço", "Material Bélico Tela / Pró-Reitoria"),
    ("Cadeira Chile Diretor Gomo com Braço Cromada", "Piumhi"),
    ("Cadeira Fixa Estofada", "Prefeitura BH"),
    ("Cadeira Universitária Escamoteável", "Pró-Reitoria"),
    ("Cadeira Universitária ISO", "Pró-Reitoria"),
    ("Poltrona Auditório Roma com Prancheta Escamoteável", "Pró-Reitoria"),
    ("Poltrona Auditório Obeso Roma com Prancheta Escamoteável", "Pró-Reitoria"),
    ("Poltrona Concha Única Anatômica", "São Roque 9"),
    ("Poltrona Auditório", "Secretaria de Segurança Pública"),
    ("Cadeira Auditório", "UNITAU"),
    ("Cadeira Ergonômica com Braços", "CRECI"),
    ("Poltrona Auditório Roma", "Londrina Auditório"),
    ("Cadeira Universitária Prefeitura BH", "Prefeitura BH Universitária"),
]
t = doc.add_table(rows=1, cols=3); t.style = "Table Grid"
for i, h in enumerate(["Nº", "Modelo consolidado", "Fonte/pasta"]):
    t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i])
for i, (model, source) in enumerate(models, 1):
    cells = t.add_row().cells; cells[0].text = str(i); cells[1].text = model; cells[2].text = source

heading("Limite da validação")
doc.add_paragraph("A contagem é uma validação da biblioteca documental e não representa aprovação técnica, conformidade com edital ou equivalência entre produtos. A consolidação por nome/modelo foi conservadora; qualquer fusão adicional deve ser confirmada por ficha técnica, desenho, amostra ou registro formal do fabricante.")
doc.add_paragraph("Fonte da contagem: PDFs de produtos localizados em Modelos de catalogos/catalogos de pedidos, excluindo editais e arquivos de apoio. Documento elaborado na pasta validação de dados.")
doc.save(OUT)
print(OUT)
