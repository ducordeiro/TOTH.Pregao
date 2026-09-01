from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos")
SOURCE = ROOT / "catalogos de pedidos" / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v4.docx"
OUTDIR = ROOT / "validação de dados"
OUT = OUTDIR / "catalogo-final-cadeira-8-executiva-lamina-ergonomica-v4.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-cadeira8-v4.docx"
OUTDIR.mkdir(parents=True, exist_ok=True)
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def title(text, size=20):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)

def heading(text, size=14):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text)

title("CADEIRA EXECUTIVA LÂMINA ERGONÔMICA")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cadeira 8 — configuração para o pedido"); r.bold = True; r.font.size = Pt(12)

heading("Apresentação")
doc.add_paragraph("Cadeira giratória executiva estofada, com assento e encosto em madeira compensada multilaminada, espuma de poliuretano e regulagens para uso profissional. Este catálogo apresenta o modelo-base Cadeira Executiva Lâmina Ergonômica e a configuração ofertada para o pedido analisado.")

heading("Identificação do modelo")
t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
for i, h in enumerate(["Campo", "Informação"]):
    t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i])
for a, b in [("Item", "8 — Cadeira Executiva Lâmina Ergonômica"), ("Categoria", "Cadeira giratória estofada"), ("Configuração de braços", "Sem braços, conforme pedido; braços são acessórios independentes"), ("Fabricante de referência", "Alpherflex Industria e Comércio de Móveis e Ferragens LTDA, conforme catálogo-base")]:
    cells = t.add_row().cells; cells[0].text, cells[1].text = a, b

heading("Características do modelo-base comprovadas no catálogo")
for x in ["Assento e encosto em madeira compensada multilaminada, moldados anatomicamente a quente.", "Espessura declarada da madeira: aproximadamente 13 mm.", "Bordas arredondadas e fixação por porcas-garras.", "Estofamento em espuma moldada/injetada de poliuretano, densidade declarada de 50 kg/m³.", "Bordas encabeçadas com perfil de PVC.", "Base giratória em aço carbono, seção 25 x 25 mm, espessura declarada de 1,2 mm, unida por solda MIG.", "Cinco rodízios de nylon com esferas de aço.", "Pistão a gás com regulagem pneumática da altura do assento e rotação de 360°.", "Mecanismo tipo relax, com travamento em posição de trabalho ou livre flutuação, ajuste de tensão e regulagem do encosto."]: bullet(x)

heading("Configuração ofertada para este pedido")
doc.add_paragraph("A configuração abaixo incorpora as exigências fornecidas para o pedido e deve ser formalizada na proposta, ficha técnica, desenho e amostra quando solicitados.")
for x in ["Sem braços.", "Assento e encosto em madeira compensada multilaminada, espessura mínima solicitada de 12 mm; o catálogo-base declara aproximadamente 13 mm.", "Revestimento em tecido 100% poliéster, cor preta.", "Espuma de poliuretano injetada, densidade especificada no pedido entre 45 e 50 kg/m³; o catálogo-base declara 50 kg/m³.", "Base com cinco patas e rodízios, em aço, com acabamento cromado na cor natural do aço, conforme configuração solicitada.", "Garantia mínima de 12 meses, aceita conforme premissa operacional do fabricante."]: bullet(x)

heading("Medidas")
t2 = doc.add_table(rows=1, cols=3); t2.style = "Table Grid"
for i, h in enumerate(["Componente", "Catálogo-base", "Configuração solicitada"]):
    t2.rows[0].cells[i].text = h; shade(t2.rows[0].cells[i])
for row in [("Encosto", "Aprox. 370 x 400 mm", "Mínimo solicitado: 472 x 335 mm; confirmar ordem das cotas no desenho"), ("Assento", "Aprox. 420 x 460 mm", "Mínimo solicitado: 472 x 435 mm"), ("Altura do assento", "450 a 560 mm", "Configuração conforme pedido e desenho técnico"), ("Rodízios", "5 unidades; nylon", "5 patas com rodízios; especificação final conforme pedido")]:
    cells = t2.add_row().cells
    for i, value in enumerate(row): cells[i].text = value

heading("Atendimento documental e validação da configuração")
for x in ["NR-17 e ABNT NBR 13962/2018 são mencionadas no catálogo-base; apresentar os documentos aplicáveis ao processo.", "A configuração sem braços e a garantia mínima de 12 meses são premissas de oferta do fabricante e devem constar expressamente na proposta.", "As medidas solicitadas diferem das medidas impressas no catálogo-base; apresentar desenho cotado e amostra da configuração ofertada.", "O acabamento cromado da base é requisito da configuração para este pedido; comprovar por ficha técnica, desenho, amostra ou documento do fabricante."]: bullet(x)

doc.save(OUT)
print(OUT)
