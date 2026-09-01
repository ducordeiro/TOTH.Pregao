from pathlib import Path
import shutil
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\marcs\OneDrive\Documents\ChatGPT\relatoriocatalogogoldflex")
TEMPLATE = ROOT / r"Modelos de catalogos\catalogos de pedidos\dados extraidos\catalogo-cadeira-universitaria-pp-prancheta.docx"
OUTDIR = TEMPLATE.parent
NAVY = RGBColor(23, 54, 93); BLUE = RGBColor(47, 117, 181); MUTED = RGBColor(102, 112, 133); INK = RGBColor(31, 41, 51)
LIGHT = "EAF2F8"; HEADER = "17365D"

MODELS = [
    {
        "file": "catalogo-modelo-fixa-estofada-bh.docx",
        "title": "Modelo Fixa Estofada",
        "subtitle": "Cadeira fixa estofada sem braços",
        "manufacturer": "ALPHERFLEX INDÚSTRIA E COMÉRCIO LTDA",
        "category": "Cadeiras estofadas",
        "type": "Cadeira fixa, quatro pés, sem braços",
        "summary": "Cadeira fixa estofada, com assento e encosto em madeira compensada multilaminada moldada anatomicamente, espuma de poliuretano, revestimento em couro ecológico preto e estrutura tubular de aço carbono.",
        "dimensions": [("Assento", "aprox. 42 cm profundidade x 43 cm largura", "Profundidade x largura"), ("Encosto", "aprox. 27 cm altura x 39 cm largura", "Altura x largura")],
        "highlights": [("Assento e encosto", "Madeira compensada multilaminada, moldagem anatômica a quente, bordas arredondadas, porcas-garras, espuma de PU e perfil de PVC."), ("Estrutura", "Tubos redondos de aço carbono, seção 7/8; suportes curvados e unidos por solda MIG."), ("Acabamento", "Tratamento antiferrugem, desengraxe, estabilização, fosforização, pintura a pó eletrostática e secagem em estufa a 250 °C."), ("Revestimento", "Couro ecológico preto."), ("Pés", "Quatro pés, ponteiras/sapatas de borracha encaixadas na estrutura.")],
        "sections": [("Estrutura", "Suportes do encosto em tubo redondo de aço carbono, seção 7/8. Travessas do assento em dois tubos redondos de aço carbono, seção 7/8. Quatro pés em dois tubos redondos de aço carbono, seção 7/8. O laudo registra reforço de 25 cm do chão, em trecho dianteiro e traseiro, com redação a conferir no desenho técnico."), ("Estofamento", "Espuma de poliuretano de alta densidade, descrita como moldada/injetada. Densidade e espessura não foram informadas no PDF analisado. Bordas arredondadas com perfil de PVC."), ("Campos não informados", "Código, capacidade, garantia, embalagem, norma, laudos/ensaios, densidade/espessura da espuma, número de parafusos, declaração de estrutura monobloco e cor do PVC não constam do laudo.")],
        "source": "Laudo completo do modelo Fixa Estofada Prefeitura de BH; catálogo-base de duas páginas.",
    },
    {
        "file": "catalogo-cadeira-fixa-universitaria-executiva-braco-escamoteavel.docx",
        "title": "Cadeira Fixa Universitária Executiva",
        "subtitle": "Com braço escamoteável e suporte para prancheta",
        "manufacturer": "Não identificado no laudo técnico",
        "category": "Cadeiras estofadas",
        "type": "Cadeira fixa universitária executiva com braço escamoteável",
        "summary": "Cadeira fixa universitária executiva com assento e encosto estofados, estrutura em aço carbono, suporte para livros e braço escamoteável com suporte para prancheta. O catálogo-base é referência aprovada/travada.",
        "dimensions": [("Assento", "aprox. 43,5 cm profundidade x 46 cm largura", "Profundidade x largura"), ("Encosto", "aprox. 37 cm altura x 42 cm largura", "Altura x largura")],
        "highlights": [("Assento e encosto", "Compensado multilaminado mínimo 12 mm, moldagem anatômica a quente, espuma de PU de alta densidade, densidade mínima 55 kg/m³ e perfil de PVC."), ("Estrutura", "Tubo redondo de aço carbono 7/8, estrutura curvada e unida por solda MIG."), ("Braços", "Par de braços fixos em aço/PU integral skin com mecanismo escamoteável, lado direito ou esquerdo."), ("Prancheta", "Suporte de prancheta incluído na configuração descrita."), ("Acabamento", "Tratamento antiferrugem, desengraxe, estabilização, fosforização, pintura a pó eletrostática e estufa a 250 °C.")],
        "sections": [("Assento e encosto", "Assento com espuma de PU de alta densidade, densidade mínima 55 kg/m³. Encosto com a mesma base construtiva e revestimento em couro sintético preto. O laudo registra moldagem anatômica, bordas arredondadas e porcas-garras."), ("Acessórios", "Suporte para livros. Ponteiras de PP/PVC flexível e acabamentos deslizantes. Braço escamoteável com suporte e parafusação da prancheta."), ("Confirmações registradas", "O laudo registra confirmações de retirada do braço quando exigido, espuma de nylon injetada, PVC conforme tipo/cor, estrutura monobloco, quatro parafusos sextavados e pintura epóxi preto-fosco, vinculadas à configuração ofertável."), ("Campos não informados", "Altura total, código, garantia, embalagem, dimensões globais completas e norma/laudos anexos não constam como documentos revisados no laudo.")],
        "source": "Laudo final completo da cadeira fixa universitária BH; catálogo-base de duas páginas.",
    },
    {
        "file": "catalogo-cadeira-ergonomica-tela-bracos-pinhais.docx",
        "title": "Cadeira Ergonômica Tela",
        "subtitle": "Giratória com braços, pistão classe 4 e Back System",
        "manufacturer": "ALPHERFLEX INDÚSTRIA E COMÉRCIO DE MÓVEIS E FERRAGENS LTDA",
        "category": "Cadeiras com encosto em tela e assento estofado",
        "type": "Cadeira de escritório ergonômica giratória com braços",
        "summary": "Cadeira de escritório ergonômica giratória com encosto em Tela Mesh, assento estofado, braços reguláveis, base de cinco rodízios, pistão classe 4 e mecanismo Back System.",
        "dimensions": [("Assento", "aprox. 48 cm profundidade x 49 cm largura x 8 cm espessura", "Profundidade x largura x espessura"), ("Encosto", "aprox. 54 cm altura x 45 cm largura x 8 cm espessura", "Altura x largura x espessura"), ("Altura do assento", "aprox. 455 a 565 mm do piso", "Faixa de regulagem"), ("Altura do encosto", "aprox. 980 a 1090 mm do piso", "Faixa indicada no catálogo"), ("Largura dos braços", "aprox. 650 mm total", "Largura total indicada")],
        "highlights": [("Encosto", "Tubo de aço revestido em Tela Mesh."), ("Assento", "Compensado multilaminado mínimo 12 mm, espuma de PU moldada/injetada, densidade mínima 50 kg/m³, espessura mínima 55 mm e perfil de PVC."), ("Base", "Base giratória de aço carbono, seção 25 x 25 mm, espessura 1,2 mm, capa de PP e solda MIG."), ("Movimento", "Pistão a gás classe 4, regulagem pneumática, rotação 360° e capa telescópica de PP."), ("Braços", "Braço digitador com seis regulagens de altura."), ("Conformidade declarada", "Carga de 140 kg; NR 17 e ABNT NBR 13962/2018 declaradas.")],
        "sections": [("Rodízios", "Cinco rodízios de nylon com esferas de aço."), ("Back System", "Regulagem de inclinação e altura do encosto e altura do assento; reclinação aproximada de -5° a 20°."), ("Revestimentos", "O encosto é Tela Mesh. As informações complementares mencionam couro ecológico/courvin preto; a aplicação de cada revestimento deve ser considerada conforme a descrição do componente."), ("Campos não informados", "Código, garantia, embalagem, peso, dimensões totais completas, tolerâncias próprias, laudos anexos e apoio lombar regulável não constam expressamente do laudo.")],
        "source": "Laudo técnico final da cadeira ergonômica em tela de São José dos Pinhais; catálogo-base de três páginas.",
    },
    {
        "file": "catalogo-cadeira-ergonomica-com-bracos-creci.docx",
        "title": "Cadeira Ergonômica com Braços",
        "subtitle": "Modelo executivo com encosto regulável",
        "manufacturer": "Não identificado no laudo técnico",
        "category": "Cadeiras estofadas",
        "type": "Cadeira de escritório giratória executiva/ergonômica com braços",
        "summary": "Cadeira de escritório giratória executiva/ergonômica com braços, encosto regulável, apoio lombar integrado, base de cinco rodízios, pistão a gás e mecanismo Back System.",
        "dimensions": [("Assento", "aprox. 46 cm profundidade x 48 cm largura", "Profundidade x largura"), ("Encosto", "aprox. 47 cm altura x 44 cm largura", "Altura x largura"), ("Ajuste do encosto", "sete posições com ajuste de 75 mm", "Curso indicado")],
        "highlights": [("Assento", "Compensado multilaminado mínimo 15 mm, moldado anatomicamente a quente, espuma de PU d55 e 50 mm, carenagem texturizada de PP injetado e bordas arredondadas."), ("Encosto", "Encosto regulável, apoio lombar integrado, espuma de PU d55 e 45 mm, concha interna e carenagem texturizada de PP."), ("Base", "Base estrela reforçada em tubo de aço carbono 25 x 25 mm, capa de PP injetado e solda MIG."), ("Movimento", "Pistão a gás com regulagem pneumática, rotação 360° e capa telescópica de PP."), ("Braços", "Braço digitador com botão de acionamento e regulagem de altura tipo T."), ("Conformidade declarada", "Carga mínima de 135 kg; NR 17 e ABNT NBR 13962/2018 declaradas.")],
        "sections": [("Rodízios", "Cinco rodízios de duplo giro, diâmetro 60 mm, injetados e descritos como PU/antirrisco."), ("Mecanismo Back System", "Regulagem de inclinação, altura do encosto e altura do assento; reclinação aproximada de -5° a 20°."), ("Revestimento", "Couro ecológico/courvin azul escuro."), ("Campos não informados", "Código, garantia, embalagem, tolerâncias, laudos anexos e dimensões globais completas não constam no PDF analisado. O número de posições dos braços não é especificado no PDF.")],
        "source": "Laudo técnico final da cadeira ergonômica CRECI; catálogo-base de três páginas.",
    },
    {
        "file": "catalogo-tela-ergonomica-relax-braco.docx",
        "title": "Tela Ergonômica Relax",
        "subtitle": "Cadeira giratória com braço regulável",
        "manufacturer": "GOLDFLEX INDÚSTRIA E COMÉRCIO DE MÓVEIS E EQUIPAMENTOS LTDA",
        "category": "Cadeiras com encosto em tela e assento estofado",
        "type": "Cadeira de escritório giratória secretária executiva",
        "summary": "Cadeira giratória com encosto em Tela Mesh, assento estofado, braços reguláveis, base giratória, mecanismo Relax e pistão a gás.",
        "dimensions": [("Assento", "aprox. 48 cm profundidade x 48 cm largura", "Profundidade x largura"), ("Encosto", "aprox. 56 cm altura x 48 cm largura", "Altura x largura"), ("Base", "seção 25 x 25 mm", "Tubo da base"), ("Rodízios", "diâmetro declarado 65 mm", "Diâmetro")],
        "highlights": [("Assento", "Madeira compensada multilaminada de 12 mm, moldada anatomicamente a quente, bordas arredondadas, porcas-garras, espuma de PU de alta densidade declarada moldada/injetada, densidade 50 kg/m³ e perfil de PVC."), ("Encosto", "Estrutura reforçada em tubo de aço revestida em Tela Mesh."), ("Base", "Base estrela em tubo de aço carbono quadrado/retangular 25 x 25 mm, capa de polipropileno injetado e solda MIG."), ("Mecanismo", "Monobloco tipo Relax, travamento na posição de trabalho ou livre flutuação, ajuste manual da tensão da mola, regulagem do encosto e inclinação sincronizada assento/encosto 2:1."), ("Braços", "Braço digitador com botão de acionamento e sete regulagens de altura."), ("Conformidade declarada", "Carga de 120 kg; NR 17 e ABNT NBR 13962/2018 declaradas; revestimento em couro ecológico/courvin preto.")],
        "sections": [("Rodízios e pistão", "Cinco rodízios com esferas de aço. Pistão a gás com regulagem pneumática da altura, rotação 360° e capa telescópica piramidal de PP injetado."), ("Tratamento da base", "Antiferrugem, desengraxe, estabilização, fosforização, pintura a pó eletrostática e secagem em estufa a 250 °C."), ("Campos não informados", "Dimensões globais, altura/largura/profundidade dos braços, altura total, profundidade do encosto, espessura do assento estofado, garantia, embalagem, material do apoia-braço e laudos anexos não constam do PDF analisado.")],
        "source": "Laudo completo da cadeira de escritório giratória com tela do Material Bélico; catálogo-base de três páginas.",
    },
]

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)

def set_cell(cell, text, bold=False, color=INK, size=8.1):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.04
    r = p.add_run(text); r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for side, value in (("top", 80), ("start", 115), ("bottom", 80), ("end", 115)):
        n = OxmlElement(f"w:{side}"); n.set(qn("w:w"), str(value)); n.set(qn("w:type"), "dxa"); mar.append(n)
    tc_pr.append(mar)

def add_table(doc, rows, widths, header=False):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"; t.autofit = False
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            c = t.cell(ri, ci); c.width = Cm(widths[ci]); set_cell(c, value, bold=(header and ri == 0) or (not header and ci == 0), color=RGBColor(255,255,255) if header and ri == 0 else (NAVY if not header and ci == 0 else INK), size=7.8 if header else 8.0)
            if header and ri == 0: shade(c, HEADER)
            elif not header and ci == 0: shade(c, LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def para(doc, text, size=8.7, color=INK, bold=False, align=None, before=0, after=3):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.07
    if align is not None: p.alignment = align
    r = p.add_run(text); r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold

def heading(doc, text):
    para(doc, text, 12.2, BLUE, True, before=5, after=3)

def build(model):
    temp = Path.home() / "AppData" / "Local" / "Temp" / ("template-" + model["file"])
    shutil.copy2(TEMPLATE, temp)
    doc = Document(temp)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"): body.remove(child)
    sec = doc.sections[0]; sec.page_width = Cm(21); sec.page_height = Cm(29.7); sec.top_margin = Cm(1.55); sec.bottom_margin = Cm(1.25); sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    para(doc, model["title"], 21, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para(doc, model["subtitle"], 15, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, model["type"], 10.2, MUTED, False, WD_ALIGN_PARAGRAPH.CENTER, after=7)
    add_table(doc, [["MARCA / FABRICANTE", "IDENTIFICAÇÃO DO MODELO"], [model["manufacturer"], model["title"] + " — " + model["subtitle"]]], [7.2, 10.8], True)
    heading(doc, "Apresentação"); para(doc, model["summary"])
    heading(doc, "Classificação do produto")
    add_table(doc, [["Categoria", model["category"]], ["Tipo construtivo", model["type"]], ["Configuração", "Modelo de referência registrado no laudo técnico"]], [4.2, 13.8])
    heading(doc, "Dimensões de referência")
    rows = [["Componente", "Medida aproximada", "Ordem da medida"]] + [list(x) for x in model["dimensions"]]
    add_table(doc, rows, [3.7, 8.3, 6.0], True)
    para(doc, "As dimensões acima reproduzem os valores aproximados registrados no laudo do modelo. Dimensões não listadas não foram acrescentadas.", 7.2, MUTED, after=2)
    heading(doc, "Destaques técnicos"); add_table(doc, [[a, b] for a, b in model["highlights"]], [4.2, 13.8])
    para(doc, "Fonte: " + model["source"], 7.0, MUTED, after=0)
    doc.add_page_break()
    heading(doc, "Registro técnico do modelo")
    for name, text in model["sections"]:
        heading(doc, name); para(doc, text)
    heading(doc, "Materiais, componentes e informações registradas")
    add_table(doc, [["Grupo", "Conteúdo do laudo"]] + [[a, b] for a, b in model["highlights"]], [4.2, 13.8], True)
    heading(doc, "Limites da evidência")
    para(doc, "Os campos abaixo permanecem como não informados ou não comprovados no laudo analisado: " + " ".join(text for name, text in model["sections"] if name == "Campos não informados"), 7.5, MUTED, after=0)
    doc.save(OUTDIR / model["file"])

for model in MODELS:
    build(model)
    print(OUTDIR / model["file"])
