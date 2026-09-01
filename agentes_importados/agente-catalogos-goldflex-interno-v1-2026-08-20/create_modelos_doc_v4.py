from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\modelos-de-cadeiras-analisados-v4.docx")
NAVY = RGBColor(31, 78, 121); BLUE = RGBColor(47, 117, 181); INK = RGBColor(31, 41, 51); MUTED = RGBColor(102, 112, 133)
rows = [
    ("Fixa", "Cadeiras em polipropileno injetado", "Braço/prancheta fixa", "Cadeira Universitária Fixa em Polipropileno com Prancheta", "aprox. 42 cm profundidade x 46 cm largura", "aprox. 28,0 cm altura x 46 cm largura", "UFRGS", "Medidas aproximadas evidenciadas"),
    ("Fixa", "Cadeiras estofadas", "Braço/prancheta fixa", "Cadeira Universitária Fixa Estofada com Prancheta", "aprox. 39 cm profundidade x 42 cm largura x 5 cm espessura", "aprox. 29 cm altura x 36 cm largura x 4 cm espessura", "UFRGS", "Medidas aproximadas evidenciadas"),
    ("Fixa", "Cadeiras estofadas", "Sem braços", "Modelo Fixa Estofada", "aprox. 42 cm profundidade x 43 cm largura", "aprox. 27 cm altura x 39 cm largura", "Prefeitura BH", "Laudo registra ausência de alguns campos"),
    ("Fixa", "Cadeiras estofadas", "Com braços", "Cadeira Fixa Universitária Executiva Com Braço Escamoteável", "aprox. 43,5 cm profundidade x 46 cm largura", "aprox. 37 cm altura x 42 cm largura", "Prefeitura BH universitária", "Braço escamoteável"),
    ("Giratória", "Cadeiras com encosto em tela e assento estofado", "Com braços", "Cadeira Ergonômica Tela Com Braços", "aprox. 48 cm profundidade x 49 cm largura x 8 cm espessura", "aprox. 54 cm altura x 45 cm largura x 8 cm espessura", "São José dos Pinhais", "Encosto em tela"),
    ("Giratória", "Cadeiras estofadas", "Com braços", "Cadeira Ergonômica Com Braços", "aprox. 46 cm profundidade x 48 cm largura", "aprox. 47 cm altura x 44 cm largura", "CRECI", "Encosto regulável"),
    ("Giratória", "Cadeiras com encosto em tela e assento estofado", "Com braços", "Modelo Tela Ergonômica Relax com Braço", "aprox. 48 cm profundidade x 48 cm largura", "aprox. 56 cm altura x 48 cm largura", "Material Bélico Tela", "Mecanismo Relax"),
    ("Rebatível", "Cadeiras de auditório estofadas", "Com braços / prancheta escamoteável", "Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável", "aprox. 48 cm profundidade x 49 cm largura", "aprox. 48 cm altura x 46 cm largura", "Almirante Alexandrino", "Modelo individual; longarina não evidenciada"),
    ("Fixa", "Cadeiras estofadas", "Não evidenciado", "Item 6 — Cadeira Secretaria Fixa", "aprox. 39 cm profundidade x 42 cm largura", "aprox. 29 cm altura x 37 cm largura", "Patrocínio Paulista", "Braços não descritos no catálogo"),
    ("Giratória", "Cadeiras estofadas", "Com braços", "Item 8 — Cadeira Executiva Lâmina Ergonômica", "aprox. 42 cm profundidade x 46 cm largura", "aprox. 37 cm altura x 40 cm largura", "Patrocínio Paulista", "Altura do assento aprox. 45–56 cm"),
    ("Fixa", "Cadeiras estofadas", "Não evidenciado", "Item 10 — Cadeira Fixa", "aprox. 39 cm profundidade x 42 cm largura", "aprox. 29 cm altura x 37 cm largura", "Patrocínio Paulista", "Braços não descritos no catálogo"),
    ("Giratória", "Bancos/cadeiras semi-sentados em PU", "Não evidenciado", "Item 2 — Banco Semi Sentado com Rodízios", "largura 35 cm x profundidade 28,5 cm", "altura 13 cm x largura 22 cm", "Material Bélico Injetado", "Altura catalogada com ordem inconsistente"),
    ("Giratória", "Cadeiras em poliuretano integral skin", "Não evidenciado", "Item 3 — Cadeira Injetada Industrial", "aprox. 41,5 cm profundidade x 43 cm largura x 4 cm espessura", "aprox. 25 cm altura x 40,5 cm largura x 4 cm espessura", "Material Bélico Injetado", "Altura catalogada aprox. 51–62 cm"),
    ("Fixa", "Cadeiras em polipropileno injetado", "Não evidenciado", "Item 1 — Cadeira Fixa Empilhável", "aprox. 43 cm profundidade x 45 cm largura", "aprox. 30 cm altura x 45 cm largura", "Assistência Jurídica", "Concha dupla e ausência de braços não explicitadas"),
]

def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcpr.append(shd)

def set_cell(cell, text, bold=False, color=INK, size=8):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, data, widths, header=True):
    t = doc.add_table(rows=len(data), cols=len(data[0])); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"; t.autofit = False
    for ri, row in enumerate(data):
        for ci, value in enumerate(row):
            c = t.cell(ri, ci); c.width = Cm(widths[ci]); set_cell(c, value, bold=(header and ri == 0), color=RGBColor(255,255,255) if header and ri == 0 else INK, size=7.2 if header else 7.5)
            if header and ri == 0: shade(c, "1F4E78")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

doc = Document(); sec = doc.sections[0]; sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5); sec.left_margin = Cm(1.4); sec.right_margin = Cm(1.4)
normal = doc.styles["Normal"]; normal.font.name = "Aptos"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); normal.font.size = Pt(9); normal.font.color.rgb = INK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3); r = p.add_run("Modelos de cadeiras analisados — registro consolidado"); r.font.name = "Aptos Display"; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = NAVY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(8); r = p.add_run("Versão atualizada com todos os modelos analisados até o momento"); r.font.size = Pt(10); r.font.color.rgb = MUTED
doc.add_paragraph("Registro de referência para consulta dos modelos disponíveis. As medidas são reproduzidas como aproximadas quando assim aparecem nos laudos e catálogos; campos não comprovados permanecem identificados como não evidenciados.")
add_table(doc, [["Resumo", "Quantidade"], ["Registros consolidados", "14"], ["Modelos fixos", "6"], ["Modelos giratórios", "6"], ["Modelo rebatível de auditório", "1"], ["Registro semi-sentado giratório", "1"]], [8, 9.5], False)
doc.add_page_break()
doc.add_heading("Registro completo", level=1)
add_table(doc, [["Tipo", "Categoria", "Braços", "Modelo", "Assento", "Encosto", "Fonte", "Observação"]] + [list(r) for r in rows], [1.8, 3.0, 3.0, 4.6, 3.5, 3.5, 2.5, 3.7], True)
doc.add_paragraph("Nota: este registro não substitui os laudos e relatórios individuais. Ele consolida nomes, classificação, braços, medidas e origem documental já analisados no projeto.")
doc.save(OUT)
print(OUT)
