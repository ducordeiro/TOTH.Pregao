from pathlib import Path
import shutil
from docx import Document

BASE = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\dados extraidos versao 2\catalogo-goldflex-catalogo-cadeiras-prefeitura-bh-v2.docx")
OUT = BASE.with_name("catalogo-goldflex-catalogo-cadeiras-prefeitura-bh-v3.docx")

shutil.copy2(BASE, OUT)
doc = Document(OUT)
doc.add_heading("Configuração confirmada pelo fabricante", level=1)
doc.add_paragraph(
    "As informações abaixo complementam o laudo e identificam a configuração ofertável. "
    "O PDF original permanece como catálogo-base travado e não foi alterado."
)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Configuração confirmada"
hdr[2].text = "Tratamento editorial"
rows = [
    ("Espuma", "Espuma de nylon injetada", "Informação confirmada pelo fabricante; manter na ficha técnica da configuração."),
    ("PVC", "Tipo e cor conforme configuração solicitada", "Registrar o tipo e a cor efetivamente ofertados; não presumir outras cores."),
    ("Estrutura", "Estrutura monobloco", "Característica confirmada; demonstrar na ficha, desenho ou amostra."),
    ("Fixação", "Quatro parafusos sextavados", "Característica confirmada; conferir na ficha, desenho ou amostra."),
    ("Pintura", "Pintura epóxi preto-fosco", "Acabamento confirmado para a configuração ofertada."),
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
doc.add_heading("Nota de aprovação e evidência", level=1)
doc.add_paragraph(
    "As medidas e características do catálogo-base foram preservadas. As confirmações do fabricante "
    "foram incorporadas como dados da configuração ofertável, sem alterar o PDF de referência. "
    "Laudos, ensaios, ficha técnica, desenho, proposta e amostra devem ser apresentados quando exigidos "
    "no edital ou necessários para comprovar a configuração ofertada."
)
doc.save(OUT)
print(OUT)
