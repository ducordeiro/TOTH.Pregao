from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos\catalogos de pedidos")
SOURCE = ROOT / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-catalogo-cadeiras-prefeitura-bh-v3.docx"
OUT = ROOT / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v3.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-mocho-estancia.docx"
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)

body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def title(text, size=20):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
    return p

def heading(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

title("MOCHO ERGONÔMICO BIPARTIDO")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Modelo Mocho Sela Bipartido | Marca Goldflex")
r.bold = True
r.font.size = Pt(12)

heading("Apresentação")
doc.add_paragraph("Mocho ergonômico bipartido, giratório, com regulagem de altura e inclinação do assento, desenvolvido para uso profissional. A configuração abaixo foi extraída do catálogo técnico destinado à Prefeitura Municipal de Estância Velha.")

heading("Identificação e classificação")
t = doc.add_table(rows=1, cols=2)
t.style = "Table Grid"
for i, h in enumerate(["Campo", "Informação"]):
    t.rows[0].cells[i].text = h
    shade(t.rows[0].cells[i])
for a, b in [
    ("Modelo", "Mocho Sela Bipartido"),
    ("Categoria construtiva", "Mocho/semi-sentado estofado"),
    ("Movimento", "Giratória 360°"),
    ("Marca", "Goldflex"),
    ("Fabricante", "Goldflex Industria e Comercio de Moveis e Equipamentos LTDA"),
]:
    cells = t.add_row().cells
    cells[0].text, cells[1].text = a, b

heading("Assento bipartido")
for x in [
    "Confeccionado em madeira compensada multilaminada.",
    "Moldado anatomicamente a quente, com bordas arredondadas.",
    "Fixação através de porcas-garras.",
    "Tapeado em espuma moldada/injetada de poliuretano de alta densidade, 45 kg/m³, isenta de CFC.",
]: bullet(x)

heading("Regulagens e desempenho funcional registrado")
for x in [
    "Regulagem pneumática da altura do assento por alavanca lateral.",
    "Altura do assento: mínima aproximada de 58 cm e máxima aproximada de 78 cm.",
    "Regulagem da inclinação do assento entre 1° e 16°, por alavanca traseira esquerda.",
    "Peso suportado declarado: 120 kg.",
]: bullet(x)

heading("Pistão, mecanismo e base")
for x in [
    "Pistão a gás classe 3, com curso/componente declarado de 200 mm.",
    "Mecanismo confeccionado em chapa de aço SAE 1006/1008 com 3 mm, fosfatizado e pintado em tinta pó epóxi.",
    "Camada de tinta declarada de aproximadamente 80 μm.",
    "Base cromada com cinco pás de apoio para fixação dos rodízios e anel central em tubo de aço carbono.",
    "Pás conformadas por estampagem e travadas por soldagem MIG.",
    "Blindagem central em polipropileno, montada por cliques de fixação, para proteção e acabamento.",
]: bullet(x)

heading("Rodízios e acabamentos")
for x in [
    "Cinco rodízios de PU de 50 mm, com esferas de aço.",
    "Capa telescópica piramidal de proteção em polipropileno injetado.",
    "Itens de acabamento em termoplástico injetado: polipropileno no acabamento frontal, montante e alavanca do pistão; poliamida 6 na alavanca de inclinação.",
    "A base recebe preparação de superfície metálica em nanocerâmica e cromagem por deposição eletrolítica.",
    "A cor pode ser definida conforme catálogo de cores do fabricante; a cor específica desta configuração não está indicada no documento analisado.",
]: bullet(x)

heading("Medidas e informações técnicas")
t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
for i, h in enumerate(["Componente", "Medida/valor", "Status"]):
    t2.rows[0].cells[i].text = h
    shade(t2.rows[0].cells[i])
for row in [
    ("Altura do assento", "58 cm a 78 cm", "Evidenciado no catálogo-base"),
    ("Pistão", "Classe 3; 200 mm declarado", "Evidenciado no catálogo-base"),
    ("Rodízios", "PU 50 mm", "Evidenciado no catálogo-base"),
    ("Mecanismo", "Chapa de aço 3 mm; inclinação 1° a 16°", "Evidenciado no catálogo-base"),
    ("Largura, profundidade e altura total", "Não evidenciadas", "Não estimar"),
    ("Dimensões do assento e encosto", "Não evidenciadas", "Confirmar em ficha/desenho"),
]:
    cells = t2.add_row().cells
    for i, value in enumerate(row): cells[i].text = value

heading("Documentação e observações")
doc.add_paragraph("O catálogo-base menciona a ABNT NBR 13962/2018 em relação aos rodízios. A menção não foi tratada como certificado ou laudo independente revisado. Normas, ensaios, garantia, embalagem, montagem e demais documentos devem ser confirmados conforme a configuração ofertada e o processo de aquisição.")
doc.add_paragraph("Este catálogo foi reorganizado a partir das informações evidenciadas no laudo técnico consolidado do modelo Estância Velha. Não foram acrescentadas especificações ausentes e o PDF de referência não foi alterado.")

doc.save(OUT)
print(OUT)
