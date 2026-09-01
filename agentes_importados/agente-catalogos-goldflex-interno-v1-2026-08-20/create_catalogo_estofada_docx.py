from pathlib import Path
import shutil
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\marcs\OneDrive\Documents\ChatGPT\relatoriocatalogogoldflex")
template = ROOT / r"Modelos de catalogos\catalogos de pedidos\dados extraidos\catalogo-cadeira-universitaria-pp-prancheta.docx"
out = ROOT / r"Modelos de catalogos\catalogos de pedidos\dados extraidos\catalogo-cadeira-universitaria-estofada-prancheta.docx"
temp = Path.home() / "AppData" / "Local" / "Temp" / "catalogo-template-estofada.docx"
shutil.copy2(template, temp)
doc = Document(temp)

NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(47, 117, 181)
MUTED = RGBColor(102, 112, 133)
INK = RGBColor(31, 41, 51)
LIGHT = "EAF2F8"
HEADER = "17365D"
GRID = "B7C9DB"

body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.55)
section.bottom_margin = Cm(1.25)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, value in (("top", 85), ("start", 120), ("bottom", 85), ("end", 120)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa"); mar.append(node)
    tc_pr.append(mar)

def set_text(cell, text, bold=False, color=INK, size=8.2):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; cell_margins(cell)

def table(rows, widths, header=False):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"; t.autofit = False
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            c = t.cell(ri, ci); c.width = Cm(widths[ci]); set_text(c, value, bold=(header and ri == 0) or (not header and ci == 0), color=RGBColor(255,255,255) if header and ri == 0 else (NAVY if not header and ci == 0 else INK), size=7.9 if header else 8.1)
            if header and ri == 0: shade(c, HEADER)
            elif not header and ci == 0: shade(c, LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def para(text, size=8.8, color=INK, bold=False, align=None, before=0, after=3):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.08
    if align is not None: p.alignment = align
    r = p.add_run(text); r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.color.rgb = color; r.font.bold = bold
    return p

def heading(text):
    return para(text, size=12.5, color=BLUE, bold=True, before=5, after=3)

# Page 1
para("Cadeira Universitária Fixa", 22, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, after=1)
para("Estofada com Prancheta Fixa", 16, NAVY, True, WD_ALIGN_PARAGRAPH.CENTER, after=2)
para("Modelo fixo universitário | Configuração com braço/prancheta fixa | Opção destra ou canhota", 10.5, MUTED, False, WD_ALIGN_PARAGRAPH.CENTER, after=8)
table([["MARCA / FABRICANTE", "IDENTIFICAÇÃO DO MODELO"], ["Marca comercial não identificada no PDF-fonte; fabricante: ALPHERFLEX", "Cadeira Universitária Fixa Estofada com Prancheta"]], [7.2, 10.8], True)
heading("Apresentação")
para("Cadeira universitária fixa com assento e encosto estofados, estrutura metálica soldada, porta-livros lateral e prancheta fixa. Esta ficha foi organizada exclusivamente a partir do catálogo analisado e do laudo técnico do modelo, preservando as medidas aproximadas e as alternativas que ainda dependem de confirmação de fabricação.")
para("Fabricante identificado no catálogo analisado", 7.3, MUTED, after=1)
para("ALPHERFLEX INDÚSTRIA E COMÉRCIO DE MÓVEIS E FERRAGENS LTDA", 8.8, INK, True, after=2)
table([["CNPJ", "40.919.354/0001-59"], ["Endereço", "Avenida Brasil, 2076, Santa Cruz, Mogi Mirim - SP, CEP 13.800-444"], ["Contato", "(19) 97126-6023 | cadeirasalphaflex@gmail.com"]], [2.8, 15.2])
heading("Classificação do produto")
table([["Categoria", "Cadeiras estofadas"], ["Tipo construtivo", "Cadeira fixa universitária"], ["Apoio lateral", "Braço/prancheta fixa"], ["Assento e encosto", "Compensado multilaminado estofado com espuma de poliuretano"], ["Configuração", "Prancheta para destro ou canhoto, conforme declaração do catálogo"]], [4.2, 13.8])
heading("Dimensões de referência")
table([["Componente", "Medida aproximada", "Ordem da medida"], ["Assento", "39 x 42 x 5 cm", "Profundidade x largura x espessura"], ["Encosto", "29 x 36 x 4 cm", "Altura x largura x espessura"], ["Prancheta", "49 x 23 x 1,8 cm", "Profundidade x largura x espessura"]], [3.7, 8.3, 6.0], True)
para("As medidas são aproximadas e reproduzem a referência registrada na análise. Altura total, altura do assento, largura total e profundidade total da cadeira não foram evidenciadas no catálogo.", 7.3, MUTED, after=2)
heading("Destaques técnicos")
table([["Estrutura", "Tubos de aço carbono redondos, seção 7/8, espessura 1,2 mm; quatro pés"], ["Estofamento", "Compensado multilaminado mínimo 12 mm e espuma de poliuretano moldada/injetada, densidade mínima 50 kg/m³"], ["Revestimento", "Tecido 100% poliéster"], ["Conformidade declarada", "Carga suportada de 130 kg e atendimento à NBR 16671/2018"]], [4.2, 13.8])
para("Referência de evidência: catálogo Cadeira Fixa Estofada com Prancheta, p. 1-2; análise técnica e laudo final dos modelos UFRGS.", 7.1, MUTED, after=0)

doc.add_page_break()
# Page 2
heading("Especificação construtiva")
table([["Elemento", "Descrição técnica evidenciada"], ["Suportes do encosto", "Tubo de aço carbono redondo, seção 7/8, espessura 1,2 mm; suportes curvados em máquinas específicas e unidos por solda MIG."], ["Travessas do assento", "Dois tubos de aço carbono redondos, seção 7/8, espessura 1,2 mm."], ["Pés", "Quatro pés em dois tubos de aço carbono redondos, seção 7/8, espessura 1,2 mm; acabamentos deslizantes."], ["Porta-livros", "Fixado nas laterais, fabricado em tubo, firme e sem terminações pontiagudas."], ["Braço/prancheta", "Estrutura em tubos de aço carbono redondos, seção 7/8, espessura 1,2 mm; prancheta fixa."]], [4.4, 13.6], True)
heading("Assento, encosto e revestimento")
para("Assento e encosto confeccionados em madeira compensada multilaminada com espessura mínima de 12 mm, tapeçados com espuma moldada/injetada de poliuretano de alta densidade, densidade mínima de 50 kg/m³. As bordas são arredondadas e recebem perfil de PVC. O revestimento é declarado como tecido 100% poliéster.")
table([["Componente", "Dimensão aproximada", "Materiais e acabamento"], ["Assento", "39 x 42 x 5 cm", "Compensado multilaminado, espuma PU e perfil de PVC"], ["Encosto", "29 x 36 x 4 cm", "Compensado multilaminado, espuma PU e perfil de PVC"]], [3.7, 5.0, 9.3], True)
heading("Tratamento e acabamento metálico")
table([["Etapa", "Informação registrada"], ["Preparação", "Desengraxe, estabilização e fosforização."], ["Proteção", "Pré-tratamento antiferrugem."], ["Pintura", "Pintura a pó pelo processo de deposição eletrostática."], ["Cura", "Secagem em estufa a 250 °C."], ["Cores", "Não informadas no catálogo analisado; confirmar na configuração ofertada."]], [4.4, 13.6], True)
heading("Prancheta de madeira")
para("Prancheta fabricada em MDP/MDF entre 18 mm, revestida em ambas as faces com laminado melamínico e acabamento em fita PVC. Medidas aproximadas: profundidade 49 cm x largura 23 cm x espessura 1,8 cm. A posição pode ser configurada para destro ou canhoto, conforme declaração do catálogo.")
heading("Informações complementares e controle")
table([["Campo", "Situação registrada"], ["Carga suportada", "130 kg, declarada no catálogo analisado."], ["Norma", "Atendimento à NBR 16671/2018, declarado no catálogo analisado."], ["Espuma", "A fonte usa a redação moldada/injetada; não escolher uma única alternativa sem confirmação da configuração."], ["Código comercial", "Não informado; atribuir somente em ficha interna/proposta."], ["Cores, garantia e embalagem", "Não informadas; confirmar na ficha técnica e na proposta."], ["Dimensões globais e tolerâncias", "Não informadas; medir e registrar na configuração ofertada."]], [5.2, 12.8], True)
heading("Orientação para proposta")
para("Usar este documento como referência do modelo aprovado. Para cada edital, apresentar separadamente a configuração ofertada, lado da prancheta, medidas efetivamente produzidas, tipo de espuma escolhido, tecido/cor, materiais, desenho, amostra e o pacote de laudos e ensaios solicitado. As medidas do catálogo-base não devem ser alteradas para eliminar diferenças; variações devem ser registradas como configuração de fabricação e confirmadas antes do fornecimento.")
para("Nota de evidência: este documento não acrescenta código, cores, dimensões globais, garantia, embalagem ou outras especificações ausentes da fonte. O histórico de aprovação é premissa fornecida pela empresa e não equivale a certificados revisados nesta pasta.", 7.4, MUTED, after=0)

doc.save(out)
print(out)
