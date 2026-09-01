from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\modelos-de-cadeiras-analisados-v6-formato-v3.docx")
groups = {
    "Cadeiras fixas": {
        "Cadeiras em polipropileno injetado": [
            ("Braço/prancheta fixa", "Cadeira Universitária Fixa em Polipropileno com Prancheta", "Assento: aprox. 42 cm profundidade x 46 cm largura", "Encosto: aprox. 28,0 cm altura x 46 cm largura", "UFRGS"),
            ("Não evidenciado", "Item 1 — Cadeira Fixa Empilhável", "Assento: aprox. 43 cm profundidade x 45 cm largura", "Encosto: aprox. 30 cm altura x 45 cm largura", "Assistência Jurídica"),
        ],
        "Cadeiras estofadas": [
            ("Braço/prancheta fixa", "Cadeira Universitária Fixa Estofada com Prancheta", "Assento: aprox. 39 x 42 x 5 cm", "Encosto: aprox. 29 x 36 x 4 cm", "UFRGS"),
            ("Sem braços", "Modelo Fixa Estofada", "Assento: aprox. 42 cm profundidade x 43 cm largura", "Encosto: aprox. 27 cm altura x 39 cm largura", "Prefeitura BH"),
            ("Com braços", "Cadeira Fixa Universitária Executiva Com Braço Escamoteável", "Assento: aprox. 43,5 cm profundidade x 46 cm largura", "Encosto: aprox. 37 cm altura x 42 cm largura", "Prefeitura BH universitária"),
            ("Não evidenciado", "Item 6 — Cadeira Secretaria Fixa", "Assento: aprox. 39 cm profundidade x 42 cm largura", "Encosto: aprox. 29 cm altura x 37 cm largura", "Patrocínio Paulista"),
            ("Não evidenciado", "Item 10 — Cadeira Fixa", "Assento: aprox. 39 cm profundidade x 42 cm largura", "Encosto: aprox. 29 cm altura x 37 cm largura", "Patrocínio Paulista"),
        ],
    },
    "Cadeiras giratórias": {
        "Cadeiras com encosto em tela e assento estofado": [
            ("Com braços", "Cadeira Ergonômica Tela Com Braços", "Assento: aprox. 48 x 49 x 8 cm", "Encosto: aprox. 54 x 45 x 8 cm", "São José dos Pinhais"),
            ("Com braços", "Modelo Tela Ergonômica Relax com Braço", "Assento: aprox. 48 cm profundidade x 48 cm largura", "Encosto: aprox. 56 cm altura x 48 cm largura", "Material Bélico Tela"),
        ],
        "Cadeiras estofadas": [
            ("Com braços", "Cadeira Ergonômica Com Braços", "Assento: aprox. 46 cm profundidade x 48 cm largura", "Encosto: aprox. 47 cm altura x 44 cm largura", "CRECI"),
            ("Com braços", "Item 8 — Cadeira Executiva Lâmina Ergonômica", "Assento: aprox. 42 cm profundidade x 46 cm largura", "Encosto: aprox. 37 cm altura x 40 cm largura", "Patrocínio Paulista"),
        ],
        "Cadeiras em poliuretano integral skin": [
            ("Não evidenciado", "Item 3 — Cadeira Injetada Industrial", "Assento: aprox. 41,5 cm profundidade x 43 cm largura x 4 cm espessura", "Encosto: aprox. 25 cm altura x 40,5 cm largura x 4 cm espessura", "Material Bélico Injetado"),
        ],
        "Bancos/cadeiras semi-sentados em PU": [
            ("Não evidenciado", "Item 2 — Banco Semi Sentado com Rodízios", "Assento: largura 35 cm x profundidade 28,5 cm", "Encosto: altura 13 cm x largura 22 cm", "Material Bélico Injetado"),
        ],
    },
    "Cadeiras rebatíveis de auditório": {
        "Cadeiras de auditório estofadas": [
            ("Com braços / prancheta escamoteável", "Cadeira de Auditório Rebatível com Braço e Prancheta Escamoteável", "Assento: aprox. 48 cm profundidade x 49 cm largura", "Encosto: aprox. 48 cm altura x 46 cm largura", "Almirante Alexandrino"),
        ],
    },
}

doc = Document()
section = doc.sections[0]
section.top_margin = Pt(54); section.bottom_margin = Pt(54); section.left_margin = Pt(60); section.right_margin = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Modelos de cadeiras analisados"); r.bold = True; r.font.size = Pt(16)
doc.add_paragraph("Classificação por tipo de movimento e categoria construtiva oficial.")
for movement, categories in groups.items():
    h = doc.add_paragraph(); rr = h.add_run(movement); rr.bold = True; rr.font.size = Pt(13)
    for category, models in categories.items():
        h = doc.add_paragraph(); rr = h.add_run(category); rr.bold = True; rr.font.size = Pt(11)
        for arms, model, seat, back, source in models:
            doc.add_paragraph(f"{arms} — {model}", style="List Bullet")
            doc.add_paragraph(f"    {seat}; {back}")
            doc.add_paragraph(f"    Fonte/registro: {source}")
doc.save(OUT)
print(OUT)
