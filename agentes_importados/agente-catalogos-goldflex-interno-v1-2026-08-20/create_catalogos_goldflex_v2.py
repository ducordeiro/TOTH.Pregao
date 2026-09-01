from pathlib import Path
import re
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos\catalogos de pedidos")
OUTDIR = ROOT / "dados extraidos" / "dados extraidos versao 2"
TEMPLATE = ROOT / "dados extraidos" / "dados extraidos versao 1" / "catalogo-cadeira-universitaria-pp-prancheta.docx"
OUTDIR.mkdir(parents=True, exist_ok=True)

def section_text(raw, heading):
    start = raw.find(heading)
    if start < 0: return ""
    tail = raw[start + len(heading):]
    nxt = tail.find("\n## ")
    return tail[:nxt if nxt >= 0 else len(tail)].strip()

def bullets(text, limit=18):
    out=[]
    for line in text.splitlines():
        line=line.strip()
        if line.startswith("-"):
            value=line[1:].strip()
            if value and value not in out: out.append(value)
    return out[:limit]

def model_title(folder, raw, pdf_lines):
    for line in pdf_lines:
        clean=line.strip()
        if re.search(r"(?i)MODELO\s*:", clean):
            value=re.split(r"(?i)MODELO\s*:", clean, maxsplit=1)[1].split("|",1)[0].strip()
            if value: return value
    for line in pdf_lines:
        clean=line.strip()
        if re.search(r"(?i)(CADEIRA|POLTRONA|MOCHO|BANCO)", clean) and len(clean) < 130:
            return clean
    return folder.name

def source_pages(folder):
    names=[]
    for f in sorted(folder.glob("*.pdf")):
        names.append(f"{f.name}")
    for f in sorted(folder.glob("*.txt")):
        names.append(f"{f.name}")
    return names

def shade(cell, fill):
    tcpr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcpr.append(shd)

def cell_text(cell, text, bold=False, color=RGBColor(31,41,51), size=8.1):
    cell.text=""; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    r=p.add_run(text); r.font.name="Aptos"; r._element.rPr.rFonts.set(qn("w:ascii"),"Aptos"); r._element.rPr.rFonts.set(qn("w:hAnsi"),"Aptos"); r.font.size=Pt(size); r.font.bold=bold; r.font.color.rgb=color

def add_table(doc, rows):
    t=doc.add_table(rows=len(rows), cols=2); t.style="Table Grid"; t.autofit=False
    for ri,row in enumerate(rows):
        for ci,value in enumerate(row):
            c=t.cell(ri,ci); c.width=Cm(4.4 if ci==0 else 13.6); cell_text(c,value,bold=(ci==0),color=RGBColor(255,255,255) if ri==0 else RGBColor(31,41,51))
            if ri==0: shade(c,"1F4E78")
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_heading(doc, text, size=12.5):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(3); r=p.add_run(text); r.font.name="Aptos Display"; r.font.size=Pt(size); r.font.bold=True; r.font.color.rgb=RGBColor(47,117,181)

def add_para(doc,text,size=8.8,color=RGBColor(31,41,51),bold=False):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); p.paragraph_format.line_spacing=1.07; r=p.add_run(text); r.font.name="Aptos"; r.font.size=Pt(size); r.font.color.rgb=color; r.font.bold=bold

def build(folder):
    laudo=folder/"laudo-tecnico-consolidado.md"
    if not laudo.exists(): return False
    raw=laudo.read_text(encoding="utf-8",errors="replace")
    pdf_lines=[]
    first_pdf=next(iter(sorted(folder.glob("*.pdf"))),None)
    if first_pdf:
        try:
            from pypdf import PdfReader
            for page in PdfReader(str(first_pdf)).pages[:2]: pdf_lines += (page.extract_text() or "").splitlines()
        except Exception: pass
    title=model_title(folder,raw,pdf_lines)
    docs=bullets(section_text(raw,"## 2. Documentos avaliados"),10)
    cls=bullets(section_text(raw,"## 3. Classificação inicial"),6)
    measures=bullets(section_text(raw,"## 4. Medidas e dimensões encontradas"),18)
    components=bullets(section_text(raw,"## 5. Materiais, componentes e acabamentos encontrados"),22)
    norms=bullets(section_text(raw,"## 6. Capacidade, normas, garantia e documentos citados"),14)
    gaps=bullets(section_text(raw,"## 7. Lacunas de evidência"),10)
    out=OUTDIR/("catalogo-goldflex-"+re.sub(r"[^A-Za-z0-9]+","-",folder.name).strip("-").lower()+"-v2.docx")
    temp=Path.home()/"AppData"/"Local"/"Temp"/("template-"+out.stem+".docx"); shutil.copy2(TEMPLATE,temp); doc=Document(temp)
    body=doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"): body.remove(child)
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.55); sec.bottom_margin=Cm(1.25); sec.left_margin=Cm(1.5); sec.right_margin=Cm(1.5)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title); r.font.name="Aptos Display"; r.font.size=Pt(19); r.font.bold=True; r.font.color.rgb=RGBColor(31,78,121)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Catálogo técnico Goldflex — versão de análise e aprovação"); r.font.size=Pt(10); r.font.color.rgb=RGBColor(102,112,133)
    add_heading(doc,"Identificação e classificação")
    add_table(doc,[["Campo","Informação registrada"],["Pasta de origem",folder.name],["Modelo",title],["Classificação","; ".join(cls) if cls else "Não classificada automaticamente"],["Documentos-base","; ".join(docs) if docs else "; ".join(source_pages(folder))]])
    add_heading(doc,"Medidas e dimensões")
    if measures:
        for x in measures: add_para(doc,"• "+x)
    else: add_para(doc,"Nenhuma medida foi localizada na extração textual.",8.2,RGBColor(102,112,133))
    add_heading(doc,"Materiais, componentes e acabamentos")
    if components:
        for x in components: add_para(doc,"• "+x)
    else: add_para(doc,"Nenhuma descrição técnica foi localizada na extração textual.",8.2,RGBColor(102,112,133))
    doc.add_page_break()
    add_heading(doc,"Capacidade, normas e documentos citados")
    if norms:
        for x in norms: add_para(doc,"• "+x)
    else: add_para(doc,"Não localizados na extração textual.",8.2,RGBColor(102,112,133))
    add_heading(doc,"Campos não evidenciados ou a confirmar")
    if gaps:
        for x in gaps: add_para(doc,"• "+x)
    else: add_para(doc,"Consultar o laudo consolidado e a ficha técnica da configuração ofertada.",8.2,RGBColor(102,112,133))
    add_heading(doc,"Nota de evidência")
    add_para(doc,"Este catálogo foi montado a partir do laudo técnico consolidado da pasta e dos documentos nele rastreados. Não foram acrescentadas especificações ausentes; a presença de uma declaração no catálogo-fonte não foi convertida automaticamente em certificado ou ensaio revisado.",8.2,RGBColor(102,112,133))
    doc.save(out); return True

created=[]
for folder in sorted(ROOT.iterdir()):
    if folder.is_dir() and folder.name.lower() != "dados extraidos" and build(folder): created.append(folder.name)
print("created",len(created))
for name in created: print(name)
