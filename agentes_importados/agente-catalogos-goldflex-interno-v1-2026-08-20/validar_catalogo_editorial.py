from pathlib import Path
from docx import Document

ROOT = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\dados extraidos versao 2")
REQUIRED = ("medid", "materia", "fonte")

files = sorted(ROOT.glob("catalogo-goldflex-*-v2.docx"))
errors = []
for path in files:
    try:
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        parts.extend(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
        text = "\n".join(parts).lower()
        if len(text.strip()) < 300:
            errors.append(f"{path.name}: conteúdo insuficiente")
        missing = [word for word in REQUIRED if word not in text]
        if missing:
            errors.append(f"{path.name}: termos ausentes: {', '.join(missing)}")
        if not doc.tables:
            errors.append(f"{path.name}: sem tabela estruturada")
    except Exception as exc:
        errors.append(f"{path.name}: {exc}")

print(f"arquivos verificados: {len(files)}")
if errors:
    print("falhas:")
    print("\n".join(errors))
    raise SystemExit(1)
print("resultado: aprovado na validação estrutural")
