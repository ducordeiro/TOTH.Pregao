from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
out = r"Modelos de catalogos\catalogos de pedidos\dados extraidos\modelos-de-cadeiras-analisados-v3.docx"
groups = {
    "Cadeiras fixas": {
        "Cadeiras em polipropileno injetado": [("Braço/prancheta fixa", "Cadeira Universitária Fixa em Polipropileno com Prancheta", "Assento: aprox. 42 cm profundidade x 46 cm largura", "Encosto: aprox. 28,0 cm altura x 46 cm largura")],
        "Cadeiras estofadas": [("Braço/prancheta fixa", "Cadeira Universitária Fixa Estofada com Prancheta", "Assento: aprox. 39 x 42 x 5 cm", "Encosto: aprox. 29 x 36 x 4 cm"), ("Sem braços", "Modelo Fixa Estofada", "Assento: aprox. 42 cm profundidade x 43 cm largura", "Encosto: aprox. 27 cm altura x 39 cm largura"), ("Com braços", "Cadeira Fixa Universitária Executiva Com Braço Escamoteavel", "Assento: aprox. 43,5 cm profundidade x 46 cm largura", "Encosto: aprox. 37 cm altura x 42 cm largura")],
    },
    "Cadeiras giratórias": {
        "Cadeiras com encosto em tela e assento estofado": [("Com braços", "Cadeira Ergonômica Tela Com braços", "Assento: aprox. 48 x 49 x 8 cm", "Encosto: aprox. 54 x 45 x 8 cm"), ("Com braços", "Modelo Tela Ergonômica Relax com Braço", "Assento: aprox. 48 cm profundidade x 48 cm largura", "Encosto: aprox. 56 cm altura x 48 cm largura")],
        "Cadeiras estofadas": [("Com braços", "Cadeira Ergonômica Com Braços", "Assento: aprox. 46 cm profundidade x 48 cm largura", "Encosto: aprox. 47 cm altura x 44 cm largura")],
    },
}
doc = Document(); s = doc.sections[0]; s.top_margin = Pt(54); s.bottom_margin = Pt(54); s.left_margin = Pt(60); s.right_margin = Pt(60)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Modelos de cadeiras analisados"); r.bold = True; r.font.size = Pt(16)
doc.add_paragraph("Classificação por tipo de movimento e categoria construtiva oficial.")
for movement, categories in groups.items():
    h = doc.add_paragraph(); rr = h.add_run(movement); rr.bold = True; rr.font.size = Pt(13)
    for category, models in categories.items():
        h = doc.add_paragraph(); rr = h.add_run(category); rr.bold = True; rr.font.size = Pt(11)
        for arms, model, seat, back in models:
            doc.add_paragraph(f"{arms} — {model}", style="List Bullet")
            doc.add_paragraph(f"    {seat}; {back}")
doc.save(out)
