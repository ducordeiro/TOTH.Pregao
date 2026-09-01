from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos")
SOURCE = ROOT / "catalogos de pedidos" / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v4.docx"
OUTDIR = ROOT / "catalogos de pedidos" / "Catalogos geral"
OUT = OUTDIR / "catalogo-goldflex-poltrona-auditorio-modular-individual-v4.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-poltrona-auditorio-v4.docx"
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def heading(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("POLTRONA DE AUDITÓRIO MODULAR INDIVIDUAL"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Prancheta escamoteável — configuração para destro"); r.bold = True; r.font.size = Pt(12)

heading("Identificação")
t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
for i, h in enumerate(["Campo", "Informação"]): t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i])
for a, b in [("Modelo", "Poltrona de Auditório Modular Individual"), ("Categoria", "Poltrona de auditório individual modular"), ("Configuração", "Destro, com prancheta escamoteável embutida no braço"), ("Fabricante", "GOLDFLEX INDUSTRIA E COMERCIO DE MOVEIS E EQUIPAMENTOS LTDA CNPJ 33.661.439/0001-14")]:
    cells = t.add_row().cells; cells[0].text, cells[1].text = a, b

heading("Características")
for x in ["Poltrona individual modular para auditório.", "Conchas em madeira laminada, com assento e encosto estofados.", "Assento rebatível com sistema antipanico.", "Braços em madeira maciça.", "Prancheta escamoteável embutida no braço, para uso destro.", "Estrutura metálica em aço SAE 1010/1020.", "Base em aço tubular SAE 1008/1020.", "Acabamento da estrutura em pintura eletrostática com tinta em pó, na cor preta.", "Assento e encosto revestidos em tecido 100% poliéster texturizado, na cor azul-marinho."]:
    bullet(x)

heading("Assento e encosto")
for x in ["Assento confeccionado em concha de madeira laminada.", "Encosto confeccionado em concha de madeira laminada.", "Estofamento em espuma de poliuretano injetado.", "Densidade da espuma: entre 50 e 60 kg/m³.", "Espessura mínima da espuma do assento: 60 mm.", "Espessura mínima da espuma do encosto: 40 mm.", "Assento rebatível com sistema antipanico."]:
    bullet(x)

heading("Medidas e dimensões")
t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
for i, h in enumerate(["Componente", "Medida mínima"]): t2.rows[0].cells[i].text = h; shade(t2.rows[0].cells[i])
for a, b in [("Assento", "460 mm de largura x 430 mm de profundidade"), ("Encosto", "455 mm de largura x 575 mm de altura"), ("Prancheta", "270 mm x 260 mm; espessura de 15 mm"), ("Espuma do assento", "60 mm"), ("Espuma do encosto", "40 mm")]:
    cells = t2.add_row().cells; cells[0].text, cells[1].text = a, b

heading("Estrutura, base e prancheta")
for x in ["Estrutura confeccionada em aço SAE 1010/1020.", "Base confeccionada em aço tubular SAE 1008/1020.", "Pintura eletrostática com tinta em pó na cor preta.", "Prancheta escamoteável embutida no braço da poltrona.", "Prancheta em madeira MDF, com 15 mm de espessura.", "Suporte e componentes metálicos da prancheta confeccionados em aço SAE 1010/1020.", "Configuração da prancheta para destro."]:
    bullet(x)

heading("Materiais e revestimentos")
for x in ["Conchas em madeira laminada.", "Braços em madeira maciça.", "Espuma de poliuretano injetado, densidade de 50 a 60 kg/m³.", "Revestimento em tecido 100% poliéster texturizado.", "Cor do revestimento: azul-marinho.", "Cor da estrutura: preta."]:
    bullet(x)

heading("Capacidade, garantia e normas")
for x in ["Garantia mínima: 12 meses.", "Fabricação conforme normas NBR/ABNT vigentes aplicáveis ao produto.", "Atendimento à NR-17."]:
    bullet(x)

heading("Observações")
doc.add_paragraph("Este catálogo técnico apresenta as características do modelo Poltrona de Auditório Modular Individual, com prancheta escamoteável para destro. As informações foram organizadas a partir do descritivo técnico fornecido para o modelo.")
doc.save(OUT)
print(OUT)
