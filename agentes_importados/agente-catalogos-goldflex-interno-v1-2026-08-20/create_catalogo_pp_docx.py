from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path(r"Modelos de catalogos\catalogos de pedidos\dados extraidos\catalogo-cadeira-universitaria-pp-prancheta.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)
NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(47, 117, 181)
MUTED = RGBColor(102, 112, 133)
INK = RGBColor(31, 41, 51)
LIGHT = "EAF2F8"
HEADER = "17365D"
GRID = "B7C9DB"

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_cell_text(cell, text, bold=False, color=INK, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)

def set_table_borders(table, color=GRID, size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)

def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            row.cells[i].width = Cm(width)
            tc_pr = row.cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 567)))
            tc_w.set(qn("w:type"), "dxa")

def add_table(doc, rows, widths_cm, header=False):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths_cm)
    set_table_borders(table)
    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            set_cell_text(table.cell(ri, ci), value, bold=(header and ri == 0) or (not header and ci == 0), color=RGBColor(255,255,255) if header and ri == 0 else (NAVY if not header and ci == 0 else INK), size=8.1 if header else 8.3)
            if header and ri == 0:
                shade(table.cell(ri, ci), HEADER)
            elif not header and ci == 0:
                shade(table.cell(ri, ci), LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

def add_para(doc, text, size=8.8, color=INK, bold=False, align=None, before=0, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    return p

def heading(doc, text):
    return add_para(doc, text, size=12.5, color=BLUE, bold=True, before=5, after=3)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(8.3)
    r.font.color.rgb = INK

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(1.55)
section.bottom_margin = Cm(1.25)
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.header_distance = Cm(0.7)
section.footer_distance = Cm(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(8.8)
normal.font.color.rgb = INK

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = header.add_run("CATALOGO TECNICO | MODELO DE REFERENCIA")
hr.font.name = "Aptos Display"; hr.font.size = Pt(8); hr.font.bold = True; hr.font.color.rgb = NAVY
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = footer.add_run("Cadeira Universitaria Fixa em Polipropileno com Prancheta  |  ")
fr.font.name = "Aptos"; fr.font.size = Pt(7.5); fr.font.color.rgb = MUTED
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
footer._p.append(fld)

add_para(doc, "Cadeira Universitaria Fixa", size=22, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add_para(doc, "em Polipropileno Injetado com Prancheta Fixa", size=16, color=NAVY, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add_para(doc, "Modelo fixo para uso universitario | Configuracao com braco/prancheta fixa", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
add_table(doc, [["MARCA / FABRICANTE", "IDENTIFICACAO DO MODELO"], ["Marca comercial nao identificada no PDF-fonte; fabricante: ALPHERFLEX", "Cadeira Universitaria Fixa em Polipropileno com Prancheta"]], [7.2, 10.8], header=True)
heading(doc, "Apresentacao")
add_para(doc, "Cadeira universitaria fixa composta por assento e encosto em polipropileno (PP) de alta resistencia, estrutura metalica soldada e prancheta fixa. Esta ficha foi organizada a partir do catalogo analisado e do laudo tecnico do modelo, preservando as medidas de referencia e separando os campos que nao foram evidenciados.")
add_para(doc, "Fabricante identificado no catalogo analisado", size=7.3, color=MUTED, after=1)
add_para(doc, "ALPHERFLEX INDUSTRIA E COMERCIO DE MOVEIS E FERRAGENS LTDA", size=8.8, bold=True, after=2)
add_table(doc, [["CNPJ", "40.919.354/0001-59"], ["Endereco", "Avenida Brasil, 2076, Santa Cruz, Mogi Mirim - SP, CEP 13.800-444"], ["Contato", "(19) 97126-6023 | cadeirasalphaflex@gmail.com"]], [2.8, 15.2])
heading(doc, "Classificacao do produto")
add_table(doc, [["Categoria", "Cadeiras em polipropileno injetado"], ["Tipo construtivo", "Cadeira fixa universitaria"], ["Apoio lateral", "Braco/prancheta fixa"], ["Assento e encosto", "Polipropileno (PP) de alta resistencia"], ["Uso documental", "Modelo de referencia aprovado/travado; medidas do catalogo nao devem ser alteradas"]], [4.2, 13.8])
heading(doc, "Dimensoes de referencia")
add_table(doc, [["Componente", "Medida aproximada", "Ordem da medida"], ["Assento", "42 cm profundidade x 46 cm largura", "Profundidade x largura"], ["Encosto", "28,0 cm altura x 46 cm largura", "Altura x largura"], ["Prancheta", "49 cm profundidade x 23 cm largura x 1,8 cm espessura", "Profundidade x largura x espessura"]], [3.7, 8.3, 6.0], header=True)
add_para(doc, "As medidas acima sao aproximadas e reproduzem a referencia registrada na analise. Altura total, altura do assento e largura total nao foram evidenciadas no catalogo analisado.", size=7.3, color=MUTED, after=2)
heading(doc, "Destaques tecnicos")
add_table(doc, [["Material estrutural", "Aco carbono com tratamento antiferrugem e pintura a po eletrostatica"], ["Montagem", "Componentes unidos por solda MIG"], ["Prancheta", "MDP/MDF entre 18 mm, laminado melaminico nas duas faces e acabamento em fita PVC"], ["Apoio de livros", "Porta-livros mencionado na referencia; detalhes dimensionais e material nao informados"]], [4.2, 13.8])
add_para(doc, "Referencia de evidencia: catalogo UFRGS de cadeira fixa em polipropileno com prancheta, p. 1; laudo-final-modelos-ufrgs.md, modelo 1.", size=7.1, color=MUTED, after=0)

doc.add_page_break()
heading(doc, "Especificacao construtiva")
add_table(doc, [["Elemento", "Descricao tecnica evidenciada"], ["Suportes do encosto", "Dois tubos oblongos de aco carbono, secao 16 x 30 mm, espessura 1,5 mm."], ["Suportes do assento", "Dois tubos redondos, secao 3/4, espessura 1,2 mm."], ["Pes", "Dois tubos oblongos, secao 16 x 30 mm, espessura 1,2 mm."], ["Uniao", "Solda MIG, conforme descricao do catalogo analisado."], ["Ponteiras/deslizantes", "Acabamentos deslizantes nos pes; tipo e composicao detalhados nao informados."]], [4.4, 13.6], header=True)
heading(doc, "Tratamento e acabamento")
add_table(doc, [["Etapa", "Informacao registrada"], ["Preparacao", "Desengraxe, estabilizacao e fosforizacao."], ["Protecao", "Tratamento antiferrugem."], ["Pintura", "Pintura a po eletrostatica."], ["Cura", "Secagem em estufa a 250 graus C."], ["Cores", "Nao informadas no catalogo analisado; confirmar na configuracao ofertada."]], [4.4, 13.6], header=True)
heading(doc, "Prancheta e ergonomia de uso")
add_para(doc, "A prancheta fixa e descrita como MDP/MDF entre 18 mm, revestida com laminado melaminico nas duas faces e com fita PVC. A medida de referencia registrada e aproximadamente 49 x 23 x 1,8 cm. O catalogo menciona porta-livros, mas nao detalha suas dimensoes, material, fixacao ou capacidade.")
heading(doc, "Controle documental e campos a confirmar")
add_table(doc, [["Campo", "Situacao no material analisado"], ["Codigo comercial", "Nao informado; atribuir somente em ficha interna/proposta."], ["Altura total, altura do assento e largura total", "Nao evidenciadas; medir e registrar na configuracao ofertada."], ["Capacidade de carga e norma", "Nao informadas neste catalogo; anexar laudos/ensaios quando exigidos no edital."], ["Cores, garantia e embalagem", "Nao informadas; confirmar comercialmente e na ficha tecnica."], ["Canhoto/destro", "Nao informado para este modelo PP; confirmar possibilidade de configuracao."]], [6.2, 11.8], header=True)
heading(doc, "Orientacao para proposta")
add_para(doc, "Usar este documento como referencia de modelo e seguranca comercial. Para cada edital, apresentar separadamente a configuracao ofertada, medidas efetivamente produzidas, materiais, acabamentos, desenhos, amostra e o pacote de laudos/ensaios solicitado. Medidas maiores ou variacoes de tecido/acabamento devem ser tratadas como configuracao a confirmar, sem modificar este catalogo-base.")
heading(doc, "Nota de evidencia")
add_para(doc, "Este documento nao acrescenta cores, capacidade, normas, garantia, dimensoes globais ou outras especificacoes ausentes da fonte. Informacoes de aprovacao historica sao premissas fornecidas pela empresa e nao equivalem a certificados revisados nesta pasta.", size=7.4, color=MUTED, after=0)

doc.save(OUT)
print(OUT.resolve())
