from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"Modelos de catalogos")
SOURCE = ROOT / "catalogos de pedidos" / "dados extraidos" / "dados extraidos versao 3" / "catalogo-goldflex-mocho-ergonomico-bipartido-estancia-velha-v4.docx"
OUTDIR = ROOT / "validação de dados"
OUT = OUTDIR / "catalogo-goldflex-cadeira-executiva-lamina-ergonomica-v4.docx"
TMP = Path.home() / "AppData" / "Local" / "Temp" / "modelo-cadeira8-generico-v4.docx"
OUTDIR.mkdir(parents=True, exist_ok=True)
shutil.copy2(SOURCE, TMP)
doc = Document(TMP)
body = doc._element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)

def shade(cell, fill="FFC000"):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill); tcPr.append(shd)

def heading(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(text)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CADEIRA EXECUTIVA LÂMINA ERGONÔMICA"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0xC0, 0x90, 0x00)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Cadeira giratória estofada"); r.bold = True; r.font.size = Pt(12)

heading("Identificação")
t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
for i, h in enumerate(["Campo", "Informação"]): t.rows[0].cells[i].text = h; shade(t.rows[0].cells[i])
for a, b in [("Modelo", "Cadeira Executiva Lâmina Ergonômica"), ("Categoria", "Cadeira giratória estofada"), ("Movimento", "Giratória 360°"), ("Fabricante", "GOLDFLEX INDUSTRIA E COMERCIO DE MOVEIS E EQUIPAMENTOS LTDA CNPJ 33.661.439/0001-14")]:
    cells = t.add_row().cells; cells[0].text, cells[1].text = a, b

heading("Características")
for x in ["Assento e encosto em madeira compensada multilaminada, moldados anatomicamente a quente.", "Espessura declarada da madeira: aproximadamente 13 mm.", "Bordas arredondadas e fixação por porcas-garras.", "Estofamento em espuma moldada/injetada de poliuretano, densidade declarada de 50 kg/m³.", "Bordas encabeçadas com perfil de PVC.", "Modelo disponível com ou sem braços, conforme a composição construtiva escolhida.", "Base giratória em aço carbono, seção 25 x 25 mm, espessura declarada de 1,2 mm, unida por solda MIG.", "Cinco rodízios de nylon com esferas de aço.", "Pistão a gás com regulagem pneumática da altura do assento e rotação de 360°.", "Mecanismo tipo relax, com travamento em posição de trabalho ou livre flutuação, ajuste de tensão e regulagem do encosto."]: bullet(x)

heading("Medidas e dimensões")
t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
for i, h in enumerate(["Componente", "Medida registrada"]): t2.rows[0].cells[i].text = h; shade(t2.rows[0].cells[i])
for a, b in [("Encosto", "Aproximadamente 370 mm de altura x 400 mm de largura"), ("Assento", "Aproximadamente 420 mm de profundidade x 460 mm de largura"), ("Altura do assento", "Aproximadamente 450 mm a 560 mm"), ("Base", "Seção aproximada de 25 x 25 mm; espessura declarada de 1,2 mm"), ("Rodízios", "5 unidades, com esferas de aço")]:
    cells = t2.add_row().cells; cells[0].text, cells[1].text = a, b

heading("Estrutura, base e regulagens")
for x in ["Estrutura da base em tubo de aço carbono, com união por solda MIG.", "Base estrela submetida a pré-tratamento antiferrugem, desengraxe, estabilização e fosforização.", "Pintura a pó por deposição eletrostática e secagem em estufa a 250 °C, conforme descrição do catálogo analisado.", "Regulagem pneumática da altura do assento por pistão a gás.", "Mecanismo relax com travamento na posição de trabalho ou livre flutuação.", "Regulagem do encosto e ajuste de tensão da mola por manípulo frontal."]: bullet(x)

heading("Materiais e acabamentos")
for x in ["Madeira compensada multilaminada no assento e encosto.", "Espuma de poliuretano moldada/injetada, densidade declarada de 50 kg/m³.", "Perfil de PVC nas bordas.", "Rodízios de nylon com esferas de aço.", "Revestimento e cores definidos conforme a composição escolhida para o modelo."]: bullet(x)

heading("Capacidade e normas citadas")
for x in ["Peso suportado declarado no catálogo analisado: 120 kg.", "Atendimento à NR-17 e à ABNT NBR 13962/2018 declarado no catálogo analisado.", "Laudos, ensaios e documentos complementares devem ser associados ao modelo quando aplicáveis."]: bullet(x)

heading("Observações")
doc.add_paragraph("Este catálogo técnico reúne as características estudadas nos laudos e no catálogo do modelo Cadeira Executiva Lâmina Ergonômica. Variações de braços, medidas, revestimentos, cores e acabamentos pertencem à composição do modelo e devem ser registradas na ficha técnica correspondente, sem alterar este catálogo-base.")
doc.save(OUT)
print(OUT)
