import re
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_PDF = Path(
    r"C:\Users\ducor\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm"
    r"\LocalState\sessions\6A86223F22A57B5C2750F0831822A5157DFFE0E7"
    r"\transfers\2026-28\19TR+-+Mobiliarios++e+eletrodomesticos.pdf"
)
OUT_DOCX = Path("Tabela_TR_Mobiliarios_Eletrodomesticos.docx")


def compact(value):
    if value is None:
        return ""
    value = str(value).replace("\r", "\n")
    lines = [" ".join(line.split()) for line in value.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def one_line(value):
    return re.sub(r"\s+", " ", value or "").strip()


def extract_rows():
    rows = []
    current_section = ""
    current_item = None

    with pdfplumber.open(str(SOURCE_PDF)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            for table in page.extract_tables():
                for raw_row in table:
                    cells = [compact(cell) for cell in raw_row]
                    # The TR table has three useful columns. Some extraction rows have
                    # extra empty cells; keep the first non-empty three-column shape.
                    while len(cells) < 3:
                        cells.append("")
                    item, spec, quant = cells[0], cells[1], cells[2]

                    if item.startswith("ANEXO"):
                        current_section = one_line(item)
                        rows.append({"type": "section", "section": current_section, "page": page_number})
                        current_item = None
                        continue

                    if item.lower() == "item" or spec.lower().startswith("especifica"):
                        continue

                    if re.fullmatch(r"\d+", item) and spec and quant:
                        current_item = {
                            "type": "item",
                            "section": current_section,
                            "item": item,
                            "spec": one_line(spec),
                            "quant": one_line(quant),
                            "page": page_number,
                        }
                        rows.append(current_item)
                        continue

                    if not item and spec and current_item:
                        current_item["spec"] = one_line(current_item["spec"] + " " + spec)

    return rows


def set_font(run, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text, *, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_font(run)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_document(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.1

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(14)
    h1.font.color.rgb = RGBColor.from_string("2E74B5")
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)


def build_docx(rows):
    items = [row for row in rows if row["type"] == "item"]
    if not items:
        raise RuntimeError("informações não encontradas")

    doc = Document()
    setup_document(doc)

    doc.add_heading("Tabela do Termo de Referência - Mobiliários e Eletrodomésticos", level=1)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    run = note.add_run(f"Fonte: {SOURCE_PDF.name}. Tabela recriada a partir do arquivo PDF.")
    set_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(89, 89, 89)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    table_width = 15400
    widths = [900, 13500, 1000]
    set_table_width(table, table_width)

    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, label, width in zip(header.cells, ["Item", "Especificação", "Quant"], widths):
        set_cell_width(cell, width)
        set_cell_margins(cell)
        shade_cell(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        write_cell(cell, label, bold=True, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in rows:
        if row["type"] == "section":
            cells = table.add_row().cells
            merged = cells[0].merge(cells[1]).merge(cells[2])
            set_cell_margins(merged)
            shade_cell(merged, "F2F4F7")
            write_cell(merged, row["section"], bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color="1F4D78")
            continue

        cells = table.add_row().cells
        values = [row["item"], row["spec"], row["quant"]]
        for idx, (cell, value, width) in enumerate(zip(cells, values, widths)):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cell, value, size=7.2, align=align)

    doc.save(OUT_DOCX)


def main():
    rows = extract_rows()
    build_docx(rows)
    items = [row for row in rows if row["type"] == "item"]
    sections = [row["section"] for row in rows if row["type"] == "section"]
    print(OUT_DOCX.resolve())
    print(f"items={len(items)}")
    print(f"sections={sections}")
    print(f"first={items[0]['item']} {items[0]['quant']} {items[0]['spec'][:80]}")
    print(f"last={items[-1]['item']} {items[-1]['quant']} {items[-1]['spec'][:80]}")


if __name__ == "__main__":
    main()
