import json
import threading
import unittest
import time
import tempfile
import zipfile
from datetime import datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch
from urllib.parse import parse_qs, urlparse

from docx import Document
from PIL import Image

import catalog
import catalog_generator
import server


def make_item(item, description, lote="", quantidade="1"):
    row = server.empty_item()
    row.update({
        "lote": lote,
        "item": item,
        "quantidade": quantidade,
        "descricao": description,
    })
    return row


class RequestBoundaryTests(unittest.TestCase):
    def test_public_urls_only_allow_http_without_credentials_or_whitespace(self):
        self.assertEqual(
            server.safe_public_url("https://pncp.gov.br/app/editais/1"),
            "https://pncp.gov.br/app/editais/1",
        )
        self.assertEqual(server.safe_public_url("javascript:alert(1)"), "")
        self.assertEqual(server.safe_public_url("https://user:secret@example.com/file"), "")
        self.assertEqual(server.safe_public_url("https://example.com/a\r\nHeader: value"), "")

    def test_request_host_and_origin_only_accept_local_addresses(self):
        for value in ("127.0.0.1:8765", "localhost:8765", "[::1]:8765", ""):
            self.assertTrue(server.local_request_host(value))
        self.assertFalse(server.local_request_host("attacker.example"))
        self.assertTrue(server.local_request_origin("http://127.0.0.1:8765"))
        self.assertTrue(server.local_request_origin("http://localhost:5173"))
        self.assertFalse(server.local_request_origin("https://attacker.example"))

    def test_json_body_rejects_oversized_and_wrong_content_type(self):
        oversized = SimpleNamespace(
            headers={"Content-Length": "11", "Content-Type": "application/json"},
            rfile=BytesIO(b"{}"),
        )
        with self.assertRaises(OverflowError):
            server.parse_json_body(oversized, maximum_size=10)

        wrong_type = SimpleNamespace(
            headers={"Content-Length": "2", "Content-Type": "text/plain"},
            rfile=BytesIO(b"{}"),
        )
        with self.assertRaisesRegex(ValueError, "application/json"):
            server.parse_json_body(wrong_type)

    def test_internal_search_maps_explicit_missing_date_option(self):
        captured = {}

        class Repository:
            def list_opportunities(self, filters):
                captured.update(filters)
                return {"items": [], "total": 0, "limit": 10, "offset": 0}

        with patch.object(server, "etl_repository", return_value=Repository()):
            server.internal_opportunities_response({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "incluirSemDataEncerramento": "1",
            })

        self.assertTrue(captured["include_missing_proposal_dates"])

    def test_internal_search_maps_selected_date_field(self):
        captured = []

        class Repository:
            def list_opportunities(self, filters):
                captured.append(filters)
                return {"items": [], "total": 0, "limit": 10, "offset": 0}

        with patch.object(server, "etl_repository", return_value=Repository()):
            server.internal_opportunities_response({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "campoData": "publicacao",
            })
            server.internal_opportunities_response({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "campoData": "abertura",
            })

        self.assertEqual(captured[0]["date_field"], "publication")
        self.assertIn("published_from", captured[0])
        self.assertNotIn("proposal_from", captured[0])
        self.assertEqual(captured[1]["date_field"], "opening")
        self.assertIn("proposal_start_from", captured[1])
        self.assertNotIn("proposal_from", captured[1])

    def test_internal_search_rejects_unknown_date_field(self):
        with self.assertRaisesRegex(ValueError, "Campo de data"):
            server.internal_opportunities_response({"campoData": "assinatura"})


class DatabaseInitializationTests(unittest.TestCase):
    def test_same_database_file_is_initialized_only_once(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "pncp.sqlite3"
            database_key = str(database_path.resolve())
            calls = []

            def initialize_once():
                calls.append(database_path)
                database_path.touch()

            server.INITIALIZED_DATABASES.pop(database_key, None)
            try:
                with (
                    patch.object(server, "DATABASE_PATH", database_path),
                    patch.object(server, "_initialize_database", side_effect=initialize_once),
                ):
                    server.init_database()
                    server.init_database()
            finally:
                server.INITIALIZED_DATABASES.pop(database_key, None)

        self.assertEqual(calls, [database_path])

    def test_replaced_database_file_is_initialized_again(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            database_path = root / "pncp.sqlite3"
            replacement_path = root / "replacement.sqlite3"
            database_key = str(database_path.resolve())
            calls = []

            def initialize_database():
                calls.append(database_path)
                database_path.touch()

            server.INITIALIZED_DATABASES.pop(database_key, None)
            try:
                with (
                    patch.object(server, "DATABASE_PATH", database_path),
                    patch.object(server, "_initialize_database", side_effect=initialize_database),
                ):
                    server.init_database()
                    replacement_path.write_bytes(b"replacement")
                    replacement_path.replace(database_path)
                    server.init_database()
            finally:
                server.INITIALIZED_DATABASES.pop(database_key, None)

        self.assertEqual(calls, [database_path, database_path])


class ItemExtractionRegressionTests(unittest.TestCase):
    def test_office_spreadsheet_inside_package_is_not_expanded_as_nested_zip(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            spreadsheet = root / "planilha.xlsx"
            with zipfile.ZipFile(spreadsheet, "w") as archive:
                for index in range(server.MAX_ARCHIVE_FILES + 1):
                    archive.writestr(f"xl/media/item-{index}.xml", "x")
            package = root / "edital.zip"
            with zipfile.ZipFile(package, "w") as archive:
                archive.write(spreadsheet, spreadsheet.name)
                archive.writestr("Termo de Referencia.pdf", b"%PDF-test")

            documents = server.downloaded_document_candidates(package)

            self.assertEqual(
                [path.name for path, _embedded in documents],
                ["planilha.xlsx", "Termo de Referencia.pdf"],
            )

    def test_proposal_preview_is_temporary_and_reused_for_unchanged_data(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            preview_dir = Path(temp_dir) / "previews"
            template_path = Path(temp_dir) / "modelo.docx"
            template_path.write_bytes(b"template")
            context = {
                "items": [
                    {
                        "item": "1",
                        "quantidade": "2",
                        "unidade": "UND",
                        "descricao": "Cadeira",
                        "marca": "Goldflex",
                        "valor_unitario": "R$ 100,00",
                        "valor_total": "R$ 200,00",
                    }
                ],
                "template_path": template_path,
                "source_name": "edital",
                "responsible_id": "1",
                "responsible": {"id": "1", "nome_completo": "Responsável"},
                "commercial_terms": {},
            }

            def fake_build(_items, _template, output_path, **_kwargs):
                output_path.write_bytes(b"docx temporario")

            def fake_convert(_docx_path, pdf_path):
                pdf_path.write_bytes(b"%PDF-1.4 preview")

            server.PROPOSAL_PREVIEW_CACHE.clear()
            with (
                patch.object(server, "PREVIEW_DIR", preview_dir),
                patch.object(server, "build_docx", side_effect=fake_build) as build_mock,
                patch.object(server, "convert_docx_to_pdf", side_effect=fake_convert),
            ):
                first = server.create_proposal_preview(context)
                second = server.create_proposal_preview(context)

                self.assertFalse(first["cached"])
                self.assertTrue(second["cached"])
                self.assertEqual(first["preview_url"], second["preview_url"])
                self.assertEqual(build_mock.call_count, 1)
                self.assertEqual(list(preview_dir.glob("*.docx")), [])
                self.assertEqual(len(list(preview_dir.glob("*.pdf"))), 1)
            server.PROPOSAL_PREVIEW_CACHE.clear()

    def test_expired_proposal_preview_is_deleted(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            preview_dir = Path(temp_dir)
            preview_path = preview_dir / "expired.pdf"
            preview_path.write_bytes(b"%PDF-1.4")
            server.PROPOSAL_PREVIEW_CACHE.clear()
            server.PROPOSAL_PREVIEW_CACHE["a" * 32] = {
                "token": "a" * 32,
                "fingerprint": "fingerprint",
                "path": preview_path,
                "created_at": 0,
                "last_access": 0,
            }

            with patch.object(server, "PREVIEW_DIR", preview_dir):
                server.cleanup_proposal_previews(server.PROPOSAL_PREVIEW_TTL + 1)

            self.assertFalse(preview_path.exists())
            self.assertEqual(server.PROPOSAL_PREVIEW_CACHE, {})

    def test_proposal_preview_falls_back_when_word_is_unavailable(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            preview_dir = Path(temp_dir) / "previews"
            template_path = Path(temp_dir) / "modelo.docx"
            template_path.write_bytes(b"template")
            context = {
                "items": [{"item": "1", "quantidade": "1", "unidade": "UND",
                           "descricao": "Mesa", "marca": "Teste",
                           "valor_unitario": "R$ 10,00", "valor_total": "R$ 10,00"}],
                "template_path": template_path,
                "source_name": "edital",
                "responsible_id": "1",
                "responsible": {"id": "1", "nome_completo": "Responsável"},
                "commercial_terms": {},
            }

            def fake_build(_items, _template, output_path, **_kwargs):
                output_path.write_bytes(b"docx temporario")

            def fake_fallback(_context, pdf_path):
                pdf_path.write_bytes(b"%PDF-1.4 compatible")

            server.PROPOSAL_PREVIEW_CACHE.clear()
            with (
                patch.object(server, "PREVIEW_DIR", preview_dir),
                patch.object(server, "build_docx", side_effect=fake_build),
                patch.object(server, "convert_docx_to_pdf", side_effect=RuntimeError("Word indisponível")),
                patch.object(server, "build_compatible_proposal_pdf", side_effect=fake_fallback) as fallback,
            ):
                result = server.create_proposal_preview(context)

            self.assertEqual(result["renderer"], "compatible")
            fallback.assert_called_once()
            self.assertEqual(len(list(preview_dir.glob("*.pdf"))), 1)
            self.assertEqual(list(preview_dir.glob("*.docx")), [])
            server.PROPOSAL_PREVIEW_CACHE.clear()

    def test_public_pncp_payload_excludes_internal_document_paths(self):
        pncp = {
            "cnpj": "12345678000199",
            "ano": 2026,
            "sequencial": 10,
            "link": "https://pncp.gov.br/app/editais/12345678000199/2026/10",
            "documento_usado": "Termo de Referencia.pdf",
            "documentos_candidatos": [
                {"path": Path(r"C:\temp\Termo de Referencia.pdf")}
            ],
        }

        public = server.public_pncp_payload(pncp)

        self.assertNotIn("documentos_candidatos", public)
        json.dumps(public)

    def test_json_serializer_handles_windows_paths_and_structured_values(self):
        payload = {
            "path": Path(r"C:\temp\edital.pdf"),
            "generated_at": datetime(2026, 7, 24, 12, 30),
            "unit_value": Decimal("1887.21"),
            "items": {"2", "1"},
        }

        serialized = json.dumps(
            payload,
            default=server.json_compatible_default,
        )
        decoded = json.loads(serialized)

        self.assertEqual(decoded["path"], r"C:\temp\edital.pdf")
        self.assertEqual(decoded["generated_at"], "2026-07-24T12:30:00")
        self.assertEqual(decoded["unit_value"], "1887.21")
        self.assertEqual(decoded["items"], ["1", "2"])

    def test_identification_is_persisted_in_sqlite(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "pncp.sqlite3"
            source_path = Path(temp_dir) / "termo.pdf"
            source_path.write_bytes(b"%PDF-test")
            source_data = {
                "source_path": source_path,
                "pncp": {
                    "cnpj": "12345678000199",
                    "ano": 2026,
                    "sequencial": 10,
                    "link": "https://pncp.gov.br/app/editais/12345678000199/2026/10",
                    "documento_usado": "Termo de Referencia.pdf",
                    "documento_tipo": "Termo de Referencia",
                    "arquivo_usado": {"sequencialDocumento": 2, "titulo": "TR"},
                    "arquivos": [
                        {"sequencialDocumento": 1, "titulo": "Edital"},
                        {"sequencialDocumento": 2, "titulo": "TR"},
                    ],
                },
            }
            identifications = server.build_item_identifications([
                make_item("1", "Cadeira giratoria completa.", quantidade="12"),
                make_item("2", "Armario de aco completo.", lote="1", quantidade="4"),
            ])

            with patch.object(server, "DATABASE_PATH", database_path), patch.object(
                server, "DATA_DIR", Path(temp_dir)
            ):
                server.persist_identification(
                    source_data,
                    identifications,
                    {"status": "ok"},
                    {"file_count": 2, "pncp_count": 2, "has_divergence": False},
                )
                with server.database_connection() as connection:
                    contract = connection.execute("SELECT * FROM contratacoes").fetchone()
                    items = connection.execute(
                        "SELECT * FROM itens ORDER BY numero_item"
                    ).fetchall()
                    files = connection.execute(
                        "SELECT * FROM arquivos_pncp ORDER BY chave_pncp"
                    ).fetchall()

            self.assertEqual(contract["total_itens"], 2)
            self.assertEqual([row["identificacao_simplificada"] for row in items], ["Cadeira", "Armário"])
            self.assertEqual(items[0]["descricao_completa"], "Cadeira giratoria completa.")
            self.assertTrue(all(row["unidade"] == "UND" for row in items))
            self.assertEqual(sum(row["selecionado"] for row in files), 1)

    def test_persisting_same_contract_replaces_stale_items(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            database_path = Path(temp_dir) / "pncp.sqlite3"
            source_path = Path(temp_dir) / "edital.pdf"
            source_path.write_bytes(b"%PDF-test")
            source_data = {
                "source_path": source_path,
                "pncp": {
                    "cnpj": "12345678000199", "ano": 2026, "sequencial": 11,
                    "link": "https://pncp.gov.br/app/editais/12345678000199/2026/11",
                    "documento_usado": "Edital.pdf", "documento_tipo": "Edital",
                    "arquivo_usado": {"sequencialDocumento": 1, "titulo": "Edital"},
                    "arquivos": [{"sequencialDocumento": 1, "titulo": "Edital"}],
                },
            }

            with patch.object(server, "DATABASE_PATH", database_path), patch.object(
                server, "DATA_DIR", Path(temp_dir)
            ):
                server.persist_identification(
                    source_data,
                    server.build_item_identifications([make_item("1", "Cadeira."), make_item("2", "Mesa.")]),
                    {"status": "ok"}, {},
                )
                server.persist_identification(
                    source_data,
                    server.build_item_identifications([make_item("1", "Cadeira atualizada.")]),
                    {"status": "ok"}, {},
                )
                with server.database_connection() as connection:
                    items = connection.execute("SELECT * FROM itens").fetchall()
                    queries = connection.execute("SELECT * FROM consultas").fetchall()

            self.assertEqual(len(items), 1)
            self.assertEqual(items[0]["descricao_completa"], "Cadeira atualizada.")
            self.assertEqual(len(queries), 2)

    def test_pdf_tables_do_not_leak_legal_text_into_last_item(self):
        tables = [
            [
                ["ITEM", "QTD", "DESCRIÇÃO"],
                ["1", "24", "Locação de equipamento de hemodiálise."],
                ["2", "24", "Locação de equipamento de osmose reversa."],
            ],
            [["", "", "Texto jurídico de outra tabela."]],
            [["9.3", "", "O prazo estabelecido no item 9.1 ficará suspenso."]],
        ]

        rows = server.normalize_pdf_tables(tables)

        self.assertEqual([row["item"] for row in rows], ["1", "2"])
        self.assertNotIn("jurídico", rows[1]["descricao"])

    def test_pdf_table_accepts_quantity_and_unit_in_the_same_cell(self):
        tables = [[
            ["ITEM", "DESCRICAO DO OBJETO", "QTDE/UND", "VALOR EM R$"],
            [
                "01",
                "AQUISICAO DE VEICULO NOVO 1.0 TURBO.",
                "01/UND",
                "R$ 117.133,33",
            ],
        ]]

        rows = server.normalize_pdf_tables(tables)

        self.assertEqual(len(rows), 1)
        self.assertEqual(rows[0]["item"], "01")
        self.assertEqual(rows[0]["quantidade"], "01")
        self.assertEqual(rows[0]["unidade"], "UND")
        self.assertEqual(rows[0]["valor_unitario"], "R$ 117.133,33")
        self.assertIn("VEICULO NOVO", rows[0]["descricao"])

    def test_legitimate_three_level_subitem_is_not_discarded(self):
        row = make_item("1.2.3", "Subitem técnico legítimo.", quantidade="5")

        self.assertFalse(server.is_spurious_document_item(row))

    def test_financial_rate_row_is_not_treated_as_an_item(self):
        row = make_item("365", "TX = Percentual da taxa anua = 6%.", quantidade="")

        self.assertTrue(server.is_spurious_document_item(row))

    def test_scanned_quantity_is_reconciled_without_replacing_description(self):
        scanned = make_item("1", "Descrição integral do arquivo.", quantidade="")
        scanned["_nome"] = "APARELHO DE AUSCULTA FETAL DF 7000"
        pncp = make_item(
            "7703811",
            "APARELHO DE AUSCULTA FETAL DF 7000 CONFORME ESPECIFICACOES DO TERMO DE REFERENCIA ANEXO.",
            quantidade="3",
        )

        report = server.reconcile_scanned_quantities([scanned], [pncp])

        self.assertEqual(scanned["quantidade"], "3")
        self.assertEqual(scanned["descricao"], "Descrição integral do arquivo.")
        self.assertEqual(report["filled_items"], ["1"])

    def test_newer_pncp_document_wins_when_type_and_revision_are_equal(self):
        path = Path("ANEXO I - Termo de Referencia.pdf")
        old = server.candidate_document_score(
            path, {"sequencialDocumento": 1}, embedded=True
        )
        new = server.candidate_document_score(
            path, {"sequencialDocumento": 2}, embedded=True
        )

        self.assertLess(new, old)

    def test_generic_embedded_pdf_can_be_classified_as_edital_by_content(self):
        with patch.object(server, "document_content_priority", return_value=1):
            score = server.candidate_document_score(
                Path("1-SEI - processo.pdf"),
                {"sequencialDocumento": 1, "tipoDocumentoNome": "Edital"},
                embedded=True,
            )

        self.assertEqual(score[0], 1)

    def test_grouped_line_peaks_keeps_strongest_position_per_line(self):
        scores = [0, 90, 120, 80, 0, 0, 140, 130, 0]

        self.assertEqual(server.grouped_line_peaks(scores, 80, maximum_gap=2), [2, 6])

    def test_missing_scanned_item_number_is_rebuilt_from_table_order(self):
        rows = [
            make_item("", "Primeiro item."),
            make_item("2", "Segundo item."),
            make_item("3", "Terceiro item."),
            make_item("", "Quarto item."),
            make_item("5", "Quinto item."),
        ]

        repaired = server.repair_scanned_item_numbers(rows)

        self.assertEqual(
            [row["item"] for row in repaired],
            ["1", "2", "3", "4", "5"],
        )

    def test_generated_table_uses_template_section_without_forcing_new_page(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template_path = Path(temp_dir) / "template.docx"
            output_path = Path(temp_dir) / "output.docx"
            template = Document()
            template.add_paragraph("Cabecalho da proposta")
            template.save(template_path)

            server.build_docx(
                [make_item("1", "Cadeira giratoria completa.")],
                template_path,
                output_path,
            )

            generated = Document(output_path)
            self.assertEqual(len(generated.sections), 1)
            self.assertEqual(len(generated.tables), 1)
            header_run = next(
                run for run in generated.tables[0].rows[0].cells[0].paragraphs[0].runs
                if run.text
            )
            body_run = next(
                run for run in generated.tables[0].rows[1].cells[0].paragraphs[0].runs
                if run.text
            )
            self.assertTrue(header_run.bold)
            self.assertEqual(header_run.font.size.pt, 10)
            self.assertEqual(body_run.font.size.pt, 9)

    def test_proposal_total_and_commercial_terms_are_added_after_table(self):
        first = make_item("1", "Cadeira completa.", quantidade="2")
        first["valor_unitario"] = "R$ 100,00"
        first["valor_total"] = "R$ 200,00"
        second = make_item("2", "Mesa completa.", quantidade="3")
        second["valor_unitario"] = "R$ 1.000,50"

        self.assertEqual(
            server.calculate_proposal_total([first, second]),
            "R$ 3.201,50",
        )

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "terms-output.docx"
            server.build_docx(
                [first, second],
                None,
                output_path,
                commercial_terms={
                    "prazo_entrega": "15 (quinze) dias úteis",
                    "prazo_pagamento": "30 (trinta) dias",
                    "validade_proposta": "60 (sessenta) dias",
                },
            )
            generated = Document(output_path)
            paragraphs = [paragraph.text for paragraph in generated.paragraphs]

            self.assertIn("VALOR TOTAL DA PROPOSTA: R$ 3.201,50", paragraphs)
            self.assertIn("Prazo de Entrega: 15 (quinze) dias úteis", paragraphs)
            self.assertIn("Prazo de pagamento: 30 (trinta) dias", paragraphs)
            self.assertIn("Validade da Proposta: 60 (sessenta) dias", paragraphs)

    def test_commercial_terms_are_read_from_official_document_text(self):
        text = """
        O prazo de entrega dos materiais será de 15 (quinze) dias úteis.
        O pagamento será efetuado no prazo máximo de 30 (trinta) dias corridos
        após o recebimento da nota fiscal.
        A validade da proposta será de 90 (noventa) dias.
        """

        terms = server.commercial_terms_from_text(text)

        self.assertEqual(terms["prazo_entrega"], "15 (quinze) dias úteis")
        self.assertEqual(terms["prazo_pagamento"], "30 (trinta) dias corridos")
        self.assertEqual(terms["validade_proposta"], "90 (noventa) dias")

    def test_missing_commercial_term_is_not_invented(self):
        terms = server.normalized_commercial_terms({})

        self.assertEqual(terms["prazo_entrega"], server.COMMERCIAL_TERM_NOT_FOUND)
        self.assertEqual(terms["prazo_pagamento"], server.COMMERCIAL_TERM_NOT_FOUND)
        self.assertEqual(terms["validade_proposta"], server.COMMERCIAL_TERM_NOT_FOUND)

    def test_payment_heading_does_not_capture_contract_signature_deadline(self):
        text = """
        CONTRATO, RECEBIMENTO E PAGAMENTO
        O adjudicatário será notificado para assinar o contrato no prazo de
        5 (cinco) dias úteis.
        """

        terms = server.commercial_terms_from_text(text)

        self.assertNotIn("prazo_pagamento", terms)
        self.assertNotIn("prazo_entrega", terms)

    def test_template_edit_link_tracks_selected_builtin_template(self):
        page = server.render_page()

        self.assertIn('id="editTemplateLink"', page)
        self.assertIn('id="openTemplateManager"', page)
        self.assertIn('id="templateManagerOverlay"', page)
        self.assertNotIn('href="/templates"', page)
        self.assertIn('/template/${encodeURIComponent(templateId)}', page)

    def test_proposal_shell_contains_sidebar_and_preserves_manager_controls(self):
        page = server.render_page()

        self.assertIn('<div class="app-brand">TOTH</div>', page)
        self.assertIn('id="navBlock1"', page)
        self.assertIn('id="navBlock2"', page)
        self.assertIn('<h1>Gerar proposta</h1>', page)
        self.assertIn('id="openResponsibleManager"', page)
        self.assertIn('id="openTemplateManager"', page)

    def test_pncp_link_rejects_unauthorized_domain(self):
        with self.assertRaisesRegex(ValueError, "domínio pncp.gov.br"):
            server.parse_pncp_link(
                "https://exemplo.com/app/editais/12345678000199/2026/10"
            )

    def test_pncp_link_accepts_common_copied_variations(self):
        expected = ("12345678000199", 2026, 10)
        self.assertEqual(
            server.parse_pncp_link(
                "www.pncp.gov.br/app/editais/12345678000199/2026/10/?pagina=1"
            ),
            expected,
        )
        self.assertEqual(
            server.parse_pncp_link(
                "/app/editais/12345678000199/2026/10#arquivos"
            ),
            expected,
        )

    def test_alere_builtin_template_accepts_generated_table(self):
        template_path = server.resolve_template("builtin:alere")
        self.assertEqual(template_path, server.ALERE_TEMPLATE)
        self.assertTrue(template_path.exists())

        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = Path(temp_dir) / "alere-output.docx"
            server.build_docx(
                [make_item("1", "Cadeira giratoria completa.")],
                template_path,
                output_path,
            )

            generated = Document(output_path)
            self.assertEqual(len(generated.tables), 1)
            self.assertEqual(generated.tables[0].rows[1].cells[0].text, "1")

    def test_accounting_codes_are_not_item_identifiers(self):
        self.assertFalse(server.is_item_identifier("33.90.39"))
        self.assertFalse(server.is_item_identifier("3.3.90.39"))
        self.assertTrue(server.is_item_identifier("01"))
        self.assertTrue(server.is_item_identifier("1.2.3"))

    def test_sanitizer_removes_unscoped_duplicate_and_accounting_row(self):
        rows = [
            make_item("01", "Cadeira completa.", lote="1", quantidade="60"),
            make_item("02", "Cadeira giratoria completa.", lote="1", quantidade="120"),
            make_item("01", "Linha espuria terminada com", quantidade=""),
            make_item("33.90.39", "Classificacao de despesa", quantidade=""),
            make_item("1.2.3", "Subitem legitimo sem lote."),
        ]

        sanitized = server.sanitize_extracted_items(rows)

        self.assertEqual(
            {server.item_lookup_key(row) for row in sanitized},
            {"1/1", "1/2", "1.2.3"},
        )

    def test_duplicate_keeps_the_most_complete_description(self):
        short = make_item("1", "Cadeira.", lote="1", quantidade="")
        complete = make_item(
            "1",
            "Cadeira giratoria com apoio para bracos e regulagem de altura.",
            lote="1",
            quantidade="60",
        )

        [item] = server.sanitize_extracted_items([short, complete])

        self.assertEqual(item["descricao"], complete["descricao"])
        self.assertEqual(item["quantidade"], "60")

    def test_unit_is_always_standardized_as_und(self):
        item = make_item("1", "Cadeira completa.")
        item["unidade"] = "KG"

        [sanitized] = server.sanitize_extracted_items([item])
        [identified] = server.build_item_identifications([item])

        self.assertEqual(server.STANDARD_UNIT, "UND")
        self.assertIn(("unidade", "UND"), server.COLUMNS)
        self.assertEqual(sanitized["unidade"], "UND")
        self.assertEqual(identified["unidade"], "UND")

    def test_identification_uses_at_most_two_simplified_words(self):
        descriptions = {
            "Cadeira giratoria com apoio para bracos.": "Cadeira",
            "Mesa para reuniao com oito lugares.": "Mesa",
            "Forno de microondas industrial.": "Micro-ondas",
            "Prestacao de servico de limpeza predial.": "Limpeza",
            "Impressora laser multifuncional monocromatica.": "Impressora multifuncional",
        }

        for description, expected in descriptions.items():
            category = server.identify_item_category(description)
            self.assertEqual(category, expected)
            self.assertLessEqual(len(category.split()), 2)
            self.assertNotIn("/", category)

    def test_optical_range_never_becomes_a_numeric_identification(self):
        descriptions = [
            "0,00 até +2,00 esférico",
            "+2,25 até +4,00 esférico / cilíndrico -2,00",
            "Bifocal Ultex com armação",
            "Progressivas / multifocais com armação",
        ]

        self.assertEqual(
            [server.identify_item_category(description) for description in descriptions],
            ["Óculos grau"] * len(descriptions),
        )

    def test_numeric_only_description_has_safe_fallback(self):
        self.assertEqual(server.identify_item_category("0,00 até +2,00"), "Indefinido")

    def test_document_cache_reuses_data_without_sharing_mutations(self):
        cache = {}
        server.cache_set(cache, "edital", {"items": [{"item": "1"}]})

        first = server.cache_get(cache, "edital", server.DOCUMENT_CACHE_TTL)
        first["items"][0]["item"] = "alterado"
        second = server.cache_get(cache, "edital", server.DOCUMENT_CACHE_TTL)

        self.assertEqual(second["items"][0]["item"], "1")
        cache["edital"]["created_at"] = time.time() - server.DOCUMENT_CACHE_TTL - 1
        self.assertIsNone(server.cache_get(cache, "edital", server.DOCUMENT_CACHE_TTL))

    def test_wide_table_text_recovers_last_item_across_pages(self):
        pages = [
            """GOVERNO DO ESTADO
Sistema de embolizacao livre
14 183008-2 30 R$ 2.999,00 R$ 89.970,00
comprimento de 4 a 14 cm.
Identificador de autenticacao: teste""",
            """GOVERNO DO ESTADO
Sistema de embolizacao controlada
15 182359-0 5 R$ 2.999,00 R$ 14.995,00
comprimento de 10 a 50 cm.
VALOR GLOBAL ESTIMADO R$ 104.965,00
DESCRICAO DA SOLUCAO""",
        ]

        rows = server.extract_from_wide_pdf_texts(pages)

        self.assertEqual([row["item"] for row in rows], ["14", "15"])
        self.assertEqual([row["quantidade"] for row in rows], ["30", "5"])
        self.assertEqual(rows[1]["descricao"], "Sistema de embolizacao controlada comprimento de 10 a 50 cm.")

    def test_merge_repairs_shifted_quantity_and_noisy_description(self):
        noisy = make_item(
            "14",
            "Sistema de embolizacao. GOVERNO DO ESTADO TERMO DE REFERENCIA Ultima Revisao",
            quantidade="R$ 89.970,00",
        )
        repaired = make_item(
            "14",
            "Sistema de embolizacao com comprimento de 4 a 14 cm.",
            quantidade="30",
        )

        [merged] = server.merge_item_lists([noisy], [repaired])

        self.assertEqual(merged["quantidade"], "30")
        self.assertEqual(merged["descricao"], repaired["descricao"])

    def test_pdf_line_break_hyphens_are_rejoined(self):
        self.assertEqual(
            server.join_pdf_description_lines(["Microesfera cali-", "brada produzida em hidrogel."]),
            "Microesfera calibrada produzida em hidrogel.",
        )


class TemplateManagementTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.template_dir = Path(self.temp_dir.name)
        self.template_dir_patch = patch.object(server, "TEMPLATE_DIR", self.template_dir)
        self.template_dir_patch.start()

    def tearDown(self):
        self.template_dir_patch.stop()
        self.temp_dir.cleanup()

    def upload_field(self, name, text="Template de teste"):
        buffer = BytesIO()
        document = Document()
        document.add_paragraph(text)
        document.save(buffer)
        buffer.seek(0)
        return SimpleNamespace(filename=name, file=buffer)

    def test_attach_list_and_delete_template(self):
        created = server.store_new_template(self.upload_field("Proposta nova.docx"))

        self.assertEqual(created["name"], "Proposta nova.docx")
        self.assertEqual([item["name"] for item in server.list_templates()], ["Proposta nova.docx"])

        server.delete_template_file(created["id"])

        self.assertEqual(server.list_templates(), [])

    def test_duplicate_attach_is_rejected(self):
        server.store_new_template(self.upload_field("Duplicado.docx", "Primeiro"))

        with self.assertRaises(FileExistsError):
            server.store_new_template(self.upload_field("Duplicado.docx", "Segundo"))

        self.assertEqual(len(server.list_templates()), 1)

    def test_replace_keeps_same_record_without_duplicate(self):
        server.store_new_template(self.upload_field("Modelo.docx", "Anterior"))

        replaced = server.replace_template_file(
            "Modelo.docx",
            self.upload_field("Outro nome.docx", "Atualizado"),
        )

        self.assertEqual(replaced["id"], "Modelo.docx")
        self.assertEqual(len(server.list_templates()), 1)
        document = Document(self.template_dir / "Modelo.docx")
        self.assertEqual(document.paragraphs[0].text, "Atualizado")

    def test_failed_replace_preserves_previous_file(self):
        server.store_new_template(self.upload_field("Seguro.docx", "Conteudo original"))
        target = self.template_dir / "Seguro.docx"
        original = target.read_bytes()
        invalid = SimpleNamespace(filename="Corrompido.docx", file=BytesIO(b"nao e um docx"))

        with self.assertRaises(ValueError):
            server.replace_template_file("Seguro.docx", invalid)

        self.assertEqual(target.read_bytes(), original)
        self.assertFalse(any(path.suffix == ".uploading" for path in self.template_dir.iterdir()))

    def test_non_docx_upload_is_rejected(self):
        field = SimpleNamespace(filename="arquivo.pdf", file=BytesIO(b"conteudo"))

        with self.assertRaises(ValueError):
            server.store_new_template(field)

    def test_oversized_upload_is_rejected_before_validation(self):
        field = SimpleNamespace(filename="grande.docx", file=BytesIO(b"x" * 11))

        with patch.object(server, "MAX_TEMPLATE_SIZE", 10):
            with self.assertRaisesRegex(ValueError, "excede"):
                server.store_new_template(field)

        self.assertEqual(list(self.template_dir.iterdir()), [])

    def test_explicit_missing_template_never_falls_back_to_default(self):
        default = server.store_new_template(
            self.upload_field("Padrao.docx", "MODELO PADRAO")
        )
        upload_dir = self.template_dir / "uploads"
        upload_dir.mkdir()

        with (
            patch.object(server, "UPLOAD_DIR", upload_dir),
            patch.object(server, "DEFAULT_TEMPLATE", self.template_dir / default["id"]),
        ):
            self.assertEqual(
                server.resolve_template(""),
                self.template_dir / default["id"],
            )
            self.assertIsNone(server.resolve_template("managed:Ausente.docx"))
            self.assertIsNone(server.resolve_template("upload:modelo_ausente.docx"))
            self.assertIsNone(server.resolve_template("modelo_ausente.docx"))

    def test_uploaded_proposal_template_is_validated_selected_and_preserved(self):
        managed = server.store_new_template(
            self.upload_field("Cadastrado.docx", "MODELO CADASTRADO")
        )
        upload_dir = self.template_dir / "uploads"
        upload_dir.mkdir()
        form = {
            "template_choice": SimpleNamespace(value=managed["id"]),
            "template_file": self.upload_field("Avulso.docx", "MARCADOR AVULSO"),
        }

        with patch.object(server, "UPLOAD_DIR", upload_dir):
            selection = server.proposal_template_selection(form)
            resolved = server.resolve_template(selection["ref"])
            output = self.template_dir / "resultado.docx"
            server.build_docx(
                [make_item("1", "Cadeira giratoria completa.")],
                resolved,
                output,
            )

        self.assertEqual(selection["source"], "upload")
        self.assertEqual(selection["name"], "Avulso.docx")
        self.assertTrue(selection["ref"].startswith("upload:modelo_"))
        self.assertEqual(resolved, selection["path"])
        self.assertEqual(Document(output).paragraphs[0].text, "MARCADOR AVULSO")

    def test_invalid_proposal_template_upload_is_rejected(self):
        upload_dir = self.template_dir / "uploads"
        upload_dir.mkdir()
        form = {
            "template_file": SimpleNamespace(
                filename="Corrompido.docx",
                file=BytesIO(b"arquivo invalido"),
            )
        }

        with patch.object(server, "UPLOAD_DIR", upload_dir):
            with self.assertRaisesRegex(ValueError, "não é um documento Word"):
                server.proposal_template_selection(form)

        self.assertEqual(list(upload_dir.iterdir()), [])


class ResponsibleManagementTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.data_dir = Path(self.temp_dir.name)
        self.database_path_patch = patch.object(
            server, "DATABASE_PATH", self.data_dir / "pncp.sqlite3"
        )
        self.data_dir_patch = patch.object(server, "DATA_DIR", self.data_dir)
        self.database_path_patch.start()
        self.data_dir_patch.start()
        server.init_database()

    def tearDown(self):
        self.data_dir_patch.stop()
        self.database_path_patch.stop()
        self.temp_dir.cleanup()

    def payload(self, **changes):
        data = {
            "nome_completo": "Ana Paula Souza",
            "empresa": "Empresa Teste LTDA",
            "cnpj": "12.345.678/0001-95",
            "rg": "12.345.678-9",
            "cpf": "123.456.789-09",
            "observacoes": "Representante legal",
        }
        data.update(changes)
        return data

    def test_initial_responsibles_are_migrated_and_resolvable(self):
        responsibles = server.list_responsibles()

        self.assertEqual(len(responsibles), 2)
        self.assertEqual(server.resolve_responsible("1")["nome_completo"], "Brendon Matheus Batista")
        self.assertIn("CPF 432.079.848-19", server.resolve_responsible("1")["document_lines"])

    def test_create_and_update_keep_the_same_record(self):
        created = server.create_responsible(self.payload())
        updated = server.update_responsible(
            created["id"], self.payload(nome_completo="Ana Paula Atualizada")
        )

        self.assertEqual(updated["id"], created["id"])
        self.assertEqual(updated["nome_completo"], "Ana Paula Atualizada")
        self.assertEqual(len(server.list_responsibles()), 3)

    def test_duplicate_cpf_is_rejected(self):
        server.create_responsible(self.payload())

        with self.assertRaises(FileExistsError):
            server.create_responsible(self.payload(nome_completo="Outra Pessoa"))

    def test_linked_responsible_cannot_be_deleted(self):
        created = server.create_responsible(self.payload())
        output_path = self.data_dir / "proposta.docx"
        output_path.write_bytes(b"documento")
        server.record_generated_document(created["id"], output_path)

        with self.assertRaisesRegex(server.ResponsibleInUseError, "1 documento"):
            server.delete_responsible(created["id"])

        self.assertIsNotNone(server.get_responsible(created["id"]))

    def test_unlinked_responsible_can_be_deleted(self):
        created = server.create_responsible(self.payload())

        server.delete_responsible(created["id"])

        self.assertIsNone(server.get_responsible(created["id"]))

    def test_initial_responsibles_are_not_recreated_after_deletion(self):
        server.delete_responsible("1")
        server.delete_responsible("2")

        self.assertEqual(server.list_responsibles(), [])


class CatalogGenerationTests(unittest.TestCase):
    def make_catalog_data(self):
        item = make_item(
            "12",
            (
                "Cadeira giratória com assento em espuma injetada. "
                "Encosto revestido em tela. Estrutura em aço. "
                "Dimensões: largura 470 mm. Conforme ABNT NBR 13962."
            ),
            quantidade="10",
        )
        pncp = {
            "ano": 2026,
            "sequencial": 43,
            "link": "https://pncp.gov.br/app/editais/45780087000103/2026/43",
            "documento_tipo": "Termo de Referência",
            "documento_usado": "termo.pdf",
            "metadata": {
                "numero_compra": "13/2026",
                "processo": "12425/2025",
                "modalidade": "Pregão - Eletrônico",
                "objeto": "Aquisição de mobiliário.",
                "orgao": "Município de Várzea Paulista",
                "orgao_cnpj": "45780087000103",
                "unidade": "Várzea Paulista",
                "municipio": "Várzea Paulista",
                "uf": "SP",
            },
        }
        data = catalog.catalog_draft_from_item(item, pncp)
        data["fabricante"].update({
            "razao_social": "Fabricante Teste Ltda.",
            "nome_fantasia": "Fabricante Teste",
            "cnpj": "33661439000114",
        })
        data["produto"].update({"marca": "Marca Teste", "modelo": "Modelo A"})
        return data

    def test_catalog_draft_preserves_source_and_groups_technical_content(self):
        data = self.make_catalog_data()

        self.assertIn("espuma injetada", data["item"]["descricao"])
        self.assertIn("assento", data["secoes"]["assento"].lower())
        self.assertIn("encosto", data["secoes"]["encosto"].lower())
        self.assertIn("470 mm", data["secoes"]["dimensoes"])
        self.assertIn("NBR 13962", data["secoes"]["normas"])

    def test_catalog_sections_keep_values_after_technical_labels(self):
        sections = catalog.section_text_from_description(
            "Dimensões assento: 450 x 490, dimensões encosto: 450 x 560, "
            "material estrutura: resina termoplástica injetada."
        )

        self.assertIn("450 x 490", sections["assento"])
        self.assertIn("450 x 560", sections["encosto"])
        self.assertIn("resina termoplástica", sections["estrutura"])
        self.assertIn("450 x 490", sections["dimensoes"])

    def test_catalog_draft_uses_structured_item_when_official_file_is_unavailable(self):
        link = "https://pncp.gov.br/app/editais/45780087000103/2026/98765"
        identification = {
            "items": [make_item("1", "Cadeira ergonômica com apoio lombar", quantidade="8")],
            "pncp": {
                "cnpj": "45780087000103",
                "ano": 2026,
                "sequencial": 98765,
                "link": link,
                "documento_tipo": "Base estruturada",
            },
        }
        server.CATALOG_DRAFT_CACHE.clear()
        with (
            patch.object(server, "identify_items_from_pncp_link", return_value=identification),
            patch.object(
                server,
                "source_from_pncp_link",
                side_effect=RuntimeError("Nenhum arquivo encontrado no PNCP"),
            ),
            patch.object(
                server,
                "pncp_purchase_metadata",
                return_value={"numero_compra": "44/2026", "orgao": "Órgão teste"},
            ),
        ):
            result = server.catalog_draft_from_pncp_link(link, "1")
        server.CATALOG_DRAFT_CACHE.clear()

        self.assertEqual(result["source"], "base_estruturada")
        self.assertIn("Nenhum arquivo", result["enrichment_warning"])
        self.assertEqual(result["draft"]["item"]["numero"], "1")
        self.assertEqual(result["draft"]["item"]["quantidade"], "8")
        self.assertEqual(result["draft"]["origem"]["tipo"], "Base estruturada local")

    def test_purchase_metadata_prefers_normalized_local_opportunity(self):
        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=lambda *_args: {
                "opportunity": {
                    "title": "Edital Pregão Eletrônico nº 204/2026",
                    "description": "Aquisição de cadeiras",
                    "process_number": "100/2026",
                    "modality": "Pregão eletrônico",
                    "buyer_name": "Órgão local",
                    "buyer_cnpj": "45780087000103",
                    "city": "Várzea Paulista",
                    "uf": "SP",
                    "uasg": "123456",
                    "status": "Aberta",
                    "pncp_control_number": "controle-local",
                    "proposal_start_at": "2026-08-20T09:00:00",
                    "proposal_end_at": "2026-08-25T09:00:00",
                    "estimated_value": 1250,
                    "origin_url": "https://origem.test/compra",
                }
            },
        )
        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "request_json") as remote_request,
        ):
            metadata = server.pncp_purchase_metadata("45780087000103", 2026, 98765)

        remote_request.assert_not_called()
        self.assertEqual(metadata["numero_compra"], "204/2026")
        self.assertEqual(metadata["objeto"], "Aquisição de cadeiras")
        self.assertEqual(metadata["codigo_unidade"], "123456")

    def test_block7_normalizes_with_traceability_and_exports_all_formats(self):
        raw = [make_item("1", "Cadeira giratória em tela", quantidade="12")]
        items = catalog_generator.normalize_items(
            raw,
            "https://pncp.gov.br/app/editais/45780087000103/2026/43",
        )
        self.assertEqual(items[0]["categoria"], "Mobiliário")
        self.assertEqual(items[0]["fontes"][0]["secao"], "Item 1")

        with tempfile.TemporaryDirectory() as temp_dir:
            exports = catalog_generator.export_catalog(
                Path(temp_dir),
                {"objeto": "Aquisição de cadeiras"},
                items,
                "a" * 32,
            )
            self.assertEqual(set(exports), {"xlsx", "csv", "json", "pdf"})
            self.assertTrue(all((Path(temp_dir) / entry["filename"]).stat().st_size for entry in exports.values()))

    def test_block7_revalidates_edited_required_fields(self):
        items = catalog_generator.normalize_items(
            [make_item("1", "Cadeira giratória", quantidade="12")],
            "https://pncp.gov.br/app/editais/45780087000103/2026/43",
        )
        edited = dict(items[0], descricao="", quantidade="", campos_ausentes=[])

        summary = catalog_generator.validation_summary([edited])
        sanitized = catalog_generator.sanitize_export_items([edited])[0]

        self.assertEqual(summary["incompletos"], 1)
        self.assertEqual(sanitized["status_evidencia"], "incompleto")
        self.assertEqual(
            sanitized["campos_ausentes"],
            ["descrição", "quantidade", "unidade"],
        )

    def test_block7_pdf_preserves_long_description_and_escapes_markup(self):
        items = catalog_generator.normalize_items(
            [make_item("1", "Cadeira giratória", quantidade="12")],
            "https://pncp.gov.br/app/editais/45780087000103/2026/43",
        )
        description = ("Descrição técnica <segura> detalhada. " * 120) + "SENTINELA_FINAL_QA"
        items[0]["descricao"] = description

        with tempfile.TemporaryDirectory() as temp_dir:
            exports = catalog_generator.export_catalog(
                Path(temp_dir), {"objeto": "Aquisição <teste>"}, items, "c" * 32
            )
            pdf_path = Path(temp_dir) / exports["pdf"]["filename"]
            with server.pdfplumber.open(pdf_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)

        self.assertIn("SENTINELA_FINAL_QA", text)
        self.assertIn("segura", text)

    def test_block7_exports_are_unique_within_the_same_second(self):
        items = catalog_generator.normalize_items(
            [make_item("1", "Cadeira giratória", quantidade="12")],
            "https://pncp.gov.br/app/editais/45780087000103/2026/43",
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir)
            with patch.object(catalog_generator, "datetime") as clock:
                clock.now.return_value = datetime(2026, 8, 31, 12, 0, 0)
                first = catalog_generator.export_catalog(output, {}, items, "d" * 32)
                original = (output / first["json"]["filename"]).read_bytes()
                items[0]["descricao"] = "Segunda versão"
                second = catalog_generator.export_catalog(output, {}, items, "d" * 32)

            self.assertNotEqual(first["json"]["filename"], second["json"]["filename"])
            self.assertEqual((output / first["json"]["filename"]).read_bytes(), original)

    def test_block7_keeps_document_fallback_when_item_api_is_unavailable(self):
        job_id = "b" * 32
        server.CATALOG_GENERATOR_JOBS[job_id] = {
            "id": job_id,
            "status": "queued",
            "stage": "validacao",
            "progress": 0,
        }
        document_item = make_item("3", "Cadeira fixa empilhável", quantidade="20")
        candidate = {
            "path": Path("termo.pdf"),
            "file_info": {"titulo": "Termo de Referência"},
            "embedded": False,
        }
        with (
            patch.object(server, "pncp_purchase_metadata", return_value={"objeto": "Cadeiras"}),
            patch.object(server, "catalog_generator_document_candidates", return_value=([candidate], [])),
            patch.object(server, "extract_items_cached", return_value=[document_item]),
            patch.object(server, "identify_items_from_pncp_link", side_effect=RuntimeError("PNCP indisponível")),
        ):
            server.run_catalog_generator_job(
                job_id,
                "https://pncp.gov.br/app/editais/45780087000103/2026/43",
            )

        job = server.catalog_generator_job(job_id)
        self.assertEqual(job["status"], "ready")
        self.assertEqual(job["result"]["items"][0]["numero"], "3")
        self.assertTrue(any("continuou pelos documentos" in warning for warning in job["result"]["warnings"]))

    def test_catalog_specification_accepts_qualifier_added_to_item_title(self):
        item = make_item("2", "Cadeira Caixa Alta Secretaria")
        source_text = (
            "b) Cadeira Caixa Alta Secretaria - Giratória: Assento anatômico "
            "com espuma D-28 e regulagem de altura por pistão a gás.\n"
            "c) Cadeira Aproximação Atendimento: Estrutura fixa."
        )
        with patch.object(server, "catalog_document_text", return_value=source_text):
            description = server.catalog_specification_from_document(
                Path("termo.pdf"),
                item,
            )

        self.assertIn("Giratória", source_text)
        self.assertIn("espuma D-28", description)
        self.assertNotIn("Estrutura fixa", description)

    def test_catalog_job_validates_selection_before_starting_a_worker(self):
        link = "https://pncp.gov.br/app/editais/45780087000103/2026/43"
        with patch.object(server.threading, "Thread") as worker:
            for selection in ([], "1", [""], [None], ["1/"]):
                with self.subTest(selection=selection), self.assertRaises(ValueError):
                    server.create_catalog_generator_job({
                        "pncp_link": link, "selected_item_keys": selection,
                    })
            worker.assert_not_called()

    def test_catalog_job_passes_normalized_selection_to_worker(self):
        link = "https://pncp.gov.br/app/editais/45780087000103/2026/43"
        with (
            patch.dict(server.CATALOG_GENERATOR_JOBS, {}, clear=True),
            patch.object(server.threading, "Thread") as worker,
        ):
            job = server.create_catalog_generator_job({
                "pncp_link": link, "selected_item_keys": ["01/02", "1/2", "3"],
            })
            self.assertEqual(job["selected_item_keys"], ["1/2", "3"])
            self.assertEqual(worker.call_args.kwargs["args"], (job["id"], link, ["1/2", "3"]))

    def test_catalog_selection_filters_document_items_and_preserves_distinct_lots(self):
        job_id = "c" * 32
        first = {**make_item("1", "Cadeira fixa", quantidade="2"), "lote": "1"}
        second = {**first, "lote": "2"}
        unselected = make_item("3", "Mesa", quantidade="4")
        with (
            patch.dict(server.CATALOG_GENERATOR_JOBS, {job_id: {"id": job_id}}, clear=True),
            patch.object(server, "pncp_purchase_metadata", return_value={"objeto": "Mobiliario"}),
            patch.object(server, "identify_items_from_pncp_link", return_value={"items": [first, unselected]}),
            patch.object(server, "catalog_generator_document_candidates", return_value=([
                {"path": Path("itens.xlsx"), "file_info": {"titulo": "Itens"}},
            ], [])),
            patch.object(server, "extract_items_cached", return_value=[second, unselected]),
        ):
            server.run_catalog_generator_job(
                job_id, "https://pncp.gov.br/app/editais/45780087000103/2026/43", ["1/1", "2/1"],
            )
            job = server.catalog_generator_job(job_id)

        self.assertEqual(job["status"], "ready")
        items = job["result"]["items"]
        self.assertEqual([(item["lote"], item["numero"]) for item in items], [("1", "1"), ("2", "1")])
        self.assertEqual(job["result"]["metadata"]["total_itens"], 2)
        self.assertEqual(catalog_generator.sanitize_export_items(items)[1]["lote"], "2")

    def test_catalog_missing_selection_fails_instead_of_processing_every_item(self):
        job_id = "d" * 32
        with (
            patch.dict(server.CATALOG_GENERATOR_JOBS, {job_id: {"id": job_id}}, clear=True),
            patch.object(server, "pncp_purchase_metadata", return_value={}),
            patch.object(server, "identify_items_from_pncp_link", return_value={"items": [make_item("1", "Cadeira")]}),
            patch.object(server, "catalog_generator_document_candidates", return_value=([], [])),
            patch.object(server.LOGGER, "exception"),
        ):
            server.run_catalog_generator_job(
                job_id, "https://pncp.gov.br/app/editais/45780087000103/2026/43", ["99"],
            )
            job = server.catalog_generator_job(job_id)

        self.assertEqual(job["status"], "failed")
        self.assertIn("99", job["error"])
        self.assertNotIn("result", job)

    def test_catalog_exports_docx_pdf_json_and_images(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            logo_path = temp_path / "logo.png"
            product_path = temp_path / "produto.png"
            Image.new("RGB", (320, 120), "white").save(logo_path)
            Image.new("RGB", (500, 500), "#dddddd").save(product_path)
            assets = [
                {"path": logo_path, "name": "logo.png", "role": "logo", "section": "", "caption": ""},
                {"path": product_path, "name": "produto.png", "role": "principal", "section": "", "caption": "Produto"},
            ]
            data = self.make_catalog_data()

            with patch.object(server, "OUTPUT_DIR", temp_path):
                response = server.generate_catalog_exports(data, assets)

            paths = {
                key: temp_path / value["filename"]
                for key, value in response["exports"].items()
            }
            self.assertTrue(all(path.exists() and path.stat().st_size for path in paths.values()))

            document = Document(paths["docx"])
            section = document.sections[0]
            self.assertLess(section.page_width, section.page_height)
            document_xml = "\n".join(
                part.blob.decode("utf-8", errors="ignore")
                for part in document.part.package.parts
                if getattr(part, "blob", None)
            )
            self.assertIn("CatalogWatermark", document_xml)

            payload = json.loads(paths["json"].read_text(encoding="utf-8"))
            self.assertEqual(payload["schema"], "catalogo-tecnico-licitacao/v1")
            self.assertEqual(payload["dados"]["item"]["numero"], "12")
            with zipfile.ZipFile(paths["images"]) as archive:
                self.assertIn("manifesto.json", archive.namelist())
                self.assertEqual(len([name for name in archive.namelist() if name.endswith(".png")]), 2)


class StructuredItemPriorityTests(unittest.TestCase):
    def test_proposal_processing_uses_structured_items_without_official_document(self):
        link = "https://pncp.gov.br/app/editais/12345678000199/2026/987653"
        identification = {
            "source": "opportunity_items",
            "items": [
                make_item("1", "Cadeira giratoria completa", lote="2", quantidade="4"),
                make_item("2", "Mesa retangular completa", lote="2", quantidade="3"),
            ],
            "pncp": {
                "cnpj": "12345678000199",
                "ano": 2026,
                "sequencial": 987653,
                "link": link,
                "documento_usado": "",
                "documento_tipo": "Base estruturada",
            },
            "pncp_items_check": {
                "structured_count": 2,
                "has_divergence": False,
            },
        }

        with (
            patch.object(server, "identify_items_from_pncp_link", return_value=identification),
            patch.object(
                server,
                "source_from_pncp_link",
                side_effect=AssertionError("arquivo oficial nao deveria ser consultado"),
            ) as official_source,
            patch.object(
                server,
                "list_pncp_items",
                side_effect=AssertionError("API de itens nao deveria ser consultada novamente"),
            ) as pncp_items,
            patch.object(
                server,
                "extract_items_cached",
                side_effect=AssertionError("PDF nao deveria ser processado"),
            ) as file_items,
        ):
            result = server.proposal_process_from_structured_items(link, "2/1")

        self.assertEqual(result["count"], 1)
        self.assertEqual(result["items"][0]["item"], "1")
        self.assertEqual(result["source_name"], "PNCP_12345678000199_2026_987653")
        self.assertEqual(result["pncp_items_check"]["selected_count"], 1)
        self.assertEqual(result["commercial_terms"]["status"], "warn")
        official_source.assert_not_called()
        pncp_items.assert_not_called()
        file_items.assert_not_called()

    def test_identification_uses_opportunity_items_before_official_files(self):
        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=lambda *_args: {
                "opportunity": {"id": "opportunity-1"},
                "items": [{
                    "lot_number": "2",
                    "item_number": "7",
                    "title": "Cadeira giratoria",
                    "description": "Cadeira giratoria com apoio de bracos",
                    "technical_object": None,
                    "quantity": 12,
                    "unit": "UN",
                }],
                "documents": [],
                "matches": [],
            }
        )
        link = "https://pncp.gov.br/app/editais/12345678000199/2026/987654"
        server.IDENTIFICATION_CACHE.clear()

        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "ALLOW_RUNTIME_PNCP_API", True),
            patch.object(server, "list_pncp_items", return_value=[{
                "lote": "2",
                "item": "7",
                "quantidade": "12",
                "unidade": "UN",
                "descricao": "Cadeira giratoria com apoio de bracos",
            }]) as api_items,
            patch.object(
                server,
                "source_from_pncp_link",
                side_effect=AssertionError("arquivo oficial não deveria ser consultado"),
            ),
        ):
            result = server.identify_items_from_pncp_link(link)

        self.assertEqual(result["source"], "opportunity_items")
        self.assertEqual(result["items"][0]["lote"], "2")
        self.assertEqual(result["items"][0]["item"], "7")
        self.assertEqual(result["items"][0]["quantidade"], "12")
        self.assertEqual(result["items"][0]["unidade"], "UN")
        self.assertFalse(result["pncp_items_check"]["api_available"])
        self.assertFalse(result["pncp_items_check"]["has_divergence"])
        api_items.assert_not_called()

    def test_structured_items_are_not_compared_with_api(self):
        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=lambda *_args: {
                "opportunity": {"id": "opportunity-2"},
                "items": [{
                    "lot_number": "",
                    "item_number": "1",
                    "title": "Item interno",
                    "description": "Descrição estruturada completa",
                    "technical_object": None,
                    "quantity": 4,
                    "unit": "UN",
                }],
                "documents": [],
                "matches": [],
            }
        )
        link = "https://pncp.gov.br/app/editais/12345678000199/2026/987655"
        server.IDENTIFICATION_CACHE.clear()
        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "ALLOW_RUNTIME_PNCP_API", True),
            patch.object(server, "list_pncp_items", return_value=[{
                "lote": "",
                "item": "2",
                "quantidade": "4",
                "unidade": "UN",
                "descricao": "Outro item na API",
            }]) as api_items,
        ):
            result = server.identify_items_from_pncp_link(link)

        self.assertEqual(result["items"][0]["item"], "1")
        self.assertEqual(result["items"][0]["descricao"], "Descrição estruturada completa")
        self.assertFalse(result["pncp_items_check"]["has_divergence"])
        self.assertEqual(result["pncp_items_check"]["only_in_file"], [])
        self.assertEqual(result["pncp_items_check"]["added_from_pncp"], [])
        api_items.assert_not_called()

    def test_missing_items_are_fetched_from_pncp_and_saved(self):
        detail = {
            "opportunity": {
                "id": "opportunity-3",
                "external_key": "12345678000199-1-987656/2026",
            },
            "items": [],
            "documents": [],
            "matches": [],
        }
        saved_items = []
        successful_audits = []

        def replace_items(opportunity_id, items):
            self.assertEqual(opportunity_id, "opportunity-3")
            saved_items.extend(items)
            detail["items"] = [{
                "source_item_id": item.source_item_id,
                "lot_number": item.lot_number,
                "item_number": item.item_number,
                "title": item.title,
                "description": item.description,
                "technical_object": item.technical_object,
                "quantity": item.quantity,
                "unit": item.unit,
            } for item in items]
            return {"persisted": True, "count": len(items)}

        def unexpected_failure(**_kwargs):
            self.fail("A consulta de itens nao deveria registrar falha")

        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=lambda *_args: detail,
            create_run=lambda *_args, **_kwargs: "run-1",
            persist_opportunity_items_enrichment=lambda **kwargs: (
                successful_audits.append(kwargs)
                or replace_items(kwargs["opportunity_id"], kwargs["items"])
            ),
            record_opportunity_items_enrichment_failure=unexpected_failure,
        )
        connector_calls = []

        def iter_items(cnpj, year, sequence, max_pages):
            connector_calls.append((cnpj, year, sequence, max_pages))
            yield SimpleNamespace(
                records=[{
                    "numeroItem": 1,
                    "descricao": "Cadeira giratoria com apoio de bracos",
                    "quantidade": 6,
                    "unidadeMedida": "UN",
                }],
                raw_payload={"data": [{"numeroItem": 1}]},
                request_url="https://pncp.test/itens?pagina=1",
            )

        link = "https://pncp.gov.br/app/editais/12345678000199/2026/987656"
        server.IDENTIFICATION_CACHE.clear()
        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "PNCPConnector", return_value=SimpleNamespace(iter_items=iter_items)),
            patch.object(server, "ALLOW_BLOCO2_ON_DEMAND_ENRICHMENT", True),
            patch.object(server, "ALLOW_RUNTIME_PNCP_API", False),
        ):
            result = server.identify_items_from_pncp_link(link)

        self.assertEqual(connector_calls, [("12345678000199", 2026, 987656, 20)])
        self.assertEqual(len(saved_items), 1)
        self.assertEqual(result["source"], "opportunity_items_enriquecido_bloco2")
        self.assertEqual(result["items"][0]["item"], "1")
        self.assertEqual(result["items"][0]["quantidade"], "6")
        self.assertEqual(successful_audits[0]["audit_summary"]["items_normalized"], 1)

    def test_missing_opportunity_is_imported_before_items_are_enriched(self):
        stored_detail = None
        completed_runs = []
        persisted_items = []

        def get_opportunity(*_args):
            return stored_detail

        def persist_record(**kwargs):
            nonlocal stored_detail
            opportunity = kwargs["opportunity"]
            stored_detail = {
                "opportunity": {
                    "id": "opportunity-imported",
                    "external_key": opportunity.external_key,
                },
                "items": [],
                "documents": [],
                "matches": [],
            }
            return "inserted", "opportunity-imported"

        def persist_items(**kwargs):
            persisted_items.extend(kwargs["items"])
            stored_detail["items"] = [{
                "source_item_id": item.source_item_id,
                "lot_number": item.lot_number,
                "item_number": item.item_number,
                "title": item.title,
                "description": item.description,
                "technical_object": item.technical_object,
                "quantity": item.quantity,
                "unit": item.unit,
            } for item in kwargs["items"]]
            return {"persisted": True, "count": len(kwargs["items"])}

        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=get_opportunity,
            create_run=lambda _source, run_type, _filters: f"run-{run_type}",
            persist_record=persist_record,
            finish_run=lambda run_id, **kwargs: completed_runs.append((run_id, kwargs)),
            persist_opportunity_items_enrichment=persist_items,
            record_opportunity_items_enrichment_failure=lambda **_kwargs: None,
        )

        class Connector:
            def fetch_detail(self, cnpj, year, sequence):
                return SimpleNamespace(
                    request_url="https://pncp.test/detalhe",
                    payload={
                        "numeroControlePNCP": f"{cnpj}-1-{sequence:06d}/{year}",
                        "numeroCnpj": cnpj,
                        "anoCompra": year,
                        "sequencialCompra": sequence,
                        "numeroCompra": "PE 123/2026",
                        "objetoCompra": "Aquisicao de mobiliarios escolares",
                    },
                )

            def iter_items(self, _cnpj, _year, _sequence, _max_pages):
                yield SimpleNamespace(
                    records=[{
                        "numeroItem": 1,
                        "descricao": "Cadeira escolar",
                        "quantidade": 20,
                        "unidadeMedida": "UN",
                    }],
                    request_url="https://pncp.test/itens?pagina=1",
                )

        link = "https://pncp.gov.br/app/editais/87613188000121/2026/219"
        server.IDENTIFICATION_CACHE.clear()
        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "PNCPConnector", return_value=Connector()),
            patch.object(server, "ALLOW_BLOCO2_ON_DEMAND_ENRICHMENT", True),
        ):
            result = server.identify_items_from_pncp_link(link)

        self.assertEqual(result["items"][0]["descricao"], "Cadeira escolar")
        self.assertEqual(result["source"], "opportunity_items_enriquecido_bloco2")
        self.assertEqual(len(persisted_items), 1)
        self.assertEqual(completed_runs[0][0], "run-opportunity_on_demand")
        self.assertEqual(completed_runs[0][1]["status"], "success")

    def test_missing_opportunity_api_failure_is_audited(self):
        completed_runs = []
        repository = SimpleNamespace(
            create_run=lambda *_args, **_kwargs: "run-import-failed",
            finish_run=lambda run_id, **kwargs: completed_runs.append((run_id, kwargs)),
        )
        connector = SimpleNamespace(
            fetch_detail=lambda *_args: (_ for _ in ()).throw(TimeoutError("timeout")),
        )

        with self.assertRaisesRegex(RuntimeError, "nao foi possivel importa-la"):
            server.import_pncp_opportunity_on_demand(
                "87613188000121",
                2026,
                219,
                repository=repository,
                connector=connector,
            )

        self.assertEqual(completed_runs[0][0], "run-import-failed")
        self.assertEqual(completed_runs[0][1]["status"], "failed")
        self.assertEqual(completed_runs[0][1]["counters"]["failed"], 1)


class DescriptionReviewRegressionTests(unittest.TestCase):
    def test_continuation_punctuation_is_a_warning_not_a_blocker(self):
        item = make_item("1", "Cadeira giratoria com estrutura em aco;")

        review = server.build_description_review([item], file_items=[item])

        self.assertEqual(review["status"], "warn")
        self.assertEqual(review["blocking_items"], [])
        self.assertEqual(review["warning_items"], ["1"])
        self.assertEqual(review["complete_count"], 1)

    def test_missing_description_remains_blocking(self):
        item = make_item("1", "")

        review = server.build_description_review([item], file_items=[item])

        self.assertEqual(review["status"], "error")
        self.assertEqual(review["blocking_items"], ["1"])
        self.assertIn("descricao ausente", server.norm(review["message"]))

    def test_document_section_inside_description_is_blocking(self):
        item = make_item("14", "Sistema de embolizacao. DESCRICAO DA SOLUCAO Texto externo.")

        review = server.build_description_review([item], file_items=[item])

        self.assertEqual(review["status"], "error")
        self.assertEqual(review["blocking_items"], ["14"])


class PncpSearchPaginationTests(unittest.TestCase):
    def test_opportunity_search_does_not_fetch_items_or_documents(self):
        row = {
            "title": "Pregao para mobiliario",
            "description": "Aquisicao de cadeiras escolares",
        }
        with (
            patch.object(
                server,
                "get_search_row_pncp_items",
                side_effect=AssertionError("itens nao devem ser consultados na busca"),
            ) as items,
            patch.object(
                server,
                "get_search_row_document_items",
                side_effect=AssertionError("documentos nao devem ser consultados na busca"),
            ) as documents,
        ):
            filtered = server.filter_opportunity_rows_by_search_term([row], "cadeiras")
            matches_type = server.row_matches_opportunity_type(row, "material")

        self.assertEqual(filtered, [row])
        self.assertTrue(matches_type)
        items.assert_not_called()
        documents.assert_not_called()

    def test_opportunity_type_does_not_discard_unclassified_rows(self):
        row = {
            "title": "Edital 42/2026",
            "description": "Atendimento das necessidades da rede municipal",
        }

        self.assertTrue(server.row_matches_opportunity_type(row, "material"))
        self.assertTrue(server.row_matches_opportunity_type(row, "servico"))

    def test_opportunity_type_rejects_a_known_opposite_type(self):
        row = {
            "title": "Edital 43/2026",
            "description": "Contratacao de empresa especializada para limpeza predial",
        }

        self.assertFalse(server.row_matches_opportunity_type(row, "material"))
        self.assertTrue(server.row_matches_opportunity_type(row, "servico"))

    def test_semicolon_keywords_are_normalized_and_deduplicated(self):
        self.assertEqual(
            server.split_search_keywords(" Cadeira de rodas; monitor;cadeira de RODAS; "),
            ["Cadeira de rodas", "monitor"],
        )

    def test_search_term_requires_complete_words_and_phrase(self):
        self.assertTrue(
            server.matches_complete_search_term(
                {"title": "Edital", "description": "Cadeira de rodas motorizada"},
                "cadeira de rodas",
            )
        )
        self.assertTrue(
            server.matches_complete_search_term(
                {"title": "Aquisição de MONITOR", "description": ""},
                "monitor",
            )
        )
        self.assertTrue(
            server.matches_complete_search_term(
                {
                    "title": "Edital",
                    "description": "Registro eletrônico integrado de ponto",
                },
                "registro de ponto",
            )
        )
        self.assertFalse(
            server.matches_complete_search_term(
                {"title": "Edital", "description": "Serviço de monitoramento"},
                "monitor",
            )
        )

        self.assertTrue(
            server.matches_complete_search_term(
                {"title": "Edital", "description": "Aquisição de monitores"},
                "monitor",
            )
        )
        self.assertTrue(
            server.matches_complete_search_term(
                {"title": "Edital", "description": "Aquisicao de cadeira de rodas"},
                "cadeiras",
            )
        )
        self.assertFalse(
            server.matches_complete_search_term(
                {"title": "Edital", "description": "Cadeira para rodas"},
                "cadeira de rodas",
            )
        )
        self.assertFalse(
            server.matches_complete_search_term(
                {
                    "title": "Edital",
                    "description": (
                        "Registro de preços para aquisição de relógios de ponto"
                    ),
                },
                "registro de ponto",
            )
        )

    def test_search_term_can_match_buyer_name(self):
        self.assertTrue(
            server.matches_complete_search_term(
                {
                    "title": "Edital 42/2026",
                    "description": "Aquisicao de materiais",
                    "orgao_nome": "Prefeitura Municipal de Campinas",
                },
                "Campinas",
            )
        )

    def test_search_term_can_match_official_item_title(self):
        self.assertTrue(
            server.items_match_search_term(
                [{"titulo": "Equipamento hospitalar", "descricao": "Modelo padrao"}],
                "hospitalar",
            )
        )

    def test_pncp_item_listing_uses_large_pages_and_extended_timeout(self):
        with patch.object(server, "request_json", return_value=[]) as request:
            payload = server.list_pncp_item_payload("12345678000199", 2026, 7)

        self.assertEqual(payload, [])
        request.assert_called_once()
        url = request.call_args.args[0]
        self.assertIn("tamanhoPagina=500", url)
        self.assertEqual(request.call_args.kwargs["timeout"], 45)

    def test_search_term_can_match_official_item_description(self):
        row = {
            "orgao_cnpj": "46422408000152",
            "ano": "2026",
            "numero_sequencial": "367",
            "title": "Edital 60/2026",
            "description": "Aquisição de Registradores Eletrônicos de Ponto",
        }
        pncp_items = [{
            "descricao": "Registro eletrônico de ponto por reconhecimento facial",
        }]

        with patch.object(server, "list_pncp_items", return_value=pncp_items) as items:
            result = server.filter_rows_by_complete_search_term(
                [row],
                "registro de ponto",
            )

        self.assertEqual(result, [row])
        items.assert_called_once_with("46422408000152", "2026", "367")
        server.SEARCH_ITEM_CACHE.clear()

    def test_search_falls_back_to_document_when_api_omits_matching_item(self):
        row = {
            "orgao_cnpj": "46422408000152",
            "ano": "2026",
            "numero_sequencial": "368",
            "title": "Edital 61/2026",
            "description": "Aquisição de equipamentos",
        }
        pncp_items = [{"descricao": "Monitor LED"}]
        document_result = {
            "items": [{"descricao": "Teclado ergonômico USB"}],
            "error": "",
        }

        with (
            patch.object(server, "list_pncp_items", return_value=pncp_items),
            patch.object(server, "get_search_row_document_items", return_value=document_result) as document,
        ):
            result = server.filter_rows_by_complete_search_term([row], "teclado")

        self.assertEqual(result, [row])
        document.assert_called_once_with(row)
        server.SEARCH_ITEM_CACHE.clear()

    def test_search_does_not_download_document_when_api_already_matches(self):
        row = {
            "orgao_cnpj": "46422408000152",
            "ano": "2026",
            "numero_sequencial": "369",
        }
        with (
            patch.object(server, "list_pncp_items", return_value=[{"descricao": "Teclado USB"}]),
            patch.object(server, "get_search_row_document_items") as document,
        ):
            result = server.filter_rows_by_complete_search_term([row], "teclado")

        self.assertEqual(result, [row])
        document.assert_not_called()
        server.SEARCH_ITEM_CACHE.clear()

    def test_single_word_search_term_can_match_official_item_description(self):
        row = {
            "orgao_cnpj": "48813638000178",
            "ano": "2026",
            "numero_sequencial": "97",
            "title": "Pregao eletronico 28/2026",
            "description": "Eventual aquisicao de equipamentos de informatica",
        }
        pncp_items = [{
            "descricao": "Kit teclado e mouse USB com fio",
        }]

        server.SEARCH_ITEM_CACHE.clear()
        with patch.object(server, "list_pncp_items", return_value=pncp_items) as items:
            result = server.filter_rows_by_complete_search_term([row], "teclado")

        self.assertEqual(result, [row])
        items.assert_called_once_with("48813638000178", "2026", "97")
        server.SEARCH_ITEM_CACHE.clear()

    def test_search_term_uses_all_opportunity_and_item_fields_with_or_semantics(self):
        row = {
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "40",
            "title": "Edital",
            "description": "Aquisicao de suporte para teclado",
        }
        pncp_items = [{
            "descricao": "Suporte para monitor",
        }]

        server.SEARCH_ITEM_CACHE.clear()
        with patch.object(server, "list_pncp_items", return_value=pncp_items):
            result = server.filter_rows_by_complete_search_term([row], "teclado")

        self.assertEqual(result, [row])
        server.SEARCH_ITEM_CACHE.clear()

    def test_full_online_search_checks_item_fields_returned_by_pncp(self):
        row = {
            "id": "item-title-match",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "44",
            "orgao_nome": "Orgao de teste",
            "title": "Edital 44/2026",
            "description": "Aquisicao de equipamentos",
            "data_fim_vigencia": "2026-08-20T10:00:00",
        }

        def fake_request(url, timeout=18):
            if url.startswith(server.PNCP_SEARCH_URL):
                return {"items": [row], "total": 1}
            if "/compras/2026/44/itens" in url:
                return [{
                    "numeroItem": 1,
                    "materialOuServicoNome": "Monitor profissional",
                    "descricao": "Tela LED de 27 polegadas",
                }]
            raise AssertionError(f"URL inesperada: {url}")

        server.PNCP_RESULT_CACHE.clear()
        server.SEARCH_ITEM_CACHE.clear()
        with patch.object(server, "request_json", side_effect=fake_request):
            response = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP",
                "palavraChave": "monitor profissional",
                "pagina": "1",
                "tamanhoPagina": "10",
            })

        self.assertEqual(response["total"], 1)
        self.assertEqual(response["results"][0]["sequencial"], "44")
        server.PNCP_RESULT_CACHE.clear()
        server.SEARCH_ITEM_CACHE.clear()

    def test_material_filter_can_use_official_item_descriptions(self):
        row = {
            "orgao_cnpj": "46068425000133",
            "ano": "2026",
            "numero_sequencial": "1396",
            "title": "Pregao eletronico 1669/2026",
            "description": "PE DGA No 90220/2026",
        }
        pncp_items = [{
            "descricao": "Microcomputador com teclado e mouse",
        }]

        server.SEARCH_ITEM_CACHE.clear()
        with patch.object(server, "list_pncp_items", return_value=pncp_items):
            self.assertTrue(server.row_matches_object_type(row, "material"))
        server.SEARCH_ITEM_CACHE.clear()

    def test_multiple_keywords_are_combined_with_or_semantics(self):
        def row(identifier, description):
            return {
                "id": identifier,
                "orgao_cnpj": "12345678000199",
                "ano": "2026",
                "numero_sequencial": identifier.removeprefix("id-"),
                "orgao_nome": "Orgao",
                "title": "Edital",
                "description": description,
                "data_fim_vigencia": "2026-08-01T10:00:00",
            }

        shared = row("id-1", "Cadeira e monitor")

        def fake_request(url, timeout=18):
            query = parse_qs(urlparse(url).query)
            keyword = query.get("q", [""])[0]
            if keyword == "cadeira de rodas":
                return {"items": [shared, row("id-2", "Cadeira de rodas")], "total": 2}
            if keyword == "monitor":
                return {"items": [shared, row("id-3", "Monitor")], "total": 2}
            return {"items": [], "total": 0}

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(server, "request_json", side_effect=fake_request) as request:
            response = server.search_pncp_open_bids({
                "dataInicial": "20260725",
                "dataFinal": "20260823",
                "uf": "SP",
                "palavraChave": "cadeira de rodas;monitor",
                "pagina": "1",
                "tamanhoPagina": "10",
            })

        self.assertEqual(response["total"], 3)
        self.assertEqual(len(response["results"]), 3)
        requested_urls = [call.args[0] for call in request.call_args_list]
        self.assertTrue(any("q=cadeira+de+rodas" in url for url in requested_urls))
        self.assertTrue(any("q=monitor" in url for url in requested_urls))
        server.PNCP_RESULT_CACHE.clear()

    def test_purchase_number_and_uasg_filters_are_applied_together(self):
        matching = {
            "id": "match",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "1",
            "orgao_nome": "Órgão",
            "title": "Pregão eletrônico 90010/2026",
            "description": "Aquisição de mobiliário",
            "unidade_codigo": "00123456",
            "data_fim_vigencia": "2026-08-20T10:00:00",
        }
        wrong_uasg = {**matching, "id": "wrong-uasg", "unidade_codigo": "654321"}
        wrong_number = {**matching, "id": "wrong-number", "title": "Pregão 90011/2026"}

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(
            server,
            "request_json",
            return_value={"items": [matching, wrong_uasg, wrong_number], "total": 3},
        ):
            response = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP",
                "numeroCompra": "90010/2026",
                "uasg": "123456",
                "pagina": "1",
                "tamanhoPagina": "10",
            })

        self.assertEqual(response["total"], 1)
        self.assertEqual(response["results"][0]["numeroCompra"], "Pregão eletrônico 90010/2026")
        self.assertEqual(response["results"][0]["codigoUnidade"], "00123456")
        server.PNCP_RESULT_CACHE.clear()

    def test_search_job_key_includes_purchase_number_and_uasg(self):
        base = {"dataInicial": "20260801", "dataFinal": "20260830"}
        self.assertNotEqual(
            server.pncp_search_job_key({**base, "numeroCompra": "1/2026"}),
            server.pncp_search_job_key({**base, "numeroCompra": "2/2026"}),
        )
        self.assertNotEqual(
            server.pncp_search_job_key({**base, "uasg": "123456"}),
            server.pncp_search_job_key({**base, "uasg": "654321"}),
        )
        self.assertNotEqual(
            server.pncp_search_job_key({**base, "campoData": "publicacao"}),
            server.pncp_search_job_key({**base, "campoData": "encerramento"}),
        )

    def test_online_search_filters_the_selected_date_field(self):
        publication_match = {
            "id": "publication-match",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "1",
            "title": "Edital publicado em agosto",
            "description": "Aquisicao de mobiliario",
            "data_publicacao_pncp": "2026-08-10T09:00:00",
            "data_inicio_vigencia": "2026-09-01T09:00:00",
            "data_fim_vigencia": "2026-09-10T18:00:00",
        }
        closing_match = {
            "id": "closing-match",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "2",
            "title": "Edital encerrado em agosto",
            "description": "Aquisicao de mobiliario",
            "data_publicacao_pncp": "2026-07-10T09:00:00",
            "data_inicio_vigencia": "2026-07-15T09:00:00",
            "data_fim_vigencia": "2026-08-10T18:00:00",
        }

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(
            server,
            "request_json",
            return_value={"items": [publication_match, closing_match], "total": 2},
        ):
            by_publication = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "campoData": "publicacao",
                "uf": "SP",
            })
            by_closing = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "campoData": "encerramento",
                "uf": "SP",
            })

        self.assertEqual([row["sequencial"] for row in by_publication["results"]], ["1"])
        self.assertEqual([row["sequencial"] for row in by_closing["results"]], ["2"])
        self.assertEqual(by_publication["campoData"], "publicacao")
        server.PNCP_RESULT_CACHE.clear()

    def test_online_search_deduplicates_cards_for_the_same_contract(self):
        base = {
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "55",
            "title": "Edital 55",
            "description": "Aquisicao de cadeiras",
            "data_fim_vigencia": "2026-08-20T18:00:00",
        }
        duplicate_rows = [
            {**base, "id": "document-a", "item_url": "https://portal.test/a"},
            {**base, "id": "document-b", "item_url": "https://portal.test/b"},
        ]

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(
            server,
            "request_json",
            return_value={"items": duplicate_rows, "total": 2},
        ):
            response = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP",
            })

        self.assertEqual(response["total"], 1)
        self.assertEqual(len(response["results"]), 1)
        server.PNCP_RESULT_CACHE.clear()

    def test_search_reconciliation_preserves_existing_items_and_documents(self):
        persisted = []
        finished = []
        repository = SimpleNamespace(
            initialize=lambda: None,
            create_run=lambda *_args: "run-1",
            persist_record=lambda **kwargs: (
                persisted.append(kwargs) or ("inserted", "opportunity-1")
            ),
            save_failed_source_record=lambda **_kwargs: None,
            finish_run=lambda run_id, **kwargs: finished.append((run_id, kwargs)),
        )
        row = {
            "id": "pncp-search-1",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "77",
            "orgao_nome": "Orgao de teste",
            "title": "Edital 77",
            "description": "Aquisicao de cadeiras",
            "uf": "SP",
            "data_fim_vigencia": "2026-08-30T10:00:00",
        }

        with patch.object(server, "etl_repository", return_value=repository):
            result = server.reconcile_pncp_search_rows(
                [row],
                "https://pncp.gov.br/api/search",
                {"dataInicial": "20260801", "dataFinal": "20260830"},
            )

        self.assertEqual(result["inserted"], 1)
        self.assertEqual(result["failed"], 0)
        self.assertFalse(persisted[0]["replace_children"])
        self.assertEqual(persisted[0]["opportunity"].sequence, 77)
        self.assertEqual(finished[0][1]["status"], "success")

    def test_search_reconciliation_persists_items_used_by_keyword_filter(self):
        persisted = []
        repository = SimpleNamespace(
            initialize=lambda: None,
            create_run=lambda *_args: "run-items",
            persist_record=lambda **kwargs: (
                persisted.append(kwargs) or ("inserted", "opportunity-items")
            ),
            save_failed_source_record=lambda **_kwargs: None,
            finish_run=lambda *_args, **_kwargs: None,
        )
        row = {
            "id": "pncp-search-items",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "79",
            "orgao_nome": "Orgao de teste",
            "title": "Edital 79",
            "description": "Aquisicao de equipamentos",
        }
        cache_key = "12345678000199:2026:79"
        server.SEARCH_ITEM_CACHE.clear()
        server.cache_set(
            server.SEARCH_ITEM_CACHE,
            cache_key,
            [{"item": "1", "titulo": "Monitor", "descricao": "Tela LED"}],
        )

        with patch.object(server, "etl_repository", return_value=repository):
            result = server.reconcile_pncp_search_rows(
                [row],
                "https://pncp.gov.br/api/search",
                {"keywords": ["monitor"]},
            )

        self.assertEqual(result["inserted"], 1)
        self.assertEqual(len(persisted[0]["opportunity"].items), 1)
        self.assertEqual(persisted[0]["opportunity"].items[0].title, "Monitor")
        server.SEARCH_ITEM_CACHE.clear()

    def test_full_online_search_reconciles_when_requested(self):
        row = {
            "id": "pncp-search-2",
            "orgao_cnpj": "12345678000199",
            "ano": "2026",
            "numero_sequencial": "78",
            "orgao_nome": "Orgao de teste",
            "title": "Edital 78",
            "description": "Aquisicao de mesas",
            "uf": "SP",
            "data_fim_vigencia": "2026-08-20T10:00:00",
        }
        summary = {
            "run_id": "run-2",
            "status": "success",
            "fetched": 1,
            "inserted": 1,
            "updated": 0,
            "skipped": 0,
            "failed": 0,
        }
        official_summary = {
            "status": "success",
            "fetched": 0,
            "inserted": 0,
            "updated": 0,
            "skipped": 0,
            "failed": 0,
            "run_ids": [],
            "endpoints": [
                {"name": "proposta", "status": "success"},
                {"name": "publicacao", "status": "success"},
                {"name": "atualizacao", "status": "success"},
            ],
        }
        server.PNCP_RESULT_CACHE.clear()
        with (
            patch.object(server, "request_json", return_value={"items": [row], "total": 1}),
            patch.object(
                server,
                "reconcile_pncp_search_rows",
                return_value=summary,
            ) as reconcile,
            patch.object(
                server,
                "sync_pncp_opportunity_endpoints",
                return_value=official_summary,
            ),
        ):
            response = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP",
                "pagina": "1",
                "tamanhoPagina": "10",
                "reconciliar": "1",
            })

        reconcile.assert_called_once()
        self.assertEqual(response["reconciliation"]["inserted"], 1)
        self.assertEqual(
            {endpoint["name"] for endpoint in response["reconciliation"]["endpoints"]},
            {"api/search", "proposta", "publicacao", "atualizacao"},
        )
        self.assertTrue(response["complete"])
        server.PNCP_RESULT_CACHE.clear()

    def test_all_opportunity_endpoints_are_synced_without_details(self):
        requests = []

        class FakeService:
            def sync(self, request):
                requests.append(request)
                return {
                    "run_id": f"run-{request.endpoint}",
                    "status": "success",
                    "fetched": 1,
                    "inserted": 1,
                    "updated": 0,
                    "skipped": 0,
                    "failed": 0,
                }

        server.PNCP_OPPORTUNITY_SYNC_CACHE.clear()
        with patch.object(server, "ETLSyncService", return_value=FakeService()):
            result = server.sync_pncp_opportunity_endpoints(
                {
                    "uf": "SP",
                    "codigoModalidadeContratacao": "6",
                },
                "20260801",
                "20260819",
            )

        self.assertEqual({request.endpoint for request in requests}, {
            "proposta", "publicacao", "atualizacao",
        })
        self.assertTrue(all(request.fetch_details is False for request in requests))
        self.assertTrue(all(request.max_pages is None for request in requests))
        self.assertEqual(result["fetched"], 3)
        self.assertEqual(result["inserted"], 3)
        server.PNCP_OPPORTUNITY_SYNC_CACHE.clear()

    def test_all_official_modalities_continue_when_one_unit_fails(self):
        requests = []

        class FakeService:
            def sync(self, request):
                requests.append(request)
                modality = request.filters.get("codigoModalidadeContratacao")
                if request.endpoint == "publicacao" and modality == 2:
                    raise RuntimeError("falha temporaria")
                return {
                    "run_id": f"run-{request.endpoint}-{modality or 'all'}",
                    "status": "success",
                    "fetched": 1,
                    "inserted": 1,
                    "updated": 0,
                    "skipped": 0,
                    "failed": 0,
                }

        server.PNCP_OPPORTUNITY_SYNC_CACHE.clear()
        with patch.object(server, "ETLSyncService", return_value=FakeService()):
            result = server.sync_pncp_opportunity_endpoints(
                {},
                "20260827",
                "20260827",
            )

        publication_modalities = {
            request.filters.get("codigoModalidadeContratacao")
            for request in requests
            if request.endpoint == "publicacao"
        }
        update_modalities = {
            request.filters.get("codigoModalidadeContratacao")
            for request in requests
            if request.endpoint == "atualizacao"
        }
        endpoints = {item["name"]: item for item in result["endpoints"]}

        self.assertEqual(publication_modalities, set(server.PNCP_MODALITY_IDS))
        self.assertEqual(update_modalities, set(server.PNCP_MODALITY_IDS))
        self.assertEqual(endpoints["publicacao"]["fetched"], len(server.PNCP_MODALITY_IDS) - 1)
        self.assertEqual(endpoints["publicacao"]["failed"], 1)
        self.assertEqual(endpoints["publicacao"]["status"], "partial")
        self.assertEqual(endpoints["atualizacao"]["fetched"], len(server.PNCP_MODALITY_IDS))
        server.PNCP_OPPORTUNITY_SYNC_CACHE.clear()

    def test_modality_and_object_type_are_applied_to_official_search(self):
        rows = [
            {
                "id": "material-1",
                "orgao_cnpj": "12345678000199",
                "ano": "2026",
                "numero_sequencial": "1",
                "orgao_nome": "Orgao",
                "title": "Edital",
                "description": "Aquisicao de cadeiras escolares",
                "data_fim_vigencia": "2026-08-01T10:00:00",
            },
            {
                "id": "servico-1",
                "orgao_cnpj": "12345678000199",
                "ano": "2026",
                "numero_sequencial": "2",
                "orgao_nome": "Orgao",
                "title": "Edital",
                "description": "Contratacao de empresa especializada para limpeza predial",
                "data_fim_vigencia": "2026-08-01T10:00:00",
            },
        ]

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(server, "request_json", return_value={"items": rows, "total": 2}) as request:
            response = server.search_pncp_open_bids({
                "dataInicial": "20260725",
                "dataFinal": "20260823",
                "uf": "SP",
                "tipoObjeto": "servico",
                "codigoModalidadeContratacao": "6",
                "pagina": "1",
                "tamanhoPagina": "10",
            })

        self.assertEqual(response["total"], 1)
        self.assertEqual(response["results"][0]["sequencial"], "2")
        requested_url = next(
            call.args[0]
            for call in request.call_args_list
            if call.args[0].startswith(server.PNCP_SEARCH_URL)
        )
        query = parse_qs(urlparse(requested_url).query)
        self.assertEqual(query["modalidades"], ["6"])
        self.assertEqual(query["ufs"], ["SP"])
        server.PNCP_RESULT_CACHE.clear()

    def test_multiple_ufs_are_queried_and_combined_without_duplicates(self):
        def fake_request(url, timeout=18):
            uf = parse_qs(urlparse(url).query)["ufs"][0]
            return {
                "items": [{
                    "id": f"item-{uf}",
                    "orgao_cnpj": "12345678000199",
                    "ano": "2026",
                    "numero_sequencial": "1" if uf == "SP" else "2",
                    "orgao_nome": f"Orgao {uf}",
                    "title": "Edital",
                    "description": "Aquisicao de mobiliario",
                    "uf": uf,
                    "data_fim_vigencia": "2026-08-20T10:00:00",
                }],
                "total": 1,
            }

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(server, "request_json", side_effect=fake_request) as request:
            response = server.search_pncp_open_bids({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP,RJ",
                "pagina": "1",
                "tamanhoPagina": "10",
            })

        requested_ufs = {
            parse_qs(urlparse(call.args[0]).query)["ufs"][0]
            for call in request.call_args_list
        }
        self.assertEqual(requested_ufs, {"SP", "RJ"})
        self.assertEqual(response["total"], 2)
        self.assertEqual({row["uf"] for row in response["results"]}, {"SP", "RJ"})
        server.PNCP_RESULT_CACHE.clear()

    def test_quick_preview_queries_each_selected_uf(self):
        def fake_request(url, timeout=12):
            uf = parse_qs(urlparse(url).query)["ufs"][0]
            return {
                "items": [{
                    "id": f"preview-{uf}",
                    "orgao_cnpj": "12345678000199",
                    "ano": "2026",
                    "numero_sequencial": "1" if uf == "SP" else "2",
                    "orgao_nome": f"Orgao {uf}",
                    "title": "Edital",
                    "description": "Aquisicao de mobiliario",
                    "uf": uf,
                    "data_fim_vigencia": "2026-08-20T10:00:00",
                }],
                "total": 1,
            }

        with patch.object(server, "request_json", side_effect=fake_request) as request:
            response = server.quick_pncp_search_preview({
                "dataInicial": "20260801",
                "dataFinal": "20260830",
                "uf": "SP,RJ",
            })

        requested_ufs = {
            parse_qs(urlparse(call.args[0]).query)["ufs"][0]
            for call in request.call_args_list
        }
        self.assertEqual(requested_ufs, {"SP", "RJ"})
        self.assertEqual({row["uf"] for row in response["results"]}, {"SP", "RJ"})

    def test_invalid_object_type_is_rejected_before_requesting_pncp(self):
        with patch.object(server, "request_json") as request:
            with self.assertRaisesRegex(ValueError, "Tipo do objeto"):
                server.search_pncp_open_bids({
                    "dataInicial": "20260725",
                    "dataFinal": "20260823",
                    "tipoObjeto": "obra",
                })

        request.assert_not_called()

    def test_invalid_date_field_is_rejected_before_requesting_pncp(self):
        with patch.object(server, "request_json") as request:
            with self.assertRaisesRegex(ValueError, "Campo de data"):
                server.search_pncp_open_bids({
                    "dataInicial": "20260725",
                    "dataFinal": "20260823",
                    "campoData": "assinatura",
                })

        request.assert_not_called()

    def test_fast_search_returns_preview_then_completed_result(self):
        preview = {
            "results": [{"processo": "preview"}],
            "total": 1,
            "pagina": 1,
            "tamanhoPagina": 10,
            "total_pages": 1,
            "complete": False,
            "searching": True,
        }
        complete = {
            "results": [{"processo": "complete"}],
            "total": 120,
            "pagina": 1,
            "tamanhoPagina": 10,
            "total_pages": 12,
            "complete": True,
        }
        params = {
            "dataInicial": "20260725",
            "dataFinal": "20260823",
            "codigoModalidadeContratacao": "6",
        }
        server.PNCP_SEARCH_JOBS.clear()
        with (
            patch.object(server, "quick_pncp_search_preview", return_value=preview),
            patch.object(server, "search_pncp_open_bids", return_value=complete),
        ):
            first = server.search_pncp_open_bids_fast(params)
            for _ in range(50):
                with server.PNCP_SEARCH_JOB_LOCK:
                    status = next(iter(server.PNCP_SEARCH_JOBS.values()))["status"]
                if status == "complete":
                    break
                time.sleep(0.01)
            second = server.search_pncp_open_bids_fast(params)

        self.assertTrue(first["searching"])
        self.assertFalse(second["searching"])
        self.assertEqual(second["total"], 120)
        server.PNCP_SEARCH_JOBS.clear()

    def test_search_returns_full_requested_page_and_pagination_metadata(self):
        rows = [
            {
                "orgao_cnpj": f"{index:014d}",
                "ano": "2026",
                "numero_sequencial": str(index),
                "orgao_nome": f"Orgao {index}",
                "title": f"Edital {index}",
                "description": f"Objeto {index}",
                "data_fim_vigencia": "2026-08-01T10:00:00",
            }
            for index in range(1, 126)
        ]

        with patch.object(server, "request_json", return_value={"items": rows, "total": 125}) as request:
            server.PNCP_RESULT_CACHE.clear()
            response = server.search_pncp_open_bids({
                "dataInicial": "20260725",
                "dataFinal": "20260823",
                "uf": "SP",
                "pagina": "2",
                "tamanhoPagina": "50",
            })

        self.assertEqual(len(response["results"]), 50)
        self.assertEqual(response["pagina"], 2)
        self.assertEqual(response["tamanhoPagina"], 50)
        self.assertEqual(response["total_pages"], 3)
        self.assertTrue(response["has_previous"])
        self.assertTrue(response["has_next"])
        requested_url = request.call_args.args[0]
        self.assertIn("pagina=1", requested_url)
        self.assertIn("tam_pagina=500", requested_url)
        server.PNCP_RESULT_CACHE.clear()

    def test_search_collects_all_source_pages_before_filtering_and_deduplicating(self):
        def make_rows(start, end):
            return [
                {
                    "id": f"id-{index}",
                    "orgao_cnpj": f"{index:014d}",
                    "ano": "2026",
                    "numero_sequencial": str(index),
                    "orgao_nome": f"Orgao {index}",
                    "title": f"Edital {index}",
                    "description": f"Aquisicao de material {index}",
                    "data_fim_vigencia": "2026-08-01T10:00:00",
                }
                for index in range(start, end)
            ]

        def fake_request(url, timeout=18):
            page = int(parse_qs(urlparse(url).query)["pagina"][0])
            if page == 1:
                return {"items": make_rows(1, 501), "total": 750}
            return {"items": make_rows(500, 751), "total": 750}

        server.PNCP_RESULT_CACHE.clear()
        with patch.object(server, "request_json", side_effect=fake_request):
            response = server.search_pncp_open_bids({
                "dataInicial": "20260725",
                "dataFinal": "20260823",
                "uf": "SP",
                "tipoObjeto": "material",
                "pagina": "1",
                "tamanhoPagina": "50",
            })

        self.assertEqual(response["source_total"], 750)
        self.assertEqual(response["total"], 750)
        self.assertEqual(response["total_pages"], 15)
        self.assertEqual(response["pages_checked"], 2)
        self.assertTrue(response["complete"])
        self.assertEqual(len(response["results"]), 50)
        server.PNCP_RESULT_CACHE.clear()


class OpportunityDetailTests(unittest.TestCase):
    def test_detail_uses_local_record_without_waiting_for_remote_metadata(self):
        local = {"opportunity": {"id": "c" * 32}, "items": [{}], "documents": []}
        expected = {"itens": [{"numero": "1", "descricao": "Item local"}]}
        repository = SimpleNamespace(
            get_opportunity_by_pncp_identity=lambda *_args: local,
        )
        with (
            patch.object(server, "etl_repository", return_value=repository),
            patch.object(server, "internal_opportunity_detail", return_value=expected) as detail_mock,
            patch.object(server, "pncp_purchase_metadata") as metadata_mock,
        ):
            result = server.opportunity_detail_from_pncp_link(
                "https://pncp.gov.br/app/editais/00394700000108/2026/250"
            )

        self.assertEqual(result, expected)
        detail_mock.assert_called_once_with("c" * 32)
        metadata_mock.assert_not_called()

    def test_detail_combines_official_pncp_metadata_files_and_items(self):
        metadata = {
            "numero_compra": "164/2026",
            "processo": "123/2026",
            "modalidade": "Dispensa Eletrônica",
            "objeto": "Aquisição de cadeiras e apoio de punho",
            "orgao": "Órgão de teste",
            "orgao_cnpj": "00394700000108",
            "unidade": "Unidade compradora",
            "municipio": "Brasília",
            "uf": "DF",
            "numero_controle_pncp": "controle",
            "abertura": "2026-08-05T08:30:00",
            "encerramento": "2026-08-05T14:30:00",
            "situacao": "Divulgada",
            "valor_total_estimado": 8283.58,
            "modo_disputa": "Dispensa com Disputa",
            "codigo_unidade": "102329",
            "link_sistema_origem": "https://compras.gov.br/compra/164",
        }
        raw_items = [{
            "numeroItem": 1,
            "descricao": "Cadeira ergonômica",
            "quantidade": 2,
            "unidadeMedida": "UNIDADE",
            "valorUnitarioEstimado": 100,
            "criterioJulgamentoNome": "Menor preço",
            "situacaoCompraItemNome": "Aberto",
        }]
        raw_files = [{
            "titulo": "Termo de Referência",
            "tipoDocumentoNome": "Termo de Referência",
            "url": "https://pncp.gov.br/arquivo.pdf",
        }]

        with (
            patch.object(server, "etl_repository", return_value=SimpleNamespace(get_opportunity_by_pncp_identity=lambda *_args: None)),
            patch.object(server, "pncp_purchase_metadata", return_value=metadata),
            patch.object(server, "identify_items_from_pncp_link", return_value={"items": raw_items}),
            patch.object(server, "list_pncp_files", return_value=raw_files),
        ):
            result = server.opportunity_detail_from_pncp_link(
                "https://pncp.gov.br/app/editais/00394700000108/2026/251"
            )

        self.assertEqual(result["oportunidade"]["portal_origem"], "Comprasnet")
        self.assertEqual(result["oportunidade"]["codigo_unidade"], "102329")
        self.assertIn("Mobiliário", result["oportunidade"]["categorias"])
        self.assertIn("Acessórios ergonômicos", result["oportunidade"]["categorias"])
        self.assertEqual(result["itens"][0]["valor_total_estimado"], 200.0)
        self.assertEqual(result["arquivos"][0]["titulo"], "Termo de Referência")

    def test_detail_prioritizes_document_items_and_reports_divergence(self):
        raw_items = [
            {
                "numeroItem": 1,
                "descricao": "Descrição resumida da API",
                "quantidade": 1,
                "unidadeMedida": "UNIDADE",
                "valorUnitarioEstimado": 50,
            },
            {
                "numeroItem": 3,
                "descricao": "Item existente somente na API",
                "quantidade": 4,
                "unidadeMedida": "UNIDADE",
            },
        ]
        file_items = [
            make_item("1", "Descrição completa do documento", quantidade="2"),
            make_item("2", "Item existente somente no documento", quantidade="3"),
        ]
        source = {
            "source_path": Path("termo.pdf"),
            "pncp": {"documento_usado": "Termo de Referência.pdf"},
        }

        with (
            patch.object(server, "etl_repository", return_value=SimpleNamespace(get_opportunity_by_pncp_identity=lambda *_args: None)),
            patch.object(server, "pncp_purchase_metadata", return_value={}),
            patch.object(server, "identify_items_from_pncp_link", return_value={"items": raw_items}),
            patch.object(server, "list_pncp_files", return_value=[]),
            patch.object(server, "source_for_opportunity_detail", return_value=source),
            patch.object(server, "extract_items_cached", return_value=file_items),
        ):
            result = server.opportunity_detail_from_pncp_link(
                "https://pncp.gov.br/app/editais/00394700000108/2026/252"
            )

        self.assertEqual(len(result["itens"]), 3)
        self.assertEqual(result["itens"][0]["descricao"], "Descrição completa do documento")
        self.assertEqual(result["itens"][0]["valor_unitario_estimado"], 50)
        self.assertEqual(result["verificacao_itens"]["file_count"], 2)
        self.assertEqual(result["verificacao_itens"]["pncp_count"], 2)
        self.assertTrue(result["verificacao_itens"]["has_divergence"])
        self.assertEqual(result["verificacao_itens"]["source"], "documento_oficial")

    def test_detail_falls_back_to_api_when_document_cannot_be_read(self):
        raw_items = [{"numeroItem": 1, "descricao": "Item da API", "quantidade": 1}]
        with (
            patch.object(server, "etl_repository", return_value=SimpleNamespace(get_opportunity_by_pncp_identity=lambda *_args: None)),
            patch.object(server, "pncp_purchase_metadata", return_value={}),
            patch.object(server, "identify_items_from_pncp_link", return_value={"items": raw_items}),
            patch.object(server, "list_pncp_files", return_value=[]),
            patch.object(server, "source_for_opportunity_detail", side_effect=RuntimeError("arquivo indisponível")),
        ):
            result = server.opportunity_detail_from_pncp_link(
                "https://pncp.gov.br/app/editais/00394700000108/2026/253"
            )

        self.assertEqual(len(result["itens"]), 1)
        self.assertEqual(result["itens"][0]["descricao"], "Item da API")
        self.assertEqual(result["verificacao_itens"]["source"], "api_pncp")
        self.assertIn("arquivo indisponível", result["verificacao_itens"]["file_error"])

    def test_document_question_returns_only_matching_official_excerpts(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            document = Path(temp_dir) / "termo.pdf"
            document.write_bytes(b"%PDF-test")
            source = {
                "source_path": document,
                "pncp": {
                    "documento_usado": "Termo de Referência.pdf",
                    "documento_tipo": "Termo de Referência",
                },
            }
            text = (
                "O prazo de entrega será de 15 dias úteis após o recebimento da ordem.\n\n"
                "A garantia mínima dos equipamentos será de 12 meses."
            )
            with (
                patch.object(server, "source_from_pncp_link", return_value=source),
                patch.object(server, "catalog_document_text", return_value=text),
            ):
                result = server.answer_opportunity_question("link", "Qual é o prazo de entrega?")

        self.assertEqual(result["tipo_documento"], "Termo de Referência")
        self.assertEqual(len(result["trechos"]), 1)
        self.assertIn("15 dias úteis", result["trechos"][0])


class BusinessItemSelectionTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.data_dir = Path(self.temp_dir.name)
        self.database_path_patch = patch.object(
            server, "DATABASE_PATH", self.data_dir / "pncp.sqlite3"
        )
        self.data_dir_patch = patch.object(server, "DATA_DIR", self.data_dir)
        self.database_path_patch.start()
        self.data_dir_patch.start()
        server.init_database()
        server.BUSINESS_FILE_CACHE.clear()
        self.link = "https://pncp.gov.br/app/editais/18428888000123/2026/138"
        self.metadata = {
            "numero_compra": "138/2026",
            "processo": "123/2026",
            "modalidade": "Pregão Eletrônico",
            "objeto": "Aquisição de mobiliário",
            "orgao": "Órgão de teste",
            "orgao_cnpj": "18428888000123",
            "unidade": "Unidade compradora",
            "municipio": "São Paulo",
            "uf": "SP",
            "numero_controle_pncp": "controle",
            "abertura": "2026-08-05T08:30:00",
            "encerramento": "2026-08-05T14:30:00",
            "situacao": "Divulgada",
        }

    def tearDown(self):
        self.data_dir_patch.stop()
        self.database_path_patch.stop()
        self.temp_dir.cleanup()

    def item(self, number, description):
        return {
            "numero": str(number),
            "lote": "1",
            "descricao": description,
            "quantidade": "2",
            "unidade": "UND",
            "valor_unitario_estimado": 100,
            "valor_total_estimado": 200,
            "criterio_julgamento": "Menor preço",
            "situacao": "Aberto",
        }

    def test_only_selected_items_are_persisted_and_returned_in_business_detail(self):
        with patch.object(server, "pncp_purchase_metadata", return_value=self.metadata):
            created = server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(1, "Cadeira"), self.item(3, "Armário")],
            })

        detail = server.get_business(created["negocio"]["id"], include_details=True)
        self.assertEqual(detail["total_itens"], 2)
        self.assertEqual([item["numero"] for item in detail["itens"]], ["1", "3"])
        self.assertEqual(detail["itens"][1]["descricao"], "Armário")

    def test_new_selection_replaces_items_of_existing_business(self):
        with patch.object(server, "pncp_purchase_metadata", return_value=self.metadata):
            first = server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(1, "Cadeira"), self.item(3, "Armário")],
            })
            server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(2, "Mesa")],
            })

        detail = server.get_business(first["negocio"]["id"], include_details=True)
        self.assertEqual(detail["total_itens"], 1)
        self.assertEqual(detail["itens"][0]["numero"], "2")

    def test_empty_item_selection_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "Selecione ao menos um item"):
            server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [],
            })


    def test_position_is_defined_on_main_business_record(self):
        with patch.object(server, "pncp_purchase_metadata", return_value=self.metadata):
            created = server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(1, "Cadeira")],
            })

        business_id = created["negocio"]["id"]
        updated = server.update_business(business_id, {"position_number": 2})
        self.assertEqual(updated["position_number"], 2)

        cleared = server.update_business(business_id, {"position_number": None})
        self.assertIsNone(cleared["position_number"])

    def test_classification_position_is_synchronized_back_to_main_screen(self):
        with patch.object(server, "pncp_purchase_metadata", return_value=self.metadata):
            created = server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(1, "Cadeira")],
            })

        business_id = created["negocio"]["id"]
        server.sync_business_position_from_proposal({
            "business_id": int(business_id),
            "position_number": 3,
        })

        self.assertEqual(server.get_business(business_id)["position_number"], 3)

    def test_remote_document_lookup_does_not_hold_the_database_lock(self):
        with patch.object(server, "pncp_purchase_metadata", return_value=self.metadata):
            created = server.import_business({
                "pncp_link": self.link,
                "empresa": "Empresa Teste",
                "itens": [self.item(1, "Cadeira")],
            })

        entered = threading.Event()

        def slow_remote_lookup(*_args):
            entered.set()
            time.sleep(0.5)
            return []

        with patch.object(server, "list_pncp_files", side_effect=slow_remote_lookup) as remote_lookup:
            worker = threading.Thread(
                target=server.get_business,
                args=(created["negocio"]["id"], True),
                daemon=True,
            )
            worker.start()
            self.assertTrue(entered.wait(2))
            started = time.perf_counter()
            server.list_responsibles()
            elapsed = time.perf_counter() - started
            worker.join(2)
            cached_started = time.perf_counter()
            server.get_business(created["negocio"]["id"], True)
            cached_elapsed = time.perf_counter() - cached_started

        self.assertFalse(worker.is_alive())
        self.assertLess(elapsed, 0.25)
        self.assertLess(cached_elapsed, 0.25)
        remote_lookup.assert_called_once()

if __name__ == "__main__":
    unittest.main()
