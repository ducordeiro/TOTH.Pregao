import csv
import html
import json
import re
import uuid
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from catalog_rules import (
    analyze_catalog_item,
    apply_user_catalog_repertoire,
    build_catalog_entries,
    catalog_repertoire_key,
    catalog_policy_summary,
    catalog_summary,
    repertoire_summary,
)


MANUFACTURER = {
    "razao_social": "GOLDFLEX INDUSTRIA E COMERCIO DE MOVEIS E EQUIPAMENTOS LTDA",
    "cnpj": "33.661.439/0001-14",
}

COMPUTED_EXPORT_KEYS = {"modelo_catalogo", "resultado_aderencia", "status_catalogo"}

EXPORT_COLUMNS = (
    ("numero", "Item"),
    ("lote", "Lote"),
    ("codigo", "Código"),
    ("produto", "Produto"),
    ("descricao", "Descrição da oportunidade"),
    ("especificacao_tecnica", "Especificação da oportunidade"),
    ("unidade", "Unidade"),
    ("quantidade", "Quantidade"),
    ("marca_referencia", "Marca/referência"),
    ("valor_estimado", "Valor estimado"),
    ("criterios_aceitacao", "Critérios de aceitação"),
    ("observacoes", "Observações"),
    ("categoria", "Categoria"),
    ("subcategoria", "Subcategoria"),
    ("status_evidencia", "Evidência da oportunidade"),
    ("modelo_catalogo", "Modelo técnico de referência"),
    ("resultado_aderencia", "Resultado da análise"),
    ("status_catalogo", "Status do catálogo"),
)


def compact(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def required_field_gaps(item):
    return [
        label
        for field, label in (
            ("descricao", "descrição"),
            ("quantidade", "quantidade"),
            ("unidade", "unidade"),
        )
        if not compact(item.get(field))
    ]


def refresh_validation(item):
    refreshed = dict(item)
    missing = required_field_gaps(refreshed)
    refreshed["campos_ausentes"] = missing
    refreshed["status_evidencia"] = "incompleto" if missing else "confirmado"
    return refreshed


def item_number(item, index):
    return compact(item.get("item") or item.get("numero") or item.get("numeroItem")) or str(index + 1)


def classify_item(description):
    text = compact(description).lower()
    rules = (
        (("cadeira", "poltrona", "longarina", "banqueta"), ("Mobiliário", "Assentos")),
        (("mesa", "estação de trabalho"), ("Mobiliário", "Mesas")),
        (("armário", "arquivo", "gaveteiro", "estante"), ("Mobiliário", "Armazenamento")),
        (("computador", "notebook", "monitor"), ("Tecnologia", "Equipamentos")),
    )
    for terms, category in rules:
        if any(term in text for term in terms):
            return category
    return "Não classificado", ""


def normalize_items(items, source_url, source_name="API oficial do PNCP"):
    normalized = []
    seen = set()
    for index, raw in enumerate(items or []):
        description = compact(
            raw.get("descricao") or raw.get("description") or raw.get("objeto")
        )
        number = item_number(raw, index)
        lot = compact(raw.get("lote"))
        key = (lot, number, description.lower())
        if key in seen:
            continue
        seen.add(key)
        category, subcategory = classify_item(description)
        item_source = compact(raw.get("_catalog_source_name")) or source_name
        item = refresh_validation({
            "id": f"item-{index + 1}",
            "numero": number,
            "lote": lot,
            "codigo": compact(raw.get("codigo") or raw.get("codigoItem") or raw.get("material")),
            "produto": compact(raw.get("produto") or raw.get("nome") or raw.get("_nome")) or description[:120],
            "descricao": description,
            "especificacao_tecnica": compact(raw.get("especificacao_tecnica") or raw.get("especificacao") or description),
            "unidade": compact(raw.get("unidade") or raw.get("unidadeMedida")),
            "quantidade": compact(raw.get("quantidade") or raw.get("quantidadeItem")),
            "marca_referencia": compact(raw.get("marca_referencia") or raw.get("marca")),
            "valor_estimado": raw.get("valor_estimado") or raw.get("valorUnitarioEstimado") or "",
            "criterios_aceitacao": compact(raw.get("criterios_aceitacao")),
            "observacoes": compact(raw.get("observacoes")),
            "categoria": category,
            "subcategoria": subcategory,
            "conflitos": [],
            "fontes": [{
                "documento": item_source,
                "pagina": None,
                "secao": f"Item {number}" + (f" · Lote {lot}" if lot else ""),
                "url": source_url,
            }],
        })
        normalized.append(analyze_catalog_item(item))
    return normalized


def validation_summary(items):
    validated = []
    for item in items:
        if not isinstance(item, dict):
            raise ValueError("Cada item do catálogo deve ser um objeto válido.")
        validated.append(refresh_validation(item))
    incomplete = sum(bool(item["campos_ausentes"]) for item in validated)
    conflicts = sum(bool(item.get("conflitos")) for item in validated)
    warnings = []
    if incomplete:
        warnings.append(f"{incomplete} item(ns) possuem campos obrigatórios ausentes.")
    if conflicts:
        warnings.append(f"{conflicts} item(ns) possuem divergências entre fontes.")
    if not items:
        warnings.append("Nenhum item foi identificado nas fontes disponíveis.")
    return {"incompletos": incomplete, "conflitos": conflicts, "avisos": warnings}


def sanitize_export_items(items):
    clean = []
    for index, raw in enumerate(items or []):
        if not isinstance(raw, dict):
            raise ValueError("Cada item do catálogo deve ser um objeto válido.")
        item = {
            key: raw.get(key, "")
            if isinstance(raw.get(key, ""), (str, int, float))
            else ""
            for key, _ in EXPORT_COLUMNS
            if key not in COMPUTED_EXPORT_KEYS
        }
        item["id"] = compact(raw.get("id")) or f"item-{index + 1}"
        item["numero"] = compact(item["numero"]) or str(index + 1)
        item["fontes"] = raw.get("fontes") if isinstance(raw.get("fontes"), list) else []
        item["conflitos"] = raw.get("conflitos") if isinstance(raw.get("conflitos"), list) else []
        refreshed = refresh_validation(item)
        user_repertoire = raw.get("repertorio_usuario")
        if (
            isinstance(user_repertoire, dict)
            and user_repertoire.get("item_key") == catalog_repertoire_key(refreshed)
        ):
            analyzed = apply_user_catalog_repertoire(refreshed, user_repertoire)
        else:
            analyzed = analyze_catalog_item(refreshed)
        reference = analyzed.get("modelo_referencia") or {}
        fit = analyzed.get("analise_aderencia") or {}
        analyzed["modelo_catalogo"] = reference.get("nome", "")
        analyzed["resultado_aderencia"] = fit.get("resultado", "")
        clean.append(analyzed)
    return clean


def prepare_catalog_items(items):
    return sanitize_export_items(items)


def format_number(value):
    number = float(value)
    if number.is_integer():
        return str(int(number))
    return f"{number:.1f}".replace(".", ",")


def format_dimension(dimension):
    values = dimension["valores_mm"]
    use_millimetres = dimension["parte"] in {"altura_assento", "altura_total"}
    if use_millimetres:
        display_values = [format_number(value) for value in values]
        unit = "mm"
    else:
        display_values = [format_number(value / 10.0) for value in values]
        unit = "cm"
    separator = " a " if dimension["parte"] in {"altura_assento", "altura_total"} else " x "
    prefix = "aprox. " if dimension["aproximada"] else ""
    return f"{dimension['rotulo']}: {prefix}{separator.join(display_values)} {unit}"


def catalog_characteristics(entry):
    characteristics = list(entry["caracteristicas"])
    characteristics.extend(format_dimension(dimension) for dimension in entry["dimensoes"])
    if entry.get("capacidade_kg") is not None:
        characteristics.append(
            f"Capacidade declarada no repertório: {format_number(entry['capacidade_kg'])} kg."
        )
    if entry.get("normas"):
        characteristics.append("Normas declaradas: " + ", ".join(entry["normas"]) + ".")
    return characteristics


def set_cell_shading(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def configure_docx(document):
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in (
        ("Title", 25, "232323"),
        ("Heading 1", 18, "232323"),
        ("Heading 2", 13, "232323"),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("GOLDFLEX")
    header_run.bold = True
    header_run.font.name = "Arial"
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor.from_string("666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Documento técnico para revisão")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string("777777")


def add_docx_band(document, text):
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFC000")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("222222")


def add_catalog_paragraph(document, text="", style=None, *, bold=False, size=None):
    paragraph = document.add_paragraph()
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        if size is not None:
            run.font.size = Pt(size)
    return paragraph


def write_catalog_docx(path, entries, template_path=None):
    document = Document(str(template_path)) if template_path else Document()
    marker = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "{CATALOGO}"),
        None,
    )
    body = document._body._element
    original_elements = set(body.iterchildren())
    template_has_content = bool(
        template_path
        and (
            any(paragraph.text.strip() for paragraph in document.paragraphs if paragraph is not marker)
            or document.tables
        )
    )

    if not template_path:
        configure_docx(document)
    elif marker is None and template_has_content:
        document.add_page_break()
    if not document.core_properties.title:
        document.core_properties.title = "Catálogo técnico Goldflex"
    if not document.core_properties.subject:
        document.core_properties.subject = "Catálogo genérico por modelo"
    if not document.core_properties.author:
        document.core_properties.author = MANUFACTURER["razao_social"]

    title = add_catalog_paragraph(document, style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Catálogo técnico Goldflex")
    title_run.bold = True
    title_run.font.size = Pt(25)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Modelos técnicos consolidados para revisão")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor.from_string("666666")
    document.add_paragraph()
    add_docx_band(document, "GOLDFLEX")
    manufacturer = document.add_paragraph()
    manufacturer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    manufacturer.add_run(MANUFACTURER["razao_social"]).bold = True
    manufacturer.add_run(f"\nCNPJ {MANUFACTURER['cnpj']}")
    status = document.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_run = status.add_run("RASCUNHO TÉCNICO — REVISÃO HUMANA OBRIGATÓRIA")
    status_run.bold = True
    status_run.font.color.rgb = RGBColor.from_string("9C6500")

    if not entries:
        document.add_page_break()
        add_docx_band(document, "MODELO NÃO IDENTIFICADO")
        warning = document.add_paragraph()
        warning.add_run(
            "Nenhuma característica técnica foi publicada porque o repertório estruturado "
            "não apresentou um modelo com evidência suficiente."
        )
        document.add_paragraph(
            "A liberação permanece bloqueada até a identificação do modelo e a validação das evidências."
        )
    for entry in entries:
        document.add_page_break()
        heading = add_catalog_paragraph(document, style="Heading 1")
        heading_run = heading.add_run(entry["nome"])
        heading_run.bold = True
        heading_run.font.size = Pt(18)
        family = document.add_paragraph()
        family_run = family.add_run(entry["familia"])
        family_run.bold = True
        family_run.font.color.rgb = RGBColor.from_string("666666")
        add_docx_band(document, "Características")
        for value in catalog_characteristics(entry):
            add_catalog_paragraph(document, value, style="List Bullet")
        add_catalog_paragraph(
            document,
            "Documentação e referência",
            style="Heading 2",
            bold=True,
            size=13,
        )
        document.add_paragraph(entry["fonte"])
        add_catalog_paragraph(
            document,
            "Dados a confirmar",
            style="Heading 2",
            bold=True,
            size=13,
        )
        for value in entry["pendencias"]:
            add_catalog_paragraph(document, value, style="List Bullet")
        review = document.add_paragraph()
        review_run = review.add_run("Status: rascunho para revisão humana.")
        review_run.bold = True
        review_run.font.color.rgb = RGBColor.from_string("9C6500")

    if marker is not None:
        generated_elements = [
            element
            for element in body.iterchildren()
            if element not in original_elements and element.tag != qn("w:sectPr")
        ]
        for element in generated_elements:
            marker._p.addprevious(element)
        marker._element.getparent().remove(marker._element)

    document.save(path)


def pdf_header_footer(canvas, document):
    canvas.saveState()
    width, height = A4
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.setFont("Helvetica-Bold", 8)
    canvas.drawRightString(width - 18 * mm, height - 10 * mm, "GOLDFLEX")
    canvas.setFont("Helvetica", 7.5)
    canvas.drawCentredString(
        width / 2,
        9 * mm,
        f"Documento técnico para revisão - Página {document.page}",
    )
    canvas.restoreState()


def write_catalog_pdf(path, entries):
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "GoldflexTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=24,
        leading=29,
        textColor=colors.HexColor("#232323"),
        alignment=1,
        spaceAfter=8 * mm,
    )
    subtitle_style = ParagraphStyle(
        "GoldflexSubtitle",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#666666"),
        alignment=1,
        spaceAfter=7 * mm,
    )
    heading_style = ParagraphStyle(
        "GoldflexHeading",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=17,
        leading=21,
        textColor=colors.HexColor("#232323"),
        spaceAfter=3 * mm,
    )
    section_style = ParagraphStyle(
        "GoldflexSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#232323"),
    )
    body_style = ParagraphStyle(
        "GoldflexBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor("#333333"),
        spaceAfter=2.2 * mm,
    )
    bullet_style = ParagraphStyle(
        "GoldflexBullet",
        parent=body_style,
        leftIndent=5 * mm,
        firstLineIndent=-3 * mm,
        bulletIndent=1 * mm,
    )
    status_style = ParagraphStyle(
        "GoldflexStatus",
        parent=body_style,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#9C6500"),
        alignment=1,
        spaceBefore=4 * mm,
    )
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title="Catálogo técnico Goldflex",
        author=MANUFACTURER["razao_social"],
    )
    story = [
        Spacer(1, 27 * mm),
        Paragraph("Catálogo técnico Goldflex", title_style),
        Paragraph("Modelos técnicos consolidados para revisão", subtitle_style),
        Table(
            [[Paragraph("<b>GOLDFLEX</b>", section_style)]],
            colWidths=[174 * mm],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFC000")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5 * mm),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5 * mm),
                ("TOPPADDING", (0, 0), (-1, -1), 3 * mm),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm),
            ]),
        ),
        Spacer(1, 7 * mm),
        Paragraph(html.escape(MANUFACTURER["razao_social"]), subtitle_style),
        Paragraph(f"CNPJ {MANUFACTURER['cnpj']}", subtitle_style),
        Paragraph("RASCUNHO TÉCNICO — REVISÃO HUMANA OBRIGATÓRIA", status_style),
    ]
    if not entries:
        story.extend([
            PageBreak(),
            Table(
                [[Paragraph("<b>MODELO NÃO IDENTIFICADO</b>", section_style)]],
                colWidths=[174 * mm],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFC000")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5 * mm),
                    ("TOPPADDING", (0, 0), (-1, -1), 3 * mm),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm),
                ]),
            ),
            Spacer(1, 6 * mm),
            Paragraph(
                "Nenhuma característica técnica foi publicada porque o repertório estruturado "
                "não apresentou um modelo com evidência suficiente.",
                body_style,
            ),
            Paragraph(
                "A liberação permanece bloqueada até a identificação do modelo e a validação das evidências.",
                body_style,
            ),
        ])
    for entry in entries:
        story.extend([
            PageBreak(),
            Paragraph(html.escape(entry["nome"]), heading_style),
            Paragraph(f"<b>{html.escape(entry['familia'])}</b>", body_style),
            Table(
                [[Paragraph("<b>Características</b>", section_style)]],
                colWidths=[174 * mm],
                style=TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFC000")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5 * mm),
                    ("TOPPADDING", (0, 0), (-1, -1), 3 * mm),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3 * mm),
                ]),
            ),
            Spacer(1, 4 * mm),
        ])
        story.extend(
            Paragraph(f"• {html.escape(value)}", bullet_style)
            for value in catalog_characteristics(entry)
        )
        story.extend([
            Spacer(1, 2 * mm),
            Paragraph("Documentação e referência", section_style),
            Paragraph(html.escape(entry["fonte"]), body_style),
            Paragraph("Dados a confirmar", section_style),
        ])
        story.extend(
            Paragraph(f"• {html.escape(value)}", bullet_style)
            for value in entry["pendencias"]
        )
        story.append(Paragraph("Status: rascunho para revisão humana.", status_style))

    document.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def write_audit_json(path, metadata, safe_items, entries):
    path.write_text(json.dumps({
        "schema_version": "2.0",
        "fabricante": MANUFACTURER,
        "regras_catalogo": catalog_policy_summary(),
        "repertorio": repertoire_summary(),
        "catalogos": entries,
        "auditoria_oportunidade": {
            "edital": metadata,
            "itens": safe_items,
            "resumo": catalog_summary(safe_items),
            "nota": (
                "A análise identifica referências e pendências. Ela não constitui declaração "
                "automática de atendimento ao edital."
            ),
        },
        "gerado_em": datetime.now().isoformat(timespec="seconds"),
    }, ensure_ascii=False, indent=2), encoding="utf-8")


def write_audit_csv(path, safe_items):
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[key for key, _ in EXPORT_COLUMNS],
            extrasaction="ignore",
            delimiter=";",
        )
        writer.writeheader()
        writer.writerows(safe_items)


def write_audit_xlsx(path, safe_items):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Auditoria"
    sheet.append([label for _, label in EXPORT_COLUMNS])
    for item in safe_items:
        sheet.append([item.get(key, "") for key, _ in EXPORT_COLUMNS])
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    header_fill = PatternFill("solid", fgColor="FFC000")
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = Font(bold=True, color="222222")
    for column in sheet.columns:
        width = min(max(len(str(cell.value or "")) for cell in column) + 2, 48)
        sheet.column_dimensions[column[0].column_letter].width = width
    workbook.save(path)


def export_catalog(output_dir, metadata, items, job_id, template_path=None):
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_items = prepare_catalog_items(items)
    entries = build_catalog_entries(safe_items)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    stem = f"catalogo_goldflex_{job_id[:8]}_{stamp}_{uuid.uuid4().hex[:8]}"

    json_path = output_dir / f"{stem}_auditoria.json"
    csv_path = output_dir / f"{stem}_auditoria.csv"
    xlsx_path = output_dir / f"{stem}_auditoria.xlsx"
    docx_path = output_dir / f"{stem}.docx"
    pdf_path = output_dir / f"{stem}.pdf"

    write_audit_json(json_path, metadata, safe_items, entries)
    write_audit_csv(csv_path, safe_items)
    write_audit_xlsx(xlsx_path, safe_items)
    write_catalog_docx(docx_path, entries, template_path=template_path)
    write_catalog_pdf(pdf_path, entries)

    return {
        kind: {"filename": path.name, "download_url": f"/download/{path.name}"}
        for kind, path in (
            ("docx", docx_path),
            ("pdf", pdf_path),
            ("xlsx", xlsx_path),
            ("csv", csv_path),
            ("json", json_path),
        )
    }
