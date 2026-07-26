import html
import json
import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape

from PIL import Image as PillowImage
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as ReportLabImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


TECHNICAL_SECTIONS = (
    ("assento", "ASSENTO"),
    ("encosto", "ENCOSTO"),
    ("estrutura", "ESTRUTURA"),
    ("base", "BASE"),
    ("pes", "PÉS"),
    ("bracos", "BRAÇOS"),
    ("rodizios", "RODÍZIOS"),
    ("mecanismos", "MECANISMOS"),
    ("acessorios", "ACESSÓRIOS"),
    ("dimensoes", "DIMENSÕES"),
    ("normas", "NORMAS E CONFORMIDADES"),
    ("complementares", "INFORMAÇÕES COMPLEMENTARES"),
    ("observacoes", "OBSERVAÇÕES"),
)

SECTION_KEYWORDS = {
    "assento": ("assento", "espuma do assento", "revestimento do assento"),
    "encosto": ("encosto", "apoio lombar"),
    "estrutura": ("estrutura", "chassi", "armação", "solda", "tubo de aço"),
    "base": ("base", "sapata", "coluna central"),
    "pes": ("pé", "pés", "longarina"),
    "bracos": ("braço", "braços", "apoia-braço", "apoio de braço"),
    "rodizios": ("rodízio", "rodízios", "rodizio", "rodizios"),
    "mecanismos": ("mecanismo", "regulagem", "ajuste", "inclinação", "reclinação"),
    "acessorios": ("acessório", "acessórios", "prancheta", "porta-copos"),
    "dimensoes": (
        "dimensão",
        "dimensões",
        "largura",
        "altura",
        "profundidade",
        "espessura",
        "diâmetro",
    ),
    "normas": (
        "abnt",
        "nbr",
        "nr ",
        "inmetro",
        "certificação",
        "certificado",
        "conformidade",
    ),
}


def compact(value):
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalized_sentences(description):
    text = str(description or "").replace("\r", "\n")
    parts = re.split(r"(?<=[.;:])\s+|\n+|,\s+", text)
    return [compact(part) for part in parts if compact(part)]


def section_text_from_description(description):
    grouped = {key: [] for key, _ in TECHNICAL_SECTIONS}
    for sentence in normalized_sentences(description):
        normalized = sentence.casefold()
        matched = False
        for key, keywords in SECTION_KEYWORDS.items():
            if any(keyword in normalized for keyword in keywords):
                grouped[key].append(sentence)
                matched = True
        if not matched:
            grouped["complementares"].append(sentence)
    return {key: "\n".join(values) for key, values in grouped.items()}


def catalog_draft_from_item(item, pncp):
    description = str(item.get("descricao") or "").strip()
    sections = section_text_from_description(description)
    metadata = pncp.get("metadata") or {}
    unit = compact(item.get("unidade")) or "UND"
    draft = {
        "documento_licitacao": {
            "titulo": "CATÁLOGO TÉCNICO",
            "numero_pregao": compact(metadata.get("numero_compra"))
            or f"{pncp.get('sequencial', '')}/{pncp.get('ano', '')}".strip("/"),
            "processo": compact(metadata.get("processo")),
            "modalidade": compact(metadata.get("modalidade")),
            "objeto": compact(metadata.get("objeto")),
            "link_pncp": compact(pncp.get("link")),
        },
        "orgao": {
            "nome": compact(metadata.get("orgao")),
            "unidade": compact(metadata.get("unidade")),
            "cnpj": compact(metadata.get("orgao_cnpj")),
            "endereco": "",
            "municipio": compact(metadata.get("municipio")),
            "uf": compact(metadata.get("uf")),
        },
        "fabricante": {
            "razao_social": "",
            "nome_fantasia": "",
            "cnpj": "",
            "inscricao_estadual": "",
            "endereco": "",
            "telefone": "",
            "email": "",
            "site": "",
        },
        "item": {
            "numero": compact(item.get("item")),
            "lote": compact(item.get("lote")),
            "quantidade": compact(item.get("quantidade")),
            "unidade": unit,
            "descricao": description,
        },
        "produto": {
            "marca": compact(item.get("marca")),
            "modelo": "",
            "cor": "",
            "revestimento": "",
            "peso": "",
            "garantia": "",
        },
        "resumo": {"caracteristicas": description},
        "secoes": sections,
        "marca_dagua": {
            "ativa": True,
            "texto_personalizado": "",
            "cor": "#C62828",
            "opacidade": 12,
        },
        "origem": {
            "tipo": compact(pncp.get("documento_tipo")) or "Arquivo oficial do PNCP",
            "arquivo": compact(pncp.get("documento_usado")),
            "link": compact(pncp.get("link")),
        },
    }
    return draft


def all_catalog_text(data):
    pieces = []
    for group in ("documento_licitacao", "orgao", "fabricante", "item", "produto", "resumo", "secoes"):
        values = data.get(group) or {}
        if isinstance(values, dict):
            pieces.extend(str(value or "") for value in values.values())
    return "\n".join(pieces)


def catalog_alerts(data, assets):
    errors = []
    warnings = []
    required = (
        ("documento_licitacao.numero_pregao", "Número do pregão"),
        ("documento_licitacao.processo", "Processo"),
        ("orgao.nome", "Órgão destinatário"),
        ("fabricante.razao_social", "Razão social do fabricante"),
        ("fabricante.cnpj", "CNPJ do fabricante"),
        ("item.numero", "Número do item"),
        ("item.descricao", "Descrição do item"),
        ("produto.marca", "Marca"),
        ("produto.modelo", "Modelo"),
    )
    for path, label in required:
        group, field = path.split(".", 1)
        if not compact((data.get(group) or {}).get(field)):
            errors.append(f"Campo ausente: {label}.")

    manufacturer_cnpj = re.sub(r"\D", "", compact((data.get("fabricante") or {}).get("cnpj")))
    if manufacturer_cnpj and len(manufacturer_cnpj) != 14:
        errors.append("O CNPJ do fabricante deve possuir 14 dígitos.")

    roles = {asset.get("role") for asset in assets}
    if "logo" not in roles:
        warnings.append("Logo da empresa não informado.")
    if "principal" not in roles:
        errors.append("Imagem principal do produto não informada.")

    combined = all_catalog_text(data).casefold()
    contradiction_pairs = (
        ("com braços", "sem braços"),
        ("com rodízios", "sem rodízios"),
        ("com rodizios", "sem rodizios"),
        ("base fixa", "base giratória"),
    )
    for positive, negative in contradiction_pairs:
        if positive in combined and negative in combined:
            warnings.append(
                f"Possível informação contraditória: “{positive}” e “{negative}”."
            )

    dimensions = compact((data.get("secoes") or {}).get("dimensoes"))
    units = {
        match.casefold()
        for match in re.findall(r"\b(?:mm|cm|m|kg|g)\b", dimensions, flags=re.IGNORECASE)
    }
    if len(units & {"mm", "cm", "m"}) > 1:
        warnings.append(
            "A seção DIMENSÕES utiliza unidades de comprimento diferentes; revise a consistência."
        )
    if re.search(r"\d[ \t]{2,}\d", all_catalog_text(data)):
        warnings.append("Há possível erro de espaçamento entre números.")
    if len(compact((data.get("resumo") or {}).get("caracteristicas"))) > 1200:
        warnings.append("O resumo excede 1.200 caracteres e pode ocupar mais de uma página.")

    empty_sections = [
        title
        for key, title in TECHNICAL_SECTIONS
        if not compact((data.get("secoes") or {}).get(key))
    ]
    if empty_sections:
        warnings.append(
            "Seções sem conteúdo não serão exportadas: " + ", ".join(empty_sections) + "."
        )
    return {"errors": errors, "warnings": warnings}


def watermark_text(data):
    custom = compact((data.get("marca_dagua") or {}).get("texto_personalizado"))
    if custom:
        return custom
    manufacturer = data.get("fabricante") or {}
    bidding = data.get("documento_licitacao") or {}
    recipient = data.get("orgao") or {}
    return " | ".join(
        value
        for value in (
            compact(manufacturer.get("razao_social")),
            compact(manufacturer.get("cnpj")),
            compact(recipient.get("nome")),
            f"Pregão {compact(bidding.get('numero_pregao'))}"
            if compact(bidding.get("numero_pregao"))
            else "",
        )
        if value
    )


def find_asset(assets, role):
    return next((asset for asset in assets if asset.get("role") == role), None)


def section_assets(assets, section):
    return [
        asset
        for asset in assets
        if asset.get("role") == "tecnica" and asset.get("section") == section
    ]


def set_cell_fill(cell, color):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, bottom=90, start=120, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def format_docx_paragraph(paragraph, size=9, bold=False, color="1F2937", alignment=None):
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.08
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_docx_watermark(document, text, color="#C62828", opacity=12):
    if not text:
        return
    hex_color = re.sub(r"[^0-9A-Fa-f]", "", color)[:6] or "C62828"
    opacity_value = max(3, min(int(opacity or 12), 30)) / 100
    for section in document.sections:
        paragraph = section.header.paragraphs[0]
        pict = parse_xml(
            f"""
            <w:pict
              xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
              xmlns:v="urn:schemas-microsoft-com:vml"
              xmlns:o="urn:schemas-microsoft-com:office:office">
              <v:shape id="CatalogWatermark" o:spid="_x0000_s2049" type="#_x0000_t136"
                style="position:absolute;margin-left:0;margin-top:0;width:520pt;height:90pt;
                rotation:315;z-index:-251654144;mso-position-horizontal:center;
                mso-position-horizontal-relative:page;mso-position-vertical:center;
                mso-position-vertical-relative:page"
                fillcolor="#{hex_color}" stroked="f">
                <v:fill opacity="{opacity_value:.2f}"/>
                <v:textpath style="font-family:Arial;font-size:1pt;font-weight:bold"
                  string="{xml_escape(text)}"/>
              </v:shape>
            </w:pict>
            """
        )
        run = paragraph.add_run()
        run._r.append(pict)


def add_docx_heading(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(153, 27, 27)
    return paragraph


def add_docx_key_value_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        row = table.add_row()
        row.cells[0].width = Mm(45)
        row.cells[1].width = Mm(125)
        row.cells[0].text = label
        row.cells[1].text = compact(value)
        set_cell_fill(row.cells[0], "E5E7EB")
        for index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                format_docx_paragraph(paragraph, bold=index == 0)
    return table


def add_docx_picture(document, asset, max_width_mm=115):
    if not asset:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(asset["path"]), width=Mm(max_width_mm))
    caption = compact(asset.get("caption") or asset.get("name"))
    if caption:
        caption_paragraph = document.add_paragraph(caption)
        format_docx_paragraph(
            caption_paragraph,
            size=8,
            color="6B7280",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )


def configure_docx(document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)


def build_catalog_docx(data, assets, output_path):
    document = Document()
    configure_docx(document)
    bidding = data.get("documento_licitacao") or {}
    recipient = data.get("orgao") or {}
    manufacturer = data.get("fabricante") or {}
    item = data.get("item") or {}
    product = data.get("produto") or {}
    summary = data.get("resumo") or {}

    watermark = data.get("marca_dagua") or {}
    if watermark.get("ativa", True):
        add_docx_watermark(
            document,
            watermark_text(data),
            watermark.get("cor", "#C62828"),
            watermark.get("opacidade", 12),
        )

    header = document.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.rows[0].cells[0].width = Mm(48)
    header.rows[0].cells[1].width = Mm(130)
    logo = find_asset(assets, "logo")
    if logo:
        logo_paragraph = header.cell(0, 0).paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_paragraph.add_run().add_picture(str(logo["path"]), width=Mm(36))
    else:
        header.cell(0, 0).text = "LOGO"
    title_paragraph = header.cell(0, 1).paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_paragraph.add_run(compact(bidding.get("titulo")) or "CATÁLOGO TÉCNICO")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(153, 27, 27)
    meta = header.cell(0, 1).add_paragraph(
        f"Pregão: {compact(bidding.get('numero_pregao'))}\n"
        f"Processo: {compact(bidding.get('processo'))}\n"
        f"Órgão: {compact(recipient.get('nome'))}"
    )
    format_docx_paragraph(meta, size=9, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    add_docx_heading(document, "IDENTIFICAÇÃO DO FABRICANTE")
    add_docx_key_value_table(
        document,
        (
            ("Razão social", manufacturer.get("razao_social")),
            ("Nome fantasia", manufacturer.get("nome_fantasia")),
            ("CNPJ", manufacturer.get("cnpj")),
            ("Inscrição estadual", manufacturer.get("inscricao_estadual")),
            ("Endereço", manufacturer.get("endereco")),
            ("Contato", " | ".join(filter(None, (
                compact(manufacturer.get("telefone")),
                compact(manufacturer.get("email")),
                compact(manufacturer.get("site")),
            )))),
        ),
    )

    add_docx_heading(document, f"ITEM {compact(item.get('numero'))}")
    description = document.add_paragraph(compact(item.get("descricao")))
    format_docx_paragraph(description, size=9)
    add_docx_key_value_table(
        document,
        (
            ("Marca", product.get("marca")),
            ("Modelo", product.get("modelo")),
            ("Quantidade", f"{compact(item.get('quantidade'))} {compact(item.get('unidade'))}".strip()),
        ),
    )
    add_docx_picture(document, find_asset(assets, "principal"), max_width_mm=92)

    add_docx_heading(document, "INFORMAÇÕES RESUMIDAS")
    add_docx_key_value_table(
        document,
        (
            ("Cor", product.get("cor")),
            ("Revestimento", product.get("revestimento")),
            ("Peso", product.get("peso")),
            ("Normas", (data.get("secoes") or {}).get("normas")),
            ("Características principais", summary.get("caracteristicas")),
        ),
    )

    secondary_assets = [asset for asset in assets if asset.get("role") == "secundaria"]
    sections = data.get("secoes") or {}
    for key, title in TECHNICAL_SECTIONS:
        content = str(sections.get(key) or "").strip()
        images = section_assets(assets, key)
        if not content and not images:
            continue
        document.add_page_break()
        add_docx_heading(document, title)
        if content:
            for block in re.split(r"\n{2,}", content):
                paragraph = document.add_paragraph(block.strip())
                format_docx_paragraph(paragraph, size=9)
        for asset in images:
            add_docx_picture(document, asset, max_width_mm=120)

    for asset in secondary_assets:
        document.add_page_break()
        add_docx_heading(document, "IMAGEM SECUNDÁRIA")
        add_docx_picture(document, asset, max_width_mm=130)

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = compact((data.get("origem") or {}).get("link"))
    format_docx_paragraph(
        footer,
        size=7,
        color="6B7280",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    document.save(output_path)


def reportlab_image(asset, max_width, max_height):
    if not asset:
        return None
    with PillowImage.open(asset["path"]) as image:
        width, height = image.size
    scale = min(max_width / width, max_height / height)
    return ReportLabImage(str(asset["path"]), width=width * scale, height=height * scale)


def pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CatalogTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=21,
            textColor=colors.HexColor("#991B1B"),
            alignment=TA_LEFT,
            spaceAfter=5,
        ),
        "heading": ParagraphStyle(
            "CatalogHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#991B1B"),
            spaceBefore=5,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "CatalogBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=11.2,
            textColor=colors.HexColor("#1F2937"),
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "CatalogSmall",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=7.6,
            leading=9.2,
            textColor=colors.HexColor("#4B5563"),
        ),
        "center": ParagraphStyle(
            "CatalogCenter",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            alignment=TA_CENTER,
        ),
    }


def pdf_paragraph(value, style):
    return Paragraph(html.escape(str(value or "")).replace("\n", "<br/>"), style)


def pdf_key_value_table(rows, styles, widths=(42 * mm, 126 * mm)):
    values = [
        [
            pdf_paragraph(label, styles["small"]),
            pdf_paragraph(compact(value), styles["body"]),
        ]
        for label, value in rows
    ]
    table = Table(values, colWidths=list(widths), hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E5E7EB")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1F2937")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#9CA3AF")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def pdf_page_decor(data):
    watermark = data.get("marca_dagua") or {}
    text = watermark_text(data)
    color = colors.HexColor(watermark.get("cor", "#C62828"))
    opacity = max(3, min(int(watermark.get("opacidade", 12) or 12), 30)) / 100
    source_link = compact((data.get("origem") or {}).get("link"))

    def draw(canvas, document):
        canvas.saveState()
        if watermark.get("ativa", True) and text:
            red, green, blue = color.rgb()
            canvas.setFillColorRGB(
                1 - ((1 - red) * opacity),
                1 - ((1 - green) * opacity),
                1 - ((1 - blue) * opacity),
            )
            canvas.setFont("Helvetica-Bold", 17)
            canvas.translate(A4[0] / 2, A4[1] / 2)
            canvas.rotate(35)
            parts = [part.strip() for part in text.split(" | ") if part.strip()]
            first_line = " | ".join(parts[:2])
            second_line = " | ".join(parts[2:])
            canvas.drawCentredString(0, 8, first_line[:76])
            if second_line:
                canvas.drawCentredString(0, -14, second_line[:76])
        canvas.restoreState()
        canvas.saveState()
        canvas.setFillColor(colors.HexColor("#6B7280"))
        canvas.setFont("Helvetica", 7)
        canvas.drawString(16 * mm, 8 * mm, source_link[:95])
        canvas.drawRightString(A4[0] - 16 * mm, 8 * mm, f"Página {document.page}")
        canvas.restoreState()

    return draw


def build_catalog_pdf(data, assets, output_path):
    styles = pdf_styles()
    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        title=compact((data.get("documento_licitacao") or {}).get("titulo"))
        or "Catálogo técnico",
        author=compact((data.get("fabricante") or {}).get("razao_social")),
    )
    bidding = data.get("documento_licitacao") or {}
    recipient = data.get("orgao") or {}
    manufacturer = data.get("fabricante") or {}
    item = data.get("item") or {}
    product = data.get("produto") or {}
    summary = data.get("resumo") or {}
    story = []

    logo_image = reportlab_image(find_asset(assets, "logo"), 38 * mm, 25 * mm)
    header_text = [
        pdf_paragraph(compact(bidding.get("titulo")) or "CATÁLOGO TÉCNICO", styles["title"]),
        pdf_paragraph(
            f"Pregão: {compact(bidding.get('numero_pregao'))}\n"
            f"Processo: {compact(bidding.get('processo'))}\n"
            f"Órgão: {compact(recipient.get('nome'))}",
            styles["small"],
        ),
    ]
    header = Table(
        [[logo_image or pdf_paragraph("LOGO", styles["center"]), header_text]],
        colWidths=[44 * mm, 124 * mm],
    )
    header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.extend([header, Spacer(1, 4 * mm)])

    story.append(pdf_paragraph("IDENTIFICAÇÃO DO FABRICANTE", styles["heading"]))
    story.append(
        pdf_key_value_table(
            (
                ("Razão social", manufacturer.get("razao_social")),
                ("Nome fantasia", manufacturer.get("nome_fantasia")),
                ("CNPJ", manufacturer.get("cnpj")),
                ("Inscrição estadual", manufacturer.get("inscricao_estadual")),
                ("Endereço", manufacturer.get("endereco")),
                ("Contato", " | ".join(filter(None, (
                    compact(manufacturer.get("telefone")),
                    compact(manufacturer.get("email")),
                    compact(manufacturer.get("site")),
                )))),
            ),
            styles,
        )
    )
    story.extend(
        [
            Spacer(1, 3 * mm),
            pdf_paragraph(f"ITEM {compact(item.get('numero'))}", styles["heading"]),
            pdf_paragraph(compact(item.get("descricao")), styles["body"]),
            pdf_key_value_table(
                (
                    ("Marca", product.get("marca")),
                    ("Modelo", product.get("modelo")),
                    ("Quantidade", f"{compact(item.get('quantidade'))} {compact(item.get('unidade'))}".strip()),
                ),
                styles,
            ),
            Spacer(1, 3 * mm),
        ]
    )
    main_image = reportlab_image(find_asset(assets, "principal"), 92 * mm, 58 * mm)
    if main_image:
        main_image.hAlign = "CENTER"
        story.extend([main_image, Spacer(1, 2 * mm)])
    story.append(pdf_paragraph("INFORMAÇÕES RESUMIDAS", styles["heading"]))
    story.append(
        pdf_key_value_table(
            (
                ("Cor", product.get("cor")),
                ("Revestimento", product.get("revestimento")),
                ("Peso", product.get("peso")),
                ("Normas", (data.get("secoes") or {}).get("normas")),
                ("Características principais", summary.get("caracteristicas")),
            ),
            styles,
        )
    )

    sections = data.get("secoes") or {}
    for key, title in TECHNICAL_SECTIONS:
        content = str(sections.get(key) or "").strip()
        images = section_assets(assets, key)
        if not content and not images:
            continue
        story.extend([PageBreak(), pdf_paragraph(title, styles["heading"])])
        if content:
            story.append(pdf_paragraph(content, styles["body"]))
        for asset in images:
            image = reportlab_image(asset, 150 * mm, 175 * mm)
            if image:
                image.hAlign = "CENTER"
                story.extend([Spacer(1, 3 * mm), image])
                if compact(asset.get("caption") or asset.get("name")):
                    story.append(
                        pdf_paragraph(asset.get("caption") or asset.get("name"), styles["center"])
                    )

    for asset in [asset for asset in assets if asset.get("role") == "secundaria"]:
        story.extend([PageBreak(), pdf_paragraph("IMAGEM SECUNDÁRIA", styles["heading"])])
        image = reportlab_image(asset, 165 * mm, 225 * mm)
        if image:
            image.hAlign = "CENTER"
            story.append(image)

    decor = pdf_page_decor(data)
    document.build(story, onFirstPage=decor, onLaterPages=decor)


def write_catalog_json(data, assets, output_path, alerts):
    payload = {
        "schema": "catalogo-tecnico-licitacao/v1",
        "dados": data,
        "imagens": [
            {
                "arquivo": asset.get("name"),
                "papel": asset.get("role"),
                "secao": asset.get("section", ""),
                "legenda": asset.get("caption", ""),
            }
            for asset in assets
        ],
        "validacao": alerts,
    }
    output_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def build_images_zip(assets, output_path):
    manifest = []
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, asset in enumerate(assets, start=1):
            source = Path(asset["path"])
            role = compact(asset.get("role")) or "imagem"
            name = f"{index:02d}_{source.name}"
            archive_name = f"{role}/{name}"
            archive.write(source, archive_name)
            manifest.append(
                {
                    "arquivo": archive_name,
                    "papel": role,
                    "secao": compact(asset.get("section")),
                    "legenda": compact(asset.get("caption")),
                }
            )
        archive.writestr(
            "manifesto.json",
            json.dumps({"imagens": manifest}, ensure_ascii=False, indent=2),
        )
