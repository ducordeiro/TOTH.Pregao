import warnings

warnings.filterwarnings("ignore", category=DeprecationWarning, message="'cgi' is deprecated.*")

import cgi
import copy
import hashlib
import html
import json
import logging
import mimetypes
import os
import re
import shutil
import sqlite3
import subprocess
import tempfile
import threading
import time
import unicodedata
import uuid
import http.client
import urllib.error
import urllib.request
import zipfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import contextmanager
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from difflib import SequenceMatcher
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, quote, unquote, urlencode, urlparse
import pdfplumber
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from catalog import (
    build_catalog_docx,
    build_catalog_pdf,
    build_images_zip,
    catalog_alerts,
    catalog_draft_from_item,
    write_catalog_json,
)

import kanban as kanban_store
from etl import ETLRepository, ETLSyncService, OpportunityClassifier, PNCPConnector, PNCPMapper, SyncRequest
from etl.connectors import HttpJsonClient
from etl.search_filters import classify_object_text


LOGGER = logging.getLogger("toth.pregao")


ROOT = Path(__file__).resolve().parent
WORKSPACE = ROOT.parent
UPLOAD_DIR = Path(os.environ.get("TOTH_UPLOAD_DIR") or ROOT / "uploads")
OUTPUT_DIR = Path(os.environ.get("TOTH_OUTPUT_DIR") or ROOT / "outputs")
TEMPLATE_DIR = ROOT / "templates"
DATA_DIR = Path(os.environ.get("TOTH_DATA_DIR") or ROOT / "data")
FRONTEND_DIST = ROOT / "frontend_dist"
DATABASE_PATH = Path(os.environ.get("TOTH_DATABASE_PATH") or DATA_DIR / "pncp.sqlite3")
PREVIEW_DIR = Path(
    os.environ.get("TOTH_PREVIEW_DIR")
    or Path(tempfile.gettempdir()) / "ocr_edital_web_previews"
)
DOCX_TO_PDF_SCRIPT = ROOT / "scripts" / "convert_docx_to_pdf.ps1"
DEFAULT_TEMPLATE = TEMPLATE_DIR / "Document 2.docx"
ALERE_TEMPLATE = TEMPLATE_DIR / "template_Alere.docx"
LEGACY_TEMPLATE_KEYS = {
    "default": DEFAULT_TEMPLATE.name,
    "alere": ALERE_TEMPLATE.name,
}
MAX_TEMPLATE_SIZE = 15 * 1024 * 1024
MAX_CATALOG_ASSET_SIZE = 12 * 1024 * 1024
MAX_CATALOG_REQUEST_SIZE = 80 * 1024 * 1024
MAX_PROPOSAL_REQUEST_SIZE = MAX_TEMPLATE_SIZE + (5 * 1024 * 1024)
MAX_JSON_REQUEST_SIZE = 8 * 1024 * 1024
LOCAL_REQUEST_HOSTS = {"127.0.0.1", "localhost", "::1"}
PORT = int(os.environ.get("TOTH_PORT") or 8765)
PNCP_API_BASE = "https://pncp.gov.br/api"
PNCP_APP_BASE = "https://pncp.gov.br/app/editais"
PNCP_SEARCH_URL = "https://pncp.gov.br/api/search"
ALLOW_RUNTIME_PNCP_API = os.environ.get("TOTH_ALLOW_RUNTIME_PNCP_API", "").strip().lower() in {
    "1",
    "true",
    "yes",
    "sim",
}
ALLOW_BLOCO2_ON_DEMAND_ENRICHMENT = os.environ.get(
    "TOTH_BLOCO2_ON_DEMAND_ENRICHMENT",
    "1",
).strip().lower() in {"1", "true", "yes", "sim"}
ALLOW_DETAIL_DOCUMENT_ON_DEMAND = os.environ.get(
    "TOTH_DETAIL_DOCUMENT_ON_DEMAND",
    "1",
).strip().lower() in {"1", "true", "yes", "sim"}
ALLOW_DETAIL_ITEMS_ON_DEMAND = os.environ.get(
    "TOTH_DETAIL_ITEMS_ON_DEMAND",
    "1",
).strip().lower() in {"1", "true", "yes", "sim"}
ENABLE_PNCP_SEARCH = os.environ.get(
    "TOTH_ENABLE_PNCP_SEARCH",
    "1",
).strip().lower() in {"1", "true", "yes", "sim"}
SEARCH_CACHE = {}
PNCP_RESULT_CACHE = {}
PNCP_OPPORTUNITY_SYNC_CACHE = {}
SEARCH_ITEM_CACHE = {}
SEARCH_DOCUMENT_ITEM_CACHE = {}
PNCP_SEARCH_JOBS = {}
SEARCH_CACHE_TTL = 300
SOURCE_CACHE = {}
EXTRACTED_ITEMS_CACHE = {}
IDENTIFICATION_CACHE = {}
BLOCO2_ENRICHMENT_LOCKS = {}
BLOCO2_ENRICHMENT_LOCKS_GUARD = threading.Lock()
DETAIL_DOCUMENT_ENRICHMENT_LOCK = threading.Lock()
CATALOG_TEXT_CACHE = {}
CATALOG_DRAFT_CACHE = {}
DOCUMENT_CACHE_TTL = 900
PROPOSAL_PREVIEW_TTL = 30 * 60
CACHE_MAX_ENTRIES = 32
CACHE_LOCK = threading.Lock()
PNCP_SEARCH_JOB_LOCK = threading.RLock()
OCR_LOCK = threading.Lock()
TEMPLATE_LOCK = threading.RLock()
DATABASE_LOCK = threading.RLock()
PROPOSAL_PREVIEW_LOCK = threading.RLock()
WORD_CONVERSION_LOCK = threading.Lock()
PROPOSAL_PREVIEW_CACHE = {}
OCR_ENGINE = None
OCR_DPI = 150
STANDARD_UNIT = "UND"
BUSINESS_STAGES = (
    "oportunidade",
    "qualificacao",
    "disputa",
    "classificacao",
    "contrato",
)
DEFAULT_BUSINESS_TASKS = (
    "Validar objeto da contratação",
    "Confirmar data e horário de abertura",
    "Revisar Termo de Referência ou Edital",
    "Analisar requisitos técnicos",
    "Analisar condições comerciais",
    "Validar prazo de entrega",
    "Validar condições de pagamento",
    "Conferir documentos de habilitação",
    "Definir responsável interno",
    "Registrar decisão de participação",
    "Preparar proposta comercial",
    "Revisar preços e margens",
    "Anexar documentação necessária",
    "Acompanhar disputa e resultado",
    "Registrar encerramento do processo",
)

COLUMNS = [
    ("item", "ITEM"),
    ("quantidade", "QTD"),
    ("unidade", STANDARD_UNIT),
    ("descricao", "DESCRIÇÃO."),
    ("marca", "MARCA"),
    ("valor_unitario", "VALOR UNITÁRIO"),
    ("valor_total", "VALOR TOTAL"),
]
LOT_COLUMN = ("lote", "LOTE")

PROPOSAL_SIGNATURE_LOCATION_DATE = "Mogi Mirim - SP 14 de Julho de 2026"
INITIAL_RESPONSIBLES = (
    (
        1,
        "Brendon Matheus Batista",
        "Goldflex Industria e Comércio de Moveis e Equipamentos LTDA",
        "33.661.439/0001-14",
        "50.630.673-2",
        "432.079.848-19",
        "",
    ),
    (
        2,
        "Marcos Sérgio Rodrigues Pereira Júnior",
        "Goldflex Industria e Comércio de Moveis e Equipamentos LTDA",
        "33.661.439/0001-14",
        "",
        "",
        "Sócio Majoritário",
    ),
)


class ResponsibleInUseError(Exception):
    pass


def ensure_dirs():
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    PREVIEW_DIR.mkdir(parents=True, exist_ok=True)


@contextmanager
def database_connection():
    connection = sqlite3.connect(DATABASE_PATH, timeout=30)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys = ON")
    try:
        with connection:
            yield connection
    finally:
        connection.close()


def init_database():
    ensure_dirs()
    kanban_store.initialize(DATABASE_PATH)
    with DATABASE_LOCK, database_connection() as connection:
        connection.execute("PRAGMA journal_mode = WAL")
        connection.executescript(
            """
            CREATE TABLE IF NOT EXISTS contratacoes (
                id INTEGER PRIMARY KEY,
                cnpj TEXT NOT NULL,
                ano INTEGER NOT NULL,
                sequencial INTEGER NOT NULL,
                link_pncp TEXT NOT NULL,
                documento_usado TEXT NOT NULL DEFAULT '',
                documento_tipo TEXT NOT NULL DEFAULT '',
                caminho_documento TEXT NOT NULL DEFAULT '',
                total_itens INTEGER NOT NULL DEFAULT 0,
                revisao_descricoes TEXT NOT NULL DEFAULT '{}',
                verificacao_pncp TEXT NOT NULL DEFAULT '{}',
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL,
                UNIQUE (cnpj, ano, sequencial)
            );

            CREATE TABLE IF NOT EXISTS arquivos_pncp (
                id INTEGER PRIMARY KEY,
                contratacao_id INTEGER NOT NULL REFERENCES contratacoes(id) ON DELETE CASCADE,
                chave_pncp TEXT NOT NULL,
                titulo TEXT NOT NULL DEFAULT '',
                tipo_documento TEXT NOT NULL DEFAULT '',
                url TEXT NOT NULL DEFAULT '',
                caminho_local TEXT NOT NULL DEFAULT '',
                selecionado INTEGER NOT NULL DEFAULT 0 CHECK (selecionado IN (0, 1)),
                metadados TEXT NOT NULL DEFAULT '{}',
                UNIQUE (contratacao_id, chave_pncp)
            );

            CREATE TABLE IF NOT EXISTS itens (
                id INTEGER PRIMARY KEY,
                contratacao_id INTEGER NOT NULL REFERENCES contratacoes(id) ON DELETE CASCADE,
                lote TEXT NOT NULL DEFAULT '',
                numero_item TEXT NOT NULL,
                quantidade TEXT NOT NULL DEFAULT '',
                unidade TEXT NOT NULL DEFAULT 'UND',
                identificacao_simplificada TEXT NOT NULL,
                descricao_completa TEXT NOT NULL,
                documento_fonte TEXT NOT NULL DEFAULT '',
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL,
                UNIQUE (contratacao_id, lote, numero_item)
            );

            CREATE TABLE IF NOT EXISTS consultas (
                id INTEGER PRIMARY KEY,
                contratacao_id INTEGER NOT NULL REFERENCES contratacoes(id) ON DELETE CASCADE,
                consultado_em TEXT NOT NULL,
                total_itens INTEGER NOT NULL,
                origem TEXT NOT NULL DEFAULT 'arquivo'
            );

            CREATE TABLE IF NOT EXISTS responsaveis (
                id INTEGER PRIMARY KEY,
                nome_completo TEXT NOT NULL,
                empresa TEXT NOT NULL,
                cnpj TEXT NOT NULL,
                rg TEXT NOT NULL DEFAULT '',
                cpf TEXT NOT NULL DEFAULT '',
                observacoes TEXT NOT NULL DEFAULT '',
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS documentos_gerados (
                id INTEGER PRIMARY KEY,
                responsavel_id INTEGER NOT NULL REFERENCES responsaveis(id) ON DELETE RESTRICT,
                nome_arquivo TEXT NOT NULL,
                caminho_arquivo TEXT NOT NULL,
                criado_em TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS negocios (
                id INTEGER PRIMARY KEY,
                contratacao_id INTEGER REFERENCES contratacoes(id) ON DELETE SET NULL,
                empresa TEXT NOT NULL,
                cnpj_orgao TEXT NOT NULL,
                ano INTEGER NOT NULL,
                sequencial INTEGER NOT NULL,
                link_pncp TEXT NOT NULL,
                titulo TEXT NOT NULL DEFAULT '',
                titulo_interno TEXT NOT NULL DEFAULT '',
                orgao TEXT NOT NULL DEFAULT '',
                municipio TEXT NOT NULL DEFAULT '',
                uf TEXT NOT NULL DEFAULT '',
                modalidade TEXT NOT NULL DEFAULT '',
                numero_compra TEXT NOT NULL DEFAULT '',
                processo TEXT NOT NULL DEFAULT '',
                plataforma TEXT NOT NULL DEFAULT 'PNCP',
                fonte_integracao TEXT NOT NULL DEFAULT 'API PNCP',
                abertura TEXT NOT NULL DEFAULT '',
                encerramento TEXT NOT NULL DEFAULT '',
                etapa TEXT NOT NULL DEFAULT 'oportunidade',
                situacao TEXT NOT NULL DEFAULT '',
                prioridade INTEGER NOT NULL DEFAULT 2 CHECK (prioridade BETWEEN 1 AND 3),
                position_number INTEGER,
                favorito INTEGER NOT NULL DEFAULT 0 CHECK (favorito IN (0, 1)),
                arquivado INTEGER NOT NULL DEFAULT 0 CHECK (arquivado IN (0, 1)),
                removido INTEGER NOT NULL DEFAULT 0 CHECK (removido IN (0, 1)),
                responsavel TEXT NOT NULL DEFAULT '',
                prazo_interno TEXT NOT NULL DEFAULT '',
                anotacoes TEXT NOT NULL DEFAULT '',
                decisao_comercial TEXT NOT NULL DEFAULT '',
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL,
                UNIQUE (empresa, cnpj_orgao, ano, sequencial)
            );

            CREATE TABLE IF NOT EXISTS negocio_historico (
                id INTEGER PRIMARY KEY,
                negocio_id INTEGER NOT NULL REFERENCES negocios(id) ON DELETE CASCADE,
                evento TEXT NOT NULL,
                etapa_anterior TEXT NOT NULL DEFAULT '',
                etapa_nova TEXT NOT NULL DEFAULT '',
                justificativa TEXT NOT NULL DEFAULT '',
                criado_em TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS negocio_tarefas (
                id INTEGER PRIMARY KEY,
                negocio_id INTEGER NOT NULL REFERENCES negocios(id) ON DELETE CASCADE,
                titulo TEXT NOT NULL,
                concluida INTEGER NOT NULL DEFAULT 0 CHECK (concluida IN (0, 1)),
                ordem INTEGER NOT NULL DEFAULT 0,
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS negocio_itens (
                id INTEGER PRIMARY KEY,
                negocio_id INTEGER NOT NULL REFERENCES negocios(id) ON DELETE CASCADE,
                ordem INTEGER NOT NULL DEFAULT 0,
                lote TEXT NOT NULL DEFAULT '',
                numero_item TEXT NOT NULL,
                descricao TEXT NOT NULL,
                quantidade TEXT NOT NULL DEFAULT '',
                unidade TEXT NOT NULL DEFAULT 'UND',
                valor_unitario_estimado TEXT NOT NULL DEFAULT '',
                valor_total_estimado TEXT NOT NULL DEFAULT '',
                criterio_julgamento TEXT NOT NULL DEFAULT '',
                situacao TEXT NOT NULL DEFAULT '',
                criado_em TEXT NOT NULL,
                atualizado_em TEXT NOT NULL,
                UNIQUE (negocio_id, lote, numero_item)
            );

            CREATE TABLE IF NOT EXISTS app_migrations (
                chave TEXT PRIMARY KEY,
                aplicado_em TEXT NOT NULL
            );

            CREATE INDEX IF NOT EXISTS idx_itens_contratacao
                ON itens (contratacao_id, lote, numero_item);
            CREATE INDEX IF NOT EXISTS idx_arquivos_contratacao
                ON arquivos_pncp (contratacao_id, selecionado);
            CREATE UNIQUE INDEX IF NOT EXISTS idx_responsaveis_cpf
                ON responsaveis (cpf) WHERE cpf <> '';
            CREATE INDEX IF NOT EXISTS idx_documentos_responsavel
                ON documentos_gerados (responsavel_id, criado_em);
            CREATE INDEX IF NOT EXISTS idx_negocios_empresa_etapa
                ON negocios (empresa, etapa, arquivado, removido);
            CREATE INDEX IF NOT EXISTS idx_negocios_abertura
                ON negocios (abertura);
            CREATE INDEX IF NOT EXISTS idx_negocio_historico
                ON negocio_historico (negocio_id, criado_em);
            CREATE INDEX IF NOT EXISTS idx_negocio_tarefas
                ON negocio_tarefas (negocio_id, ordem, id);
            CREATE INDEX IF NOT EXISTS idx_negocio_itens
                ON negocio_itens (negocio_id, ordem, id);
            """
        )
        business_columns = {
            row[1] for row in connection.execute("PRAGMA table_info(negocios)")
        }
        if "position_number" not in business_columns:
            connection.execute("ALTER TABLE negocios ADD COLUMN position_number INTEGER")
        connection.execute(
            """
            UPDATE negocios SET position_number = (
                SELECT p.position_number FROM proposals p WHERE p.business_id = negocios.id LIMIT 1
            ) WHERE position_number IS NULL AND EXISTS (
                SELECT 1 FROM proposals p WHERE p.business_id = negocios.id AND p.position_number IS NOT NULL
            )
            """
        )
        now = datetime.now().astimezone().isoformat(timespec="seconds")
        seed_migration = "seed_responsaveis_v1"
        migration_applied = connection.execute(
            "SELECT 1 FROM app_migrations WHERE chave = ?", (seed_migration,)
        ).fetchone()
        if not migration_applied:
            connection.executemany(
                """
                INSERT OR IGNORE INTO responsaveis (
                    id, nome_completo, empresa, cnpj, rg, cpf, observacoes,
                    criado_em, atualizado_em
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                [(*responsible, now, now) for responsible in INITIAL_RESPONSIBLES],
            )
            connection.execute(
                "INSERT INTO app_migrations (chave, aplicado_em) VALUES (?, ?)",
                (seed_migration, now),
            )
    ETLRepository(DATABASE_PATH).initialize()


def responsible_record(row):
    return {
        "id": str(row["id"]),
        "nome_completo": row["nome_completo"],
        "empresa": row["empresa"],
        "cnpj": row["cnpj"],
        "rg": row["rg"],
        "cpf": row["cpf"],
        "observacoes": row["observacoes"],
        "criado_em": row["criado_em"],
        "atualizado_em": row["atualizado_em"],
    }


def responsible_id_value(responsible_id):
    value = str(responsible_id or "").strip()
    if not value.isdigit() or int(value) <= 0:
        raise ValueError("Responsável inválido.")
    return int(value)


def validate_responsible_payload(payload):
    if not isinstance(payload, dict):
        raise ValueError("Dados do responsável inválidos.")
    data = {
        "nome_completo": compact(payload.get("nome_completo")),
        "empresa": compact(payload.get("empresa")),
        "cnpj": compact(payload.get("cnpj")),
        "rg": compact(payload.get("rg")),
        "cpf": compact(payload.get("cpf")),
        "observacoes": str(payload.get("observacoes") or "").strip(),
    }
    if not data["nome_completo"]:
        raise ValueError("Informe o nome completo do responsável.")
    if not data["empresa"]:
        raise ValueError("Informe a empresa do responsável.")
    if not data["cnpj"]:
        raise ValueError("Informe o CNPJ da empresa.")
    if len(re.sub(r"\D", "", data["cnpj"])) != 14:
        raise ValueError("Informe um CNPJ com 14 dígitos.")
    if data["cpf"] and len(re.sub(r"\D", "", data["cpf"])) != 11:
        raise ValueError("Informe um CPF com 11 dígitos.")
    limits = {
        "nome_completo": 200,
        "empresa": 200,
        "cnpj": 24,
        "rg": 30,
        "cpf": 20,
        "observacoes": 1000,
    }
    for field, limit in limits.items():
        if len(data[field]) > limit:
            raise ValueError(f"O campo {field.replace('_', ' ')} excede {limit} caracteres.")
    return data


def list_responsibles():
    init_database()
    with DATABASE_LOCK, database_connection() as connection:
        rows = connection.execute(
            "SELECT * FROM responsaveis ORDER BY nome_completo COLLATE NOCASE, id"
        ).fetchall()
    return [responsible_record(row) for row in rows]


def get_responsible(responsible_id):
    init_database()
    responsible_id = responsible_id_value(responsible_id)
    with DATABASE_LOCK, database_connection() as connection:
        row = connection.execute(
            "SELECT * FROM responsaveis WHERE id = ?", (responsible_id,)
        ).fetchone()
    return responsible_record(row) if row else None


def create_responsible(payload):
    data = validate_responsible_payload(payload)
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    try:
        with DATABASE_LOCK, database_connection() as connection:
            cursor = connection.execute(
                """
                INSERT INTO responsaveis (
                    nome_completo, empresa, cnpj, rg, cpf, observacoes,
                    criado_em, atualizado_em
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (*data.values(), now, now),
            )
            responsible_id = cursor.lastrowid
    except sqlite3.IntegrityError as exc:
        raise FileExistsError("Já existe um responsável cadastrado com este CPF.") from exc
    return get_responsible(responsible_id)


def update_responsible(responsible_id, payload):
    responsible_id = responsible_id_value(responsible_id)
    data = validate_responsible_payload(payload)
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    try:
        with DATABASE_LOCK, database_connection() as connection:
            cursor = connection.execute(
                """
                UPDATE responsaveis SET
                    nome_completo = ?, empresa = ?, cnpj = ?, rg = ?, cpf = ?,
                    observacoes = ?, atualizado_em = ?
                WHERE id = ?
                """,
                (*data.values(), now, responsible_id),
            )
            if not cursor.rowcount:
                raise FileNotFoundError("Responsável não encontrado.")
    except sqlite3.IntegrityError as exc:
        raise FileExistsError("Já existe um responsável cadastrado com este CPF.") from exc
    return get_responsible(responsible_id)


def delete_responsible(responsible_id):
    responsible_id = responsible_id_value(responsible_id)
    init_database()
    with DATABASE_LOCK, database_connection() as connection:
        linked_documents = connection.execute(
            "SELECT COUNT(*) FROM documentos_gerados WHERE responsavel_id = ?",
            (responsible_id,),
        ).fetchone()[0]
        if linked_documents:
            raise ResponsibleInUseError(
                f"Este responsável possui vínculo com {linked_documents} documento(s) gerado(s) e não pode ser excluído."
            )
        cursor = connection.execute(
            "DELETE FROM responsaveis WHERE id = ?", (responsible_id,)
        )
        if not cursor.rowcount:
            raise FileNotFoundError("Responsável não encontrado.")


def record_generated_document(responsible_id, output_path):
    responsible_id = responsible_id_value(responsible_id)
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    with DATABASE_LOCK, database_connection() as connection:
        connection.execute(
            """
            INSERT INTO documentos_gerados (
                responsavel_id, nome_arquivo, caminho_arquivo, criado_em
            ) VALUES (?, ?, ?, ?)
            """,
            (responsible_id, output_path.name, str(output_path.resolve()), now),
        )


def pncp_file_key(file_info):
    return compact(
        file_info.get("id")
        or file_info.get("sequencialDocumento")
        or file_info.get("url")
        or file_info.get("uri")
        or file_info.get("titulo")
    )


def persist_identification(source_data, identifications, description_review, pncp_items_check):
    init_database()
    pncp = source_data["pncp"]
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    source_path = Path(source_data["source_path"])
    chosen = pncp.get("arquivo_usado") or {}
    chosen_key = pncp_file_key(chosen)

    with DATABASE_LOCK, database_connection() as connection:
        connection.execute("BEGIN IMMEDIATE")
        connection.execute(
            """
            INSERT INTO contratacoes (
                cnpj, ano, sequencial, link_pncp, documento_usado, documento_tipo,
                caminho_documento, total_itens, revisao_descricoes, verificacao_pncp,
                criado_em, atualizado_em
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(cnpj, ano, sequencial) DO UPDATE SET
                link_pncp = excluded.link_pncp,
                documento_usado = excluded.documento_usado,
                documento_tipo = excluded.documento_tipo,
                caminho_documento = excluded.caminho_documento,
                total_itens = excluded.total_itens,
                revisao_descricoes = excluded.revisao_descricoes,
                verificacao_pncp = excluded.verificacao_pncp,
                atualizado_em = excluded.atualizado_em
            """,
            (
                pncp["cnpj"], pncp["ano"], pncp["sequencial"], pncp["link"],
                pncp.get("documento_usado", ""), pncp.get("documento_tipo", ""),
                str(source_path.resolve()), len(identifications),
                json.dumps(description_review, ensure_ascii=False),
                json.dumps(pncp_items_check, ensure_ascii=False), now, now,
            ),
        )
        contratacao_id = connection.execute(
            "SELECT id FROM contratacoes WHERE cnpj = ? AND ano = ? AND sequencial = ?",
            (pncp["cnpj"], pncp["ano"], pncp["sequencial"]),
        ).fetchone()["id"]

        connection.execute("DELETE FROM arquivos_pncp WHERE contratacao_id = ?", (contratacao_id,))
        for index, file_info in enumerate(pncp.get("arquivos") or []):
            file_key = pncp_file_key(file_info) or f"arquivo-{index + 1}"
            selected = int(file_key == chosen_key)
            connection.execute(
                """
                INSERT INTO arquivos_pncp (
                    contratacao_id, chave_pncp, titulo, tipo_documento, url,
                    caminho_local, selecionado, metadados
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    contratacao_id, file_key, compact(file_info.get("titulo")),
                    compact(file_info.get("tipoDocumentoNome") or file_info.get("tipoDocumentoDescricao")),
                    compact(file_info.get("url") or file_info.get("uri")),
                    str(source_path.resolve()) if selected else "", selected,
                    json.dumps(file_info, ensure_ascii=False, default=str),
                ),
            )

        connection.execute("DELETE FROM itens WHERE contratacao_id = ?", (contratacao_id,))
        for item in identifications:
            connection.execute(
                """
                INSERT INTO itens (
                    contratacao_id, lote, numero_item, quantidade, unidade,
                    identificacao_simplificada, descricao_completa, documento_fonte,
                    criado_em, atualizado_em
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    contratacao_id, item["lote"], item["item"], item["quantidade"],
                    STANDARD_UNIT, item["categoria"], item["descricao"],
                    pncp.get("documento_usado", ""), now, now,
                ),
            )

        connection.execute(
            """
            INSERT INTO consultas (contratacao_id, consultado_em, total_itens, origem)
            VALUES (?, ?, ?, 'arquivo')
            """,
            (contratacao_id, now, len(identifications)),
        )
    return contratacao_id


def template_path_from_name(name):
    filename = unquote(compact(name))
    if not filename or Path(filename).name != filename or Path(filename).suffix.lower() != ".docx":
        return None
    candidate = TEMPLATE_DIR / filename
    if candidate.parent.resolve() != TEMPLATE_DIR.resolve():
        return None
    return candidate


def template_display_name(filename):
    stem = Path(filename).stem
    if filename == DEFAULT_TEMPLATE.name:
        return "Modelo padrão"
    if filename == ALERE_TEMPLATE.name:
        return "Alere Solar"
    return stem.replace("_", " ").strip() or stem


def template_record(path):
    stat = path.stat()
    return {
        "id": path.name,
        "name": path.name,
        "display_name": template_display_name(path.name),
        "size": stat.st_size,
        "updated_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds"),
        "download_url": f"/template/{quote(path.name)}",
    }


def list_templates():
    ensure_dirs()
    with TEMPLATE_LOCK:
        paths = sorted(
            (
                path
                for path in TEMPLATE_DIR.iterdir()
                if path.is_file() and path.suffix.lower() == ".docx"
            ),
            key=lambda path: path.name.casefold(),
        )
        return [template_record(path) for path in paths]


def default_template_name(templates=None):
    templates = templates if templates is not None else list_templates()
    names = {template["id"] for template in templates}
    if DEFAULT_TEMPLATE.name in names:
        return DEFAULT_TEMPLATE.name
    return templates[0]["id"] if templates else ""


def safe_template_filename(filename):
    raw_name = Path(filename or "").name.strip()
    if Path(raw_name).suffix.lower() != ".docx":
        raise ValueError("Selecione um arquivo no formato .docx.")
    stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", Path(raw_name).stem).strip(" .")
    if not stem:
        raise ValueError("O arquivo precisa ter um nome válido.")
    return f"{stem[:120]}.docx"


def save_template_candidate(field):
    filename = safe_template_filename(field.filename)
    temp_path = TEMPLATE_DIR / f".{uuid.uuid4().hex}.uploading"
    size = 0
    try:
        with temp_path.open("wb") as handle:
            while True:
                chunk = field.file.read(1024 * 1024)
                if not chunk:
                    break
                size += len(chunk)
                if size > MAX_TEMPLATE_SIZE:
                    raise ValueError("O arquivo excede o limite de 15 MB.")
                handle.write(chunk)
        if size == 0:
            raise ValueError("O arquivo enviado está vazio.")
        if not zipfile.is_zipfile(temp_path):
            raise ValueError("O arquivo não é um documento Word .docx válido.")
        with zipfile.ZipFile(temp_path) as package:
            members = set(package.namelist())
            if "[Content_Types].xml" not in members or "word/document.xml" not in members:
                raise ValueError("O arquivo não é um documento Word .docx válido.")
        Document(str(temp_path))
        return temp_path, filename
    except Exception:
        if temp_path.exists():
            temp_path.unlink()
        raise


def store_new_template(field):
    temp_path, filename = save_template_candidate(field)
    target = template_path_from_name(filename)
    try:
        with TEMPLATE_LOCK:
            if target.exists():
                raise FileExistsError("Já existe um template com esse nome. Use Substituir.")
            os.replace(temp_path, target)
            return template_record(target)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def replace_template_file(template_id, field):
    target = template_path_from_name(template_id)
    if target is None:
        raise FileNotFoundError("Template não encontrado.")
    temp_path, _ = save_template_candidate(field)
    try:
        with TEMPLATE_LOCK:
            if not target.exists():
                raise FileNotFoundError("Template não encontrado.")
            os.replace(temp_path, target)
            return template_record(target)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def delete_template_file(template_id):
    target = template_path_from_name(template_id)
    if target is None:
        raise FileNotFoundError("Template não encontrado.")
    with TEMPLATE_LOCK:
        if not target.exists():
            raise FileNotFoundError("Template não encontrado.")
        target.unlink()


def compact(value):
    if value is None:
        return ""
    value = str(value).replace("\r", "\n")
    lines = [" ".join(line.split()) for line in value.splitlines()]
    return " ".join(line for line in lines if line).strip()


def safe_public_url(value):
    text = str(value or "").strip()
    if not text:
        return ""
    if any(character.isspace() for character in text):
        return ""
    parsed = urlparse(text)
    if (
        parsed.scheme.lower() not in {"http", "https"}
        or not parsed.hostname
        or parsed.username
        or parsed.password
    ):
        return ""
    return text


def local_request_host(value):
    text = compact(value)
    if not text:
        return True
    parsed = urlparse(f"//{text}")
    return (parsed.hostname or "").lower() in LOCAL_REQUEST_HOSTS


def local_request_origin(value):
    text = compact(value)
    if not text:
        return True
    parsed = urlparse(text)
    return (
        parsed.scheme.lower() in {"http", "https"}
        and (parsed.hostname or "").lower() in LOCAL_REQUEST_HOSTS
    )


def cache_get(cache, key, ttl):
    with CACHE_LOCK:
        cached = cache.get(key)
        if not cached:
            return None
        if time.time() - cached["created_at"] >= ttl:
            cache.pop(key, None)
            return None
        return copy.deepcopy(cached["data"])


def cache_set(cache, key, data):
    with CACHE_LOCK:
        if key not in cache and len(cache) >= CACHE_MAX_ENTRIES:
            oldest_key = min(cache, key=lambda item: cache[item]["created_at"])
            cache.pop(oldest_key, None)
        cache[key] = {"created_at": time.time(), "data": copy.deepcopy(data)}


def cache_discard(cache, key):
    with CACHE_LOCK:
        cache.pop(key, None)


@contextmanager
def bloco2_enrichment_scope(cnpj, ano, sequencial):
    key = (str(cnpj), int(ano), int(sequencial))
    with BLOCO2_ENRICHMENT_LOCKS_GUARD:
        entry = BLOCO2_ENRICHMENT_LOCKS.get(key)
        if entry is None:
            entry = {"lock": threading.Lock(), "users": 0}
            BLOCO2_ENRICHMENT_LOCKS[key] = entry
        entry["users"] += 1
    lock = entry["lock"]
    lock.acquire()
    try:
        yield
    finally:
        lock.release()
        with BLOCO2_ENRICHMENT_LOCKS_GUARD:
            entry["users"] -= 1
            if entry["users"] == 0 and BLOCO2_ENRICHMENT_LOCKS.get(key) is entry:
                BLOCO2_ENRICHMENT_LOCKS.pop(key, None)


def norm(value):
    value = compact(value).lower()
    value = unicodedata.normalize("NFKD", value)
    value = "".join(ch for ch in value if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", " ", value).strip()


def empty_item():
    return {
        "lote": "",
        "item": "",
        "quantidade": "",
        "unidade": "",
        "descricao": "",
        "marca": "",
        "valor_unitario": "",
        "valor_total": "",
    }


def header_map(cells):
    mapping = {}
    for idx, cell in enumerate(cells):
        text = norm(cell)
        # Cabeçalhos são células curtas. Sem esse limite, uma especificação
        # extensa contendo palavras como "item" e "marca" vira um falso cabeçalho.
        if not text or len(text) > 80:
            continue
        if (
            re.search(r"\blote\b", text)
            or re.search(r"\bgrupo\b", text)
        ) and "lote" not in mapping:
            mapping["lote"] = idx
        if re.search(r"\bitem\b", text) and "item" not in mapping:
            mapping["item"] = idx
        if ("qtd" in text or "quant" in text) and "quantidade" not in mapping:
            mapping["quantidade"] = idx
        if (
            re.search(r"\bun\b", text)
            or "unid" in text
            or "unidade" in text
            or "medida" in text
        ) and "unidade" not in mapping:
            mapping["unidade"] = idx
        if (
            "descr" in text
            or "especific" in text
            or "objeto" in text
        ) and "descricao" not in mapping:
            mapping["descricao"] = idx
        if "marca" in text and "marca" not in mapping:
            mapping["marca"] = idx
        if (
            "valor unit" in text
            or "vlr unit" in text
            or "unitario" in text
        ) and "valor_unitario" not in mapping:
            mapping["valor_unitario"] = idx
        if (
            "valor total" in text
            or "vlr total" in text
            or ("total" in text and "valor_unitario" not in mapping)
        ) and "valor_total" not in mapping:
            mapping["valor_total"] = idx
    if "item" in mapping and len(set(mapping.values())) >= 2 and (
        "descricao" in mapping
        or "quantidade" in mapping
        or "unidade" in mapping
        or "valor_unitario" in mapping
    ):
        return mapping
    return None


def is_header_row(cells):
    return header_map(cells) is not None


def is_skip_row(cells):
    joined = norm(" ".join(cells))
    if not joined:
        return True
    return joined.startswith("anexo") or joined in {"item especificacao quant", "item descricao qtd"}


def is_item_table_end_row(cells):
    nonempty = [norm(cell) for cell in cells if norm(cell)]
    if not nonempty:
        return False
    return bool(re.match(r"^(?:total geral|subtotal|total do lote|valor total geral)\b", nonempty[0]))


def format_lote_value(value):
    text = norm(value)
    if not text:
        return ""
    if text.isdigit():
        return str(int(text))
    if text == "unico":
        return "Único"
    return text.upper()


def extract_lote_from_text(value):
    text = norm(value)
    if not text:
        return ""
    match = re.match(
        r"^(?:lote|grupo)(?:\s+(?:n|no|numero))?\s*(?:0+)?([0-9]+|[a-z]+|unico)\b",
        text,
    )
    if not match:
        return ""
    if match.group(1) in {"item", "descricao", "descr", "qtd", "quantidade", "unidade", "valor"}:
        return ""
    return format_lote_value(match.group(1))


def extract_lote_from_row(cells):
    nonempty = [cell for cell in cells if compact(cell)]
    if not nonempty:
        return ""
    joined = " ".join(nonempty)
    if len(joined) > 140:
        return ""
    return extract_lote_from_text(joined)


ACCOUNTING_CODE_PATTERN = re.compile(
    r"^(?:\d{2}\.\d{2}\.\d{2}|\d\.\d\.\d{2}\.\d{2})$"
)


def is_item_identifier(value):
    text = compact(value)
    if not text or ACCOUNTING_CODE_PATTERN.fullmatch(text):
        return False
    return bool(re.fullmatch(r"\d+(?:[\.-]\d+)*", text))


def get_cell(cells, idx):
    if idx is None or idx >= len(cells):
        return ""
    return compact(cells[idx])


def item_from_mapping(cells, mapping):
    row = empty_item()
    for key in row:
        row[key] = get_cell(cells, mapping.get(key))
    return row


def compact_header_mapping(header_cells, mapping):
    used_indexes = [idx for idx, cell in enumerate(header_cells) if compact(cell)]
    if not used_indexes:
        return None
    compacted = {}
    for key, idx in mapping.items():
        if idx in used_indexes:
            compacted[key] = used_indexes.index(idx)
    return compacted if compacted else None


def choose_row_mapping(cells, mapping, compact_mapping):
    if not mapping:
        return None
    if compact_mapping and max(mapping.values(), default=0) >= len(cells):
        return compact_mapping
    return mapping


def is_money_value(value):
    text = compact(value)
    return bool(re.fullmatch(r"(?:R\$\s*)?\d{1,3}(?:\.\d{3})*,\d{2}", text))


def parse_brazilian_number(value):
    text = compact(value)
    if not text:
        return None
    text = re.sub(r"[^0-9,.-]", "", text)
    if not text or text in {"-", ",", "."}:
        return None
    if "," in text:
        text = text.replace(".", "").replace(",", ".")
    elif text.count(".") > 1:
        text = text.replace(".", "")
    try:
        return Decimal(text)
    except InvalidOperation:
        return None


def format_brazilian_money(value):
    formatted = f"{value.quantize(Decimal('0.01')):,.2f}"
    return "R$ " + formatted.translate(str.maketrans({",": ".", ".": ","}))


def calculate_proposal_total(items):
    total = Decimal("0")
    for item in items:
        item_total = parse_brazilian_number(item.get("valor_total"))
        if item_total is None:
            quantity = parse_brazilian_number(item.get("quantidade"))
            unit_value = parse_brazilian_number(item.get("valor_unitario"))
            if quantity is not None and unit_value is not None:
                item_total = quantity * unit_value
        if item_total is not None:
            total += item_total
    return format_brazilian_money(total)


def is_quantity_value(value):
    return bool(re.fullmatch(r"\d+(?:[\.,]\d+)?", compact(value)))


def is_unit_value(value):
    return norm(value) in {
        "un", "und", "unid", "unidade", "kg", "g", "mg", "l", "ml", "m", "m2", "m3",
        "cm", "mm", "caixa", "cx", "pacote", "pct", "par", "jogo", "conjunto",
    }


def is_description_value(value):
    text = compact(value)
    if not text or is_item_identifier(text) or is_money_value(text) or is_unit_value(text):
        return False
    return norm(text) not in {
        "item", "descricao", "especificacao", "quantidade", "qtd", "valor unitario", "valor total",
    }


def item_from_sparse_cells(cells):
    row = empty_item()
    values = [compact(cell) for cell in cells if compact(cell)]
    if len(values) < 2 or not is_item_identifier(values[0]):
        return row

    row["item"] = values[0]
    description_candidates = [
        (idx, value)
        for idx, value in enumerate(values[1:], start=1)
        if is_description_value(value) and not is_quantity_value(value)
    ]
    if not description_candidates:
        return row
    description_idx, row["descricao"] = max(description_candidates, key=lambda candidate: len(candidate[1]))

    remaining = [value for idx, value in enumerate(values[1:], start=1) if idx != description_idx]
    money_values = [value for value in remaining if is_money_value(value)]
    if money_values:
        row["valor_unitario"] = money_values[0]
        if len(money_values) > 1:
            row["valor_total"] = money_values[1]

    quantity_values = [value for value in remaining if is_quantity_value(value) and not is_money_value(value)]
    if quantity_values:
        # Tabelas com quantidade mínima/máxima usam a última quantidade como total do item.
        row["quantidade"] = quantity_values[-1]

    unit_values = [value for value in remaining if is_unit_value(value)]
    if unit_values:
        row["unidade"] = unit_values[0]
    return row


def description_from_cells(cells, row_mapping=None):
    mapped = get_cell(cells, row_mapping.get("descricao") if row_mapping else None)
    if is_description_value(mapped) and not is_quantity_value(mapped):
        return mapped
    candidates = [
        compact(cell)
        for cell in cells
        if is_description_value(cell) and not is_quantity_value(cell)
    ]
    return max(candidates, key=len) if candidates else ""


def item_from_heuristic(cells):
    return item_from_sparse_cells(cells)


def normalize_rows(rows):
    items = []
    current = None
    mapping = None
    compact_mapping = None
    current_lote = ""

    for raw_row in rows:
        cells = [compact(cell) for cell in raw_row]
        if not any(cells):
            continue
        if is_item_table_end_row(cells):
            current = None
            mapping = None
            compact_mapping = None
            current_lote = ""
            continue
        if is_skip_row(cells):
            current = None
            continue

        detected = header_map(cells)
        if detected and is_header_row(cells):
            mapping = detected
            compact_mapping = compact_header_mapping(cells, detected)
            continue

        row_mapping = choose_row_mapping(cells, mapping, compact_mapping)
        row_lote = extract_lote_from_row(cells)
        if row_lote:
            current_lote = row_lote
            current = None
            continue

        item_cell = get_cell(cells, row_mapping.get("item") if row_mapping else 0)
        sparse_row = item_from_sparse_cells(cells)
        if not is_item_identifier(item_cell) and is_item_identifier(sparse_row.get("item")):
            item_cell = sparse_row["item"]
        if is_item_identifier(item_cell):
            row = item_from_mapping(cells, row_mapping) if row_mapping else item_from_heuristic(cells)
            if sparse_row.get("item") == item_cell:
                for field in row:
                    if not compact(row.get(field)) and compact(sparse_row.get(field)):
                        row[field] = sparse_row[field]
                if not is_description_value(row.get("descricao")) or is_quantity_value(row.get("descricao")):
                    row["descricao"] = sparse_row.get("descricao", "")
            if not row["item"]:
                row["item"] = item_cell
            elif not is_item_identifier(row["item"]):
                row["item"] = item_cell
            if compact(row.get("lote")):
                current_lote = compact(row.get("lote"))
            elif current_lote:
                row["lote"] = current_lote
            if row["descricao"] or row["quantidade"] or row["unidade"] or row["valor_unitario"]:
                items.append(row)
                current = row
            continue

        if current:
            desc = description_from_cells(cells, row_mapping)
            if desc:
                current["descricao"] = compact(current["descricao"] + " " + desc)

    return sanitize_extracted_items(items)


MONEY_PATTERN = r"(?:R\$\s*)?\d{1,3}(?:\.\d{3})*,\d{2}"
TEXT_ITEM_PATTERN = re.compile(
    r"^(?P<item>0*\d{1,4})\s+"
    r"(?P<quantidade>\d+(?:[\.,]\d+)*)\s+"
    r"(?P<unidade>\S{1,12})\s+"
    r"(?P<descricao>.+)$"
)
WIDE_TEXT_ITEM_PATTERN = re.compile(
    rf"^(?P<item>0*\d{{1,4}})\s+"
    rf"(?P<catalogo>\d[\d.-]{{2,}})\s+"
    rf"(?P<descricao>.*?)\s*"
    rf"(?P<quantidade>\d+(?:[\.,]\d+)?)\s+"
    rf"(?P<valor_unitario>{MONEY_PATTERN})\s+"
    rf"(?P<valor_total>{MONEY_PATTERN})$"
)
WIDE_TEXT_NOISE_PREFIXES = (
    "governo do estado",
    "fundacao hospital",
    "assessoria juridica",
    "formulario",
    "termo de referencia",
    "codigo",
    "pres asjur",
    "identificador de autenticacao",
    "n do protocolo",
    "pae n",
    "o que sera contratado",
    "quanti",
    "dade",
    "valor uni",
    "item simas",
    "tario",
    "para 12",
    "meses",
)
DESCRIPTION_DOCUMENT_NOISE = (
    "governo do estado",
    "fundacao hospital",
    "assessoria juridica",
    "termo de referencia ultima revisao",
    "valor global estimado",
    "descricao da solucao",
    "qual o motivo da contratacao",
    "natureza do bem",
    "criterios de selecao",
)


def identifier_sort_key(value):
    text = compact(value)
    if not text:
        return (1, (), "")
    parts = re.findall(r"\d+", text)
    if parts:
        return (0, tuple(int(part) for part in parts), text)
    return (1, (), text)


def normalize_item_reference(value):
    text = compact(value)
    if re.fullmatch(r"\d+", text):
        return str(int(text))
    return text


def item_sort_key(row):
    return (
        identifier_sort_key(row.get("lote")),
        identifier_sort_key(row.get("item")),
    )


def item_lookup_key(row):
    item = normalize_item_reference(row.get("item"))
    lote = normalize_item_reference(row.get("lote"))
    if not item:
        return ""
    return f"{lote}/{item}" if lote else item


def extracted_item_quality(row):
    description = compact(row.get("descricao"))
    return (
        bool(description),
        len(description),
        bool(compact(row.get("quantidade"))),
        bool(compact(row.get("unidade"))),
    )


def quantity_from_combined_value(value):
    text = compact(value)
    match = re.fullmatch(
        r"(?P<quantidade>\d+(?:[\.,]\d+)?)\s*(?:/|\s+)\s*"
        r"(?P<unidade>[A-Za-z]{1,12})\.?",
        text,
    )
    if not match or not is_unit_value(match.group("unidade")):
        return ""
    return match.group("quantidade")


def sanitize_extracted_items(items):
    valid_items = []
    for row in items:
        if not is_item_identifier(row.get("item")):
            continue
        normalized_row = dict(row)
        combined_quantity = quantity_from_combined_value(
            normalized_row.get("quantidade")
        )
        if combined_quantity:
            normalized_row["quantidade"] = combined_quantity
        normalized_row["unidade"] = STANDARD_UNIT
        valid_items.append(normalized_row)
    scoped_item_numbers = {
        normalize_item_reference(row.get("item"))
        for row in valid_items
        if compact(row.get("lote"))
    }

    # Em tabelas divididas por lotes, linhas soltas repetindo um item já lotado
    # normalmente vieram de outra tabela da mesma página (dotação, cronograma etc.).
    valid_items = [
        row for row in valid_items
        if compact(row.get("lote"))
        or normalize_item_reference(row.get("item")) not in scoped_item_numbers
    ]

    deduplicated = {}
    for row in valid_items:
        key = item_lookup_key(row)
        existing = deduplicated.get(key)
        if existing is None:
            deduplicated[key] = row
            continue

        preferred, fallback = (
            (row, existing)
            if extracted_item_quality(row) > extracted_item_quality(existing)
            else (existing, row)
        )
        merged = dict(preferred)
        for field in merged:
            if not compact(merged.get(field)) and compact(fallback.get(field)):
                merged[field] = fallback[field]
        deduplicated[key] = merged

    return sorted(deduplicated.values(), key=item_sort_key)


def is_spurious_document_item(row):
    item = compact(row.get("item"))
    quantity = compact(row.get("quantidade"))
    description = compact(row.get("descricao"))
    normalized = norm(description)
    numeric_parts = re.findall(r"\d+", item)

    if len(numeric_parts) > 3:
        return True
    if item.isdigit() and 1900 <= int(item) <= 2100:
        return True
    if not description or re.fullmatch(r"[-+]?\d+(?:[.,]\d+)?", description):
        return True
    if quantity and not is_quantity_value(quantity):
        return True

    legal_prefixes = (
        "nos termos do art",
        "o prazo estabelecido no item",
        "tx percentual da taxa",
        "documentos de habilitacao",
        "assinatura avancada",
    )
    if not is_quantity_value(quantity) and any(
        normalized.lstrip(".0123456789 ").startswith(prefix)
        for prefix in legal_prefixes
    ):
        return True
    return False


def remove_spurious_document_items(items):
    return [row for row in items if not is_spurious_document_item(row)]


def normalize_pdf_tables(tables):
    merged = []
    for table in tables:
        normalized = normalize_rows(table)
        merged = merge_item_lists(merged, normalized)
    return remove_spurious_document_items(merged)


def strip_price_columns(text):
    text = compact(text)
    text = re.sub(rf"\s+{MONEY_PATTERN}\s+{MONEY_PATTERN}\s*$", "", text)
    text = re.sub(rf"\s+{MONEY_PATTERN}\s*$", "", text)
    return compact(text)


def clean_text_item_line(line):
    line = compact(line)
    line = re.sub(r"^(?:[A-Z]\s+)+(?=0*\d{1,4}\s+\d)", "", line)
    line = re.sub(r"^[01]\s+(?=0\d{1,3}\s+\d)", "", line)
    return line


def clean_text_continuation(line):
    line = compact(line)
    line = re.sub(r"^(?:[A-Z]\s+)+(?=[A-Z0-9])", "", line)
    return strip_price_columns(line)


def is_pdf_text_noise(line):
    text = norm(line)
    if not text:
        return True
    if re.fullmatch(r"\d+", text):
        return True
    prefixes = (
        "municipio de",
        "prefeitura municipal",
        "departamento de",
        "rua ",
        "av ",
        "cnpj ",
        "cep ",
        "e mail",
        "telefone",
        "praca ",
        "pagina ",
        "www ",
        "pregao eletronico",
        "processo digital",
    )
    return any(text.startswith(prefix) for prefix in prefixes)


def is_text_table_stop(line):
    text = norm(line)
    stops = (
        "valor total",
        "valor total por extenso",
        "forma de solicitacao",
        "especificacao",
        "quantitativo",
        "estimativa de precos",
        "fundamentacao",
        "descricao da necessidade",
    )
    return any(text.startswith(stop) for stop in stops)


def parse_pdf_text_item_line(line):
    line = clean_text_item_line(line)
    match = TEXT_ITEM_PATTERN.match(line)
    if not match:
        return None
    row = empty_item()
    row["item"] = match.group("item")
    row["quantidade"] = match.group("quantidade")
    row["unidade"] = match.group("unidade").rstrip(".")
    row["descricao"] = strip_price_columns(match.group("descricao"))
    if not row["descricao"]:
        return None
    return row


def extract_from_pdf_text(path):
    items = []
    current = None
    current_lote = ""
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            for raw_line in text.splitlines():
                line = compact(raw_line)
                if is_pdf_text_noise(line):
                    continue
                if is_text_table_stop(line):
                    current = None
                    continue

                line_lote = extract_lote_from_text(line)
                if line_lote:
                    current_lote = line_lote
                    current = None
                    continue

                row = parse_pdf_text_item_line(line)
                if row:
                    if current_lote:
                        row["lote"] = current_lote
                    items.append(row)
                    current = row
                    continue

                if current:
                    desc = clean_text_continuation(line)
                    if desc and not is_header_row([desc]) and not is_text_table_stop(desc):
                        current["descricao"] = compact(current["descricao"] + " " + desc)

    return sanitize_extracted_items(items)


def join_pdf_description_lines(lines):
    description = ""
    for raw_line in lines:
        line = compact(raw_line)
        if not line:
            continue
        if description.endswith("-") and re.match(r"^[a-zà-ÿ]", line):
            description = description[:-1] + line
        else:
            description = compact(description + " " + line)
    return description


def extract_from_wide_pdf_texts(page_texts):
    lines = []
    stopped = False
    for page_text in page_texts:
        for raw_line in (page_text or "").splitlines():
            line = compact(raw_line)
            normalized = norm(line)
            if normalized.startswith("valor global estimado"):
                stopped = True
                break
            if not line or any(normalized.startswith(prefix) for prefix in WIDE_TEXT_NOISE_PREFIXES):
                continue
            lines.append(line)
        if stopped:
            break

    anchors = []
    for index, line in enumerate(lines):
        match = WIDE_TEXT_ITEM_PATTERN.match(line)
        if match:
            anchors.append((index, match))
    if len(anchors) < 2:
        return []

    items = []
    for position, (line_index, match) in enumerate(anchors):
        start = 0 if position == 0 else (anchors[position - 1][0] + line_index) // 2 + 1
        end = (
            len(lines) - 1
            if position + 1 == len(anchors)
            else (line_index + anchors[position + 1][0]) // 2
        )
        description_lines = [
            match.group("descricao") if index == line_index else lines[index]
            for index in range(start, end + 1)
        ]
        row = empty_item()
        row.update({
            "item": match.group("item"),
            "quantidade": match.group("quantidade"),
            "unidade": STANDARD_UNIT,
            "descricao": join_pdf_description_lines(description_lines),
            "valor_unitario": match.group("valor_unitario"),
            "valor_total": match.group("valor_total"),
        })
        items.append(row)
    return sanitize_extracted_items(items)


def description_has_document_noise(description):
    text = norm(description)
    return any(marker in text for marker in DESCRIPTION_DOCUMENT_NOISE)


def merge_item_lists(primary, fallback):
    merged = {}
    for row in primary:
        key = item_lookup_key(row)
        if not key:
            continue
        merged[key] = dict(row)

    for row in fallback:
        key = item_lookup_key(row)
        if not key:
            continue
        existing = merged.get(key)
        if existing:
            for field in existing:
                if not compact(existing.get(field)) and compact(row.get(field)):
                    existing[field] = row.get(field, "")
            if is_quantity_value(row.get("quantidade")) and not is_quantity_value(existing.get("quantidade")):
                existing["quantidade"] = row["quantidade"]
            if is_money_value(row.get("valor_unitario")) and not is_money_value(existing.get("valor_unitario")):
                existing["valor_unitario"] = row["valor_unitario"]
            if is_money_value(row.get("valor_total")) and not is_money_value(existing.get("valor_total")):
                existing["valor_total"] = row["valor_total"]
            if (
                compact(row.get("descricao"))
                and description_has_document_noise(existing.get("descricao"))
                and not description_has_document_noise(row.get("descricao"))
            ):
                existing["descricao"] = row["descricao"]
            continue
        merged[key] = dict(row)
    return sanitize_extracted_items(merged.values())


def item_key_set(items):
    return {item_lookup_key(row) for row in items if item_lookup_key(row)}


def build_pncp_items_check(file_items, pncp_items):
    file_keys = item_key_set(file_items)
    pncp_keys = item_key_set(pncp_items)
    added_from_pncp = sorted(pncp_keys - file_keys, key=identifier_sort_key)
    only_in_file = sorted(file_keys - pncp_keys, key=identifier_sort_key)
    return {
        "file_count": len(file_items),
        "pncp_count": len(pncp_items),
        "added_from_pncp": added_from_pncp,
        "only_in_file": only_in_file,
        "has_divergence": bool(added_from_pncp or only_in_file or len(file_items) != len(pncp_items)),
    }


DESCRIPTION_END_CONNECTORS = {
    "a", "ao", "aos", "as", "atraves", "com", "da", "das", "de", "do", "dos",
    "e", "em", "entre", "mediante", "ou", "para", "pela", "pelas", "pelo", "pelos",
    "por", "que", "sem", "sob", "sobre",
}
DESCRIPTION_START_CONNECTORS = {
    "com", "da", "das", "de", "do", "dos", "e", "ou", "para", "pela", "pelas",
    "pelo", "pelos", "por", "que", "sem", "sob", "sobre",
}


def description_integrity_findings(description):
    description = compact(description)
    blocking = []
    warnings_found = []
    if not description:
        return ["descrição ausente"], []
    if description_has_document_noise(description):
        blocking.append("descrição contém cabeçalho ou seção externa ao item")

    words = norm(description).split()
    if words and words[-1] in DESCRIPTION_END_CONNECTORS:
        warnings_found.append(f'possível fim interrompido em "{words[-1]}"')
    elif description.rstrip().endswith((",", ";", ":", "-", "/")):
        warnings_found.append("possível fim interrompido por pontuação de continuidade")

    if words and words[0] in DESCRIPTION_START_CONNECTORS:
        warnings_found.append(f'início aparenta continuação em "{words[0]}"')

    for opening, closing, label in (("(", ")", "parênteses"), ("[", "]", "colchetes")):
        if description.count(opening) != description.count(closing):
            warnings_found.append(f"{label} não balanceados")

    return blocking, warnings_found


def build_description_review(items, file_items=None, pncp_items=None):
    file_items = file_items or []
    pncp_items = pncp_items or []
    final_lookup = {item_lookup_key(row): row for row in items if item_lookup_key(row)}
    file_lookup = {item_lookup_key(row): row for row in file_items if item_lookup_key(row)}
    pncp_lookup = {item_lookup_key(row): row for row in pncp_items if item_lookup_key(row)}
    expected_keys = set(pncp_lookup) or set(final_lookup)

    def sort_keys(values):
        return sorted(values, key=identifier_sort_key)

    missing_items = sort_keys(expected_keys - set(final_lookup))
    blocking_items = []
    warning_items = []
    pncp_fallback_items = []
    different_source_items = []
    details = []

    for key in sort_keys(set(final_lookup)):
        description = compact(final_lookup[key].get("descricao"))
        file_description = compact(file_lookup.get(key, {}).get("descricao"))
        pncp_description = compact(pncp_lookup.get(key, {}).get("descricao"))
        blocking, warnings_found = description_integrity_findings(description)

        if blocking:
            blocking_items.append(key)
        if warnings_found:
            warning_items.append(key)
        if file_lookup and not file_description and pncp_description:
            pncp_fallback_items.append(key)
        if file_description and pncp_description and file_description != pncp_description:
            different_source_items.append(key)

        if description and description == file_description and description == pncp_description:
            source = "arquivo oficial e aba PNCP"
        elif description and description == file_description:
            source = "arquivo oficial"
        elif description and description == pncp_description:
            source = "aba oficial PNCP"
        else:
            source = "conciliação das fontes"

        details.append({
            "item": key,
            "status": "erro" if blocking else ("atenção" if warnings_found else "ok"),
            "source": source,
            "characters": len(description),
            "blocking": blocking,
            "warnings": warnings_found,
        })

    problem_items = sort_keys(set(missing_items + blocking_items))
    if problem_items:
        status = "error"
        joined = ", ".join(problem_items)
        message = f"Revisão das descrições reprovada: descrição ausente no(s) item(ns) {joined}."
    elif warning_items or pncp_fallback_items:
        status = "warn"
        notices = []
        if warning_items:
            notices.append(f"item(ns) {', '.join(sort_keys(set(warning_items)))} exigem conferência")
        if pncp_fallback_items:
            notices.append(f"item(ns) {', '.join(sort_keys(set(pncp_fallback_items)))} vieram da aba oficial PNCP")
        message = "Revisão das descrições concluída com atenção: " + "; ".join(notices) + "."
    else:
        status = "ok"
        message = (
            f"Revisão das descrições concluída: {len(details)} de {len(details)} item(ns) "
            "com início, continuidade e fim, sem sinais de corte."
        )

    return {
        "status": status,
        "message": message,
        "reviewed_count": len(details),
        "complete_count": len(details) - len(set(blocking_items)),
        "missing_items": missing_items,
        "blocking_items": sort_keys(set(blocking_items)),
        "warning_items": sort_keys(set(warning_items)),
        "pncp_fallback_items": sort_keys(set(pncp_fallback_items)),
        "different_source_items": sort_keys(set(different_source_items)),
        "items": details,
    }


ITEM_CATEGORY_RULES = [
    (r"\bcadeira\b", "Cadeira"),
    (r"\bsofa\b", "Sofá"),
    (r"\bpoltrona\b", "Poltrona"),
    (r"\barmario\b", "Armário"),
    (r"\barquivo\b", "Arquivo"),
    (r"\bmesa\b", "Mesa"),
    (r"\bestacao\b", "Estação"),
    (r"\bestante\b", "Estante"),
    (r"\bsuporte\b", "Suporte"),
    (r"\bquadro\b", "Quadro"),
    (r"\baspirador\b", "Aspirador"),
    (r"\bbalanca\b", "Balança"),
    (r"\bbebedouro\b", "Bebedouro"),
    (r"\bcoifa\b", "Coifa"),
    (r"\bespremedor\b|extrator de frutas", "Espremedor"),
    (r"\bexaustor\b", "Exaustor"),
    (r"\bfogao\b", "Fogão"),
    (r"\bmicroondas\b", "Micro-ondas"),
    (r"\bforno\b", "Forno"),
    (r"\bfreezer\b", "Freezer"),
    (r"\bgeladeira\b|refrigerador", "Geladeira"),
    (r"\bliquidificador\b", "Liquidificador"),
    (r"\bpanela\b", "Panela"),
    (r"\bpicador\b", "Picador"),
    (r"\bprocessador\b", "Processador"),
    (r"\bventilador\b", "Ventilador"),
    (r"\bimpressora\b.*\bmultifuncional\b", "Impressora multifuncional"),
    (r"\bimpressora\b", "Impressora"),
    (
        r"\b(?:oculos|lentes?|esferic\w*|cilindric\w*|bifoc\w*|multifoc\w*|progressiv\w*)\b",
        "Óculos grau",
    ),
]


CATEGORY_FALLBACK_STOPWORDS = {
    "a", "ao", "aos", "as", "ate", "com", "contratacao", "cor", "da", "das", "de", "do", "dos",
    "empresa", "em", "e", "especificacao", "fornecimento", "item", "marca", "material", "medida",
    "modelo", "o", "objeto", "para", "por", "prestacao", "sem", "servico", "tipo", "unidade",
    "cm", "g", "kg", "l", "m", "mg", "ml", "mm", "und",
}


def fallback_item_category(description):
    text = compact(description)
    words = re.findall(r"[A-Za-zÀ-ÿ]+(?:-[A-Za-zÀ-ÿ]+)?", text)
    for word in words:
        if norm(word) not in CATEGORY_FALLBACK_STOPWORDS:
            return word.title()
    return "Indefinido"


def identify_item_category(description):
    text = norm(description)
    for pattern, label in ITEM_CATEGORY_RULES:
        if re.search(pattern, text):
            return label
    return fallback_item_category(description)


def build_item_identifications(items):
    identifications = []
    for row in sorted(items, key=item_sort_key):
        description = compact(row.get("descricao"))
        identifications.append({
            "lote": compact(row.get("lote")),
            "item": compact(row.get("item")),
            "quantidade": compact(row.get("quantidade")),
            "unidade": STANDARD_UNIT,
            "categoria": identify_item_category(description),
            "descricao": description,
        })
    return identifications


def nullable_compact(value):
    return None if value is None else compact(value)


def valid_bloco2_cache_entry(cached, cnpj, ano, sequencial):
    if not isinstance(cached, dict):
        return False
    pncp = cached.get("pncp")
    items = cached.get("items")
    if not isinstance(pncp, dict) or not isinstance(items, list) or not items:
        return False
    if (
        compact(pncp.get("cnpj")) != str(cnpj)
        or pncp.get("ano") != int(ano)
        or pncp.get("sequencial") != int(sequencial)
    ):
        return False
    if not str(cached.get("source") or "").startswith("opportunity_items"):
        return False
    return all(
        isinstance(item, dict)
        and compact(item.get("item"))
        and compact(item.get("descricao"))
        for item in items
    )


def raw_bloco2_item_identity(payload):
    if not isinstance(payload, dict):
        return None
    item_number = compact(
        payload.get("numeroItem")
        or payload.get("numero")
        or payload.get("item")
        or payload.get("item_number")
    )
    lot_number = compact(
        payload.get("numeroGrupo")
        or payload.get("grupo")
        or payload.get("lote")
        or payload.get("numeroLote")
    )
    description = compact(
        payload.get("descricao")
        or payload.get("descricaoItem")
        or payload.get("description")
        or payload.get("itemDescricao")
        or payload.get("descricaoCompleta")
        or payload.get("descricaoDetalhada")
    )
    if not item_number or not description:
        return None
    return lot_number, item_number


def normalize_bloco2_api_items(raw_items):
    if not isinstance(raw_items, list) or not raw_items:
        raise RuntimeError("A API do PNCP nao retornou itens validos para esta oportunidade.")
    identities = []
    for payload in raw_items:
        identity = raw_bloco2_item_identity(payload)
        if identity is None:
            raise RuntimeError("A API do PNCP retornou um item incompleto ou malformado.")
        identities.append(identity)
    if len(set(identities)) != len(identities):
        raise RuntimeError("A API do PNCP retornou itens duplicados para esta oportunidade.")

    mapped_items = PNCPMapper().map_items(raw_items)
    if len(mapped_items) != len(raw_items):
        raise RuntimeError("A normalizacao dos itens do PNCP ficou incompleta.")
    for payload, mapped in zip(raw_items, mapped_items):
        has_explicit_total = any(
            payload.get(key) is not None and payload.get(key) != ""
            for key in ("valorTotal", "valorTotalEstimado", "estimated_total_value")
        )
        if not has_explicit_total:
            mapped.estimated_total_value = None
    return mapped_items


def identify_items_from_opportunity_store(cnpj, ano, sequencial):
    detail = etl_repository().get_opportunity_by_pncp_identity(cnpj, ano, sequencial)
    if not detail or not detail.get("items"):
        return None

    items = []
    for stored in detail["items"]:
        description = compact(
            stored.get("description")
            or stored.get("technical_object")
            or stored.get("title")
        )
        item_number = compact(stored.get("item_number"))
        if not item_number or not description:
            continue
        items.append({
            "lote": compact(stored.get("lot_number")),
            "item": item_number,
            "quantidade": (
                None
                if stored.get("quantity") is None
                else format_pncp_quantity(stored.get("quantity"))
            ),
            "unidade": nullable_compact(stored.get("unit")),
            "categoria": identify_item_category(description),
            "descricao": description,
            "valor_unitario_estimado": stored.get("estimated_unit_value"),
            "valor_total_estimado": stored.get("estimated_total_value"),
        })
    if not items:
        return None

    items.sort(key=item_sort_key)
    description_review = build_description_review(items)
    for review in description_review.get("items", []):
        review["source"] = "opportunity_items"
    description_review["message"] = (
        f"{len(items)} item(ns) carregado(s) da base estruturada opportunity_items."
    )
    return {
        "count": len(items),
        "source": "opportunity_items",
        "items": items,
        "description_review": description_review,
        "pncp_items_check": {
            "file_count": 0,
            "pncp_count": len(items),
            "structured_count": len(items),
            "has_divergence": False,
            "source": "opportunity_items",
        },
        "quantity_reconciliation": {
            "source": "opportunity_items",
            "matched_count": len(items),
            "unmatched_file_count": 0,
            "unmatched_pncp_count": 0,
        },
        "pncp": {
            "cnpj": cnpj,
            "ano": ano,
            "sequencial": sequencial,
            "link": pncp_app_link(cnpj, ano, sequencial),
            "documento_usado": "",
            "documento_tipo": "Base estruturada",
        },
    }


def import_pncp_opportunity_on_demand(
    cnpj,
    ano,
    sequencial,
    *,
    repository=None,
    connector=None,
):
    repository = repository or etl_repository()
    connector = connector or PNCPConnector()
    request_url = pncp_app_link(cnpj, ano, sequencial)
    counters = {"fetched": 0, "inserted": 0, "updated": 0, "skipped": 0, "failed": 0}
    run_id = repository.create_run("pncp", "opportunity_on_demand", {
        "cnpj": cnpj,
        "year": ano,
        "sequence": sequencial,
    })
    try:
        fetched = connector.fetch_detail(cnpj, ano, sequencial)
        request_url = fetched.request_url
        payload = fetched.payload
        if not isinstance(payload, dict) or not payload:
            raise RuntimeError("A API do PNCP nao retornou os dados da contratacao.")

        mapping_payload = dict(payload)
        mapping_payload.setdefault("numeroCnpj", cnpj)
        mapping_payload.setdefault("anoCompra", ano)
        mapping_payload.setdefault("sequencialCompra", sequencial)
        mapping_payload.setdefault(
            "numeroControlePNCP",
            f"{cnpj}-1-{int(sequencial):06d}/{int(ano)}",
        )
        opportunity = PNCPMapper().map(mapping_payload)
        received_identity = (
            opportunity.source_cnpj,
            opportunity.year,
            opportunity.sequence,
        )
        expected_identity = (str(cnpj), int(ano), int(sequencial))
        if received_identity != expected_identity:
            raise RuntimeError("A API do PNCP retornou uma contratacao diferente da solicitada.")

        outcome, _ = repository.persist_record(
            run_id=run_id,
            source_endpoint="detail",
            request_url=request_url,
            raw_payload={"detail": payload},
            opportunity=opportunity,
            match=OpportunityClassifier().classify(opportunity, {}),
            replace_children=False,
        )
        counters["fetched"] = 1
        counters[outcome] += 1
        repository.finish_run(run_id, status="success", counters=counters)
    except Exception as exc:
        counters["failed"] = 1
        repository.finish_run(
            run_id,
            status="failed",
            counters=counters,
            error_message=str(exc),
        )
        raise RuntimeError(
            "Esta contratacao nao esta no banco local e nao foi possivel importa-la "
            "diretamente da API do PNCP agora. Tente novamente mais tarde."
        ) from exc

    detail = repository.get_opportunity_by_pncp_identity(cnpj, ano, sequencial)
    if detail is None:
        raise RuntimeError("A contratacao foi consultada, mas nao foi gravada no banco local.")
    return detail


def enrich_opportunity_items(cnpj, ano, sequencial):
    with bloco2_enrichment_scope(cnpj, ano, sequencial):
        structured = identify_items_from_opportunity_store(cnpj, ano, sequencial)
        if structured is not None:
            structured["cache_status"] = "hit"
            return structured

        repository = etl_repository()
        connector = PNCPConnector()
        detail = repository.get_opportunity_by_pncp_identity(cnpj, ano, sequencial)
        if detail is None:
            detail = import_pncp_opportunity_on_demand(
                cnpj,
                ano,
                sequencial,
                repository=repository,
                connector=connector,
            )

        opportunity_id = detail["opportunity"]["id"]
        run_id = repository.create_run("pncp", "opportunity_item_enrichment", {
            "opportunity_id": opportunity_id,
            "cnpj": cnpj,
            "year": ano,
            "sequence": sequencial,
        })
        request_url = pncp_app_link(cnpj, ano, sequencial)
        page_count = 0
        raw_item_count = 0
        try:
            raw_items = []
            for page in connector.iter_items(cnpj, ano, sequencial, 20):
                raw_items.extend(page.records)
                page_count += 1
                raw_item_count += len(page.records)
                request_url = page.request_url
            mapped_items = normalize_bloco2_api_items(raw_items)
            persistence = repository.persist_opportunity_items_enrichment(
                run_id=run_id,
                opportunity_id=opportunity_id,
                items=mapped_items,
                request_url=request_url,
                audit_summary={
                    "pages_received": page_count,
                    "items_received": raw_item_count,
                    "items_normalized": len(mapped_items),
                },
                external_key=detail["opportunity"].get("external_key"),
            )
        except Exception as exc:
            try:
                repository.record_opportunity_items_enrichment_failure(
                    run_id=run_id,
                    opportunity_id=opportunity_id,
                    request_url=request_url,
                    audit_summary={
                        "pages_received": page_count,
                        "items_received": raw_item_count,
                    },
                    error_message=f"{type(exc).__name__}: {exc}",
                    external_key=detail["opportunity"].get("external_key"),
                )
            except Exception:
                pass
            raise RuntimeError(
                "Nao foi possivel carregar os itens desta oportunidade agora. "
                "A oportunidade permanece salva no banco e os itens continuam pendentes. "
                "Tente novamente mais tarde."
            ) from exc

        structured = identify_items_from_opportunity_store(cnpj, ano, sequencial)
        if structured is None:
            raise RuntimeError(
                "O enriquecimento concluiu, mas nenhum item foi gravado para esta oportunidade."
            )

        structured["source"] = "opportunity_items_enriquecido"
        structured["cache_status"] = (
            "miss_enriched" if persistence["persisted"] else "hit_concurrent"
        )
        structured["description_review"]["message"] = (
            f"{len(structured['items'])} item(ns) enriquecido(s), salvo(s) no banco "
            "e carregado(s) da base local."
        )
        return structured


def enrich_items_for_bloco2(cnpj, ano, sequencial):
    if not ALLOW_BLOCO2_ON_DEMAND_ENRICHMENT:
        raise RuntimeError("Enriquecimento sob demanda do Bloco 2 esta desativado.")
    structured = enrich_opportunity_items(cnpj, ano, sequencial)
    if structured.get("cache_status") == "miss_enriched":
        structured["source"] = "opportunity_items_enriquecido_bloco2"
        structured["description_review"]["message"] = (
            f"{len(structured['items'])} item(ns) enriquecido(s), salvo(s) no banco "
            "e carregado(s) para o Bloco 2."
        )
    return structured


def identify_items_from_pncp_link(link):
    cnpj, ano, sequencial = parse_pncp_link(link)
    cache_key = pncp_app_link(cnpj, ano, sequencial)
    cached = cache_get(IDENTIFICATION_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if valid_bloco2_cache_entry(cached, cnpj, ano, sequencial):
        return cached
    if cached is not None:
        cache_discard(IDENTIFICATION_CACHE, cache_key)

    structured = identify_items_from_opportunity_store(cnpj, ano, sequencial)
    if structured is not None:
        check = {
            "file_count": 0,
            "pncp_count": 0,
            "structured_count": len(structured["items"]),
            "added_from_pncp": [],
            "only_in_file": [],
            "has_divergence": False,
            "api_available": False,
            "api_error": "Itens encontrados no banco local; consulta ao PNCP nao necessaria.",
            "source": "opportunity_items_base_local",
        }
        structured["description_review"]["message"] = (
            f"{len(structured['items'])} item(ns) carregado(s) da base estruturada local."
        )
        structured["pncp_items_check"] = check
        cache_set(IDENTIFICATION_CACHE, cache_key, structured)
        return structured

    if ALLOW_BLOCO2_ON_DEMAND_ENRICHMENT:
        structured = enrich_items_for_bloco2(cnpj, ano, sequencial)
        cache_set(IDENTIFICATION_CACHE, cache_key, structured)
        return structured

    raise RuntimeError(
        "Esta contratacao ainda nao tem itens no banco local e o preenchimento "
        "sob demanda do Bloco 2 esta desativado."
    )


def pncp_purchase_metadata(cnpj, ano, sequencial):
    url = (
        f"{PNCP_API_BASE}/consulta/v1/orgaos/{cnpj}/compras/"
        f"{ano}/{sequencial}"
    )
    payload = request_json(url)
    if not isinstance(payload, dict) or payload.get("timeout"):
        return {}
    orgao = payload.get("orgaoEntidade") or {}
    unidade = payload.get("unidadeOrgao") or {}
    return {
        "numero_compra": compact(payload.get("numeroCompra")),
        "processo": compact(payload.get("processo")),
        "modalidade": compact(payload.get("modalidadeNome")),
        "objeto": compact(payload.get("objetoCompra")),
        "orgao": compact(orgao.get("razaoSocial")),
        "orgao_cnpj": compact(orgao.get("cnpj")) or compact(cnpj),
        "unidade": compact(unidade.get("nomeUnidade")),
        "municipio": compact(unidade.get("municipioNome")),
        "uf": compact(unidade.get("ufSigla")),
        "numero_controle_pncp": compact(payload.get("numeroControlePNCP")),
        "abertura": compact(payload.get("dataAberturaProposta")),
        "encerramento": compact(payload.get("dataEncerramentoProposta")),
        "situacao": compact(
            payload.get("situacaoCompraNome")
            or payload.get("situacaoCompra")
        ),
        "valor_total_estimado": payload.get("valorTotalEstimado"),
        "modo_disputa": compact(
            payload.get("modoDisputaNome")
            or payload.get("modoDisputa")
        ),
        "codigo_unidade": compact(
            unidade.get("codigoUnidade")
            or unidade.get("codigo")
        ),
        "link_sistema_origem": compact(
            payload.get("linkSistemaOrigem")
            or payload.get("urlSistemaOrigem")
        ),
    }


def business_id_value(value, label="Negócio"):
    text = str(value or "").strip()
    if not text.isdigit() or int(text) <= 0:
        raise ValueError(f"{label} inválido.")
    return int(text)


def business_record(row):
    if not row:
        return None
    return {
        "id": str(row["id"]),
        "contratacao_id": str(row["contratacao_id"] or ""),
        "empresa": row["empresa"],
        "cnpj_orgao": row["cnpj_orgao"],
        "ano": row["ano"],
        "sequencial": row["sequencial"],
        "link_pncp": row["link_pncp"],
        "titulo": row["titulo_interno"] or row["titulo"],
        "titulo_oficial": row["titulo"],
        "orgao": row["orgao"],
        "municipio": row["municipio"],
        "uf": row["uf"],
        "modalidade": row["modalidade"],
        "numero_compra": row["numero_compra"],
        "processo": row["processo"],
        "plataforma": row["plataforma"],
        "fonte_integracao": row["fonte_integracao"],
        "abertura": row["abertura"],
        "encerramento": row["encerramento"],
        "etapa": row["etapa"],
        "situacao": row["situacao"],
        "prioridade": row["prioridade"],
        "favorito": bool(row["favorito"]),
        "arquivado": bool(row["arquivado"]),
        "removido": bool(row["removido"]),
        "responsavel": row["responsavel"],
        "prazo_interno": row["prazo_interno"],
        "anotacoes": row["anotacoes"],
        "decisao_comercial": row["decisao_comercial"],
        "checklist_concluido": row["checklist_concluido"],
        "checklist_total": row["checklist_total"],
        "total_itens": row["total_itens"],
        "criado_em": row["criado_em"],
        "atualizado_em": row["atualizado_em"],
        "position_number": row["effective_position_number"],
        "pode_mover": True,
    }


BUSINESS_SELECT = """
    SELECT n.*,
        (SELECT COUNT(*) FROM negocio_tarefas t
         WHERE t.negocio_id = n.id AND t.concluida = 1) AS checklist_concluido,
        (SELECT COUNT(*) FROM negocio_tarefas t
         WHERE t.negocio_id = n.id) AS checklist_total,
        (SELECT COUNT(*) FROM negocio_itens i
         WHERE i.negocio_id = n.id) AS total_itens,
        COALESCE(n.position_number, (SELECT p.position_number FROM proposals p
         WHERE p.business_id = n.id LIMIT 1)) AS effective_position_number
    FROM negocios n
"""


def list_businesses(include_archived=False):
    init_database()
    where = "WHERE removido = 0"
    if not include_archived:
        where += " AND arquivado = 0"
    with DATABASE_LOCK, database_connection() as connection:
        rows = connection.execute(
            f"{BUSINESS_SELECT} {where} ORDER BY atualizado_em DESC, id DESC"
        ).fetchall()
    return [business_record(row) for row in rows]


def business_files(connection, row):
    files = []
    if row["contratacao_id"]:
        stored = connection.execute(
            """
            SELECT titulo, tipo_documento, url, selecionado
            FROM arquivos_pncp
            WHERE contratacao_id = ?
            ORDER BY selecionado DESC, id
            """,
            (row["contratacao_id"],),
        ).fetchall()
        files = [
            {
                "titulo": item["titulo"] or item["tipo_documento"] or "Arquivo oficial",
                "tipo": item["tipo_documento"],
                "url": item["url"],
                "selecionado": bool(item["selecionado"]),
            }
            for item in stored
        ]
    if files:
        return files
    try:
        remote = list_pncp_files(row["cnpj_orgao"], row["ano"], row["sequencial"])
    except Exception:
        remote = []
    return [
        {
            "titulo": compact(
                item.get("titulo")
                or item.get("tipoDocumentoNome")
                or "Arquivo oficial"
            ),
            "tipo": compact(
                item.get("tipoDocumentoNome")
                or item.get("tipoDocumentoDescricao")
            ),
            "url": compact(item.get("url") or item.get("uri")),
            "selecionado": False,
        }
        for item in remote
    ]


def get_business(business_id, include_details=False):
    business_id = business_id_value(business_id)
    init_database()
    with DATABASE_LOCK, database_connection() as connection:
        row = connection.execute(
            f"{BUSINESS_SELECT} WHERE n.id = ?",
            (business_id,),
        ).fetchone()
        if not row:
            raise FileNotFoundError("Negócio não encontrado.")
        result = business_record(row)
        if include_details:
            result["historico"] = [
                {
                    "id": str(item["id"]),
                    "evento": item["evento"],
                    "etapa_anterior": item["etapa_anterior"],
                    "etapa_nova": item["etapa_nova"],
                    "justificativa": item["justificativa"],
                    "criado_em": item["criado_em"],
                }
                for item in connection.execute(
                    """
                    SELECT * FROM negocio_historico
                    WHERE negocio_id = ?
                    ORDER BY criado_em DESC, id DESC
                    """,
                    (business_id,),
                ).fetchall()
            ]
            result["tarefas"] = [
                {
                    "id": str(item["id"]),
                    "titulo": item["titulo"],
                    "concluida": bool(item["concluida"]),
                    "ordem": item["ordem"],
                }
                for item in connection.execute(
                    """
                    SELECT * FROM negocio_tarefas
                    WHERE negocio_id = ?
                    ORDER BY ordem, id
                    """,
                    (business_id,),
                ).fetchall()
            ]
            result["arquivos"] = business_files(connection, row)
            result["itens"] = [
                {
                    "id": str(item["id"]),
                    "ordem": item["ordem"],
                    "lote": item["lote"],
                    "numero": item["numero_item"],
                    "descricao": item["descricao"],
                    "quantidade": item["quantidade"],
                    "unidade": item["unidade"],
                    "valor_unitario_estimado": item["valor_unitario_estimado"],
                    "valor_total_estimado": item["valor_total_estimado"],
                    "criterio_julgamento": item["criterio_julgamento"],
                    "situacao": item["situacao"],
                }
                for item in connection.execute(
                    """
                    SELECT * FROM negocio_itens
                    WHERE negocio_id = ?
                    ORDER BY ordem, id
                    """,
                    (business_id,),
                ).fetchall()
            ]
    return result


def default_business_company():
    responsibles = list_responsibles()
    if responsibles:
        return responsibles[0]["empresa"]
    return "Empresa principal"


def ensure_business_tasks(connection, business_id, now):
    current = connection.execute(
        "SELECT COUNT(*) FROM negocio_tarefas WHERE negocio_id = ?",
        (business_id,),
    ).fetchone()[0]
    if current:
        return
    connection.executemany(
        """
        INSERT INTO negocio_tarefas (
            negocio_id, titulo, concluida, ordem, criado_em, atualizado_em
        ) VALUES (?, ?, 0, ?, ?, ?)
        """,
        [
            (business_id, title, index, now, now)
            for index, title in enumerate(DEFAULT_BUSINESS_TASKS, start=1)
        ],
    )


def validate_business_items(payload):
    if "itens" not in payload:
        return None
    raw_items = payload.get("itens")
    if not isinstance(raw_items, list) or not raw_items:
        raise ValueError("Selecione ao menos um item para adicionar ao negócio.")
    if len(raw_items) > 5000:
        raise ValueError("A seleção excede o limite de 5.000 itens por negócio.")
    items = []
    seen = set()
    for index, raw in enumerate(raw_items, start=1):
        if not isinstance(raw, dict):
            raise ValueError("A seleção de itens contém dados inválidos.")
        item = {
            "ordem": index,
            "lote": compact(raw.get("lote")),
            "numero": compact(raw.get("numero") or raw.get("item")),
            "descricao": compact(raw.get("descricao")),
            "quantidade": compact(raw.get("quantidade")),
            "unidade": compact(raw.get("unidade")) or STANDARD_UNIT,
            "valor_unitario_estimado": compact(raw.get("valor_unitario_estimado")),
            "valor_total_estimado": compact(raw.get("valor_total_estimado")),
            "criterio_julgamento": compact(raw.get("criterio_julgamento")),
            "situacao": compact(raw.get("situacao")),
        }
        if not item["numero"] or not item["descricao"]:
            raise ValueError("Todo item selecionado deve possuir número e descrição.")
        key = (item["lote"], item["numero"])
        if key in seen:
            raise ValueError(f"O item {item['numero']} foi selecionado mais de uma vez.")
        seen.add(key)
        items.append(item)
    return items


def replace_business_items(connection, business_id, items, now):
    if items is None:
        return
    connection.execute("DELETE FROM negocio_itens WHERE negocio_id = ?", (business_id,))
    connection.executemany(
        """
        INSERT INTO negocio_itens (
            negocio_id, ordem, lote, numero_item, descricao, quantidade,
            unidade, valor_unitario_estimado, valor_total_estimado,
            criterio_julgamento, situacao, criado_em, atualizado_em
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        [
            (
                business_id, item["ordem"], item["lote"], item["numero"],
                item["descricao"], item["quantidade"], item["unidade"],
                item["valor_unitario_estimado"], item["valor_total_estimado"],
                item["criterio_julgamento"], item["situacao"], now, now,
            )
            for item in items
        ],
    )


def import_business(payload):
    if not isinstance(payload, dict):
        raise ValueError("Dados do negócio inválidos.")
    link = compact(payload.get("pncp_link"))
    cnpj, ano, sequencial = parse_pncp_link(link)
    empresa = compact(payload.get("empresa")) or default_business_company()
    selected_items = validate_business_items(payload)
    if len(empresa) > 200:
        raise ValueError("O nome da empresa excede 200 caracteres.")
    fallback = payload.get("oportunidade")
    fallback = fallback if isinstance(fallback, dict) else {}
    metadata = {
        "numero_compra": compact(fallback.get("numero_compra")),
        "processo": compact(fallback.get("processo")),
        "modalidade": compact(fallback.get("modalidade")),
        "objeto": compact(fallback.get("objeto")),
        "orgao": compact(fallback.get("orgao")),
        "orgao_cnpj": cnpj,
        "unidade": compact(fallback.get("unidade")),
        "municipio": compact(fallback.get("municipio")),
        "uf": compact(fallback.get("uf")),
        "numero_controle_pncp": compact(fallback.get("numero_controle_pncp")),
        "abertura": compact(fallback.get("abertura")),
        "encerramento": compact(fallback.get("encerramento")),
        "situacao": compact(fallback.get("situacao")),
    }
    remote_metadata = (
        {}
        if metadata["objeto"] and metadata["orgao"]
        else pncp_purchase_metadata(cnpj, ano, sequencial)
    )
    metadata.update({
        key: value
        for key, value in remote_metadata.items()
        if value not in (None, "")
    })
    if not metadata["objeto"] and not metadata["orgao"]:
        raise RuntimeError("A contratação não foi localizada na API oficial do PNCP.")
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    official_link = pncp_app_link(cnpj, ano, sequencial)
    with DATABASE_LOCK, database_connection() as connection:
        contract = connection.execute(
            """
            SELECT id FROM contratacoes
            WHERE cnpj = ? AND ano = ? AND sequencial = ?
            """,
            (cnpj, ano, sequencial),
        ).fetchone()
        existing = connection.execute(
            """
            SELECT id FROM negocios
            WHERE empresa = ? AND cnpj_orgao = ? AND ano = ? AND sequencial = ?
            """,
            (empresa, cnpj, ano, sequencial),
        ).fetchone()
        if existing:
            business_id = existing["id"]
            connection.execute(
                """
                UPDATE negocios SET
                    contratacao_id = COALESCE(?, contratacao_id),
                    link_pncp = ?, titulo = ?, orgao = ?, municipio = ?, uf = ?,
                    modalidade = ?, numero_compra = ?, processo = ?,
                    abertura = ?, encerramento = ?, situacao = ?,
                    removido = 0, atualizado_em = ?
                WHERE id = ?
                """,
                (
                    contract["id"] if contract else None,
                    official_link,
                    metadata["objeto"],
                    metadata["orgao"],
                    metadata["municipio"],
                    metadata["uf"],
                    metadata["modalidade"],
                    metadata["numero_compra"],
                    metadata["processo"],
                    metadata["abertura"],
                    metadata["encerramento"],
                    metadata["situacao"],
                    now,
                    business_id,
                ),
            )
            created = False
        else:
            cursor = connection.execute(
                """
                INSERT INTO negocios (
                    contratacao_id, empresa, cnpj_orgao, ano, sequencial,
                    link_pncp, titulo, orgao, municipio, uf, modalidade,
                    numero_compra, processo, abertura, encerramento, situacao,
                    criado_em, atualizado_em
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    contract["id"] if contract else None,
                    empresa,
                    cnpj,
                    ano,
                    sequencial,
                    official_link,
                    metadata["objeto"],
                    metadata["orgao"],
                    metadata["municipio"],
                    metadata["uf"],
                    metadata["modalidade"],
                    metadata["numero_compra"],
                    metadata["processo"],
                    metadata["abertura"],
                    metadata["encerramento"],
                    metadata["situacao"],
                    now,
                    now,
                ),
            )
            business_id = cursor.lastrowid
            connection.execute(
                """
                INSERT INTO negocio_historico (
                    negocio_id, evento, etapa_nova, criado_em
                ) VALUES (?, ?, 'oportunidade', ?)
                """,
                (business_id, "Negócio adicionado a partir do PNCP", now),
            )
            created = True
        replace_business_items(connection, business_id, selected_items, now)
        if selected_items is not None:
            connection.execute(
                """
                INSERT INTO negocio_historico (
                    negocio_id, evento, criado_em
                ) VALUES (?, ?, ?)
                """,
                (
                    business_id,
                    f"{len(selected_items)} item(ns) selecionado(s) para o negócio",
                    now,
                ),
            )
        ensure_business_tasks(connection, business_id, now)
    return {"negocio": get_business(business_id), "criado": created}


def classification_portal_name(detected_portal, source_link=""):
    """Map source-system names to the columns used by the classification sub-screen."""
    detected = compact(detected_portal)
    aliases = {
        "BLL Compras": "BLL",
        "BNC Compras": "BNC",
        "Portal de Compras Publicas": "Portal de Compras Públicas",
        "Licitações-e": "Licitações-e",
    }
    if detected and detected != "PNCP":
        return aliases.get(detected, detected)
    hostname = (urlparse(source_link or "").hostname or "").removeprefix("www.")
    return hostname or "PNCP"


def classify_business_into_source_portal(business_id):
    """Verify edital/API data and upsert the business in its source portal column."""
    business = get_business(business_id, include_details=True)
    fallback = {
        "objeto": business["titulo_oficial"],
        "orgao": business["orgao"],
        "municipio": business["municipio"],
        "uf": business["uf"],
        "modalidade": business["modalidade"],
        "numero_compra": business["numero_compra"],
        "processo": business["processo"],
        "abertura": business["abertura"],
        "encerramento": business["encerramento"],
        "situacao": business["situacao"],
    }
    detail = opportunity_detail_from_pncp_link(business["link_pncp"], fallback)
    opportunity = detail["oportunidade"]
    source_link = opportunity.get("link_origem") or business["link_pncp"]
    portal = classification_portal_name(opportunity.get("portal_origem"), source_link)
    column = kanban_store.ensure_column(DATABASE_PATH, portal)
    verification = detail.get("verificacao_itens") or {}
    verification_note = (
        f"Conferência edital/API: {verification.get('file_count', 0)} item(ns) no documento e "
        f"{verification.get('pncp_count', 0)} na API."
    )
    if verification.get("has_divergence"):
        verification_note += " Foram identificadas divergências entre as fontes."
    if verification.get("file_error"):
        verification_note += f" Documento indisponível: {verification['file_error']}"
    proposal = kanban_store.upsert_business_proposal(DATABASE_PATH, business_id, {
        "column_id": str(column["id"]),
        "portal": portal,
        "position_number": str(business.get("position_number") or ""),
        "modality": business["modalidade"],
        "agency_name": business["orgao"],
        "notice_number": business["numero_compra"],
        "uasg": opportunity.get("codigo_unidade", ""),
        "pncp_control_number": opportunity.get("numero_controle_pncp", ""),
        "opening_at": business["abertura"],
        "critical_deadline": business["encerramento"],
        "internal_identifier": f"negocio-{business_id}",
        "title": business["titulo"],
        "object_description": business["titulo_oficial"],
        "phase_status": business["situacao"],
        "priority": {1: "alta", 2: "normal", 3: "baixa"}.get(business["prioridade"], "normal"),
        "pending_documents": "",
        "estimated_value": str(opportunity.get("valor_total_estimado") or ""),
        "responsible": business["responsavel"],
        "next_review_at": business["prazo_interno"],
        "notes": verification_note,
        "source_link": source_link,
    })
    return {"portal": portal, "proposal": proposal, "verification": verification}


def sync_business_position_from_proposal(proposal):
    business_id = proposal.get("business_id") if proposal else None
    if not business_id:
        return
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    with DATABASE_LOCK, database_connection() as connection:
        connection.execute(
            "UPDATE negocios SET position_number = ?, atualizado_em = ? WHERE id = ?",
            (proposal.get("position_number"), now, business_id),
        )


def clear_business_position_from_deleted_proposal(proposal):
    business_id = proposal.get("business_id") if proposal else None
    if not business_id:
        return
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    with DATABASE_LOCK, database_connection() as connection:
        connection.execute(
            "UPDATE negocios SET position_number = NULL, atualizado_em = ? WHERE id = ?",
            (now, business_id),
        )


def validate_business_update(payload):
    if not isinstance(payload, dict):
        raise ValueError("Dados do negócio inválidos.")
    allowed = {
        "titulo_interno",
        "etapa",
        "prioridade",
        "favorito",
        "position_number",
        "responsavel",
        "prazo_interno",
        "anotacoes",
        "decisao_comercial",
        "arquivado",
        "removido",
    }
    return {key: payload[key] for key in allowed if key in payload}


def update_business(business_id, payload):
    business_id = business_id_value(business_id)
    changes = validate_business_update(payload)
    if not changes:
        raise ValueError("Nenhuma alteração foi informada.")
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    justification = compact(payload.get("justificativa"))
    classification = None
    requested_stage = compact(changes.get("etapa")).lower() if "etapa" in changes else ""
    if requested_stage == "classificacao":
        with DATABASE_LOCK, database_connection() as connection:
            stage_row = connection.execute(
                "SELECT etapa FROM negocios WHERE id = ?", (business_id,)
            ).fetchone()
        if not stage_row:
            raise FileNotFoundError("Negócio não encontrado.")
        if stage_row["etapa"] != "classificacao":
            classification = classify_business_into_source_portal(business_id)

    with DATABASE_LOCK, database_connection() as connection:
        current = connection.execute(
            "SELECT * FROM negocios WHERE id = ?", (business_id,)
        ).fetchone()
        if not current:
            raise FileNotFoundError("Negócio não encontrado.")
        if "etapa" in changes:
            stage = compact(changes["etapa"]).lower()
            if stage not in BUSINESS_STAGES:
                raise ValueError("Etapa do negócio inválida.")
            previous = current["etapa"]
            if stage != previous:
                backward = BUSINESS_STAGES.index(stage) < BUSINESS_STAGES.index(previous)
                if (stage == "contrato" or backward) and not justification:
                    raise ValueError("Informe uma justificativa para esta movimentação.")
                connection.execute(
                    """
                    INSERT INTO negocio_historico (
                        negocio_id, evento, etapa_anterior, etapa_nova,
                        justificativa, criado_em
                    ) VALUES (?, ?, ?, ?, ?, ?)
                    """,
                    (business_id, "Etapa alterada", previous, stage, justification, now),
                )
            changes["etapa"] = stage
        if "prioridade" in changes:
            priority = int(changes["prioridade"])
            if priority not in {1, 2, 3}:
                raise ValueError("Prioridade inválida.")
            changes["prioridade"] = priority
        if "position_number" in changes:
            raw_position = changes["position_number"]
            if raw_position in (None, ""):
                changes["position_number"] = None
            else:
                try:
                    position = int(raw_position)
                except (TypeError, ValueError):
                    raise ValueError("A posição deve ser um número inteiro maior que zero.")
                if position < 1:
                    raise ValueError("A posição deve ser um número inteiro maior que zero.")
                changes["position_number"] = position
        for boolean_field in ("favorito", "arquivado", "removido"):
            if boolean_field in changes:
                changes[boolean_field] = 1 if bool(changes[boolean_field]) else 0
        for text_field in (
            "titulo_interno", "responsavel", "prazo_interno", "anotacoes", "decisao_comercial",
        ):
            if text_field in changes:
                changes[text_field] = str(changes[text_field] or "").strip()
        assignments = ", ".join(f"{field} = ?" for field in changes)
        connection.execute(
            f"UPDATE negocios SET {assignments}, atualizado_em = ? WHERE id = ?",
            (*changes.values(), now, business_id),
        )
        if "position_number" in changes:
            connection.execute(
                "UPDATE proposals SET position_number = ?, updated_at = ? WHERE business_id = ?",
                (changes["position_number"], now, business_id),
            )
        if "arquivado" in changes or "removido" in changes:
            event = "Negócio arquivado" if changes.get("arquivado") else "Negócio removido"
            connection.execute(
                """
                INSERT INTO negocio_historico (
                    negocio_id, evento, justificativa, criado_em
                ) VALUES (?, ?, ?, ?)
                """,
                (business_id, event, justification, now),
            )
        if classification:
            connection.execute(
                "UPDATE negocios SET plataforma = ?, fonte_integracao = ? WHERE id = ?",
                (classification["portal"], "Edital oficial + API PNCP", business_id),
            )
            check = classification["verification"]
            detail_text = (
                f"Portal: {classification['portal']}; documento: {check.get('file_count', 0)} item(ns); "
                f"API PNCP: {check.get('pncp_count', 0)} item(ns); "
                f"divergência: {'sim' if check.get('has_divergence') else 'não'}"
            )
            connection.execute(
                "INSERT INTO negocio_historico (negocio_id, evento, justificativa, criado_em) "
                "VALUES (?, ?, ?, ?)",
                (business_id, "Classificado automaticamente no portal de origem", detail_text, now),
            )
    return get_business(business_id)


def create_business_task(business_id, payload):
    business_id = business_id_value(business_id)
    title = compact((payload or {}).get("titulo"))
    if not title:
        raise ValueError("Informe o título da tarefa.")
    if len(title) > 300:
        raise ValueError("O título da tarefa excede 300 caracteres.")
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    with DATABASE_LOCK, database_connection() as connection:
        if not connection.execute(
            "SELECT 1 FROM negocios WHERE id = ?", (business_id,)
        ).fetchone():
            raise FileNotFoundError("Negócio não encontrado.")
        next_order = connection.execute(
            """
            SELECT COALESCE(MAX(ordem), 0) + 1
            FROM negocio_tarefas WHERE negocio_id = ?
            """,
            (business_id,),
        ).fetchone()[0]
        cursor = connection.execute(
            """
            INSERT INTO negocio_tarefas (
                negocio_id, titulo, concluida, ordem, criado_em, atualizado_em
            ) VALUES (?, ?, 0, ?, ?, ?)
            """,
            (business_id, title, next_order, now, now),
        )
        task_id = cursor.lastrowid
        connection.execute(
            "UPDATE negocios SET atualizado_em = ? WHERE id = ?",
            (now, business_id),
        )
    detail = get_business(business_id, include_details=True)
    task = next(item for item in detail["tarefas"] if int(item["id"]) == task_id)
    return {"tarefa": task, "negocio": business_record_from_detail(detail)}


def business_record_from_detail(detail):
    return {
        key: value
        for key, value in detail.items()
        if key not in {"historico", "tarefas", "arquivos"}
    }


def update_business_task(business_id, task_id, payload):
    business_id = business_id_value(business_id)
    task_id = business_id_value(task_id, "Tarefa")
    if not isinstance(payload, dict):
        raise ValueError("Dados da tarefa inválidos.")
    changes = {}
    if "concluida" in payload:
        changes["concluida"] = 1 if bool(payload["concluida"]) else 0
    if "titulo" in payload:
        title = compact(payload["titulo"])
        if not title:
            raise ValueError("Informe o título da tarefa.")
        changes["titulo"] = title
    if not changes:
        raise ValueError("Nenhuma alteração foi informada.")
    init_database()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    with DATABASE_LOCK, database_connection() as connection:
        assignments = ", ".join(f"{field} = ?" for field in changes)
        cursor = connection.execute(
            f"""
            UPDATE negocio_tarefas
            SET {assignments}, atualizado_em = ?
            WHERE id = ? AND negocio_id = ?
            """,
            (*changes.values(), now, task_id, business_id),
        )
        if not cursor.rowcount:
            raise FileNotFoundError("Tarefa não encontrada.")
        connection.execute(
            "UPDATE negocios SET atualizado_em = ? WHERE id = ?",
            (now, business_id),
        )
    return get_business(business_id, include_details=True)


def catalog_document_text(path):
    cache_key = f"{path.resolve()}:{path.stat().st_mtime_ns}"
    cached = cache_get(CATALOG_TEXT_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None:
        return cached
    if path.suffix.lower() == ".pdf":
        with pdfplumber.open(str(path)) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
    elif path.suffix.lower() == ".docx":
        document = Document(path)
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend(
                " | ".join(cell.text for cell in row.cells)
                for row in table.rows
            )
        text = "\n".join(blocks)
    else:
        text = ""
    cache_set(CATALOG_TEXT_CACHE, cache_key, text)
    return text


COMMERCIAL_TERM_NOT_FOUND = "Não localizado no documento oficial"
COMMERCIAL_TERM_FIELDS = (
    "prazo_entrega",
    "prazo_pagamento",
    "validade_proposta",
)
COMMERCIAL_DURATION_PATTERN = (
    r"(?:at[eé]\s+)?\d{1,4}(?:º|°)?"
    r"(?:\s*\([^) \n][^)\n]{1,45}\))?\s*"
    r"(?:dias?|horas?|mes(?:es)?)"
    r"(?:\s+(?:[uú]teis|corridos|consecutivos|calend[aá]rio))?"
)


def commercial_terms_from_text(text):
    searchable = re.sub(r"(?<=\w)-\s+(?=\w)", "", text or "")
    searchable = re.sub(r"\s+", " ", searchable).strip()
    patterns = {
        "prazo_entrega": [
            (
                rf"prazo\s+(?:m[aá]ximo\s+)?de\s+entrega"
                rf".{{0,180}}?(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
            (
                rf"(?:entrega|fornecimento)\s+"
                rf"(?:dever[aá]\s+(?:ser\s+)?(?:realizad[ao]|efetuad[ao]|ocorrer)"
                rf"|ser[aá]\s+(?:realizad[ao]|efetuad[ao])).{{0,120}}?"
                rf"(?:no|em)\s+prazo\s+(?:m[aá]ximo\s+)?de\s+"
                rf"(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
        ],
        "prazo_pagamento": [
            (
                rf"prazo\s+(?:m[aá]ximo\s+)?de\s+pagamento"
                rf".{{0,180}}?(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
            (
                rf"pagamento\s+(?:dever[aá]\s+ser|ser[aá])\s+"
                rf"(?:efetuado|realizado|processado).{{0,120}}?"
                rf"(?:no|em)\s+prazo\s+(?:m[aá]ximo\s+)?de\s+"
                rf"(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
        ],
        "validade_proposta": [
            (
                rf"validade\s+(?:m[ií]nima\s+)?da\s+proposta"
                rf".{{0,180}}?(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
            (
                rf"proposta.{{0,100}}?validade"
                rf".{{0,100}}?(?P<value>{COMMERCIAL_DURATION_PATTERN})"
            ),
        ],
    }
    values = {}
    for field, field_patterns in patterns.items():
        for pattern in field_patterns:
            match = re.search(pattern, searchable, flags=re.IGNORECASE)
            if match:
                values[field] = compact(match.group("value")).strip(" .;,:")
                break
    return values


def extract_commercial_terms(source_data):
    source_path = Path(source_data["source_path"])
    candidates = [source_path]
    candidates.extend(
        Path(candidate["path"])
        for candidate in source_data.get("documentos_candidatos") or []
        if candidate.get("score", (2,))[0] <= 1
    )
    values = {}
    sources = {}
    visited = set()
    for path in candidates:
        resolved = str(path.resolve())
        if resolved in visited or not path.is_file():
            continue
        visited.add(resolved)
        try:
            found = commercial_terms_from_text(catalog_document_text(path))
        except Exception:
            continue
        for field in COMMERCIAL_TERM_FIELDS:
            if field not in values and found.get(field):
                values[field] = found[field]
                sources[field] = path.name
        if len(values) == len(COMMERCIAL_TERM_FIELDS):
            break

    missing = [field for field in COMMERCIAL_TERM_FIELDS if not values.get(field)]
    for field in missing:
        values[field] = COMMERCIAL_TERM_NOT_FOUND
    values["fontes"] = sources
    values["status"] = "ok" if not missing else "warn"
    values["campos_nao_localizados"] = missing
    return values


def normalized_commercial_terms(terms):
    terms = terms if isinstance(terms, dict) else {}
    return {
        field: compact(terms.get(field))[:200] or COMMERCIAL_TERM_NOT_FOUND
        for field in COMMERCIAL_TERM_FIELDS
    }


def commercial_terms_without_document():
    missing = list(COMMERCIAL_TERM_FIELDS)
    return {
        **{field: COMMERCIAL_TERM_NOT_FOUND for field in missing},
        "fontes": {},
        "status": "warn",
        "campos_nao_localizados": missing,
    }


def normalized_wanted_item_key(value):
    parts = [normalize_item_reference(part) for part in compact(value).split("/", 1)]
    if not parts or not parts[0]:
        return ""
    return "/".join(parts)


def proposal_process_from_structured_items(pncp_link, wanted_items):
    cnpj, ano, sequencial = parse_pncp_link(pncp_link)
    identification = identify_items_from_pncp_link(pncp_link)
    available_items = identification.get("items") or []
    wanted_keys = {
        normalized_wanted_item_key(value)
        for value in compact(wanted_items).split(",")
        if normalized_wanted_item_key(value)
    }
    if not wanted_keys:
        raise ValueError("Selecione pelo menos um item para a proposta.")

    selected_items = [
        item for item in available_items
        if item_lookup_key(item) in wanted_keys
    ]
    selected_keys = {item_lookup_key(item) for item in selected_items}
    missing_keys = sorted(wanted_keys - selected_keys, key=identifier_sort_key)
    if missing_keys:
        raise ValueError(
            "Os itens selecionados nao foram encontrados na base estruturada: "
            + ", ".join(missing_keys)
            + ". Atualize a identificacao dos itens e tente novamente."
        )

    description_review = build_description_review(selected_items)
    for review in description_review.get("items", []):
        review["source"] = "base estruturada local"
    description_review["message"] = (
        f"{len(selected_items)} item(ns) selecionado(s) e validado(s) na base estruturada local."
    )
    if description_review["status"] == "error":
        raise ValueError(description_review["message"])

    pncp_items_check = dict(identification.get("pncp_items_check") or {})
    pncp_items_check["selected_count"] = len(selected_items)
    pncp_items_check["source"] = identification.get("source") or "opportunity_items"
    pncp_info = identification.get("pncp") or {
        "cnpj": cnpj,
        "ano": ano,
        "sequencial": sequencial,
        "link": pncp_app_link(cnpj, ano, sequencial),
        "documento_usado": "",
        "documento_tipo": "Base estruturada",
    }
    return {
        "count": len(selected_items),
        "items": selected_items,
        "source_name": f"PNCP_{cnpj}_{ano}_{sequencial}",
        "pncp": public_pncp_payload(pncp_info),
        "pncp_items_check": pncp_items_check,
        "description_review": description_review,
        "commercial_terms": commercial_terms_without_document(),
    }


def catalog_specification_from_document(path, item):
    title = compact(item.get("descricao"))
    item_number = normalize_item_reference(item.get("item"))
    if not title:
        return ""
    try:
        text = catalog_document_text(path)
    except Exception:
        return ""
    if not text:
        return ""
    flexible_title = r"\s+".join(re.escape(part) for part in title.split())
    title_with_qualifier = (
        rf"{flexible_title}(?:\s*[-–]\s*[^:\n]{{1,100}})?"
    )
    patterns = [
        rf"(?ims)^\s*[a-z]\)\s*{title_with_qualifier}\s*:\s*(.+?)(?=^\s*[a-z]\)\s+|\Z)",
        (
            rf"(?ims)^\s*(?:ITEM\s*)?0*{re.escape(item_number)}\s*[-:]\s*"
            rf"{title_with_qualifier}\s*[:.-]\s*(.+?)"
            rf"(?=^\s*(?:ITEM\s*)?0*\d+\s*[-:]|\Z)"
        ),
    ]
    matches = []
    for pattern in patterns:
        matches.extend(re.findall(pattern, text))
    if not matches:
        return ""
    detail = max(matches, key=lambda value: len(compact(value)))
    detail = re.sub(r"(?i)P[áa]gina\s+\d+\s+de\s+\d+", " ", detail)
    detail = compact(detail)
    if len(detail) < 40:
        return ""
    return compact(f"{title}: {detail}")


def richest_catalog_item(source_data, selected_key, fallback):
    best_item = dict(fallback)
    best_document = compact(
        (source_data.get("pncp") or {}).get("documento_usado")
    )
    best_score = len(compact(best_item.get("descricao")))
    candidates = (
        source_data.get("documentos_candidatos")
        or (source_data.get("pncp") or {}).get("documentos_candidatos")
        or []
    )
    reviewed = 0
    for candidate in candidates:
        score = candidate.get("score") or (2,)
        if score[0] > 1 or reviewed >= 8:
            continue
        path = Path(candidate.get("path") or "")
        if not path.is_file():
            continue
        reviewed += 1
        try:
            rows = extract_items_cached(path)
        except Exception:
            continue
        matching = next(
            (
                row
                for row in rows
                if item_lookup_key(row) == selected_key
                or normalize_item_reference(row.get("item")) == selected_key
            ),
            None,
        )
        if not matching:
            continue
        detailed_description = catalog_specification_from_document(path, matching)
        candidate_item = dict(matching)
        if detailed_description:
            candidate_item["descricao"] = detailed_description
        description_score = len(compact(candidate_item.get("descricao")))
        if description_score > best_score:
            best_item = candidate_item
            best_document = path.name
            best_score = description_score
        if detailed_description and score[0] == 0:
            break
    return best_item, best_document


def catalog_draft_from_pncp_link(link, selected_key):
    cnpj, ano, sequencial = parse_pncp_link(link)
    selected_key = normalize_item_reference(selected_key)
    cache_key = f"{pncp_app_link(cnpj, ano, sequencial)}#{selected_key}"
    cached = cache_get(CATALOG_DRAFT_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None:
        return cached
    identification = identify_items_from_pncp_link(link)
    items = identification.get("items") or []
    selected = next(
        (
            item
            for item in items
            if item_lookup_key(item) == selected_key
            or normalize_item_reference(item.get("item")) == selected_key
        ),
        None,
    )
    if not selected:
        raise ValueError("Selecione um item válido do edital.")
    source_data = source_from_pncp_link(link)
    selected, selected_document = richest_catalog_item(
        source_data,
        selected_key,
        selected,
    )
    pncp = dict(identification.get("pncp") or {})
    if selected_document:
        pncp["documento_usado"] = selected_document
    pncp["metadata"] = pncp_purchase_metadata(
        pncp.get("cnpj"),
        pncp.get("ano"),
        pncp.get("sequencial"),
    )
    result = {
        "draft": catalog_draft_from_item(selected, pncp),
        "items": items,
        "pncp": pncp,
    }
    cache_set(CATALOG_DRAFT_CACHE, cache_key, result)
    return result


def catalog_asset_field(form, key):
    if key not in form:
        return None
    field = form[key]
    if isinstance(field, list):
        return field[0] if field else None
    return field


def save_catalog_assets(form, payload):
    specs = payload.get("assets") or []
    if not isinstance(specs, list):
        raise ValueError("A lista de imagens do catálogo é inválida.")
    saved = []
    allowed_suffixes = {".png", ".jpg", ".jpeg", ".webp"}
    try:
        for index, spec in enumerate(specs):
            if not isinstance(spec, dict):
                continue
            upload_key = compact(spec.get("upload_key"))
            field = catalog_asset_field(form, upload_key)
            if field is None or not getattr(field, "filename", ""):
                continue
            filename = Path(field.filename).name
            suffix = Path(filename).suffix.lower()
            if suffix not in allowed_suffixes:
                raise ValueError(
                    f"A imagem {filename} deve estar em PNG, JPG, JPEG ou WEBP."
                )
            content = field.file.read(MAX_CATALOG_ASSET_SIZE + 1)
            if not content:
                raise ValueError(f"A imagem {filename} está vazia.")
            if len(content) > MAX_CATALOG_ASSET_SIZE:
                raise OverflowError(
                    f"A imagem {filename} excede o limite de 12 MB."
                )
            target = UPLOAD_DIR / (
                f"catalogo_{safe_name(Path(filename).stem)}_"
                f"{uuid.uuid4().hex}{suffix}"
            )
            target.write_bytes(content)
            try:
                from PIL import Image

                with Image.open(target) as image:
                    image.verify()
            except Exception as exc:
                target.unlink(missing_ok=True)
                raise ValueError(f"O arquivo {filename} não é uma imagem válida.") from exc
            saved.append(
                {
                    "path": target,
                    "name": filename,
                    "role": compact(spec.get("role")) or "secundaria",
                    "section": compact(spec.get("section")),
                    "caption": compact(spec.get("caption")),
                    "order": index,
                }
            )
    except Exception:
        for asset in saved:
            Path(asset["path"]).unlink(missing_ok=True)
        raise
    return saved


def generate_catalog_exports(data, assets):
    if not isinstance(data, dict):
        raise ValueError("Os dados estruturados do catálogo são inválidos.")
    alerts = catalog_alerts(data, assets)
    if alerts["errors"]:
        error = ValueError("Revise os campos obrigatórios antes de exportar o catálogo.")
        error.catalog_alerts = alerts
        raise error

    item_number = safe_name((data.get("item") or {}).get("numero") or "item")
    stamp = f"{int(time.time())}_{uuid.uuid4().hex[:8]}"
    base_name = f"Catalogo_Tecnico_Item_{item_number}_{stamp}"
    paths = {
        "docx": OUTPUT_DIR / f"{base_name}.docx",
        "pdf": OUTPUT_DIR / f"{base_name}.pdf",
        "json": OUTPUT_DIR / f"{base_name}.json",
        "images": OUTPUT_DIR / f"{base_name}_imagens.zip",
    }
    build_catalog_docx(data, assets, paths["docx"])
    build_catalog_pdf(data, assets, paths["pdf"])
    write_catalog_json(data, assets, paths["json"], alerts)
    build_images_zip(assets, paths["images"])
    return {
        "alerts": alerts,
        "exports": {
            key: {
                "filename": path.name,
                "download_url": f"/download/{quote(path.name)}",
            }
            for key, path in paths.items()
        },
    }


def has_item_number_gaps(items):
    numbers = sorted({
        int(compact(row.get("item")))
        for row in items
        if re.fullmatch(r"\d+", compact(row.get("item")))
    })
    if len(numbers) < 2:
        return False
    return numbers[-1] - numbers[0] + 1 != len(numbers)


def grouped_line_peaks(scores, minimum_score, maximum_gap=6):
    groups = []
    for position, score in enumerate(scores):
        if score < minimum_score:
            continue
        if not groups or position - groups[-1][-1][0] > maximum_gap:
            groups.append([])
        groups[-1].append((position, score))
    return [max(group, key=lambda pair: pair[1])[0] for group in groups]


def scanned_table_grid(image):
    try:
        import cv2
        import numpy as np
    except ImportError as exc:
        raise RuntimeError(
            "O suporte de OCR para PDFs escaneados não está instalado."
        ) from exc

    gray = cv2.cvtColor(np.asarray(image), cv2.COLOR_RGB2GRAY)
    binary = cv2.threshold(gray, 180, 255, cv2.THRESH_BINARY_INV)[1]
    height, width = gray.shape

    vertical = cv2.morphologyEx(
        binary,
        cv2.MORPH_OPEN,
        cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(50, height // 25))),
    )
    vertical_scores = (vertical > 0).sum(axis=0)
    vertical_lines = grouped_line_peaks(
        vertical_scores,
        max(80, int(height * 0.06)),
    )
    vertical_lines = [
        x for x in vertical_lines if width * 0.03 <= x <= width * 0.97
    ]
    if len(vertical_lines) < 6:
        return None
    if len(vertical_lines) > 10:
        strongest = sorted(
            vertical_lines,
            key=lambda x: vertical_scores[x],
            reverse=True,
        )[:10]
        vertical_lines = sorted(strongest)

    left, right = vertical_lines[0], vertical_lines[-1]
    horizontal = cv2.morphologyEx(
        binary,
        cv2.MORPH_OPEN,
        cv2.getStructuringElement(cv2.MORPH_RECT, (max(100, width // 10), 1)),
    )
    horizontal_scores = (horizontal[:, left : right + 1] > 0).sum(axis=1)
    horizontal_lines = grouped_line_peaks(
        horizontal_scores,
        max(90, int((right - left) * 0.12)),
        maximum_gap=10,
    )
    horizontal_lines = [
        y for y in horizontal_lines if height * 0.01 < y < height * 0.985
    ]
    if len(horizontal_lines) < 2:
        return None
    return vertical_lines, horizontal_lines


def get_ocr_engine():
    global OCR_ENGINE
    if OCR_ENGINE is None:
        try:
            from rapidocr import RapidOCR
        except ImportError as exc:
            raise RuntimeError(
                "O suporte de OCR para PDFs escaneados não está instalado."
            ) from exc
        OCR_ENGINE = RapidOCR()
    return OCR_ENGINE


def ocr_page_tokens(image):
    import numpy as np

    with OCR_LOCK:
        result = get_ocr_engine()(
            np.asarray(image),
            use_cls=False,
            text_score=0.5,
        )
    if not result or result.boxes is None or not result.txts:
        return []

    tokens = []
    for box, text, score in zip(result.boxes, result.txts, result.scores):
        if score < 0.5 or not compact(text):
            continue
        xs = [point[0] for point in box]
        ys = [point[1] for point in box]
        tokens.append({
            "text": compact(text),
            "left": min(xs),
            "top": min(ys),
            "center_x": sum(xs) / len(xs),
            "center_y": sum(ys) / len(ys),
        })
    return sorted(tokens, key=lambda token: (token["top"], token["left"]))


def tokens_in_scanned_cell(tokens, left, right, top, bottom):
    return [
        token["text"]
        for token in tokens
        if left < token["center_x"] < right
        and top < token["center_y"] < bottom
    ]


def scanned_row_item_number(values):
    text = " ".join(values)
    match = re.search(r"\b(\d{1,4})\b", text)
    return match.group(1) if match else ""


def scanned_row_quantity(values):
    text = " ".join(values)
    match = re.search(r"\b\d+(?:[.,]\d+)?\b", text)
    return match.group(0) if match else ""


def append_scanned_description(item, lines):
    addition = join_pdf_description_lines(
        line for line in lines if not is_pdf_text_noise(line)
    )
    if not addition:
        return
    item["descricao"] = compact(f"{item.get('descricao', '')} {addition}")


def is_scanned_header_row(name_values, description_values):
    name = norm(" ".join(name_values))
    description = norm(" ".join(description_values))
    return name in {"nome", "item"} or description.startswith("descricao")


def repair_scanned_item_numbers(items):
    index = 0
    while index < len(items):
        if compact(items[index].get("item")).isdigit():
            index += 1
            continue
        run_start = index
        while index < len(items) and not compact(items[index].get("item")).isdigit():
            index += 1
        run_end = index
        previous_number = None
        next_number = None
        if run_start > 0:
            previous = compact(items[run_start - 1].get("item"))
            previous_number = int(previous) if previous.isdigit() else None
        if run_end < len(items):
            following = compact(items[run_end].get("item"))
            next_number = int(following) if following.isdigit() else None

        missing_count = run_end - run_start
        expected_start = 1 if previous_number is None else previous_number + 1
        expected_end = expected_start + missing_count
        if next_number == expected_end:
            for offset, row_index in enumerate(range(run_start, run_end)):
                items[row_index]["item"] = str(expected_start + offset)

    known = []
    for position, row in enumerate(items, start=1):
        item = compact(row.get("item"))
        if item.isdigit():
            known.append((position, int(item)))
    if not known:
        return items
    aligned = sum(position == number for position, number in known)
    if aligned / len(known) < 0.8:
        return items
    for position, row in enumerate(items, start=1):
        row["item"] = str(position)
    return items


def extract_from_scanned_pdf(path):
    items = []
    current_item = None
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            image = page.to_image(resolution=OCR_DPI, antialias=True).original.convert("RGB")
            grid = scanned_table_grid(image)
            if not grid:
                continue
            vertical_lines, horizontal_lines = grid
            if len(vertical_lines) < 6:
                continue
            crop_left = vertical_lines[0]
            crop_top = horizontal_lines[0]
            crop_right = vertical_lines[5]
            crop_bottom = horizontal_lines[-1]
            table_image = image.crop((crop_left, crop_top, crop_right, crop_bottom))
            tokens = ocr_page_tokens(table_image)
            vertical_lines = [x - crop_left for x in vertical_lines[:6]]
            horizontal_lines = [y - crop_top for y in horizontal_lines]

            for top, bottom in zip(horizontal_lines, horizontal_lines[1:]):
                item_values = tokens_in_scanned_cell(
                    tokens, vertical_lines[0], vertical_lines[1], top, bottom
                )
                name_values = tokens_in_scanned_cell(
                    tokens, vertical_lines[1], vertical_lines[2], top, bottom
                )
                description_values = tokens_in_scanned_cell(
                    tokens, vertical_lines[2], vertical_lines[3], top, bottom
                )
                quantity_values = tokens_in_scanned_cell(
                    tokens, vertical_lines[4], vertical_lines[5], top, bottom
                )
                item_number = scanned_row_item_number(item_values)

                starts_item = bool(item_number or name_values)
                if starts_item and not is_scanned_header_row(name_values, description_values):
                    row = empty_item()
                    row["item"] = item_number
                    row["quantidade"] = scanned_row_quantity(quantity_values)
                    row["unidade"] = STANDARD_UNIT
                    row["_nome"] = join_pdf_description_lines(name_values)
                    append_scanned_description(
                        row, description_values or name_values
                    )
                    items.append(row)
                    current_item = row
                elif current_item and description_values:
                    append_scanned_description(current_item, description_values)

    return sanitize_extracted_items(repair_scanned_item_numbers(items))


def extract_from_pdf(path):
    tables = []
    page_texts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                tables.append([
                    [
                        join_pdf_description_lines(str(cell).splitlines())
                        if cell is not None
                        else None
                        for cell in row
                    ]
                    for row in table
                ])
            page_texts.append(page.extract_text() or "")
    table_items = normalize_pdf_tables(tables)
    wide_text_items = extract_from_wide_pdf_texts(page_texts)
    if wide_text_items:
        table_items = merge_item_lists(table_items, wide_text_items)
    if not table_items and sum(len(compact(text)) for text in page_texts) < 100:
        return extract_from_scanned_pdf(path)
    if table_items and not items_without_description(table_items) and not has_item_number_gaps(table_items):
        return table_items
    text_items = remove_spurious_document_items(extract_from_pdf_text(path))
    return merge_item_lists(table_items, text_items)


def items_without_description(items):
    return [compact(row.get("item")) for row in items if not compact(row.get("descricao"))]


def extract_from_docx(path):
    doc = Document(str(path))
    rows = []
    for table in doc.tables:
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
    return normalize_rows(rows)


def extract_items(path):
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_from_pdf(path)
    if suffix == ".docx":
        return extract_from_docx(path)
    raise ValueError("Formato não suportado. Envie PDF ou DOCX.")


def extract_items_cached(path):
    path = Path(path)
    stat = path.stat()
    cache_key = (str(path.resolve()), stat.st_mtime_ns, stat.st_size)
    cached = cache_get(EXTRACTED_ITEMS_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None:
        return cached
    items = extract_items(path)
    cache_set(EXTRACTED_ITEMS_CACHE, cache_key, items)
    return items


def set_font(run, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_grid_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=80, bottom=80, start=100, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_column_cell_margins(cell, key):
    if key in {"lote", "item", "quantidade", "unidade"}:
        set_cell_margins(cell, start=20, end=20)
    else:
        set_cell_margins(cell)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def write_cell(cell, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text or "")
    set_font(run)
    run.bold = bold
    run.font.size = Pt(size)


def resolve_responsible(responsible_id):
    try:
        responsible = get_responsible(responsible_id)
    except ValueError:
        return None
    if not responsible:
        return None
    lines = [PROPOSAL_SIGNATURE_LOCATION_DATE]
    if responsible["empresa"]:
        lines.append(responsible["empresa"])
    if responsible["cnpj"]:
        lines.append(f'CNPJ {responsible["cnpj"]}')
    lines.append(responsible["nome_completo"])
    if responsible["rg"]:
        lines.append(f'RG {responsible["rg"]}')
    if responsible["cpf"]:
        lines.append(f'CPF {responsible["cpf"]}')
    lines.extend(line for line in responsible["observacoes"].splitlines() if line.strip())
    return {**responsible, "label": responsible["nome_completo"], "document_lines": lines}


def add_responsible_block(doc, responsible):
    if not responsible:
        return

    doc.add_paragraph()
    doc.add_paragraph()
    for line in responsible["document_lines"]:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_font(run)
        run.font.size = Pt(10)


def add_commercial_terms(doc, items, commercial_terms=None):
    commercial_terms = normalized_commercial_terms(commercial_terms)
    lines = [
        ("VALOR TOTAL DA PROPOSTA:", calculate_proposal_total(items), True),
        ("Prazo de Entrega:", commercial_terms["prazo_entrega"], False),
        ("Prazo de pagamento:", commercial_terms["prazo_pagamento"], False),
        ("Validade da Proposta:", commercial_terms["validade_proposta"], False),
    ]
    for index, (label, value, emphasize) in enumerate(lines):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(8 if index == 0 else 0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = index < len(lines) - 1
        label_run = paragraph.add_run(f"{label} ")
        set_font(label_run)
        label_run.bold = True
        label_run.font.size = Pt(10)
        value_run = paragraph.add_run(value)
        set_font(value_run)
        value_run.bold = emphasize
        value_run.font.size = Pt(10)


def columns_for_items(items):
    if any(compact(item.get("lote")) for item in items):
        return [LOT_COLUMN] + COLUMNS
    return COLUMNS


def widths_for_columns(columns):
    if columns and columns[0][0] == "lote":
        return [750, 650, 600, 600, 3850, 1250, 1550, 1850]
    return [650, 650, 600, 4500, 1200, 1600, 1900]


def fit_widths_to_section(widths, section):
    available_width = int(
        (section.page_width - section.left_margin - section.right_margin) / 635
    )
    total_width = min(sum(widths), available_width)
    scale = total_width / sum(widths)
    fitted = [max(1, round(width * scale)) for width in widths]
    fitted[-1] += total_width - sum(fitted)
    return total_width, fitted


def build_docx(items, template_path, output_path, responsible=None, commercial_terms=None):
    if template_path and template_path.exists():
        with TEMPLATE_LOCK:
            shutil.copyfile(template_path, output_path)
        doc = Document(str(output_path))
    else:
        doc = Document()
        doc.add_paragraph("PROPOSTA FINAL")

    columns = columns_for_items(items)
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_grid_borders(table)
    table.autofit = False

    total_width, widths = fit_widths_to_section(
        widths_for_columns(columns), doc.sections[-1]
    )
    set_table_width(table, total_width)

    header_row = table.rows[0]
    repeat_header(header_row)
    for cell, (key, header), width in zip(header_row.cells, columns, widths):
        set_cell_width(cell, width)
        set_column_cell_margins(cell, key)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        write_cell(cell, header, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    for item in items:
        row = table.add_row()
        for idx, (cell, (key, _), width) in enumerate(zip(row.cells, columns, widths)):
            set_cell_width(cell, width)
            set_column_cell_margins(cell, key)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            align = WD_ALIGN_PARAGRAPH.LEFT if key == "descricao" else WD_ALIGN_PARAGRAPH.CENTER
            write_cell(cell, item.get(key, ""), bold=False, size=9, align=align)

    add_commercial_terms(doc, items, commercial_terms)
    add_responsible_block(doc, responsible)
    doc.save(str(output_path))


def safe_name(name):
    stem = Path(name or "arquivo").stem
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", stem).strip("._")
    return stem or "arquivo"


def save_upload(field, prefix):
    filename = safe_name(field.filename)
    suffix = Path(field.filename or "").suffix.lower()
    path = UPLOAD_DIR / f"{prefix}_{uuid.uuid4().hex}{suffix}"
    with path.open("wb") as handle:
        shutil.copyfileobj(field.file, handle)
    return path, filename


def request_json(url, timeout=18):
    cached = SEARCH_CACHE.get(url)
    if cached and time.time() - cached["created_at"] < SEARCH_CACHE_TTL:
        cached_data = cached["data"]
        if isinstance(cached_data, dict):
            data = dict(cached_data)
            data["cache_hit"] = True
            return data
        if isinstance(cached_data, list):
            return list(cached_data)
        return cached_data

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126 Safari/537.36",
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "pt-BR,pt;q=0.9,en-US;q=0.8,en;q=0.7",
        "Referer": "https://pncp.gov.br/app/editais?pagina=1",
        "Connection": "close",
    }
    last_error = ""
    for attempt in range(2):
        req = urllib.request.Request(url, headers=headers)
        try:
            with urllib.request.urlopen(req, timeout=timeout) as response:
                data = response.read()
                if not data:
                    return {}
                payload = json.loads(data.decode("utf-8"))
                SEARCH_CACHE[url] = {"created_at": time.time(), "data": payload}
                return payload
        except urllib.error.HTTPError as exc:
            if exc.code == 429:
                return {"data": [], "rate_limited": True, "message": "Limite temporário de consultas do PNCP atingido."}
            last_error = str(exc)
        except (TimeoutError, urllib.error.URLError, ConnectionError, http.client.RemoteDisconnected) as exc:
            last_error = str(exc)
        if attempt == 0:
            time.sleep(0.35)
    return {"data": [], "timeout": True, "message": last_error or "O PNCP não respondeu dentro do tempo limite."}


def parse_pncp_link(link):
    value = compact(link).strip("\"'")
    if value.startswith("/app/editais/"):
        value = f"https://pncp.gov.br{value}"
    elif not re.match(r"^https?://", value, flags=re.IGNORECASE):
        value = f"https://{value}"

    parsed = urlparse(value)
    hostname = (parsed.hostname or "").lower()
    if hostname.startswith("www."):
        hostname = hostname[4:]
    if parsed.scheme not in {"http", "https"} or hostname != "pncp.gov.br":
        raise ValueError("Link PNCP inválido. Use um endereço público do domínio pncp.gov.br.")
    match = re.search(
        r"^/app/editais/(\d{14})/(\d{4})/(\d+)/?$",
        parsed.path,
        flags=re.IGNORECASE,
    )
    if not match:
        raise ValueError("Link PNCP inválido. Use um link no formato https://pncp.gov.br/app/editais/CNPJ/ANO/SEQUENCIAL")
    cnpj, ano, sequencial = match.groups()
    return cnpj, int(ano), int(sequencial)


def pncp_app_link(cnpj, ano, sequencial):
    return f"{PNCP_APP_BASE}/{cnpj}/{ano}/{sequencial}"


def list_pncp_files(cnpj, ano, sequencial):
    url = f"{PNCP_API_BASE}/pncp/v1/orgaos/{cnpj}/compras/{ano}/{sequencial}/arquivos"
    files = request_json(url)
    if isinstance(files, dict):
        files = files.get("data") or files.get("items") or files.get("content") or []
    if not isinstance(files, list):
        return []
    return [item for item in files if item.get("statusAtivo", True)]


def format_pncp_quantity(value):
    if value is None:
        return ""
    text = str(value)
    if re.fullmatch(r"\d+\.0+", text):
        return text.split(".", 1)[0]
    return text.replace(".", ",")


def list_pncp_item_payload(cnpj, ano, sequencial):
    base_url = f"{PNCP_API_BASE}/pncp/v1/orgaos/{cnpj}/compras/{ano}/{sequencial}/itens"
    page_size = 100
    payload = []
    for page in range(1, 101):
        url = f"{base_url}?{urlencode({'pagina': page, 'tamanhoPagina': page_size})}"
        page_payload = request_json(url)
        if isinstance(page_payload, dict):
            page_items = page_payload.get("data") or page_payload.get("items") or page_payload.get("content") or []
        else:
            page_items = page_payload or []
        if not page_items:
            break
        payload.extend(page_items)
        if len(page_items) < page_size:
            break
        time.sleep(0.15)
    return payload


def list_pncp_items(cnpj, ano, sequencial):
    payload = list_pncp_item_payload(cnpj, ano, sequencial)

    rows = []
    for item in payload or []:
        row = empty_item()
        row["lote"] = compact(
            item.get("numeroLote")
            or item.get("lote")
            or item.get("grupo")
            or item.get("numeroGrupo")
            or item.get("numeroLoteCompra")
        )
        row["item"] = compact(item.get("numeroItem"))
        row["quantidade"] = format_pncp_quantity(item.get("quantidade"))
        row["unidade"] = STANDARD_UNIT
        row["descricao"] = compact(item.get("descricao"))
        if row["item"] and row["descricao"]:
            rows.append(row)
    return rows


OPPORTUNITY_CATEGORY_RULES = (
    ("Mobiliário", ("cadeira", "mesa", "armario", "sofa", "estante", "mobiliario", "poltrona", "banco")),
    ("Equipamentos de escritório", ("impressora", "fragmentadora", "projetor", "scanner", "copiadora", "escritorio")),
    ("Acessórios ergonômicos", ("ergonom", "apoio de punho", "mouse pad", "suporte para monitor", "apoio para pes")),
    ("Informática", ("computador", "notebook", "monitor", "teclado", "mouse", "servidor", "software", "informatica")),
    ("Artigos para ginástica", ("ginastica", "academia", "halter", "colchonete", "esteira ergometrica")),
    ("Saúde", ("hospital", "medicamento", "cirurg", "odontolog", "saude", "enfermagem")),
    ("Limpeza", ("limpeza", "higiene", "detergente", "desinfetante", "saneante")),
    ("EPI e segurança", ("capacete", "luva", "equipamento de protecao", "epi", "seguranca")),
    ("Climatização", ("ar condicionado", "climatizador", "ventilador", "refrigeracao")),
    ("Alimentos", ("alimento", "genero alimenticio", "refeicao", "cafe", "leite")),
)


def opportunity_categories(object_text, items):
    searchable = norm(" ".join(
        [object_text] + [compact(item.get("descricao")) for item in items]
    ))
    categories = [
        label
        for label, keywords in OPPORTUNITY_CATEGORY_RULES
        if any(norm(keyword) in searchable for keyword in keywords)
    ]
    return categories[:4] or ["Outros"]


def opportunity_source_portal(link):
    hostname = (urlparse(link or "").hostname or "").lower()
    if "compras.gov.br" in hostname or "comprasnet" in hostname:
        return "Comprasnet"
    if "licitanet" in hostname:
        return "Licitanet"
    if "bll" in hostname:
        return "BLL Compras"
    if "bnc" in hostname:
        return "BNC Compras"
    if "comprasbr" in hostname:
        return "ComprasBR"
    if "licitapp" in hostname or "licita-pp" in hostname:
        return "Licita PP"
    if "licitardigital" in hostname:
        return "Licitar Digital"
    if "novobbmnet" in hostname or "bbmnet" in hostname:
        return "NovoBBMNet"
    if "portaldecompraspublicas" in hostname:
        return "Portal de Compras Públicas"
    if "compras.rs.gov.br" in hostname:
        return "Portal de Compras RS"
    if "sislog" in hostname:
        return "SISLOG"
    if "licitacoes-e" in hostname:
        return "Licitações-e"
    if hostname:
        return hostname.removeprefix("www.")
    return "PNCP"


def opportunity_file_record(item):
    return {
        "titulo": compact(
            item.get("titulo")
            or item.get("tipoDocumentoNome")
            or "Arquivo oficial"
        ),
        "tipo": compact(
            item.get("tipoDocumentoNome")
            or item.get("tipoDocumentoDescricao")
        ),
        "url": safe_public_url(item.get("url") or item.get("uri")),
    }


def opportunity_item_record(item):
    quantity = item.get("quantidade")
    unit_value = item.get("valorUnitarioEstimado")
    total_value = item.get("valorTotal")
    if total_value is None and quantity is not None and unit_value is not None:
        try:
            total_value = float(quantity) * float(unit_value)
        except (TypeError, ValueError):
            total_value = None
    return {
        "numero": compact(item.get("numeroItem") or item.get("item")),
        "lote": compact(
            item.get("numeroLote")
            or item.get("lote")
            or item.get("grupo")
        ),
        "descricao": compact(item.get("descricao")),
        "quantidade": format_pncp_quantity(quantity),
        "unidade": compact(
            item.get("unidadeMedida")
            or item.get("unidadeFornecimento")
        ) or STANDARD_UNIT,
        "valor_unitario_estimado": unit_value,
        "valor_total_estimado": total_value,
        "criterio_julgamento": compact(
            item.get("criterioJulgamentoNome")
            or item.get("criterioJulgamento")
        ),
        "situacao": compact(
            item.get("situacaoCompraItemNome")
            or item.get("situacaoItemNome")
            or item.get("situacaoCompraItem")
        ),
        "tipo": compact(
            item.get("materialOuServicoNome")
            or item.get("tipoItemNome")
        ),
    }


def opportunity_detail_from_pncp_link(link, fallback=None):
    cnpj, ano, sequencial = parse_pncp_link(link)
    fallback = fallback if isinstance(fallback, dict) else {}
    metadata = {
        "numero_compra": compact(fallback.get("numero_compra")),
        "processo": compact(fallback.get("processo")),
        "modalidade": compact(fallback.get("modalidade")),
        "objeto": compact(fallback.get("objeto")),
        "orgao": compact(fallback.get("orgao")),
        "orgao_cnpj": cnpj,
        "unidade": compact(fallback.get("unidade")),
        "municipio": compact(fallback.get("municipio")),
        "uf": compact(fallback.get("uf")),
        "numero_controle_pncp": compact(fallback.get("numero_controle_pncp")),
        "abertura": compact(fallback.get("abertura")),
        "encerramento": compact(fallback.get("encerramento")),
        "situacao": compact(fallback.get("situacao")),
        "valor_total_estimado": fallback.get("valor_total_estimado"),
        "modo_disputa": compact(fallback.get("modo_disputa")),
        "codigo_unidade": compact(fallback.get("codigo_unidade")),
        "link_sistema_origem": compact(fallback.get("link_sistema_origem")),
    }
    with ThreadPoolExecutor(max_workers=3) as executor:
        metadata_future = executor.submit(
            pncp_purchase_metadata, cnpj, ano, sequencial
        )
        items_future = executor.submit(
            list_pncp_item_payload, cnpj, ano, sequencial
        )
        files_future = executor.submit(list_pncp_files, cnpj, ano, sequencial)
        remote_metadata = metadata_future.result()
        raw_items = items_future.result()
        files = files_future.result()
    metadata.update({
        key: value
        for key, value in remote_metadata.items()
        if value not in (None, "")
    })
    if not metadata:
        raise RuntimeError("A contratação não foi localizada na API oficial do PNCP.")
    api_items = [
        opportunity_item_record(item)
        for item in raw_items
        if compact(item.get("numeroItem") or item.get("item"))
    ]
    official_link = pncp_app_link(cnpj, ano, sequencial)
    file_items = []
    file_error = ""
    source_data = None
    try:
        source_data = source_from_pncp_link(official_link)
        file_items = extract_items_cached(source_data["source_path"])
    except Exception as exc:
        file_error = str(exc) or "Documento oficial não pôde ser lido."

    api_by_key = {
        (compact(item.get("lote")), compact(item.get("numero"))): item
        for item in api_items
    }
    if file_items:
        items = []
        for file_item in file_items:
            key = (compact(file_item.get("lote")), compact(file_item.get("item")))
            api_item = api_by_key.get(key, {})
            document_item = opportunity_item_record({
                "numeroItem": file_item.get("item"),
                "numeroLote": file_item.get("lote"),
                "descricao": file_item.get("descricao"),
                "quantidade": file_item.get("quantidade"),
                "unidadeMedida": file_item.get("unidade"),
            })
            merged = {**api_item, **{
                field: value
                for field, value in document_item.items()
                if value not in (None, "")
            }}
            items.append(merged)
    else:
        items = api_items

    comparison_api_items = [
        {"item": item.get("numero"), "lote": item.get("lote"), "descricao": item.get("descricao")}
        for item in api_items
    ]
    items_check = build_pncp_items_check(file_items, comparison_api_items)
    if file_error:
        items_check["file_error"] = file_error
    items_check["source"] = "documento_oficial" if file_items else "api_pncp"
    items_check["documento"] = (
        source_data.get("pncp", {}).get("documento_usado", "") if source_data else ""
    )
    source_link = safe_public_url(metadata.get("link_sistema_origem", ""))
    return {
        "oportunidade": {
            **metadata,
            "cnpj": cnpj,
            "ano": ano,
            "sequencial": sequencial,
            "link_pncp": official_link,
            "link_origem": source_link,
            "portal_origem": opportunity_source_portal(source_link),
            "categorias": opportunity_categories(metadata.get("objeto", ""), items),
        },
        "arquivos": [opportunity_file_record(item) for item in files],
        "itens": items,
        "verificacao_itens": items_check,
        "fontes": {
            "oportunidade": "API oficial do PNCP - contratação e busca pública",
            "arquivos": "API oficial do PNCP - arquivos da contratação",
            "itens": (
                "Documento oficial conferido com a API do PNCP"
                if file_items
                else "API oficial do PNCP (documento indisponível para conferência)"
            ),
        },
    }


OPPORTUNITY_QUESTION_STOPWORDS = {
    "a", "ao", "aos", "as", "com", "como", "da", "das", "de", "do", "dos",
    "e", "em", "esta", "este", "foi", "ha", "no", "nos", "o", "os", "para",
    "por", "qual", "que", "sao", "se", "sobre", "um", "uma",
}


def answer_opportunity_question(link, question):
    question = compact(question)
    if len(question) < 3:
        raise ValueError("Digite uma pergunta sobre o edital.")
    if len(question) > 500:
        raise ValueError("A pergunta deve ter no máximo 500 caracteres.")
    source = source_from_pncp_link(link)
    source_path = Path(source["source_path"])
    document_text = catalog_document_text(source_path)
    chunks = [
        compact(chunk)
        for chunk in re.split(r"\n\s*\n|(?<=[.;:])\s+(?=[A-ZÁÉÍÓÚÂÊÔÃÕÇ0-9])", document_text)
        if len(compact(chunk)) >= 40
    ]
    terms = {
        term for term in norm(question).split()
        if len(term) >= 3 and term not in OPPORTUNITY_QUESTION_STOPWORDS
    }
    ranked = []
    for index, chunk in enumerate(chunks):
        normalized = norm(chunk)
        score = sum(1 for term in terms if term in normalized)
        if score:
            ranked.append((score, -index, chunk[:1200]))
    excerpts = [row[2] for row in sorted(ranked, reverse=True)[:3]]
    return {
        "resposta": (
            "Encontrei os trechos abaixo no documento oficial que mais se relacionam à pergunta."
            if excerpts
            else "Não localizei no documento oficial um trecho relacionado à pergunta."
        ),
        "trechos": excerpts,
        "documento": source["pncp"].get("documento_usado") or source_path.name,
        "tipo_documento": source["pncp"].get("documento_tipo", ""),
    }


def pncp_item_match_text(item):
    description = norm(item.get("descricao"))
    return description.split("conforme especificacoes", 1)[0].strip(" .;:-")


def reconcile_scanned_quantities(file_items, pncp_items):
    available = {
        index: pncp_item_match_text(item)
        for index, item in enumerate(pncp_items)
        if pncp_item_match_text(item)
    }
    matched_indexes = set()
    filled_items = []

    for row in file_items:
        name = norm(row.get("_nome"))
        if not name or not available:
            continue
        best_index = None
        best_score = 0.0
        for index, candidate in available.items():
            if index in matched_indexes:
                continue
            score = SequenceMatcher(None, name, candidate).ratio()
            if name in candidate or candidate in name:
                score = max(score, 0.92)
            if score > best_score:
                best_index = index
                best_score = score
        if best_index is None or best_score < 0.62:
            continue
        matched_indexes.add(best_index)
        if not compact(row.get("quantidade")):
            row["quantidade"] = compact(pncp_items[best_index].get("quantidade"))
            if row["quantidade"]:
                filled_items.append(compact(row.get("item")))

    return {
        "matched_count": len(matched_indexes),
        "filled_items": filled_items,
        "unmatched_file_count": len(file_items) - len(matched_indexes),
        "unmatched_pncp_count": len(pncp_items) - len(matched_indexes),
    }


def document_name_priority(value):
    text = norm(value)
    if "termo de referencia" in text or re.search(r"\btr\b", text):
        return 0
    if "edital" in text:
        return 1
    return 2


def pncp_file_score(item):
    text = " ".join([
        str(item.get("titulo", "")),
        str(item.get("tipoDocumentoNome", "")),
        str(item.get("tipoDocumentoDescricao", "")),
    ])
    return (document_name_priority(text), norm(text))


def choose_pncp_file(files):
    if not files:
        raise RuntimeError("Nenhum arquivo encontrado no PNCP para este edital.")
    return sorted(files, key=pncp_file_score)[0]


def filename_from_disposition(disposition):
    if not disposition:
        return ""
    match = re.search(r'filename\*?=(?:UTF-8\'\')?"?([^";]+)"?', disposition)
    if not match:
        return ""
    return unquote(match.group(1)).strip()


def download_pncp_file(file_info):
    url = file_info.get("url") or file_info.get("uri")
    if not url:
        raise RuntimeError("Arquivo do PNCP sem URL de download.")

    req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(req, timeout=60) as response:
        data = response.read()
        headers = response.headers

    filename = filename_from_disposition(headers.get("Content-Disposition", ""))
    suffix = Path(filename).suffix.lower()
    if data.startswith(b"PK\x03\x04") and suffix != ".docx":
        suffix = ".zip"
    elif data.startswith(b"%PDF"):
        suffix = ".pdf"
    elif suffix not in {".pdf", ".docx", ".zip"}:
        title_suffix = Path(str(file_info.get("titulo", ""))).suffix.lower()
        suffix = title_suffix if title_suffix in {".pdf", ".docx", ".zip"} else ".pdf"

    stem = safe_name(file_info.get("titulo") or filename or "arquivo_pncp")
    path = UPLOAD_DIR / f"pncp_{stem}_{uuid.uuid4().hex}{suffix}"
    path.write_bytes(data)
    return path, file_info


MAX_ARCHIVE_DEPTH = 4
MAX_ARCHIVE_FILES = 250
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 300 * 1024 * 1024


def is_zip_archive(path):
    if path.suffix.lower() in {".docx", ".xlsx", ".pptx"}:
        return False
    try:
        with path.open("rb") as handle:
            return handle.read(4).startswith(b"PK")
    except OSError:
        return False


def safe_archive_target(destination, member_name):
    parts = [part for part in member_name.replace("\\", "/").split("/") if part not in {"", "."}]
    if not parts or any(part == ".." for part in parts):
        return None
    target = (destination / Path(*parts)).resolve()
    try:
        target.relative_to(destination.resolve())
    except ValueError:
        return None
    return target


def extract_archive_documents(archive_path, destination, depth=0):
    if depth > MAX_ARCHIVE_DEPTH:
        raise RuntimeError("O pacote oficial possui arquivos compactados além do limite suportado.")

    documents = []
    with zipfile.ZipFile(archive_path) as archive:
        members = [member for member in archive.infolist() if not member.is_dir()]
        if len(members) > MAX_ARCHIVE_FILES:
            raise RuntimeError("O pacote oficial possui arquivos demais para uma extração segura.")
        if sum(member.file_size for member in members) > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
            raise RuntimeError("O pacote oficial excede o tamanho máximo de extração segura.")

        for member in members:
            target = safe_archive_target(destination, member.filename)
            if target is None:
                continue
            target.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, target.open("wb") as output:
                shutil.copyfileobj(source, output)

            suffix = target.suffix.lower()
            if suffix in {".pdf", ".docx"} and not is_zip_archive(target):
                documents.append(target)
                continue
            if suffix == ".zip" or is_zip_archive(target):
                nested_destination = target.parent / f"{safe_name(target.name)}_extraido"
                nested_destination.mkdir(parents=True, exist_ok=True)
                documents.extend(extract_archive_documents(target, nested_destination, depth + 1))
    return documents


def downloaded_document_candidates(path):
    if path.suffix.lower() in {".pdf", ".docx"} and not is_zip_archive(path):
        return [(path, False)]
    if path.suffix.lower() == ".zip" or is_zip_archive(path):
        destination = path.parent / f"{safe_name(path.name)}_extraido"
        destination.mkdir(parents=True, exist_ok=True)
        return [(document, True) for document in extract_archive_documents(path, destination)]
    return []


def document_content_priority(path):
    try:
        if path.suffix.lower() == ".pdf":
            with pdfplumber.open(str(path)) as pdf:
                text = "\n".join(
                    page.extract_text() or ""
                    for page in pdf.pages[:3]
                )
        elif path.suffix.lower() == ".docx":
            document = Document(str(path))
            text = "\n".join(
                paragraph.text
                for paragraph in document.paragraphs[:80]
            )
        else:
            return 2
    except Exception:
        return 2

    heading = norm(text[:12000])
    beginning = heading[:2000]
    if "estudo tecnico preliminar" in beginning:
        return 2
    if (
        re.search(r"(?:^|anexo\s+[ivx0-9]+\s*[-:])\s*termo de referencia\b", beginning)
        or beginning.startswith("termo de referencia")
    ):
        return 0
    if (
        "edital de licitacao" in beginning
        or re.search(r"\bedital\b", beginning)
        or "pregao eletronico" in beginning
    ):
        return 1
    return 2


def candidate_document_score(path, file_info, embedded):
    priority = document_name_priority(path.name)
    if priority == 2:
        priority = document_content_priority(path)
    if not embedded and priority == 2:
        priority = pncp_file_score(file_info)[0]
    revision_text = norm(path.name)
    revision_priority = 0 if any(
        word in revision_text for word in ("retificado", "republicado", "atualizado", "versao final")
    ) else 1
    publication_priority = -int(file_info.get("sequencialDocumento") or 0)
    return (priority, revision_priority, publication_priority, norm(path.name))


def source_from_pncp_link(link):
    cnpj, ano, sequencial = parse_pncp_link(link)
    cache_key = pncp_app_link(cnpj, ano, sequencial)
    cached = cache_get(SOURCE_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None and Path(cached["source_path"]).is_file():
        return cached

    files = list_pncp_files(cnpj, ano, sequencial)
    if not files:
        raise RuntimeError("Nenhum arquivo encontrado no PNCP para este edital.")

    best = None
    document_candidates = []
    download_errors = []
    for file_info in sorted(files, key=pncp_file_score):
        try:
            downloaded_path, _ = download_pncp_file(file_info)
            candidates = downloaded_document_candidates(downloaded_path)
        except Exception as exc:
            download_errors.append(str(exc))
            continue

        for candidate_path, embedded in candidates:
            score = candidate_document_score(candidate_path, file_info, embedded)
            candidate = (score, candidate_path, file_info, embedded)
            document_candidates.append(candidate)
            if best is None or score < best[0]:
                best = candidate
    if best is None or best[0][0] > 1:
        detail = f" Detalhe: {'; '.join(download_errors)}" if download_errors else ""
        raise RuntimeError(
            "Nenhum Termo de Referência ou Edital legível foi encontrado nos arquivos oficiais do PNCP." + detail
        )

    score, source_path, chosen, embedded = best
    source_kind = "Termo de Referência" if score[0] == 0 else "Edital"
    result = {
        "source_path": source_path,
        "source_stem": f"PNCP_{cnpj}_{ano}_{sequencial}",
        "documentos_candidatos": [
            {
                "score": candidate[0],
                "path": candidate[1],
                "file_info": candidate[2],
                "embedded": candidate[3],
            }
            for candidate in sorted(document_candidates, key=lambda row: row[0])
        ],
        "pncp": {
            "cnpj": cnpj,
            "ano": ano,
            "sequencial": sequencial,
            "link": pncp_app_link(cnpj, ano, sequencial),
            "arquivo_usado": chosen,
            "documento_usado": source_path.name,
            "documento_tipo": source_kind,
            "documento_interno": embedded,
            "arquivos": files,
        },
    }
    cache_set(SOURCE_CACHE, cache_key, result)
    return result


def yyyymmdd(value):
    value = re.sub(r"\D", "", value or "")
    if len(value) != 8:
        raise ValueError("Data inválida. Use o formato AAAAMMDD.")
    return value


def optional_yyyymmdd(value):
    value = re.sub(r"\D", "", value or "")
    if not value:
        return ""
    if len(value) != 8:
        raise ValueError("Data invalida. Use o formato AAAAMMDD.")
    return value


def default_date_range():
    start = datetime.now()
    end = start + timedelta(days=29)
    return start.strftime("%Y%m%d"), end.strftime("%Y%m%d")


def date_key(value):
    return re.sub(r"\D", "", str(value or ""))[:8]


def date_in_range(value, data_inicial="", data_final=""):
    key = date_key(value)
    if not key:
        return True
    if data_inicial and key < data_inicial:
        return False
    if data_final and key > data_final:
        return False
    return True


def date_range_days(data_inicial, data_final):
    if not data_inicial or not data_final:
        return 0
    start = datetime.strptime(data_inicial, "%Y%m%d")
    end = datetime.strptime(data_final, "%Y%m%d")
    return (end - start).days + 1


def is_historical_search_period(data_inicial, data_final):
    if not data_inicial or not data_final:
        return False
    today = datetime.now().strftime("%Y%m%d")
    return data_final <= today


def bounded_int(value, default, minimum=1, maximum=50):
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = default
    return max(minimum, min(number, maximum))


def map_search_item(row):
    cnpj = row.get("orgao_cnpj", "")
    ano = row.get("ano", "")
    sequencial = row.get("numero_sequencial", "")
    return {
        "orgao": row.get("orgao_nome", ""),
        "cnpj": cnpj,
        "ano": ano,
        "sequencial": sequencial,
        "numeroCompra": row.get("title", ""),
        "processo": row.get("numero_controle_pncp", ""),
        "modalidade": row.get("modalidade_nome", "") or row.get("modalidade_licitacao_nome", ""),
        "objeto": row.get("description", ""),
        "uf": row.get("uf", ""),
        "municipio": row.get("municipio_nome", ""),
        "unidade": row.get("unidade_nome", "") or row.get("unidade_orgao_nome", ""),
        "codigoUnidade": row.get("unidade_codigo", "") or row.get("codigo_unidade", ""),
        "valorTotalEstimado": row.get("valor_total_estimado"),
        "modoDisputa": row.get("modo_disputa_nome", ""),
        "situacao": row.get("situacao_nome", "") or row.get("situacao_compra_nome", ""),
        "linkOrigem": row.get("item_url", "") or row.get("link_sistema_origem", ""),
        "abertura": row.get("data_inicio_vigencia", ""),
        "encerramento": row.get("data_fim_vigencia", ""),
        "link": pncp_app_link(cnpj, ano, sequencial) if cnpj and ano and sequencial else "",
    }


BRAZILIAN_UFS = (
    "AC", "AL", "AP", "AM", "BA", "CE", "DF", "ES", "GO",
    "MA", "MT", "MS", "MG", "PA", "PB", "PR", "PE", "PI",
    "RJ", "RN", "RS", "RO", "RR", "SC", "SP", "SE", "TO",
)
PNCP_MODALITY_IDS = (1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13)


def parse_search_ufs(value):
    selected = []
    for part in re.split(r"[,;\s]+", str(value or "").upper()):
        uf = part.strip()
        if not uf or uf in selected:
            continue
        if uf not in BRAZILIAN_UFS:
            raise ValueError(f"UF invalida: {uf}.")
        selected.append(uf)
    return tuple(selected)


def split_search_keywords(value, maximum=20):
    keywords = []
    seen = set()
    for part in str(value or "").split(";"):
        keyword = compact(part)
        if not keyword:
            continue
        normalized = norm(keyword)
        if normalized in seen:
            continue
        if len(keyword) > 120:
            raise ValueError("Cada palavra-chave deve possuir no maximo 120 caracteres.")
        seen.add(normalized)
        keywords.append(keyword)
    if len(keywords) > maximum:
        raise ValueError(f"Informe no maximo {maximum} palavras-chave.")
    return keywords


MAX_SEARCH_TERM_GAP = 2


def search_word_variants(word):
    word = norm(word)
    if not word:
        return set()
    variants = {word}
    if len(word) > 3:
        if word.endswith("oes"):
            variants.add(f"{word[:-3]}ao")
        if word.endswith("ais"):
            variants.add(f"{word[:-3]}al")
        if word.endswith("eis"):
            variants.add(f"{word[:-3]}el")
        if word.endswith("is"):
            variants.add(f"{word[:-2]}il")
        if word.endswith("es"):
            variants.add(word[:-2])
        if word.endswith("s"):
            variants.add(word[:-1])
        variants.add(f"{word}s")
        if word[-1] in {"r", "z", "n"}:
            variants.add(f"{word}es")
        if word.endswith("al"):
            variants.add(f"{word[:-2]}ais")
        if word.endswith("el"):
            variants.add(f"{word[:-2]}eis")
        if word.endswith("ao"):
            variants.add(f"{word[:-2]}oes")
    return {variant for variant in variants if variant}


def search_words_match(actual, expected):
    return bool(search_word_variants(actual) & search_word_variants(expected))


def matches_complete_words(text, search_term):
    expected = norm(search_term).split()
    if not expected:
        return True
    words = norm(text).split()
    for start, word in enumerate(words):
        if not search_words_match(word, expected[0]):
            continue
        position = start
        matched = True
        for expected_word in expected[1:]:
            stop = min(len(words), position + MAX_SEARCH_TERM_GAP + 2)
            next_position = next(
                (
                    index
                    for index in range(position + 1, stop)
                    if search_words_match(words[index], expected_word)
                ),
                None,
            )
            if next_position is None:
                matched = False
                break
            position = next_position
        if matched:
            return True
    return False


def matches_complete_search_term(row, search_term):
    return matches_complete_words(
        f"{row.get('title', '')} {row.get('description', '')}",
        search_term,
    )


def normalized_numeric_code(value):
    digits = re.sub(r"\D", "", compact(value))
    return digits.lstrip("0") or ("0" if digits else "")


def row_matches_purchase_filters(row, purchase_number="", uasg=""):
    purchase_number = norm(purchase_number)
    if purchase_number:
        purchase_text = norm(" ".join(compact(value) for value in (
            row.get("numeroCompra"),
            row.get("numero_compra"),
            row.get("numero"),
            row.get("title"),
        ) if value not in (None, "")))
        if purchase_number not in purchase_text:
            return False

    uasg = normalized_numeric_code(uasg)
    if uasg:
        unit = row.get("unidadeOrgao") or {}
        candidate_codes = (
            row.get("codigoUnidade"),
            row.get("codigo_unidade"),
            row.get("unidade_codigo"),
            unit.get("codigoUnidade") if isinstance(unit, dict) else "",
            unit.get("codigo") if isinstance(unit, dict) else "",
        )
        if not any(normalized_numeric_code(value) == uasg for value in candidate_codes):
            return False
    return True


def search_row_contract_key(row):
    cnpj = compact(row.get("orgao_cnpj"))
    ano = compact(row.get("ano"))
    sequencial = compact(row.get("numero_sequencial"))
    if not cnpj or not ano or not sequencial:
        return None
    return cnpj, ano, sequencial


def get_search_row_pncp_items(row):
    contract_key = search_row_contract_key(row)
    if contract_key is None:
        return []
    cnpj, ano, sequencial = contract_key
    cache_key = f"{cnpj}:{ano}:{sequencial}"
    cached = cache_get(SEARCH_ITEM_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None:
        return cached
    items = list_pncp_items(cnpj, ano, sequencial)
    cache_set(SEARCH_ITEM_CACHE, cache_key, items)
    return items


def get_search_row_document_items(row):
    contract_key = search_row_contract_key(row)
    if contract_key is None:
        return {"items": [], "error": "Identificação PNCP incompleta."}
    cnpj, ano, sequencial = contract_key
    cache_key = f"{cnpj}:{ano}:{sequencial}"
    cached = cache_get(SEARCH_DOCUMENT_ITEM_CACHE, cache_key, DOCUMENT_CACHE_TTL)
    if cached is not None:
        return cached
    try:
        source = source_from_pncp_link(pncp_app_link(cnpj, ano, sequencial))
        items = extract_items_cached(source["source_path"])
        result = {
            "items": items,
            "error": "",
            "documento": source.get("pncp", {}).get("documento_usado", ""),
        }
    except Exception as exc:
        result = {"items": [], "error": str(exc) or "Documento oficial não pôde ser lido."}
    cache_set(SEARCH_DOCUMENT_ITEM_CACHE, cache_key, result)
    return result


def items_match_search_term(items, search_term):
    return any(
        matches_complete_words(item.get("descricao", ""), search_term)
        for item in items
    )


def matches_search_term_after_item_identification(row, search_term, verify_document=True):
    try:
        pncp_items = get_search_row_pncp_items(row)
    except Exception:
        pncp_items = []
    if items_match_search_term(pncp_items, search_term):
        return True

    # A prévia permanece rápida. A busca completa, executada em segundo plano,
    # consulta o documento antes de rejeitar uma oportunidade.
    if verify_document:
        document_result = get_search_row_document_items(row)
        document_items = document_result.get("items") or []
        if document_items:
            return items_match_search_term(document_items, search_term)
    if pncp_items:
        return False
    return matches_complete_search_term(row, search_term)


def filter_rows_by_complete_search_term(rows, search_term, verify_documents=True):
    if not norm(search_term):
        return list(rows)

    accepted = []
    rows = list(rows)
    if not rows:
        return accepted
    with ThreadPoolExecutor(max_workers=min(6, len(rows))) as executor:
        futures = {
            executor.submit(
                matches_search_term_after_item_identification,
                row,
                search_term,
                verify_documents,
            ): row
            for row in rows
        }
        for future in as_completed(futures):
            try:
                if future.result():
                    accepted.append(futures[future])
            except Exception:
                continue
    return accepted


def classify_search_text(text):
    return classify_object_text(text)


def classify_search_object(row):
    return classify_search_text(f"{row.get('title', '')} {row.get('description', '')}")


def classify_search_items(row):
    return classify_search_text(
        " ".join(item.get("descricao", "") for item in get_search_row_pncp_items(row))
    )


def row_matches_object_type(row, object_type):
    if not object_type:
        return True
    row_type = classify_search_object(row)
    if row_type == object_type:
        return True
    item_type = classify_search_items(row)
    return item_type == object_type


def row_matches_opportunity_type(row, object_type):
    if not object_type:
        return True
    classified = classify_search_object(row)
    return not classified or classified == object_type


def filter_opportunity_rows_by_search_term(rows, search_term):
    if not norm(search_term):
        return list(rows)
    return [row for row in rows if matches_complete_search_term(row, search_term)]


def reconcile_pncp_search_rows(rows, request_url, filters):
    repository = etl_repository()
    repository.initialize()
    run_id = repository.create_run("pncp", "search_reconciliation", filters)
    mapper = PNCPMapper()
    classifier = OpportunityClassifier()
    counters = {"fetched": 0, "inserted": 0, "updated": 0, "skipped": 0, "failed": 0}
    for raw_row in rows:
        counters["fetched"] += 1
        try:
            opportunity = mapper.map(raw_row)
            match = classifier.classify(opportunity, {})
            outcome, _ = repository.persist_record(
                run_id=run_id,
                source_endpoint="api/search",
                request_url=request_url,
                raw_payload=raw_row,
                opportunity=opportunity,
                match=match,
                replace_children=False,
            )
            counters[outcome] += 1
        except Exception as exc:
            counters["failed"] += 1
            try:
                repository.save_failed_source_record(
                    run_id=run_id,
                    source="pncp",
                    source_endpoint="api/search",
                    request_url=request_url,
                    raw_payload=raw_row,
                    error_message=f"{type(exc).__name__}: {exc}",
                )
            except Exception:
                pass

    status = "partial" if counters["failed"] else "success"
    repository.finish_run(run_id, status=status, counters=counters)
    return {"run_id": run_id, "status": status, **counters}


def search_pncp_app_editais(params, data_inicial, data_final):
    keywords = split_search_keywords(params.get("palavraChave"))
    search_terms = tuple(keywords) or ("",)
    object_type = str(params.get("tipoObjeto") or "").strip().lower()
    if object_type not in {"", "material", "servico"}:
        raise ValueError("Tipo do objeto invalido.")
    purchase_number = compact(params.get("numeroCompra"))
    if len(purchase_number) > 80:
        raise ValueError("Número da compra excede 80 caracteres.")
    uasg = re.sub(r"\D", "", compact(params.get("uasg")))
    if len(uasg) > 20:
        raise ValueError("UASG excede 20 dígitos.")
    page_size = bounded_int(params.get("tamanhoPagina"), 50, 1, 500)
    start_page = bounded_int(params.get("pagina"), 1, 1, 1000)
    source_page_size = 500
    base_query = {
        "tipos_documento": "edital",
        "status": "recebendo_proposta",
        "ordenacao": str(params.get("ordenacao") or "-data"),
        "tam_pagina": source_page_size,
    }
    modalidade = str(params.get("codigoModalidadeContratacao") or "").strip()
    if modalidade:
        base_query["modalidades"] = modalidade
    selected_ufs = parse_search_ufs(params.get("uf"))
    reconcile_requested = str(params.get("reconciliar") or "").strip().lower() in {
        "1", "true", "yes", "sim",
    }
    result_cache_key = json.dumps({
        **base_query,
        "keywords": keywords,
        "ufs": selected_ufs,
        "dataInicial": data_inicial,
        "dataFinal": data_final,
        "tipoObjeto": object_type,
        "numeroCompra": purchase_number,
        "uasg": uasg,
        "reconciliar": reconcile_requested,
    }, ensure_ascii=True, sort_keys=True)
    consolidated = cache_get(PNCP_RESULT_CACHE, result_cache_key, SEARCH_CACHE_TTL)
    aggregate_cache_hit = consolidated is not None

    if consolidated is None:
        partition_ufs = selected_ufs or BRAZILIAN_UFS
        partitions = tuple(
            (search_term, partition_uf)
            for search_term in search_terms
            for partition_uf in partition_ufs
        )

        def fetch_source_page(search_term, partition_uf, page, timeout=18):
            query = dict(base_query)
            if search_term:
                query["q"] = search_term
            query["ufs"] = partition_uf
            query["pagina"] = page
            url = f"{PNCP_SEARCH_URL}?{urlencode(query)}"
            try:
                payload = request_json(url, timeout=timeout)
            except Exception as exc:
                payload = {"items": [], "timeout": True, "message": str(exc)}
            return (search_term, partition_uf, page), url, payload

        payloads = {}
        urls_by_page = {}
        with ThreadPoolExecutor(max_workers=min(6, len(partitions))) as executor:
            futures = {
                executor.submit(fetch_source_page, search_term, partition_uf, 1): partition
                for partition in partitions
                for search_term, partition_uf in (partition,)
            }
            for future in as_completed(futures):
                page_key, url, payload = future.result()
                payloads[page_key] = payload
                urls_by_page[page_key] = url

        for retry_delay in (0.75, 1.5, 3.0):
            failed_keys = [
                page_key for page_key, payload in payloads.items()
                if payload.get("rate_limited") or payload.get("timeout")
            ]
            if not failed_keys:
                break
            time.sleep(retry_delay)
            with ThreadPoolExecutor(max_workers=min(3, len(failed_keys))) as executor:
                futures = {
                    executor.submit(fetch_source_page, search_term, partition_uf, page, 30): page_key
                    for page_key in failed_keys
                    for search_term, partition_uf, page in (page_key,)
                }
                for future in as_completed(futures):
                    page_key, url, payload = future.result()
                    payloads[page_key] = payload
                    urls_by_page[page_key] = url

        partition_pages = {}
        source_total = 0
        for search_term, partition_uf in partitions:
            payload = payloads.get((search_term, partition_uf, 1), {})
            if payload.get("rate_limited") or payload.get("timeout"):
                continue
            try:
                partition_total = int(payload.get("total", len(payload.get("items", []))))
            except (TypeError, ValueError):
                partition_total = len(payload.get("items", []))
            source_total += partition_total
            partition_pages[(search_term, partition_uf)] = (
                max(1, (partition_total + source_page_size - 1) // source_page_size)
                if partition_total else 1
            )

        remaining_keys = [
            (search_term, partition_uf, page)
            for (search_term, partition_uf), pages in partition_pages.items()
            for page in range(2, pages + 1)
        ]
        if remaining_keys:
            with ThreadPoolExecutor(max_workers=min(6, len(remaining_keys))) as executor:
                futures = {
                    executor.submit(fetch_source_page, search_term, partition_uf, page): page_key
                    for page_key in remaining_keys
                    for search_term, partition_uf, page in (page_key,)
                }
                for future in as_completed(futures):
                    page_key, url, payload = future.result()
                    payloads[page_key] = payload
                    urls_by_page[page_key] = url

        for retry_delay in (0.75, 1.5, 3.0):
            failed_keys = [
                page_key for page_key, payload in payloads.items()
                if payload.get("rate_limited") or payload.get("timeout")
            ]
            if not failed_keys:
                break
            time.sleep(retry_delay)
            with ThreadPoolExecutor(max_workers=min(2, len(failed_keys))) as executor:
                futures = {
                    executor.submit(fetch_source_page, search_term, partition_uf, page, 30): page_key
                    for page_key in failed_keys
                    for search_term, partition_uf, page in (page_key,)
                }
                for future in as_completed(futures):
                    page_key, url, payload = future.result()
                    payloads[page_key] = payload
                    urls_by_page[page_key] = url

        failed_keys = [
            page_key for page_key, payload in payloads.items()
            if payload.get("rate_limited") or payload.get("timeout")
        ]
        for search_term, partition_uf, page in failed_keys:
            time.sleep(0.35)
            page_key, url, payload = fetch_source_page(
                search_term, partition_uf, page, timeout=45
            )
            payloads[page_key] = payload
            urls_by_page[page_key] = url

        candidates_by_term = {
            search_term: {}
            for search_term in search_terms
        }
        rate_limited = False
        timed_out = False
        page_cache_hit = False
        for page_key in sorted(payloads):
            search_term, _partition_uf, _page = page_key
            payload = payloads[page_key]
            rate_limited = rate_limited or bool(payload.get("rate_limited"))
            timed_out = timed_out or bool(payload.get("timeout"))
            page_cache_hit = page_cache_hit or bool(payload.get("cache_hit"))
            for row in payload.get("items", []):
                if not date_in_range(row.get("data_fim_vigencia"), data_inicial, data_final):
                    continue
                if not row_matches_purchase_filters(row, purchase_number, uasg):
                    continue
                if not row_matches_opportunity_type(row, object_type):
                    continue
                row_key = (
                    row.get("id")
                    or row.get("numero_controle_pncp")
                    or row.get("item_url")
                    or hashlib.sha256(
                        json.dumps(row, ensure_ascii=True, sort_keys=True, default=str).encode("utf-8")
                    ).hexdigest()
                )
                candidates_by_term[search_term].setdefault(row_key, row)

        filtered_rows = []
        seen = set()
        for search_term, candidates in candidates_by_term.items():
            for row in filter_opportunity_rows_by_search_term(
                candidates.values(), search_term
            ):
                row_key = (
                    row.get("id")
                    or row.get("numero_controle_pncp")
                    or row.get("item_url")
                    or hashlib.sha256(
                        json.dumps(
                            row,
                            ensure_ascii=True,
                            sort_keys=True,
                            default=str,
                        ).encode("utf-8")
                    ).hexdigest()
                )
                if row_key in seen:
                    continue
                seen.add(row_key)
                filtered_rows.append(row)

        filtered_rows.sort(
            key=lambda row: str(
                row.get("data_atualizacao_pncp")
                or row.get("createdAt")
                or row.get("data_publicacao_pncp")
                or ""
            ),
            reverse=True,
        )
        rows = [map_search_item(row) for row in filtered_rows]

        source_pages = sum(partition_pages.values())
        complete = (
            len(partition_pages) == len(partitions)
            and len(payloads) == source_pages
            and not rate_limited
            and not timed_out
        )
        reconciliation = None
        if reconcile_requested:
            try:
                reconciliation = reconcile_pncp_search_rows(
                    filtered_rows,
                    next(iter(urls_by_page.values()), PNCP_SEARCH_URL),
                    {
                        "dataInicial": data_inicial,
                        "dataFinal": data_final,
                        "ufs": selected_ufs,
                        "keywords": keywords,
                        "tipoObjeto": object_type,
                        "numeroCompra": purchase_number,
                        "uasg": uasg,
                    },
                )
                complete = complete and reconciliation["failed"] == 0
            except Exception as exc:
                reconciliation = {
                    "status": "failed",
                    "fetched": len(filtered_rows),
                    "inserted": 0,
                    "updated": 0,
                    "skipped": 0,
                    "failed": len(filtered_rows),
                    "error": str(exc) or "Nao foi possivel reconciliar a busca com a base local.",
                }
                complete = False
        failed_pages = [
            f"{search_term or '*'}:{partition_uf}:{page}"
            for (search_term, partition_uf, page), payload in payloads.items()
            if payload.get("rate_limited") or payload.get("timeout")
        ]
        consolidated = {
            "results": rows,
            "source_total": source_total,
            "source_pages": source_pages,
            "pages_checked": len(payloads),
            "source_url": next(iter(urls_by_page.values()), ""),
            "rate_limited": rate_limited,
            "timed_out": timed_out,
            "failed_pages": failed_pages,
            "complete": complete,
            "page_cache_hit": page_cache_hit,
            "reconciliation": reconciliation,
        }
        if complete:
            cache_set(PNCP_RESULT_CACHE, result_cache_key, consolidated)

    all_results = consolidated["results"]
    total = len(all_results)
    total_pages = max(1, (total + page_size - 1) // page_size) if total else 0
    if total_pages:
        start_page = min(start_page, total_pages)
    offset = (start_page - 1) * page_size
    page_results = all_results[offset:offset + page_size]

    return {
        "results": page_results,
        "total": total,
        "source_total": consolidated["source_total"],
        "pagina": start_page,
        "tamanhoPagina": page_size,
        "total_pages": total_pages,
        "has_previous": start_page > 1,
        "has_next": start_page < total_pages,
        "source_url": consolidated["source_url"],
        "source": "api/search",
        "pages_checked": consolidated["pages_checked"],
        "source_pages": consolidated["source_pages"],
        "rate_limited": consolidated["rate_limited"],
        "timed_out": consolidated["timed_out"],
        "failed_pages": consolidated["failed_pages"],
        "complete": consolidated["complete"],
        "dataInicial": data_inicial,
        "dataFinal": data_final,
        "tipoObjeto": object_type,
        "cache_hit": aggregate_cache_hit or consolidated["page_cache_hit"],
        "reconciliation": consolidated.get("reconciliation"),
    }


def search_pncp_historical_bids(params, data_inicial, data_final):
    keywords = split_search_keywords(params.get("palavraChave"))
    search_terms = tuple(keywords) or ("",)
    object_type = str(params.get("tipoObjeto") or "").strip().lower()
    if object_type not in {"", "material", "servico"}:
        raise ValueError("Tipo do objeto invalido.")
    purchase_number = compact(params.get("numeroCompra"))
    if len(purchase_number) > 80:
        raise ValueError("Número da compra excede 80 caracteres.")
    uasg = re.sub(r"\D", "", compact(params.get("uasg")))
    if len(uasg) > 20:
        raise ValueError("UASG excede 20 dígitos.")

    page_size = bounded_int(params.get("tamanhoPagina"), 50, 10, 50)
    start_page = bounded_int(params.get("pagina"), 1, 1, 1000)
    source_page_size = 50
    selected_ufs = parse_search_ufs(params.get("uf"))
    ufs = selected_ufs or ("",)
    modalidade = str(params.get("codigoModalidadeContratacao") or "").strip()
    modalities = (int(modalidade),) if modalidade else PNCP_MODALITY_IDS
    max_source_pages = 6 if (keywords or purchase_number or uasg or object_type) else 3

    rows = []
    seen = set()
    source_total = 0
    pages_checked = 0
    source_urls = []
    rate_limited = False
    timed_out = False
    failed_pages = []

    def fetch_page(modality, uf, page):
        query = {
            "dataInicial": data_inicial,
            "dataFinal": data_final,
            "codigoModalidadeContratacao": modality,
            "pagina": page,
            "tamanhoPagina": source_page_size,
        }
        if uf:
            query["uf"] = uf
        if uasg:
            query["codigoUnidadeAdministrativa"] = uasg
        url = f"{PNCP_API_BASE}/consulta/v1/contratacoes/publicacao?{urlencode(query)}"
        return url, request_json(url, timeout=30)

    for modality in modalities:
        for uf in ufs:
            for page in range(1, max_source_pages + 1):
                url, payload = fetch_page(modality, uf, page)
                if not source_urls:
                    source_urls.append(url)
                pages_checked += 1
                if payload.get("rate_limited"):
                    rate_limited = True
                    failed_pages.append(f"{modality}:{uf or '*'}:{page}")
                    break
                if payload.get("timeout"):
                    timed_out = True
                    failed_pages.append(f"{modality}:{uf or '*'}:{page}")
                    break
                page_rows = payload.get("data") or []
                if page == 1:
                    try:
                        source_total += int(payload.get("totalRegistros", len(page_rows)))
                    except (TypeError, ValueError):
                        source_total += len(page_rows)
                for source_row in page_rows:
                    row = consulta_row_to_search_item(source_row)
                    row_key = (
                        row.get("numero_controle_pncp")
                        or row.get("item_url")
                        or hashlib.sha256(
                            json.dumps(row, ensure_ascii=True, sort_keys=True, default=str).encode("utf-8")
                        ).hexdigest()
                    )
                    if row_key in seen:
                        continue
                    if not row_matches_purchase_filters(row, purchase_number, uasg):
                        continue
                    if not row_matches_object_type(row, object_type):
                        continue
                    search_text = norm(
                        f"{row.get('title', '')} {row.get('description', '')} {row.get('orgao_nome', '')}"
                    )
                    if keywords and not any(norm(term) in search_text for term in search_terms):
                        continue
                    seen.add(row_key)
                    rows.append(row)
                if not payload.get("paginasRestantes") or len(page_rows) < source_page_size:
                    break
                if len(rows) >= start_page * page_size and not (keywords or object_type):
                    break
                time.sleep(0.2)

    rows.sort(
        key=lambda row: str(
            row.get("data_publicacao_pncp")
            or row.get("data_atualizacao_pncp")
            or row.get("data_fim_vigencia")
            or ""
        ),
        reverse=True,
    )
    total = len(rows)
    total_pages = max(1, (total + page_size - 1) // page_size) if total else 0
    if total_pages:
        start_page = min(start_page, total_pages)
    offset = (start_page - 1) * page_size
    page_results = [map_search_item(row) for row in rows[offset:offset + page_size]]
    return {
        "results": page_results,
        "total": total,
        "source_total": source_total,
        "pagina": start_page,
        "tamanhoPagina": page_size,
        "total_pages": total_pages,
        "has_previous": start_page > 1,
        "has_next": start_page < total_pages,
        "source_url": source_urls[0] if source_urls else "",
        "source": "api/consulta/publicacao",
        "pages_checked": pages_checked,
        "source_pages": None,
        "rate_limited": rate_limited,
        "timed_out": timed_out,
        "failed_pages": failed_pages,
        "complete": not rate_limited and not timed_out,
        "dataInicial": data_inicial,
        "dataFinal": data_final,
        "tipoObjeto": object_type,
        "cache_hit": False,
    }


def combined_reconciliation_summary(*summaries):
    valid = [summary for summary in summaries if isinstance(summary, dict)]
    counters = {
        key: sum(int(summary.get(key) or 0) for summary in valid)
        for key in ("fetched", "inserted", "updated", "skipped", "failed")
    }
    statuses = {compact(summary.get("status")) for summary in valid}
    status = "failed" if statuses == {"failed"} else (
        "partial" if statuses & {"partial", "failed"} or counters["failed"] else "success"
    )
    run_ids = []
    endpoints = []
    for summary in valid:
        run_ids.extend(summary.get("run_ids") or ([summary["run_id"]] if summary.get("run_id") else []))
        endpoints.extend(summary.get("endpoints") or [])
    return {
        "status": status,
        **counters,
        "run_ids": run_ids,
        "endpoints": endpoints,
    }


def sync_pncp_opportunity_endpoints(params, data_inicial, data_final):
    selected_ufs = parse_search_ufs(params.get("uf"))
    selected_modality = compact(params.get("codigoModalidadeContratacao"))
    today = datetime.now().strftime("%Y%m%d")
    publication_start = data_inicial or today
    publication_end = min(data_final or today, today)
    cache_key = json.dumps({
        "dataInicial": publication_start,
        "dataFinal": data_final,
        "publicationEnd": publication_end,
        "ufs": selected_ufs,
        "modalidade": selected_modality,
    }, ensure_ascii=True, sort_keys=True)
    cached = cache_get(PNCP_OPPORTUNITY_SYNC_CACHE, cache_key, SEARCH_CACHE_TTL)
    if cached is not None:
        cached["cache_hit"] = True
        return cached

    endpoint_specs = [
        ("proposta", {"dataFinal": data_final or today}),
    ]
    if publication_start <= publication_end:
        modalities = [int(selected_modality)] if selected_modality else list(range(1, 20))
        endpoint_specs.extend([
            ("publicacao", {
                "dataInicial": publication_start,
                "dataFinal": publication_end,
                "modality_codes": modalities,
            }),
            ("atualizacao", {
                "dataInicial": publication_start,
                "dataFinal": publication_end,
                "modality_codes": modalities,
            }),
        ])
    if selected_modality:
        endpoint_specs[0][1]["codigoModalidadeContratacao"] = int(selected_modality)

    endpoint_results_by_name = {
        endpoint: {
            "name": endpoint,
            "fetched": 0, "inserted": 0, "updated": 0,
            "skipped": 0, "failed": 0,
            "run_ids": [], "errors": [],
        }
        for endpoint, _ in endpoint_specs
    }
    sync_units = []
    for endpoint, base_filters in endpoint_specs:
        for uf in selected_ufs or ("",):
            filters = dict(base_filters)
            if uf:
                filters["uf"] = uf
            sync_units.append((endpoint, filters))

    def sync_unit(endpoint, filters):
        service = ETLSyncService(
            etl_repository(),
            PNCPConnector(client=HttpJsonClient(timeout=12, retries=0)),
            PNCPMapper(),
            OpportunityClassifier(),
        )
        return service.sync(SyncRequest(
            endpoint=endpoint,
            filters=filters,
            run_type=f"search_reconciliation_{endpoint}",
            dry_run=False,
            max_pages=None,
            max_records=None,
            fetch_details=False,
            company_profile={},
        ))

    with ThreadPoolExecutor(max_workers=min(4, len(sync_units))) as executor:
        futures = {
            executor.submit(sync_unit, endpoint, filters): endpoint
            for endpoint, filters in sync_units
        }
        for future in as_completed(futures):
            endpoint = futures[future]
            endpoint_result = endpoint_results_by_name[endpoint]
            try:
                result = future.result()
                endpoint_result["run_ids"].append(result["run_id"])
                for key in ("fetched", "inserted", "updated", "skipped", "failed"):
                    endpoint_result[key] += int(result.get(key) or 0)
            except Exception as exc:
                endpoint_result["failed"] += 1
                endpoint_result["errors"].append(
                    str(exc) or f"Falha no endpoint {endpoint}."
                )

    endpoint_results = []
    for endpoint, _ in endpoint_specs:
        result = endpoint_results_by_name[endpoint]
        result["status"] = (
            "partial" if result["errors"] or result["failed"] else "success"
        )
        endpoint_results.append(result)

    summary = combined_reconciliation_summary(*endpoint_results)
    summary["endpoints"] = endpoint_results
    summary["cache_hit"] = False
    if summary["status"] == "success":
        cache_set(PNCP_OPPORTUNITY_SYNC_CACHE, cache_key, summary)
    return summary


def search_pncp_open_bids(params):
    data_inicial = optional_yyyymmdd(params.get("dataInicial"))
    data_final = optional_yyyymmdd(params.get("dataFinal"))
    if data_inicial and data_final and data_inicial > data_final:
        raise ValueError("Data inicial nao pode ser maior que a data final.")
    if date_range_days(data_inicial, data_final) > 30:
        raise ValueError("O periodo maximo e de 30 dias corridos.")
    result = search_pncp_app_editais(params, data_inicial, data_final)
    reconcile_requested = str(params.get("reconciliar") or "").strip().lower() in {
        "1", "true", "yes", "sim",
    }
    if not reconcile_requested:
        return result

    app_summary = result.get("reconciliation")
    official_summary = sync_pncp_opportunity_endpoints(params, data_inicial, data_final)
    result["reconciliation"] = combined_reconciliation_summary(
        app_summary, official_summary
    )
    result["reconciliation"]["endpoints"] = [
        {
            "name": "api/search",
            "status": (app_summary or {}).get("status", "success"),
        },
        *official_summary.get("endpoints", []),
    ]
    result["complete"] = bool(result.get("complete")) and official_summary["status"] == "success"
    return result


def pncp_search_job_key(params):
    filters = {
        key: str(params.get(key) or "").strip()
        for key in (
            "dataInicial",
            "dataFinal",
            "uf",
            "palavraChave",
            "tipoObjeto",
            "codigoModalidadeContratacao",
            "numeroCompra",
            "uasg",
            "ordenacao",
            "reconciliar",
        )
    }
    filters["palavraChave"] = split_search_keywords(filters["palavraChave"])
    return hashlib.sha256(
        json.dumps(filters, ensure_ascii=True, sort_keys=True).encode("utf-8")
    ).hexdigest()


def quick_pncp_search_preview(params):
    data_inicial = optional_yyyymmdd(params.get("dataInicial"))
    data_final = optional_yyyymmdd(params.get("dataFinal"))
    if data_inicial and data_final and data_inicial > data_final:
        raise ValueError("Data inicial nao pode ser maior que a data final.")
    if date_range_days(data_inicial, data_final) > 30:
        raise ValueError("O periodo maximo e de 30 dias corridos.")
    object_type = str(params.get("tipoObjeto") or "").strip().lower()
    if object_type not in {"", "material", "servico"}:
        raise ValueError("Tipo do objeto invalido.")
    purchase_number = compact(params.get("numeroCompra"))
    uasg = re.sub(r"\D", "", compact(params.get("uasg")))
    base_query = {
        "tipos_documento": "edital",
        "status": "recebendo_proposta",
        "ordenacao": str(params.get("ordenacao") or "-data"),
        "pagina": 1,
        "tam_pagina": 100,
    }
    keywords = split_search_keywords(params.get("palavraChave"))
    search_terms = tuple(keywords) or ("",)
    modalidade = str(params.get("codigoModalidadeContratacao") or "").strip()
    if modalidade:
        base_query["modalidades"] = modalidade
    selected_ufs = parse_search_ufs(params.get("uf"))
    preview_ufs = selected_ufs or ("",)

    def fetch_preview(search_term, uf):
        query = dict(base_query)
        if search_term:
            query["q"] = search_term
        if uf:
            query["ufs"] = uf
        url = f"{PNCP_SEARCH_URL}?{urlencode(query)}"
        return search_term, uf, url, request_json(url, timeout=12)

    previews = []
    partitions = tuple(
        (search_term, uf)
        for search_term in search_terms
        for uf in preview_ufs
    )
    with ThreadPoolExecutor(max_workers=min(6, len(partitions))) as executor:
        futures = {
            executor.submit(fetch_preview, search_term, uf): (search_term, uf)
            for search_term, uf in partitions
        }
        for future in as_completed(futures):
            previews.append(future.result())

    rows = []
    seen = set()
    source_total = 0
    rate_limited = False
    timed_out = False
    cache_hit = False
    source_urls = []
    for search_term, _uf, url, payload in previews:
        source_urls.append(url)
        try:
            source_total += int(payload.get("total", len(payload.get("items", []))))
        except (TypeError, ValueError):
            source_total += len(payload.get("items", []))
        rate_limited = rate_limited or bool(payload.get("rate_limited"))
        timed_out = timed_out or bool(payload.get("timeout"))
        cache_hit = cache_hit or bool(payload.get("cache_hit"))
        candidates = []
        for row in payload.get("items", []):
            if not date_in_range(row.get("data_fim_vigencia"), data_inicial, data_final):
                continue
            if not row_matches_purchase_filters(row, purchase_number, uasg):
                continue
            if not row_matches_opportunity_type(row, object_type):
                continue
            candidates.append(row)
        for row in filter_opportunity_rows_by_search_term(candidates, search_term):
            row_key = (
                row.get("id")
                or row.get("numero_controle_pncp")
                or row.get("item_url")
            )
            if row_key and row_key in seen:
                continue
            if row_key:
                seen.add(row_key)
            rows.append(row)
    rows.sort(
        key=lambda row: str(
            row.get("data_atualizacao_pncp")
            or row.get("createdAt")
            or row.get("data_publicacao_pncp")
            or ""
        ),
        reverse=True,
    )
    results = [map_search_item(row) for row in rows[:10]]
    return {
        "results": results,
        "total": len(results),
        "source_total": source_total,
        "pagina": 1,
        "tamanhoPagina": 10,
        "total_pages": 1 if results else 0,
        "has_previous": False,
        "has_next": False,
        "source_url": source_urls[0] if source_urls else "",
        "source": "api/search",
        "pages_checked": 1,
        "source_pages": None,
        "rate_limited": rate_limited,
        "timed_out": timed_out,
        "failed_pages": [],
        "complete": False,
        "searching": True,
        "dataInicial": data_inicial,
        "dataFinal": data_final,
        "tipoObjeto": object_type,
        "cache_hit": cache_hit,
    }


def search_pncp_open_bids_fast(params):
    job_key = pncp_search_job_key(params)
    now = time.time()
    with PNCP_SEARCH_JOB_LOCK:
        expired = [
            key for key, job in PNCP_SEARCH_JOBS.items()
            if now - job["created_at"] >= SEARCH_CACHE_TTL
        ]
        for key in expired:
            PNCP_SEARCH_JOBS.pop(key, None)
        job = PNCP_SEARCH_JOBS.get(job_key)
        if job and job["status"] == "complete":
            if job["result"].get("complete"):
                result = search_pncp_open_bids(params)
            else:
                result = copy.deepcopy(job["result"])
            result["searching"] = False
            return result
        if job:
            return copy.deepcopy(job.get("preview") or {
                "results": [],
                "total": 0,
                "pagina": 1,
                "tamanhoPagina": 10,
                "total_pages": 0,
                "complete": False,
                "searching": True,
            })
        PNCP_SEARCH_JOBS[job_key] = {
            "created_at": now,
            "status": "starting",
            "preview": None,
            "result": None,
        }

    try:
        preview = quick_pncp_search_preview(params)
    except Exception:
        with PNCP_SEARCH_JOB_LOCK:
            PNCP_SEARCH_JOBS.pop(job_key, None)
        raise
    with PNCP_SEARCH_JOB_LOCK:
        PNCP_SEARCH_JOBS[job_key]["preview"] = copy.deepcopy(preview)
        PNCP_SEARCH_JOBS[job_key]["status"] = "running"

    def complete_search():
        try:
            full_params = dict(params)
            full_params["pagina"] = "1"
            full_params["tamanhoPagina"] = "10"
            result = search_pncp_open_bids(full_params)
        except Exception as exc:
            result = copy.deepcopy(preview)
            result.update({
                "complete": False,
                "searching": False,
                "error": str(exc) or "Nao foi possivel concluir a busca no PNCP.",
            })
        with PNCP_SEARCH_JOB_LOCK:
            job = PNCP_SEARCH_JOBS.get(job_key)
            if job:
                job["status"] = "complete"
                job["result"] = result

    threading.Thread(target=complete_search, daemon=True).start()
    return preview


PUBLIC_PNCP_FIELDS = (
    "cnpj",
    "ano",
    "sequencial",
    "link",
    "arquivo_usado",
    "documento_usado",
    "documento_tipo",
    "documento_interno",
    "arquivos",
)


def public_pncp_payload(pncp):
    pncp = pncp or {}
    return {
        field: pncp[field]
        for field in PUBLIC_PNCP_FIELDS
        if field in pncp
    }


def json_compatible_default(value):
    if isinstance(value, Path):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, set):
        return sorted(value, key=str)
    raise TypeError(
        f"Object of type {value.__class__.__name__} is not JSON serializable"
    )


def json_response(handler, status, payload, headers=None):
    body = json.dumps(
        payload,
        ensure_ascii=False,
        default=json_compatible_default,
    ).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json; charset=utf-8")
    handler.send_header("Cache-Control", "no-store")
    handler.send_header("Content-Length", str(len(body)))
    for name, value in (headers or {}).items():
        handler.send_header(str(name), str(value))
    handler.end_headers()
    handler.wfile.write(body)


def html_response(handler, body):
    encoded = body.encode("utf-8")
    handler.send_response(200)
    handler.send_header("Content-Type", "text/html; charset=utf-8")
    handler.send_header("Content-Length", str(len(encoded)))
    handler.end_headers()
    handler.wfile.write(encoded)


def frontend_response(handler, request_path="/"):
    index_path = FRONTEND_DIST / "index.html"
    if not index_path.exists():
        html_response(handler, render_page())
        return

    path = urlparse(request_path).path
    if path.startswith("/assets/"):
        candidate = (FRONTEND_DIST / path.lstrip("/")).resolve()
        frontend_root = FRONTEND_DIST.resolve()
        if frontend_root not in candidate.parents or not candidate.is_file():
            handler.send_error(404)
            return
        data = candidate.read_bytes()
        content_type = mimetypes.guess_type(candidate.name)[0] or "application/octet-stream"
        cache_control = "public, max-age=31536000, immutable"
    else:
        data = index_path.read_bytes()
        content_type = "text/html; charset=utf-8"
        cache_control = "no-store, max-age=0, must-revalidate"

    handler.send_response(200)
    handler.send_header("Content-Type", content_type)
    handler.send_header("Cache-Control", cache_control)
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def file_response(handler, path):
    if not path.exists() or path.parent.resolve() != OUTPUT_DIR.resolve():
        handler.send_error(404)
        return
    data = path.read_bytes()
    handler.send_response(200)
    content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
    handler.send_header("Content-Type", content_type)
    handler.send_header("Content-Disposition", f'attachment; filename="{path.name}"')
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def template_response(handler, template_key):
    template_path = template_path_from_name(template_key)
    if not template_path or not template_path.exists():
        handler.send_error(404)
        return
    with TEMPLATE_LOCK:
        data = template_path.read_bytes()
    handler.send_response(200)
    handler.send_header(
        "Content-Type",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    handler.send_header("Content-Disposition", f'attachment; filename="{safe_name(template_path.name)}.docx"')
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def resolve_template(ref):
    if ref.startswith("managed:"):
        template = template_path_from_name(ref.removeprefix("managed:"))
        return template if template and template.exists() else None
    if ref.startswith("builtin:"):
        filename = LEGACY_TEMPLATE_KEYS.get(ref.removeprefix("builtin:"), "")
        template = template_path_from_name(filename)
        return template if template and template.exists() else None
    if ref:
        candidate = UPLOAD_DIR / Path(ref).name
        if candidate.exists() and candidate.parent.resolve() == UPLOAD_DIR.resolve():
            return candidate
    fallback = template_path_from_name(default_template_name())
    return fallback if fallback and fallback.exists() else None


def proposal_generation_context(payload):
    if not isinstance(payload, dict):
        raise ValueError("Os dados da proposta são inválidos.")
    items = payload.get("items") or []
    if (
        not isinstance(items, list)
        or not items
        or not all(isinstance(item, dict) for item in items)
    ):
        raise ValueError("informações não encontradas")

    items = [dict(item, unidade=STANDARD_UNIT) for item in items]
    if any(len(compact(item.get("marca"))) > 120 for item in items):
        raise ValueError("A marca deve possuir no máximo 120 caracteres.")
    for item in items:
        item["marca"] = compact(item.get("marca"))

    template_ref = compact(payload.get("template_ref"))
    template_path = resolve_template(template_ref)
    if template_ref and not template_path:
        raise ValueError("O modelo Word selecionado não está disponível.")

    responsible_id = compact(payload.get("responsible_id"))
    responsible = resolve_responsible(responsible_id) if responsible_id else None
    if not responsible:
        raise ValueError("Selecione um responsável válido pela proposta.")

    invalid_values = [
        item.get("item", "")
        for item in items
        if (
            parse_brazilian_number(item.get("valor_unitario")) is None
            or parse_brazilian_number(item.get("valor_unitario")) < 0
        )
    ]
    if invalid_values:
        raise ValueError(
            "Informe valores unitários válidos e não negativos para todos os itens."
        )

    return {
        "items": items,
        "template_path": template_path,
        "source_name": safe_name(payload.get("source_name", "arquivo")),
        "responsible_id": responsible_id,
        "responsible": responsible,
        "commercial_terms": payload.get("commercial_terms"),
    }


def proposal_preview_fingerprint(context):
    template_path = context.get("template_path")
    template_signature = None
    if template_path and template_path.exists():
        stat = template_path.stat()
        template_signature = {
            "name": template_path.name,
            "size": stat.st_size,
            "modified": stat.st_mtime_ns,
        }
    signature = {
        "items": context.get("items"),
        "template": template_signature,
        "responsible": context.get("responsible"),
        "commercial_terms": context.get("commercial_terms"),
    }
    encoded = json.dumps(
        signature,
        ensure_ascii=False,
        sort_keys=True,
        default=json_compatible_default,
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def cleanup_proposal_previews(now=None):
    ensure_dirs()
    now = time.time() if now is None else now
    with PROPOSAL_PREVIEW_LOCK:
        expired_tokens = [
            token
            for token, record in PROPOSAL_PREVIEW_CACHE.items()
            if now - record.get("last_access", record.get("created_at", 0))
            >= PROPOSAL_PREVIEW_TTL
        ]
        for token in expired_tokens:
            record = PROPOSAL_PREVIEW_CACHE.pop(token, {})
            preview_path = record.get("path")
            if preview_path:
                Path(preview_path).unlink(missing_ok=True)

        active_paths = {
            Path(record["path"]).resolve()
            for record in PROPOSAL_PREVIEW_CACHE.values()
            if record.get("path")
        }
        for path in PREVIEW_DIR.glob("*"):
            if not path.is_file() or path.resolve() in active_paths:
                continue
            try:
                if now - path.stat().st_mtime >= PROPOSAL_PREVIEW_TTL:
                    path.unlink(missing_ok=True)
            except OSError:
                continue


def convert_docx_to_pdf(docx_path, pdf_path):
    if os.name != "nt":
        raise RuntimeError(
            "A pré-visualização fiel requer o Microsoft Word neste servidor."
        )
    powershell = shutil.which("powershell.exe") or shutil.which("powershell")
    if not powershell or not DOCX_TO_PDF_SCRIPT.is_file():
        raise RuntimeError(
            "O conversor da pré-visualização não está disponível neste servidor."
        )

    flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        with WORD_CONVERSION_LOCK:
            result = subprocess.run(
                [
                    powershell,
                    "-NoProfile",
                    "-NonInteractive",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(DOCX_TO_PDF_SCRIPT),
                    "-InputPath",
                    str(docx_path.resolve()),
                    "-OutputPath",
                    str(pdf_path.resolve()),
                ],
                capture_output=True,
                text=True,
                timeout=120,
                creationflags=flags,
                check=False,
            )
    except subprocess.TimeoutExpired as exc:
        raise RuntimeError(
            "A geração da pré-visualização excedeu o limite de 2 minutos."
        ) from exc
    if result.returncode or not pdf_path.is_file() or not pdf_path.stat().st_size:
        detail = compact(result.stderr) or compact(result.stdout)
        raise RuntimeError(
            "Não foi possível converter o Word para a pré-visualização em PDF."
            + (f" Detalhe: {detail[:300]}" if detail else "")
        )


def build_compatible_proposal_pdf(context, pdf_path):
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Title"],
        alignment=TA_CENTER,
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#233254"),
    )
    body_style = ParagraphStyle(
        "ProposalBody",
        parent=styles["BodyText"],
        fontSize=8,
        leading=10,
    )
    document = SimpleDocTemplate(
        str(pdf_path),
        pagesize=landscape(A4),
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
        title="Pré-visualização da proposta",
    )
    responsible = context.get("responsible") or {}
    story = [
        Paragraph("PROPOSTA COMERCIAL", title_style),
        Spacer(1, 4 * mm),
        Paragraph(
            f"<b>Empresa:</b> {html.escape(compact(responsible.get('empresa')) or 'Não informada')}",
            body_style,
        ),
        Paragraph(
            f"<b>Responsável:</b> {html.escape(compact(responsible.get('nome_completo')) or 'Não informado')}",
            body_style,
        ),
        Spacer(1, 4 * mm),
    ]
    rows = [[
        Paragraph("Item", body_style),
        Paragraph("Qtd.", body_style),
        Paragraph("UND", body_style),
        Paragraph("Descrição", body_style),
        Paragraph("Marca", body_style),
        Paragraph("Valor unitário", body_style),
        Paragraph("Valor total", body_style),
    ]]
    for item in context.get("items") or []:
        rows.append([
            Paragraph(html.escape(compact(item.get("item"))), body_style),
            Paragraph(html.escape(compact(item.get("quantidade"))), body_style),
            Paragraph(html.escape(compact(item.get("unidade")) or STANDARD_UNIT), body_style),
            Paragraph(html.escape(compact(item.get("descricao"))), body_style),
            Paragraph(html.escape(compact(item.get("marca"))), body_style),
            Paragraph(html.escape(compact(item.get("valor_unitario"))), body_style),
            Paragraph(html.escape(compact(item.get("valor_total"))), body_style),
        ])
    table = Table(
        rows,
        repeatRows=1,
        colWidths=[14 * mm, 18 * mm, 15 * mm, 100 * mm, 32 * mm, 34 * mm, 34 * mm],
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#233254")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9CA3AF")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F4F6")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([table, Spacer(1, 4 * mm)])
    terms = normalized_commercial_terms(context.get("commercial_terms"))
    for label, field in (
        ("Prazo de entrega", "prazo_entrega"),
        ("Prazo de pagamento", "prazo_pagamento"),
        ("Validade da proposta", "validade_proposta"),
    ):
        story.append(Paragraph(f"<b>{label}:</b> {html.escape(terms[field])}", body_style))
    document.build(story)


def proposal_preview_payload(record, cache_hit):
    expires_at = datetime.fromtimestamp(
        record["last_access"] + PROPOSAL_PREVIEW_TTL
    ).astimezone()
    return {
        "preview_url": f"/proposal-preview/{record['token']}.pdf",
        "expires_at": expires_at.isoformat(timespec="seconds"),
        "cached": cache_hit,
        "renderer": record.get("renderer", "word"),
    }


def create_proposal_preview(context):
    cleanup_proposal_previews()
    fingerprint = proposal_preview_fingerprint(context)
    now = time.time()
    with PROPOSAL_PREVIEW_LOCK:
        cached = next(
            (
                record
                for record in PROPOSAL_PREVIEW_CACHE.values()
                if record.get("fingerprint") == fingerprint
                and Path(record.get("path") or "").is_file()
            ),
            None,
        )
        if cached:
            cached["last_access"] = now
            return proposal_preview_payload(cached, True)

        token = uuid.uuid4().hex
        preview_docx = PREVIEW_DIR / f"{token}.docx"
        preview_pdf = PREVIEW_DIR / f"{token}.pdf"
        renderer = "word"
        try:
            build_docx(
                context["items"],
                context["template_path"],
                preview_docx,
                responsible=context["responsible"],
                commercial_terms=context["commercial_terms"],
            )
            try:
                convert_docx_to_pdf(preview_docx, preview_pdf)
            except RuntimeError as exc:
                LOGGER.warning("Word preview unavailable; using compatible PDF: %s", exc)
                build_compatible_proposal_pdf(context, preview_pdf)
                renderer = "compatible"
        except Exception:
            preview_pdf.unlink(missing_ok=True)
            raise
        finally:
            preview_docx.unlink(missing_ok=True)

        record = {
            "token": token,
            "fingerprint": fingerprint,
            "path": preview_pdf,
            "created_at": now,
            "last_access": now,
            "renderer": renderer,
        }
        PROPOSAL_PREVIEW_CACHE[token] = record
        while len(PROPOSAL_PREVIEW_CACHE) > CACHE_MAX_ENTRIES:
            oldest_token = min(
                PROPOSAL_PREVIEW_CACHE,
                key=lambda key: PROPOSAL_PREVIEW_CACHE[key]["last_access"],
            )
            oldest = PROPOSAL_PREVIEW_CACHE.pop(oldest_token)
            Path(oldest["path"]).unlink(missing_ok=True)
        return proposal_preview_payload(record, False)


def proposal_preview_response(handler, token):
    cleanup_proposal_previews()
    with PROPOSAL_PREVIEW_LOCK:
        record = PROPOSAL_PREVIEW_CACHE.get(token)
        if not record:
            handler.send_error(404, "Pré-visualização expirada")
            return
        path = Path(record["path"])
        if not path.is_file() or path.parent.resolve() != PREVIEW_DIR.resolve():
            PROPOSAL_PREVIEW_CACHE.pop(token, None)
            handler.send_error(404, "Pré-visualização expirada")
            return
        record["last_access"] = time.time()
        data = path.read_bytes()

    handler.send_response(200)
    handler.send_header("Content-Type", "application/pdf")
    handler.send_header(
        "Content-Disposition",
        f'inline; filename="Pre_visualizacao_{token[:8]}.pdf"',
    )
    handler.send_header("Cache-Control", "private, no-store, max-age=0")
    handler.send_header("X-Content-Type-Options", "nosniff")
    handler.send_header("Content-Length", str(len(data)))
    handler.end_headers()
    handler.wfile.write(data)


def render_page():
    templates = list_templates()
    responsibles = list_responsibles()
    selected_template = default_template_name(templates)
    template_options = "".join(
        f'<option value="{html.escape(template["id"])}"'
        f'{" selected" if template["id"] == selected_template else ""}>'
        f'{html.escape(template["display_name"])}</option>'
        for template in templates
    )
    if not template_options:
        template_options = '<option value="" selected>Nenhum template cadastrado</option>'
    responsible_options = "".join(
        f'<option value="{html.escape(responsible["id"])}"'
        f'{" selected" if index == 0 else ""}>'
        f'{html.escape(responsible["nome_completo"])}</option>'
        for index, responsible in enumerate(responsibles)
    )
    if not responsible_options:
        responsible_options = '<option value="" selected>Nenhum responsável cadastrado</option>'
    default_msg = (
        f"{len(templates)} template(s) Word cadastrado(s)."
        if templates
        else "Nenhum template Word cadastrado."
    )
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Gerador de Proposta por Edital</title>
  <style>
    :root {{
      color-scheme: dark;
      --ink: #f1f3f5;
      --panel-ink: #f1f3f5;
      --muted: #a7adb7;
      --panel-muted: #a7adb7;
      --line: #343a40;
      --panel-line: #3f464e;
      --panel: #1f2328;
      --surface: #111315;
      --surface-soft: #191c20;
      --table: #181b1f;
      --table-head: #2b3036;
      --field: #15181c;
      --accent: #6f7782;
      --accent-strong: #858e99;
      --ok: #77d49d;
      --warn: #f0bd6a;
      --danger: #ff858f;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      min-height: 100vh;
      background: var(--surface);
      color: var(--ink);
      font-family: Arial, Helvetica, sans-serif;
      font-size: 15px;
    }}
    main {{
      width: min(1180px, calc(100vw - 32px));
      margin: 0 auto;
      padding: 28px 0 36px;
    }}
    header {{
      display: flex;
      align-items: flex-end;
      justify-content: space-between;
      gap: 20px;
      padding-bottom: 18px;
      border-bottom: 1px solid var(--line);
    }}
    h1 {{
      margin: 0 0 5px;
      font-size: 25px;
      line-height: 1.2;
      letter-spacing: 0;
    }}
    .sub {{
      margin: 0;
      color: var(--muted);
      max-width: 760px;
      line-height: 1.45;
    }}
    .status {{
      min-width: 240px;
      text-align: right;
      color: var(--muted);
      font-size: 13px;
    }}
    .workflow {{
      display: grid;
      grid-template-columns: minmax(320px, 500px) 1fr;
      gap: 18px;
      margin-top: 22px;
      align-items: start;
    }}
    .proposal-panel {{
      display: grid;
      grid-template-rows: 0fr;
      margin-top: 0;
      overflow: hidden;
      opacity: 0;
      transform: translateY(-8px);
      transition:
        grid-template-rows .32s ease,
        opacity .24s ease,
        transform .32s ease,
        margin-top .32s ease;
    }}
    .proposal-panel.is-open {{
      grid-template-rows: 1fr;
      margin-top: 22px;
      opacity: 1;
      transform: translateY(0);
    }}
    .proposal-panel-inner {{
      min-height: 0;
      overflow: hidden;
    }}
    .proposal-panel.is-open .proposal-panel-inner {{
      overflow: visible;
    }}
    .proposal-tools {{
      display: flex;
      justify-content: flex-end;
      margin-bottom: 10px;
    }}
    .proposal-tools button {{
      min-height: 36px;
      padding: 8px 12px;
    }}
    .proposal-panel .workflow {{
      margin-top: 0;
    }}
    .search-panel {{
      margin-top: 18px;
      background: var(--panel);
      border: 1px solid var(--panel-line);
      border-radius: 8px;
      color: var(--panel-ink);
      box-shadow: 0 18px 42px rgba(0, 0, 0, .28);
      overflow: visible;
    }}
    .search-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 14px;
      padding: 14px 16px;
      border-bottom: 1px solid var(--panel-line);
      background: var(--surface-soft);
    }}
    .search-head h2 {{
      margin: 0;
      font-size: 16px;
      letter-spacing: 0;
    }}
    .search-form {{
      grid-template-columns: repeat(6, minmax(120px, 1fr));
      align-items: end;
      box-shadow: none;
      border: 0;
      border-radius: 0;
      border-bottom: 1px solid var(--panel-line);
    }}
    .search-form button {{
      width: 100%;
    }}
    .search-actions {{
      display: grid;
      gap: 8px;
    }}
    .period-field {{
      grid-column: span 2;
      display: grid;
      gap: 7px;
      font-weight: 700;
    }}
    .date-range-picker {{
      position: relative;
    }}
    .date-range-input {{
      width: 100%;
      border: 1px solid #555d66;
      border-radius: 6px;
      padding: 10px 42px 10px 11px;
      background: var(--field);
      color: var(--panel-ink);
      font: inherit;
      font-weight: 400;
      justify-content: flex-start;
      text-align: left;
      position: relative;
    }}
    .date-range-input:hover {{
      background: #1b1f24;
    }}
    .date-range-input:focus {{
      outline: 2px solid rgba(160, 168, 178, .22);
      border-color: var(--accent);
    }}
    .date-range-input .placeholder {{
      color: var(--panel-muted);
    }}
    .calendar-icon {{
      position: absolute;
      right: 12px;
      top: 50%;
      width: 18px;
      height: 18px;
      transform: translateY(-50%);
      border: 2px solid #a7adb7;
      border-radius: 4px;
      pointer-events: none;
    }}
    .calendar-icon::before {{
      content: "";
      position: absolute;
      left: -2px;
      right: -2px;
      top: 4px;
      border-top: 2px solid #a7adb7;
    }}
    .calendar-icon::after {{
      content: "";
      position: absolute;
      left: 3px;
      right: 3px;
      top: -4px;
      height: 4px;
      border-left: 2px solid #a7adb7;
      border-right: 2px solid #a7adb7;
    }}
    .date-range-panel {{
      position: absolute;
      z-index: 80;
      top: calc(100% + 8px);
      left: 0;
      width: min(340px, calc(100vw - 48px));
      border: 1px solid var(--panel-line);
      border-radius: 8px;
      background: #20252b;
      box-shadow: 0 22px 48px rgba(0, 0, 0, .42);
      padding: 12px;
      pointer-events: auto;
    }}
    .date-range-panel[hidden] {{
      display: none;
    }}
    .calendar-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 8px;
      margin-bottom: 10px;
    }}
    .calendar-title {{
      flex: 1;
      min-width: 0;
      font-size: 14px;
      font-weight: 700;
      text-align: center;
      white-space: nowrap;
    }}
    .calendar-nav {{
      width: 34px;
      min-height: 34px;
      padding: 0;
      border-radius: 6px;
      background: #343a40;
      color: #ffffff;
      font-size: 18px;
      line-height: 1;
    }}
    .search-form .calendar-nav {{
      flex: 0 0 34px;
      width: 34px;
    }}
    .calendar-weekdays,
    .calendar-grid {{
      display: grid;
      grid-template-columns: repeat(7, minmax(0, 1fr));
      gap: 4px;
    }}
    .calendar-weekdays {{
      margin-bottom: 5px;
      color: var(--panel-muted);
      font-size: 11px;
      text-align: center;
      text-transform: uppercase;
    }}
    .calendar-day {{
      min-width: 0;
      min-height: 36px;
      aspect-ratio: 1;
      padding: 0;
      border-radius: 6px;
      background: transparent;
      color: var(--panel-ink);
      font-size: 13px;
      font-weight: 700;
    }}
    .calendar-day:hover {{
      background: #3a424b;
    }}
    .calendar-day.outside-month {{
      color: #747b84;
      font-weight: 400;
    }}
    .calendar-day.in-range {{
      background: rgba(133, 142, 153, .28);
    }}
    .calendar-day.range-start,
    .calendar-day.range-end {{
      background: var(--accent-strong);
      color: #ffffff;
    }}
    .calendar-day:disabled {{
      cursor: not-allowed;
      opacity: .34;
      background: transparent;
      color: #777f88;
    }}
    .calendar-help {{
      margin-top: 10px;
      color: var(--panel-muted);
      font-size: 12px;
      line-height: 1.35;
    }}
    form, .result {{
      background: var(--panel);
      border: 1px solid var(--panel-line);
      border-radius: 8px;
      color: var(--panel-ink);
      box-shadow: 0 18px 42px rgba(0, 0, 0, .34);
    }}
    form {{
      padding: 18px;
      display: grid;
      gap: 16px;
    }}
    label {{
      display: grid;
      gap: 7px;
      font-weight: 700;
    }}
    .hint {{
      color: var(--panel-muted);
      font-size: 13px;
      font-weight: 400;
      line-height: 1.35;
    }}
    input[type="file"] {{
      width: 100%;
      border: 1px dashed #5b636d;
      background: var(--field);
      border-radius: 6px;
      padding: 12px;
      color: var(--panel-ink);
    }}
    input[type="url"], input[type="date"], select {{
      width: 100%;
      border: 1px solid #555d66;
      border-radius: 6px;
      padding: 10px 11px;
      background: var(--field);
      color: var(--panel-ink);
      font: inherit;
    }}
    form > label input[type="text"] {{
      width: 100%;
      border: 1px solid #555d66;
      border-radius: 6px;
      padding: 10px 11px;
      background: var(--field);
      color: var(--panel-ink);
      font: inherit;
    }}
    button, .download {{
      appearance: none;
      border: 0;
      border-radius: 6px;
      background: var(--accent);
      color: #ffffff;
      font-weight: 700;
      padding: 12px 14px;
      cursor: pointer;
      text-decoration: none;
      display: inline-flex;
      align-items: center;
      justify-content: center;
      min-height: 42px;
    }}
    button:hover, .download:hover {{ background: var(--accent-strong); }}
    button:disabled {{
      cursor: wait;
      opacity: .72;
    }}
    .result {{
      min-height: 470px;
      overflow: hidden;
    }}
    .result-head {{
      display: flex;
      justify-content: space-between;
      align-items: center;
      gap: 16px;
      padding: 14px 16px;
      border-bottom: 1px solid var(--panel-line);
      background: var(--surface-soft);
    }}
    .result-title {{
      font-weight: 700;
    }}
    .message {{
      padding: 14px 16px;
      color: var(--panel-muted);
      line-height: 1.45;
    }}
    .message.ok {{ color: var(--ok); }}
    .message.warn {{ color: var(--warn); }}
    .message.danger {{ color: var(--danger); }}
    .table-wrap {{
      overflow: auto;
      max-height: 620px;
    }}
    .identification-panel {{
      border: 1px solid var(--panel-line);
      border-radius: 6px;
      background: var(--surface-soft);
      overflow: hidden;
    }}
    .identification-head {{
      display: flex;
      justify-content: space-between;
      gap: 12px;
      align-items: center;
      padding: 10px 12px;
      border-bottom: 1px solid var(--panel-line);
    }}
    .identification-title {{
      font-weight: 700;
    }}
    .identification-status {{
      color: var(--panel-muted);
      font-size: 13px;
      text-align: right;
    }}
    .identification-status.ok {{ color: var(--ok); }}
    .identification-status.warn {{ color: var(--warn); }}
    .identification-status.danger {{ color: var(--danger); }}
    .identification-table {{
      overflow: auto;
      max-height: 260px;
    }}
    .identification-table table {{
      min-width: 760px;
    }}
    .identification-table input[type="checkbox"] {{
      width: 18px;
      height: 18px;
      accent-color: var(--accent);
    }}
    .identification-table .identify-value {{
      min-width: 115px;
    }}
    .search-results {{
      padding: 0;
      overflow: auto;
      max-height: 390px;
    }}
    .search-results table {{
      min-width: 980px;
    }}
    .link-button {{
      min-height: 34px;
      padding: 8px 10px;
      font-size: 13px;
      white-space: nowrap;
    }}
    table {{
      border-collapse: collapse;
      width: 100%;
      min-width: 1120px;
      background: var(--table);
      color: var(--panel-ink);
    }}
    th, td {{
      border: 1px solid var(--panel-line);
      padding: 8px 9px;
      vertical-align: top;
      text-align: left;
      line-height: 1.3;
    }}
    th {{
      position: sticky;
      top: 0;
      z-index: 1;
      background: var(--table-head);
      color: #ffffff;
      font-size: 13px;
      text-transform: uppercase;
    }}
    td.short, th.short {{ width: 72px; text-align: center; }}
    td.money, th.money {{ width: 130px; text-align: center; }}
    td.desc {{ min-width: 430px; }}
    .actions {{
      display: flex;
      gap: 10px;
      align-items: center;
      flex-wrap: wrap;
    }}
    .bulk-fields {{
      display: flex;
      gap: 8px;
      align-items: center;
      flex-wrap: wrap;
    }}
    .bulk-field {{
      display: grid;
      gap: 4px;
    }}
    .bulk-field label {{
      display: block;
      font-size: 12px;
      color: var(--panel-muted);
      font-weight: 700;
    }}
    .bulk-field input,
    .bulk-field select {{
      width: 210px;
      border: 1px solid #555d66;
      border-radius: 4px;
      padding: 9px 10px;
      font: inherit;
      background: var(--field);
      color: var(--panel-ink);
    }}
    .bulk-field input.money-input {{
      width: 150px;
    }}
    .secondary {{
      background: #343a40;
      color: #ffffff;
    }}
    .secondary:hover {{
      background: #454c54;
    }}
    .template-edit {{
      display: inline-flex;
      align-items: center;
      justify-content: center;
      width: fit-content;
      min-height: 34px;
      margin-top: 2px;
      padding: 7px 10px;
      border: 1px solid #555d66;
      border-radius: 4px;
      background: #343a40;
      color: #ffffff;
      font-size: 12px;
      font-weight: 700;
      text-decoration: none;
    }}
    .template-edit:hover {{
      background: #454c54;
    }}
    .template-links {{
      display: flex;
      gap: 8px;
      align-items: center;
      flex-wrap: wrap;
    }}
    .template-manager-overlay {{
      position: fixed;
      inset: 0;
      z-index: 1000;
      display: grid;
      place-items: center;
      padding: 20px;
      background: rgba(0, 0, 0, .72);
    }}
    .template-manager-overlay[hidden] {{ display: none; }}
    .template-manager-dialog {{
      width: min(760px, 100%);
      max-height: min(780px, calc(100vh - 40px));
      overflow: hidden;
      border: 1px solid #f4a45d;
      border-radius: 6px;
      background: #ee8627;
      color: #102844;
      box-shadow: 0 22px 60px rgba(0, 0, 0, .45);
    }}
    .template-manager-head {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      padding: 20px 22px;
    }}
    .template-manager-head h2 {{
      margin: 0;
      color: #ffffff;
      font-size: 28px;
      letter-spacing: 0;
    }}
    .template-manager-close {{
      width: 38px;
      min-width: 38px;
      height: 38px;
      padding: 0;
      border: 1px solid rgba(255, 255, 255, .65);
      border-radius: 4px;
      background: transparent;
      color: #ffffff;
      font-size: 26px;
      line-height: 1;
    }}
    .template-manager-body {{
      max-height: calc(min(780px, 100vh - 40px) - 78px);
      overflow-y: auto;
      margin: 0 12px 12px;
      border-radius: 5px;
      background: #ffffff;
      color: #102844;
    }}
    .template-manager-toolbar {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      padding: 18px;
      border-bottom: 1px solid #d9dde2;
    }}
    .template-manager-toolbar strong {{ font-size: 17px; }}
    .template-manager-button {{
      min-height: 40px;
      border: 0;
      border-radius: 4px;
      padding: 10px 15px;
      background: #102844;
      color: #ffffff;
      font-size: 14px;
      font-weight: 700;
    }}
    .template-manager-button:hover {{ background: #1a3c63; }}
    .template-manager-button.delete {{ background: #cc3438; }}
    .template-manager-button.delete:hover {{ background: #ae292d; }}
    .template-manager-button:disabled,
    .template-manager-close:disabled {{ cursor: wait; opacity: .58; }}
    .template-manager-status {{
      min-height: 22px;
      margin: 0;
      padding: 0 18px;
      color: #56616d;
      font-size: 13px;
      line-height: 22px;
    }}
    .template-manager-status.success {{ color: #176b3a; }}
    .template-manager-status.error {{ color: #b3262a; }}
    .template-manager-list {{ border-top: 1px solid #e1e4e8; }}
    .template-manager-row {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      align-items: center;
      gap: 18px;
      padding: 18px;
      border-bottom: 1px solid #e1e4e8;
    }}
    .template-manager-row:last-child {{ border-bottom: 0; }}
    .template-manager-name {{
      display: block;
      overflow-wrap: anywhere;
      color: #102844;
      font-size: 15px;
    }}
    .template-manager-meta {{
      display: block;
      margin-top: 4px;
      color: #68737e;
      font-size: 12px;
    }}
    .template-manager-actions {{ display: flex; gap: 8px; }}
    .template-manager-empty {{
      padding: 42px 18px;
      color: #68737e;
      text-align: center;
    }}
    .responsible-manager-dialog {{ width: min(880px, 100%); }}
    .responsible-manager-form {{
      padding: 18px;
      border-bottom: 1px solid #d9dde2;
      background: #f7f8fa;
    }}
    .responsible-manager-form[hidden] {{ display: none; }}
    .responsible-manager-form h3 {{
      margin: 0 0 16px;
      color: #102844;
      font-size: 18px;
    }}
    .responsible-form-grid {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 13px 16px;
    }}
    .responsible-form-field {{ display: grid; gap: 5px; }}
    .responsible-form-field.full {{ grid-column: 1 / -1; }}
    .responsible-form-field label {{
      color: #26394f;
      font-size: 12px;
      font-weight: 700;
    }}
    .responsible-form-field input,
    .responsible-form-field textarea {{
      width: 100%;
      border: 1px solid #aeb7c1;
      border-radius: 4px;
      padding: 9px 10px;
      background: #ffffff;
      color: #102844;
      font: inherit;
    }}
    .responsible-form-field textarea {{ min-height: 78px; resize: vertical; }}
    .responsible-form-actions {{
      display: flex;
      justify-content: flex-end;
      gap: 8px;
      margin-top: 16px;
    }}
    .template-manager-button.cancel {{ background: #59636e; }}
    .responsible-manager-row {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      align-items: center;
      gap: 18px;
      padding: 18px;
      border-bottom: 1px solid #e1e4e8;
    }}
    .responsible-manager-row:last-child {{ border-bottom: 0; }}
    .responsible-manager-row.is-selected {{ background: #f0f5fa; }}
    .responsible-manager-details {{ min-width: 0; }}
    .responsible-manager-actions {{ display: flex; gap: 7px; }}
    .app-shell {{
      display: grid;
      grid-template-columns: 220px minmax(0, 1fr);
      min-height: 100vh;
      background: #071521;
    }}
    .app-sidebar {{
      position: sticky;
      top: 0;
      z-index: 90;
      display: flex;
      flex-direction: column;
      min-height: 100vh;
      padding: 30px 18px 22px;
      background: #ef7f1a;
      color: #ffffff;
      box-shadow: 8px 0 26px rgba(0, 0, 0, .2);
    }}
    .app-brand {{
      padding: 0 10px 28px;
      font-size: 30px;
      font-weight: 800;
      letter-spacing: 0;
    }}
    .block-navigation {{ display: grid; gap: 8px; }}
    .block-nav-item {{
      display: grid;
      grid-template-columns: 34px minmax(0, 1fr);
      justify-content: stretch;
      gap: 10px;
      width: 100%;
      min-height: 62px;
      padding: 10px;
      border: 1px solid transparent;
      border-radius: 6px;
      background: transparent;
      color: rgba(255, 255, 255, .82);
      text-align: left;
    }}
    .block-nav-item:hover {{ background: rgba(9, 30, 48, .16); }}
    .block-nav-item.is-active {{
      border-color: rgba(255, 255, 255, .24);
      background: #102b42;
      color: #ffffff;
    }}
    .block-nav-number {{
      display: grid;
      place-items: center;
      width: 34px;
      height: 34px;
      border: 1px solid currentColor;
      border-radius: 50%;
      font-size: 11px;
      font-weight: 800;
    }}
    .block-nav-item strong,
    .block-nav-item small {{ display: block; letter-spacing: 0; }}
    .block-nav-item strong {{ margin-bottom: 3px; font-size: 14px; }}
    .block-nav-item small {{ color: inherit; font-size: 11px; font-weight: 400; }}
    .sidebar-context {{
      margin-top: auto;
      padding: 18px 10px 0;
      border-top: 1px solid rgba(255, 255, 255, .3);
      color: rgba(255, 255, 255, .78);
      font-size: 12px;
    }}
    .app-content {{ min-width: 0; background: #071521; }}
    .app-content main {{
      width: min(1360px, calc(100vw - 252px));
      padding: 30px 0 42px;
    }}
    .app-content header {{
      display: grid;
      grid-template-columns: minmax(180px, 1fr) minmax(320px, 680px) minmax(180px, 1fr);
      align-items: center;
      padding-bottom: 22px;
      border-color: #234055;
    }}
    .header-copy {{ grid-column: 2; text-align: center; }}
    .header-kicker,
    .section-kicker {{
      color: #ef8a2f;
      font-size: 11px;
      font-weight: 800;
      letter-spacing: 0;
      text-transform: uppercase;
    }}
    .app-content header h1 {{ margin: 5px 0 7px; font-size: 30px; }}
    .app-content header .sub {{ max-width: none; color: #9db0bf; }}
    .app-content header > .status {{ grid-column: 3; justify-self: end; color: #8fa3b2; }}
    .search-panel {{
      border-color: #29455a;
      background: #0c2031;
      box-shadow: 0 18px 42px rgba(0, 0, 0, .2);
    }}
    .search-head,
    .result-head {{ background: #10283b; border-color: #29455a; }}
    .proposal-page-head {{
      display: flex;
      align-items: flex-start;
      justify-content: space-between;
      gap: 18px;
      margin-bottom: 14px;
      padding: 0 2px 14px;
      border-bottom: 1px solid #29455a;
    }}
    .proposal-page-head h2 {{ margin: 4px 0 5px; font-size: 22px; letter-spacing: 0; }}
    .proposal-page-head p {{ margin: 0; color: #94a9b8; font-size: 13px; }}
    .compact-button {{ min-height: 36px; padding: 8px 12px; }}
    .proposal-panel .workflow {{ display: grid; grid-template-columns: 1fr; gap: 18px; }}
    #uploadForm {{
      gap: 18px;
      padding: 20px;
      border-color: #29495f;
      background: #0c2031;
      box-shadow: 0 18px 40px rgba(0, 0, 0, .22);
    }}
    .proposal-url-field {{ color: #f3f6f8; }}
    .proposal-url-field .hint {{ color: #8fa5b5; }}
    .proposal-fields {{
      display: grid;
      grid-template-columns: minmax(160px, .8fr) minmax(250px, 1.1fr) minmax(250px, 1.1fr);
      gap: 14px;
      align-items: end;
    }}
    .proposal-fields .bulk-field {{ display: grid; gap: 7px; min-width: 0; }}
    .proposal-fields .bulk-field > label {{ color: #c9d4dc; font-size: 12px; }}
    .proposal-fields .bulk-field input,
    .proposal-fields .bulk-field select {{
      width: 100%;
      min-height: 42px;
      border-color: #3c5c72;
      background: #081925;
      color: #f3f6f8;
    }}
    .field-with-action {{
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      gap: 7px;
    }}
    .field-action-button {{
      min-height: 42px;
      padding: 9px 12px;
      border: 1px solid #ef8a2f;
      border-radius: 5px;
      background: transparent;
      color: #ffad62;
      font-size: 12px;
    }}
    .field-action-button:hover {{ background: #ef7f1a; color: #ffffff; }}
    .template-utilities {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 12px;
      padding: 12px;
      border: 1px solid #29495f;
      border-radius: 5px;
      background: #091b29;
    }}
    .custom-template-field {{
      display: flex;
      align-items: center;
      justify-content: flex-end;
      gap: 10px;
      color: #9eb0bd;
      font-size: 12px;
      font-weight: 600;
    }}
    .custom-template-field input[type="file"] {{
      width: min(360px, 46vw);
      padding: 8px;
      border-style: solid;
      border-color: #38586e;
      background: #071722;
      font-size: 12px;
    }}
    .identification-panel {{ border-color: #31526a; background: #091b29; }}
    .identification-head {{ padding: 13px 14px; border-color: #31526a; background: #10283b; }}
    .identification-caption,
    .result-caption {{ margin-top: 3px; color: #8ea3b2; font-size: 11px; }}
    .identification-table {{ max-height: 420px; }}
    .identification-table table {{ min-width: 980px; background: #0a1c29; }}
    .identification-table th {{ background: #17364d; }}
    .identification-table tbody tr:has([data-identify-select]:checked) {{ background: rgba(239, 127, 26, .09); }}
    .identification-table input[type="checkbox"] {{ accent-color: #ef7f1a; }}
    .proposal-submit-row {{
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 18px;
      padding-top: 2px;
    }}
    #submitBtn {{
      min-width: 280px;
      background: #ef7f1a;
      color: #ffffff;
    }}
    #submitBtn:hover {{ background: #ff9638; }}
    .result {{
      min-height: 230px;
      border-color: #29495f;
      background: #0c2031;
      box-shadow: 0 18px 40px rgba(0, 0, 0, .18);
    }}
    .result-head {{ padding: 14px 18px; }}
    .download {{ background: #ef7f1a; }}
    .download:hover {{ background: #ff9638; }}
    .nav-focused {{ animation: navFocus .55s ease; }}
    @keyframes navFocus {{
      0% {{ box-shadow: 0 0 0 0 rgba(239, 127, 26, .5); }}
      100% {{ box-shadow: 0 0 0 12px rgba(239, 127, 26, 0); }}
    }}
    .empty {{
      color: #8e959e;
      font-style: italic;
    }}
    .cell-input {{
      width: 100%;
      min-width: 110px;
      border: 1px solid #555d66;
      border-radius: 4px;
      padding: 7px 8px;
      font: inherit;
      background: var(--field);
      color: var(--panel-ink);
    }}
    .cell-input:focus {{
      outline: 2px solid rgba(160, 168, 178, .22);
      border-color: var(--accent);
    }}
    .money-input {{
      text-align: center;
    }}
    @media (max-width: 860px) {{
      .app-shell {{ display: block; }}
      .app-sidebar {{
        min-height: auto;
        padding: 10px 12px;
        flex-direction: row;
        align-items: center;
        gap: 12px;
      }}
      .app-brand {{ padding: 0; font-size: 22px; }}
      .block-navigation {{ flex: 1; grid-template-columns: 1fr 1fr; gap: 6px; }}
      .block-nav-item {{ min-height: 50px; padding: 7px; grid-template-columns: 28px minmax(0, 1fr); gap: 7px; }}
      .block-nav-number {{ width: 28px; height: 28px; }}
      .block-nav-item small {{ display: none; }}
      .sidebar-context {{ display: none; }}
      .app-content main {{ width: min(100vw - 20px, 1180px); padding-top: 18px; }}
      .proposal-panel {{ scroll-margin-top: 82px; }}
      .app-content header {{ display: block; }}
      .header-copy {{ text-align: left; }}
      .app-content header > .status {{ text-align: left; margin-top: 10px; }}
      .workflow {{ grid-template-columns: 1fr; }}
      .search-form {{ grid-template-columns: 1fr; }}
      .period-field {{ grid-column: auto; }}
      .date-range-panel {{ width: 100%; }}
      h1 {{ font-size: 22px; }}
      .proposal-page-head {{ align-items: center; }}
      .proposal-fields {{ grid-template-columns: 1fr; }}
      .template-utilities {{ align-items: stretch; flex-direction: column; }}
      .custom-template-field {{ align-items: stretch; flex-direction: column; }}
      .custom-template-field input[type="file"] {{ width: 100%; }}
      .proposal-submit-row {{ align-items: stretch; flex-direction: column; }}
      #submitBtn {{ width: 100%; min-width: 0; }}
      .template-manager-overlay {{ padding: 10px; }}
      .template-manager-dialog {{ max-height: calc(100vh - 20px); }}
      .template-manager-body {{ max-height: calc(100vh - 94px); }}
      .template-manager-toolbar,
      .template-manager-row {{ align-items: stretch; grid-template-columns: 1fr; flex-direction: column; }}
      .template-manager-toolbar {{ display: flex; }}
      .template-manager-toolbar .template-manager-button {{ width: 100%; }}
      .template-manager-actions {{ display: grid; grid-template-columns: 1fr 1fr; }}
      .template-manager-actions .template-manager-button {{ width: 100%; }}
      .responsible-form-grid {{ grid-template-columns: 1fr; }}
      .responsible-form-field.full {{ grid-column: auto; }}
      .responsible-manager-row {{ grid-template-columns: 1fr; align-items: stretch; }}
      .responsible-manager-actions {{ display: grid; grid-template-columns: repeat(3, 1fr); }}
      .responsible-manager-actions .template-manager-button {{ width: 100%; padding-inline: 8px; }}
    }}
    @media (max-width: 460px) {{
      .responsible-manager-actions {{ grid-template-columns: 1fr; }}
      .responsible-form-actions {{ display: grid; grid-template-columns: 1fr; }}
      .responsible-form-actions .template-manager-button {{ width: 100%; }}
    }}
  </style>
</head>
<body>
  <div class="app-shell">
    <aside class="app-sidebar" aria-label="Navegação principal">
      <div class="app-brand">TOTH</div>
      <nav class="block-navigation" aria-label="Etapas da proposta">
        <button id="navBlock1" class="block-nav-item is-active" type="button">
          <span class="block-nav-number">01</span>
          <span><strong>Bloco 1</strong><small>Consulta PNCP</small></span>
        </button>
        <button id="navBlock2" class="block-nav-item" type="button">
          <span class="block-nav-number">02</span>
          <span><strong>Bloco 2</strong><small>Gerar proposta</small></span>
        </button>
      </nav>
      <div class="sidebar-context">Licitações públicas</div>
    </aside>
    <div class="app-content">
  <main>
    <header>
      <div class="header-copy">
        <div class="header-kicker">Propostas comerciais</div>
        <h1>Gerar proposta</h1>
        <p class="sub">Consulte o edital, selecione os itens e gere o documento Word com os dados da proposta.</p>
      </div>
      <div class="status">{html.escape(default_msg)}</div>
    </header>

    <section id="block1" class="search-panel">
      <div class="search-head">
        <h2>Consulta PNCP - Filtro oficial de editais</h2>
        <div id="searchStatus" class="status">Usa o filtro oficial de editais do PNCP</div>
      </div>
      <form id="pncpSearchForm" class="search-form">
        <div class="period-field">
          <label id="periodLabel" for="dateRangeButton">Período</label>
          <div id="dateRangePicker" class="date-range-picker">
            <button id="dateRangeButton" class="date-range-input" type="button" aria-haspopup="dialog" aria-expanded="false" aria-labelledby="periodLabel dateRangeText" aria-describedby="dateRangeHint">
              <span id="dateRangeText" class="placeholder">Selecione o período</span>
              <span class="calendar-icon" aria-hidden="true"></span>
            </button>
            <input id="searchStartDate" name="dataInicial" type="hidden">
            <input id="searchEndDate" name="dataFinal" type="hidden">
            <div id="dateRangePanel" class="date-range-panel" role="dialog" aria-label="Selecionar período" hidden>
              <div class="calendar-head">
                <button id="calendarPrev" class="calendar-nav" type="button" aria-label="Mês anterior">&lt;</button>
                <div id="calendarTitle" class="calendar-title"></div>
                <button id="calendarNext" class="calendar-nav" type="button" aria-label="Próximo mês">&gt;</button>
              </div>
              <div class="calendar-weekdays" aria-hidden="true">
                <span>Dom</span><span>Seg</span><span>Ter</span><span>Qua</span><span>Qui</span><span>Sex</span><span>Sáb</span>
              </div>
              <div id="calendarGrid" class="calendar-grid"></div>
              <div id="calendarHelp" class="calendar-help">Escolha a data inicial.</div>
            </div>
          </div>
          <span id="dateRangeHint" class="hint">Selecione um período de até 30 dias.</span>
        </div>
        <label>
          UF
          <input id="searchUf" name="uf" type="text" maxlength="2" placeholder="Ex.: SP">
        </label>
        <label>
          Palavra-chave
          <span class="hint">Opcional. Com palavra-chave a busca oficial fica mais rápida e precisa.</span>
          <input id="searchKeyword" name="palavraChave" type="text" placeholder="Ex.: mobiliário">
        </label>
        <label>
          Tipo do objeto
          <select id="searchObjectType" name="tipoObjeto">
            <option value="" selected>Materiais e serviços</option>
            <option value="material">Materiais</option>
            <option value="servico">Serviços</option>
          </select>
        </label>
        <label>
          Modalidade
          <select id="searchModality" name="codigoModalidadeContratacao">
            <option value="6" selected>Pregão eletrônico</option>
            <option value="8">Dispensa eletrônica</option>
            <option value="">Todas</option>
          </select>
        </label>
        <div class="search-actions">
          <button id="searchBtn" type="submit">Buscar contratações</button>
          <button id="showProposalBtn" class="secondary" type="button" aria-controls="proposalPanel" aria-expanded="false">Gerar proposta</button>
          <button id="clearSearchBtn" class="secondary" type="button">Limpar filtros</button>
        </div>
      </form>
      <div id="searchResults" class="search-results">
        <div class="message">Busque no PNCP para copiar ou usar automaticamente um link de edital.</div>
      </div>
    </section>

    <section id="proposalPanel" class="proposal-panel" aria-hidden="true" inert>
      <div class="proposal-panel-inner">
        <div class="proposal-page-head">
          <div>
            <span class="section-kicker">Bloco 2</span>
            <h2>Gerar proposta</h2>
            <p>Informe os dados comerciais e selecione os itens que farão parte do documento.</p>
          </div>
          <button id="hideProposalBtn" class="secondary compact-button" type="button">Ocultar</button>
        </div>
        <section class="workflow">
          <form id="uploadForm">
            <label id="pncpSourcePanel" class="source-panel proposal-url-field">
              URL do edital no PNCP
              <span class="hint">Cole o link público do edital. A consulta e a extração dos arquivos serão feitas pelo backend.</span>
              <input id="pncpLink" name="pncp_link" type="url" maxlength="500" autocomplete="url" placeholder="https://pncp.gov.br/app/editais/45780087000103/2026/43">
            </label>

            <div class="proposal-fields">
              <div class="bulk-field">
                <label for="presetBrand">Marca</label>
                <input id="presetBrand" name="preset_brand" type="text" maxlength="120" autocomplete="off" placeholder="Ex.: Goldflex">
              </div>
              <div class="bulk-field">
                <label for="responsibleId">Responsável</label>
                <div class="field-with-action">
                  <select id="responsibleId" name="responsible_id">
                    {responsible_options}
                  </select>
                  <button id="openResponsibleManager" class="field-action-button" type="button" title="Selecionar ou editar responsáveis">Gerenciar</button>
                </div>
              </div>
              <div class="bulk-field">
                <label for="templateChoice">Template</label>
                <div class="field-with-action">
                  <select id="templateChoice" name="template_choice">
                    {template_options}
                  </select>
                  <button id="openTemplateManager" class="field-action-button" type="button" title="Selecionar ou gerenciar templates">Gerenciar</button>
                </div>
              </div>
            </div>

            <div class="template-utilities">
              <a id="editTemplateLink" class="template-edit" href="#" target="_blank" rel="noopener">Abrir template selecionado</a>
              <label class="custom-template-field" for="templateFile">
                <span>Usar outro arquivo nesta proposta</span>
                <input id="templateFile" name="template_file" type="file" accept=".docx">
              </label>
            </div>

            <input id="wantedItems" name="wanted_items" type="hidden">
            <section class="identification-panel" aria-live="polite">
              <div class="identification-head">
                <div>
                  <div class="identification-title">Itens encontrados no edital</div>
                  <div class="identification-caption">Selecione os itens e informe o valor unitário.</div>
                </div>
                <div id="itemIdentificationStatus" class="identification-status">Aguardando link PNCP.</div>
              </div>
              <div id="itemIdentificationPreview" class="identification-table">
                <div class="message">Informe o link para identificar o que cada item representa.</div>
              </div>
            </section>

            <div class="proposal-submit-row">
              <span class="hint">Somente os itens selecionados serão incluídos na proposta.</span>
              <button id="submitBtn" type="submit">Extrair tabelas e aplicar valores</button>
            </div>
          </form>

          <section class="result" aria-live="polite">
            <div class="result-head">
              <div>
                <div class="result-title">Resultado da proposta</div>
                <div class="result-caption">Pré-visualização, validações e documento gerado</div>
              </div>
              <div id="actions" class="actions"></div>
            </div>
            <div id="message" class="message">Aguardando processamento.</div>
            <div id="preview" class="table-wrap"></div>
          </section>
        </section>
      </div>
    </section>
  </main>
    </div>
  </div>

  <div id="templateManagerOverlay" class="template-manager-overlay" hidden>
    <section class="template-manager-dialog" role="dialog" aria-modal="true" aria-labelledby="templateManagerTitle">
      <div class="template-manager-head">
        <h2 id="templateManagerTitle">TEMPLATES</h2>
        <button id="closeTemplateManager" class="template-manager-close" type="button" aria-label="Fechar gerenciamento de templates">&times;</button>
      </div>
      <div class="template-manager-body">
        <div class="template-manager-toolbar">
          <strong>Arquivos cadastrados</strong>
          <button id="attachTemplateButton" class="template-manager-button" type="button">Anexar novo template</button>
        </div>
        <p id="templateManagerStatus" class="template-manager-status" role="status" aria-live="polite"></p>
        <div id="templateManagerList" class="template-manager-list" aria-live="polite"></div>
      </div>
      <input id="templateManagerFile" type="file" accept=".docx,application/vnd.openxmlformats-officedocument.wordprocessingml.document" hidden>
    </section>
  </div>

  <div id="responsibleManagerOverlay" class="template-manager-overlay" hidden>
    <section class="template-manager-dialog responsible-manager-dialog" role="dialog" aria-modal="true" aria-labelledby="responsibleManagerTitle">
      <div class="template-manager-head">
        <h2 id="responsibleManagerTitle">RESPONSÁVEIS</h2>
        <button id="closeResponsibleManager" class="template-manager-close" type="button" aria-label="Fechar gerenciamento de responsáveis">&times;</button>
      </div>
      <div class="template-manager-body">
        <div class="template-manager-toolbar">
          <strong>Responsáveis cadastrados</strong>
          <button id="newResponsibleButton" class="template-manager-button" type="button">Cadastrar responsável</button>
        </div>
        <p id="responsibleManagerStatus" class="template-manager-status" role="status" aria-live="polite"></p>
        <form id="responsibleManagerForm" class="responsible-manager-form" hidden>
          <h3 id="responsibleFormTitle">Novo responsável</h3>
          <div class="responsible-form-grid">
            <div class="responsible-form-field full">
              <label for="responsibleName">Nome completo *</label>
              <input id="responsibleName" name="nome_completo" maxlength="200" required autocomplete="name">
            </div>
            <div class="responsible-form-field full">
              <label for="responsibleCompany">Empresa *</label>
              <input id="responsibleCompany" name="empresa" maxlength="200" required autocomplete="organization">
            </div>
            <div class="responsible-form-field">
              <label for="responsibleCnpj">CNPJ *</label>
              <input id="responsibleCnpj" name="cnpj" maxlength="24" required inputmode="numeric" placeholder="00.000.000/0000-00">
            </div>
            <div class="responsible-form-field">
              <label for="responsibleRg">RG</label>
              <input id="responsibleRg" name="rg" maxlength="30">
            </div>
            <div class="responsible-form-field">
              <label for="responsibleCpf">CPF</label>
              <input id="responsibleCpf" name="cpf" maxlength="20" inputmode="numeric" placeholder="000.000.000-00">
            </div>
            <div class="responsible-form-field full">
              <label for="responsibleNotes">Observações</label>
              <textarea id="responsibleNotes" name="observacoes" maxlength="1000"></textarea>
            </div>
          </div>
          <div class="responsible-form-actions">
            <button id="cancelResponsibleButton" class="template-manager-button cancel" type="button">Cancelar</button>
            <button id="saveResponsibleButton" class="template-manager-button" type="submit">Salvar responsável</button>
          </div>
        </form>
        <div id="responsibleManagerList" class="template-manager-list" aria-live="polite"></div>
      </div>
    </section>
  </div>

  <script>
    const form = document.getElementById('uploadForm');
    const pncpSearchForm = document.getElementById('pncpSearchForm');
    const btn = document.getElementById('submitBtn');
    const searchBtn = document.getElementById('searchBtn');
    const showProposalBtn = document.getElementById('showProposalBtn');
    const hideProposalBtn = document.getElementById('hideProposalBtn');
    const proposalPanel = document.getElementById('proposalPanel');
    const navBlock1 = document.getElementById('navBlock1');
    const navBlock2 = document.getElementById('navBlock2');
    const block1 = document.getElementById('block1');
    const clearSearchBtn = document.getElementById('clearSearchBtn');
    const message = document.getElementById('message');
    const preview = document.getElementById('preview');
    const actions = document.getElementById('actions');
    const searchResults = document.getElementById('searchResults');
    const searchStatus = document.getElementById('searchStatus');
    const searchStartDate = document.getElementById('searchStartDate');
    const searchEndDate = document.getElementById('searchEndDate');
    const dateRangePicker = document.getElementById('dateRangePicker');
    const dateRangeButton = document.getElementById('dateRangeButton');
    const dateRangeText = document.getElementById('dateRangeText');
    const dateRangePanel = document.getElementById('dateRangePanel');
    const calendarTitle = document.getElementById('calendarTitle');
    const calendarGrid = document.getElementById('calendarGrid');
    const calendarHelp = document.getElementById('calendarHelp');
    const calendarPrev = document.getElementById('calendarPrev');
    const calendarNext = document.getElementById('calendarNext');
    const searchUf = document.getElementById('searchUf');
    const searchKeyword = document.getElementById('searchKeyword');
    const searchObjectType = document.getElementById('searchObjectType');
    const searchModality = document.getElementById('searchModality');
    const pncpLink = document.getElementById('pncpLink');
    const itemIdentificationStatus = document.getElementById('itemIdentificationStatus');
    const itemIdentificationPreview = document.getElementById('itemIdentificationPreview');
    const presetBrand = document.getElementById('presetBrand');
    const responsibleId = document.getElementById('responsibleId');
    const templateChoice = document.getElementById('templateChoice');
    const templateFile = document.getElementById('templateFile');
    const editTemplateLink = document.getElementById('editTemplateLink');
    const openTemplateManager = document.getElementById('openTemplateManager');
    const templateManagerOverlay = document.getElementById('templateManagerOverlay');
    const closeTemplateManager = document.getElementById('closeTemplateManager');
    const attachTemplateButton = document.getElementById('attachTemplateButton');
    const templateManagerFile = document.getElementById('templateManagerFile');
    const templateManagerList = document.getElementById('templateManagerList');
    const templateManagerStatus = document.getElementById('templateManagerStatus');
    const openResponsibleManager = document.getElementById('openResponsibleManager');
    const responsibleManagerOverlay = document.getElementById('responsibleManagerOverlay');
    const closeResponsibleManager = document.getElementById('closeResponsibleManager');
    const newResponsibleButton = document.getElementById('newResponsibleButton');
    const responsibleManagerStatus = document.getElementById('responsibleManagerStatus');
    const responsibleManagerList = document.getElementById('responsibleManagerList');
    const responsibleManagerForm = document.getElementById('responsibleManagerForm');
    const responsibleFormTitle = document.getElementById('responsibleFormTitle');
    const cancelResponsibleButton = document.getElementById('cancelResponsibleButton');
    const saveResponsibleButton = document.getElementById('saveResponsibleButton');
    const wantedItems = document.getElementById('wantedItems');
    let currentItems = [];
    let currentTemplateRef = '';
    let currentSourceName = '';
    let currentCommercialTerms = {{}};
    let managedTemplates = [];
    let templateManagerBusy = false;
    let pendingTemplateAction = null;
    let managedResponsibles = [];
    let responsibleManagerBusy = false;
    let editingResponsibleId = null;

    function syncTemplateEditLink() {{
      const templateId = templateChoice.value;
      editTemplateLink.hidden = !templateId;
      editTemplateLink.href = templateId ? `/template/${{encodeURIComponent(templateId)}}` : '#';
    }}

    templateChoice.addEventListener('change', syncTemplateEditLink);
    syncTemplateEditLink();

    function formatTemplateSize(bytes) {{
      if (bytes < 1024) return `${{bytes}} B`;
      if (bytes < 1024 * 1024) return `${{(bytes / 1024).toFixed(1)}} KB`;
      return `${{(bytes / (1024 * 1024)).toFixed(1)}} MB`;
    }}

    function setTemplateManagerMessage(text, kind = '') {{
      templateManagerStatus.textContent = text;
      templateManagerStatus.className = `template-manager-status ${{kind}}`.trim();
    }}

    function setTemplateManagerBusy(busy) {{
      templateManagerBusy = busy;
      closeTemplateManager.disabled = busy;
      attachTemplateButton.disabled = busy;
      templateManagerList.querySelectorAll('button').forEach(button => {{ button.disabled = busy; }});
    }}

    function syncTemplateChoice() {{
      const selected = templateChoice.value;
      templateChoice.replaceChildren();
      managedTemplates.forEach(template => {{
        const option = document.createElement('option');
        option.value = template.id;
        option.textContent = template.display_name || template.name;
        templateChoice.appendChild(option);
      }});
      if (managedTemplates.some(template => template.id === selected)) templateChoice.value = selected;
      syncTemplateEditLink();
    }}

    function renderManagedTemplates() {{
      templateManagerList.replaceChildren();
      if (!managedTemplates.length) {{
        const empty = document.createElement('div');
        empty.className = 'template-manager-empty';
        empty.textContent = 'Nenhum template cadastrado.';
        templateManagerList.appendChild(empty);
        return;
      }}
      managedTemplates.forEach(template => {{
        const row = document.createElement('article');
        row.className = 'template-manager-row';
        const details = document.createElement('div');
        const name = document.createElement('strong');
        name.className = 'template-manager-name';
        name.textContent = template.name;
        const meta = document.createElement('span');
        meta.className = 'template-manager-meta';
        meta.textContent = formatTemplateSize(template.size);
        details.append(name, meta);

        const actions = document.createElement('div');
        actions.className = 'template-manager-actions';
        const replace = document.createElement('button');
        replace.type = 'button';
        replace.className = 'template-manager-button';
        replace.textContent = 'Substituir';
        replace.disabled = templateManagerBusy;
        replace.addEventListener('click', () => chooseTemplateFile({{type: 'replace', template}}));
        const remove = document.createElement('button');
        remove.type = 'button';
        remove.className = 'template-manager-button delete';
        remove.textContent = 'Excluir';
        remove.disabled = templateManagerBusy;
        remove.addEventListener('click', () => deleteManagedTemplate(template));
        actions.append(replace, remove);
        row.append(details, actions);
        templateManagerList.appendChild(row);
      }});
    }}

    async function loadManagedTemplates(messageText = '', messageKind = '') {{
      const response = await fetch('/api/templates', {{headers: {{'Accept': 'application/json'}}}});
      const payload = await response.json().catch(() => ({{}}));
      if (!response.ok) throw new Error(payload.error || 'Não foi possível carregar os templates.');
      managedTemplates = payload.templates || [];
      syncTemplateChoice();
      renderManagedTemplates();
      setTemplateManagerMessage(messageText, messageKind);
    }}

    function openTemplateManagerModal() {{
      templateManagerOverlay.hidden = false;
      document.body.style.overflow = 'hidden';
      setTemplateManagerMessage('Carregando templates...');
      setTemplateManagerBusy(true);
      loadManagedTemplates()
        .catch(error => setTemplateManagerMessage(error.message, 'error'))
        .finally(() => {{ setTemplateManagerBusy(false); renderManagedTemplates(); closeTemplateManager.focus(); }});
    }}

    function closeTemplateManagerModal() {{
      if (templateManagerBusy) return;
      templateManagerOverlay.hidden = true;
      document.body.style.overflow = '';
      pendingTemplateAction = null;
      templateManagerFile.value = '';
      openTemplateManager.focus();
    }}

    function chooseTemplateFile(action) {{
      if (templateManagerBusy) return;
      pendingTemplateAction = action;
      templateManagerFile.value = '';
      templateManagerFile.click();
    }}

    function validateManagedTemplate(file) {{
      if (!file.name.toLowerCase().endsWith('.docx')) throw new Error('Selecione um arquivo Word no formato .docx.');
      if (!file.size) throw new Error('O arquivo selecionado está vazio.');
      if (file.size > {MAX_TEMPLATE_SIZE}) throw new Error('O arquivo excede o limite de 15 MB.');
    }}

    async function submitManagedTemplate(file) {{
      const action = pendingTemplateAction;
      pendingTemplateAction = null;
      validateManagedTemplate(file);
      setTemplateManagerBusy(true);
      setTemplateManagerMessage(action.type === 'replace' ? 'Substituindo template...' : 'Anexando template...');
      try {{
        const data = new FormData();
        data.append('template_file', file);
        const url = action.type === 'replace'
          ? `/api/templates/${{encodeURIComponent(action.template.id)}}/replace`
          : '/api/templates';
        const response = await fetch(url, {{method: 'POST', body: data}});
        const payload = await response.json().catch(() => ({{}}));
        if (!response.ok) throw new Error(payload.error || 'Não foi possível salvar o template.');
        await loadManagedTemplates(action.type === 'replace' ? 'Template substituído com sucesso.' : 'Template anexado com sucesso.', 'success');
      }} finally {{
        setTemplateManagerBusy(false);
        renderManagedTemplates();
      }}
    }}

    async function deleteManagedTemplate(template) {{
      if (templateManagerBusy || !window.confirm(`Excluir o template "${{template.name}}"?`)) return;
      setTemplateManagerBusy(true);
      setTemplateManagerMessage('Excluindo template...');
      try {{
        const response = await fetch(`/api/templates/${{encodeURIComponent(template.id)}}`, {{method: 'DELETE'}});
        const payload = await response.json().catch(() => ({{}}));
        if (!response.ok) throw new Error(payload.error || 'Não foi possível excluir o template.');
        await loadManagedTemplates('Template excluído com sucesso.', 'success');
      }} catch (error) {{
        setTemplateManagerMessage(error.message, 'error');
      }} finally {{
        setTemplateManagerBusy(false);
        renderManagedTemplates();
      }}
    }}

    openTemplateManager.addEventListener('click', openTemplateManagerModal);
    closeTemplateManager.addEventListener('click', closeTemplateManagerModal);
    attachTemplateButton.addEventListener('click', () => chooseTemplateFile({{type: 'create'}}));
    templateManagerFile.addEventListener('change', () => {{
      const file = templateManagerFile.files && templateManagerFile.files[0];
      if (!file) {{ pendingTemplateAction = null; return; }}
      submitManagedTemplate(file).catch(error => {{
        setTemplateManagerBusy(false);
        renderManagedTemplates();
        setTemplateManagerMessage(error.message, 'error');
      }});
    }});
    templateManagerOverlay.addEventListener('click', event => {{
      if (event.target === templateManagerOverlay) closeTemplateManagerModal();
    }});
    document.addEventListener('keydown', event => {{
      if (event.key === 'Escape' && !templateManagerOverlay.hidden) closeTemplateManagerModal();
    }});

    function setResponsibleManagerMessage(text, kind = '') {{
      responsibleManagerStatus.textContent = text;
      responsibleManagerStatus.className = `template-manager-status ${{kind}}`.trim();
    }}

    function setResponsibleManagerBusy(busy) {{
      responsibleManagerBusy = busy;
      closeResponsibleManager.disabled = busy;
      newResponsibleButton.disabled = busy;
      responsibleManagerForm.querySelectorAll('input, textarea, button').forEach(control => {{ control.disabled = busy; }});
      responsibleManagerList.querySelectorAll('button').forEach(button => {{ button.disabled = busy; }});
    }}

    function syncResponsibleChoice(preferredId = '') {{
      const selected = preferredId || responsibleId.value;
      responsibleId.replaceChildren();
      managedResponsibles.forEach(responsible => {{
        const option = document.createElement('option');
        option.value = responsible.id;
        option.textContent = responsible.nome_completo;
        responsibleId.appendChild(option);
      }});
      if (managedResponsibles.some(responsible => responsible.id === selected)) responsibleId.value = selected;
    }}

    function renderManagedResponsibles() {{
      responsibleManagerList.replaceChildren();
      if (!managedResponsibles.length) {{
        const empty = document.createElement('div');
        empty.className = 'template-manager-empty';
        empty.textContent = 'Nenhum responsável cadastrado.';
        responsibleManagerList.appendChild(empty);
        return;
      }}
      managedResponsibles.forEach(responsible => {{
        const row = document.createElement('article');
        row.className = 'responsible-manager-row';
        if (responsible.id === responsibleId.value) row.classList.add('is-selected');
        const details = document.createElement('div');
        details.className = 'responsible-manager-details';
        const name = document.createElement('strong');
        name.className = 'template-manager-name';
        name.textContent = responsible.nome_completo;
        const meta = document.createElement('span');
        meta.className = 'template-manager-meta';
        meta.textContent = [responsible.empresa, responsible.cpf ? `CPF ${{responsible.cpf}}` : ''].filter(Boolean).join(' · ');
        details.append(name, meta);

        const actions = document.createElement('div');
        actions.className = 'responsible-manager-actions';
        const select = document.createElement('button');
        select.type = 'button';
        select.className = 'template-manager-button';
        select.textContent = responsible.id === responsibleId.value ? 'Selecionado' : 'Selecionar';
        select.disabled = responsibleManagerBusy;
        select.addEventListener('click', () => {{
          responsibleId.value = responsible.id;
          renderManagedResponsibles();
          closeResponsibleManagerModal();
        }});
        const edit = document.createElement('button');
        edit.type = 'button';
        edit.className = 'template-manager-button cancel';
        edit.textContent = 'Editar';
        edit.disabled = responsibleManagerBusy;
        edit.addEventListener('click', () => showResponsibleForm(responsible));
        const remove = document.createElement('button');
        remove.type = 'button';
        remove.className = 'template-manager-button delete';
        remove.textContent = 'Excluir';
        remove.disabled = responsibleManagerBusy;
        remove.addEventListener('click', () => deleteManagedResponsible(responsible));
        actions.append(select, edit, remove);
        row.append(details, actions);
        responsibleManagerList.appendChild(row);
      }});
    }}

    async function loadManagedResponsibles(messageText = '', messageKind = '', preferredId = '') {{
      const response = await fetch('/api/responsaveis', {{headers: {{'Accept': 'application/json'}}}});
      const payload = await response.json().catch(() => ({{}}));
      if (!response.ok) throw new Error(payload.error || 'Não foi possível carregar os responsáveis.');
      managedResponsibles = payload.responsaveis || [];
      syncResponsibleChoice(preferredId);
      renderManagedResponsibles();
      setResponsibleManagerMessage(messageText, messageKind);
    }}

    function resetResponsibleForm() {{
      editingResponsibleId = null;
      responsibleManagerForm.reset();
      responsibleManagerForm.hidden = true;
      responsibleFormTitle.textContent = 'Novo responsável';
    }}

    function showResponsibleForm(responsible = null) {{
      if (responsibleManagerBusy) return;
      editingResponsibleId = responsible ? responsible.id : null;
      responsibleFormTitle.textContent = responsible ? 'Editar responsável' : 'Novo responsável';
      responsibleManagerForm.reset();
      if (responsible) {{
        Object.entries(responsible).forEach(([field, value]) => {{
          const control = responsibleManagerForm.elements.namedItem(field);
          if (control) control.value = value || '';
        }});
      }}
      responsibleManagerForm.hidden = false;
      setResponsibleManagerMessage('');
      document.getElementById('responsibleName').focus();
    }}

    function validateResponsibleData(data) {{
      const cnpjDigits = data.cnpj.replace(/\\D/g, '');
      const cpfDigits = data.cpf.replace(/\\D/g, '');
      if (cnpjDigits.length !== 14) throw new Error('Informe um CNPJ com 14 dígitos.');
      if (data.cpf && cpfDigits.length !== 11) throw new Error('Informe um CPF com 11 dígitos.');
    }}

    async function saveManagedResponsible(event) {{
      event.preventDefault();
      if (responsibleManagerBusy || !responsibleManagerForm.reportValidity()) return;
      const data = Object.fromEntries(new FormData(responsibleManagerForm).entries());
      try {{
        validateResponsibleData(data);
      }} catch (error) {{
        setResponsibleManagerMessage(error.message, 'error');
        return;
      }}
      const editing = Boolean(editingResponsibleId);
      setResponsibleManagerBusy(true);
      setResponsibleManagerMessage(editing ? 'Atualizando responsável...' : 'Cadastrando responsável...');
      try {{
        const url = editing ? `/api/responsaveis/${{encodeURIComponent(editingResponsibleId)}}` : '/api/responsaveis';
        const response = await fetch(url, {{
          method: editing ? 'PUT' : 'POST',
          headers: {{'Content-Type': 'application/json'}},
          body: JSON.stringify(data)
        }});
        const payload = await response.json().catch(() => ({{}}));
        if (!response.ok) throw new Error(payload.error || 'Não foi possível salvar o responsável.');
        const savedId = payload.responsavel.id;
        resetResponsibleForm();
        await loadManagedResponsibles(
          editing ? 'Responsável atualizado com sucesso.' : 'Responsável cadastrado com sucesso.',
          'success',
          savedId
        );
      }} catch (error) {{
        setResponsibleManagerMessage(error.message, 'error');
      }} finally {{
        setResponsibleManagerBusy(false);
        renderManagedResponsibles();
      }}
    }}

    async function deleteManagedResponsible(responsible) {{
      if (responsibleManagerBusy || !window.confirm(`Excluir o responsável "${{responsible.nome_completo}}"?`)) return;
      setResponsibleManagerBusy(true);
      setResponsibleManagerMessage('Verificando vínculos e excluindo responsável...');
      try {{
        const response = await fetch(`/api/responsaveis/${{encodeURIComponent(responsible.id)}}`, {{method: 'DELETE'}});
        const payload = await response.json().catch(() => ({{}}));
        if (!response.ok) throw new Error(payload.error || 'Não foi possível excluir o responsável.');
        await loadManagedResponsibles('Responsável excluído com sucesso.', 'success');
      }} catch (error) {{
        setResponsibleManagerMessage(error.message, 'error');
      }} finally {{
        setResponsibleManagerBusy(false);
        renderManagedResponsibles();
      }}
    }}

    function openResponsibleManagerModal() {{
      responsibleManagerOverlay.hidden = false;
      document.body.style.overflow = 'hidden';
      resetResponsibleForm();
      setResponsibleManagerMessage('Carregando responsáveis...');
      setResponsibleManagerBusy(true);
      loadManagedResponsibles()
        .catch(error => setResponsibleManagerMessage(error.message, 'error'))
        .finally(() => {{ setResponsibleManagerBusy(false); renderManagedResponsibles(); closeResponsibleManager.focus(); }});
    }}

    function closeResponsibleManagerModal() {{
      if (responsibleManagerBusy) return;
      responsibleManagerOverlay.hidden = true;
      document.body.style.overflow = '';
      resetResponsibleForm();
      openResponsibleManager.focus();
    }}

    openResponsibleManager.addEventListener('click', openResponsibleManagerModal);
    closeResponsibleManager.addEventListener('click', closeResponsibleManagerModal);
    newResponsibleButton.addEventListener('click', () => showResponsibleForm());
    cancelResponsibleButton.addEventListener('click', resetResponsibleForm);
    responsibleManagerForm.addEventListener('submit', saveManagedResponsible);
    responsibleManagerOverlay.addEventListener('click', event => {{
      if (event.target === responsibleManagerOverlay) closeResponsibleManagerModal();
    }});
    document.addEventListener('keydown', event => {{
      if (event.key === 'Escape' && !responsibleManagerOverlay.hidden) closeResponsibleManagerModal();
    }});

    function setProposalVisible(visible) {{
      proposalPanel.classList.toggle('is-open', visible);
      proposalPanel.setAttribute('aria-hidden', visible ? 'false' : 'true');
      proposalPanel.inert = !visible;
      showProposalBtn.setAttribute('aria-expanded', visible ? 'true' : 'false');
      navBlock1.classList.toggle('is-active', !visible);
      navBlock2.classList.toggle('is-active', visible);
      if (visible) {{
        requestAnimationFrame(() => {{
          proposalPanel.classList.add('nav-focused');
          window.setTimeout(() => proposalPanel.classList.remove('nav-focused'), 600);
        }});
        window.setTimeout(
          () => proposalPanel.scrollIntoView({{ behavior: 'smooth', block: 'start' }}),
          360
        );
      }}
    }}

    navBlock1.addEventListener('click', () => {{
      setProposalVisible(false);
      block1.scrollIntoView({{behavior: 'smooth', block: 'start'}});
    }});
    navBlock2.addEventListener('click', () => setProposalVisible(true));

    const baseLabels = [
      ['item', 'ITEM'],
      ['quantidade', 'QTD'],
      ['unidade', 'UND'],
      ['descricao', 'DESCRIÇÃO'],
      ['marca', 'MARCA'],
      ['valor_unitario', 'VALOR UNITÁRIO'],
      ['valor_total', 'VALOR TOTAL']
    ];

    function labelsForItems(items) {{
      if (items.some(item => String(item.lote || '').trim())) {{
        return [['lote', 'LOTE'], ...baseLabels];
      }}
      return baseLabels;
    }}

    const MAX_RANGE_DAYS = 30;
    let rangeStart = null;
    let rangeEnd = null;
    let choosingRangeEnd = false;
    let hoverDate = null;
    let calendarMonth = startOfMonth(new Date());

    function setDefaultDateRange(start = new Date()) {{
      const cleanStart = startOfDay(start);
      setRange(cleanStart, addDays(cleanStart, MAX_RANGE_DAYS - 1));
      calendarMonth = startOfMonth(cleanStart);
      renderCalendar();
    }}

    function startOfDay(date) {{
      return new Date(date.getFullYear(), date.getMonth(), date.getDate());
    }}

    function startOfMonth(date) {{
      return new Date(date.getFullYear(), date.getMonth(), 1);
    }}

    function addDays(date, days) {{
      const next = new Date(date);
      next.setDate(next.getDate() + days);
      return startOfDay(next);
    }}

    function dateStamp(date) {{
      return Date.UTC(date.getFullYear(), date.getMonth(), date.getDate());
    }}

    function dateDiffDays(start, end) {{
      return Math.round((dateStamp(end) - dateStamp(start)) / 86400000);
    }}

    function sameDay(a, b) {{
      return a && b && dateStamp(a) === dateStamp(b);
    }}

    function parseInputDate(value) {{
      const match = String(value || '').match(/^(\\d{{4}})-(\\d{{2}})-(\\d{{2}})$/);
      if (!match) return null;
      return new Date(Number(match[1]), Number(match[2]) - 1, Number(match[3]));
    }}

    function dateToInputValue(date) {{
      const year = date.getFullYear();
      const month = String(date.getMonth() + 1).padStart(2, '0');
      const day = String(date.getDate()).padStart(2, '0');
      return `${{year}}-${{month}}-${{day}}`;
    }}

    function formatDateBR(date) {{
      return date.toLocaleDateString('pt-BR', {{ day: '2-digit', month: '2-digit', year: 'numeric' }});
    }}

    function updateDateRangeText() {{
      if (rangeStart && rangeEnd) {{
        dateRangeText.textContent = `${{formatDateBR(rangeStart)}} — ${{formatDateBR(rangeEnd)}}`;
        dateRangeText.classList.remove('placeholder');
      }} else if (rangeStart) {{
        dateRangeText.textContent = `${{formatDateBR(rangeStart)}} — escolha a data final`;
        dateRangeText.classList.remove('placeholder');
      }} else {{
        dateRangeText.textContent = 'Selecione o período';
        dateRangeText.classList.add('placeholder');
      }}
    }}

    function setRange(start, end = start) {{
      const first = startOfDay(start);
      const second = startOfDay(end || start);
      rangeStart = first <= second ? first : second;
      rangeEnd = first <= second ? second : first;
      if (dateDiffDays(rangeStart, rangeEnd) > MAX_RANGE_DAYS - 1) {{
        rangeEnd = addDays(rangeStart, MAX_RANGE_DAYS - 1);
      }}
      searchStartDate.value = dateToInputValue(rangeStart);
      searchEndDate.value = dateToInputValue(rangeEnd);
      updateDateRangeText();
    }}

    function setRangeStart(start) {{
      rangeStart = startOfDay(start);
      rangeEnd = null;
      searchStartDate.value = dateToInputValue(rangeStart);
      searchEndDate.value = '';
      updateDateRangeText();
    }}

    function clearDateRange() {{
      rangeStart = null;
      rangeEnd = null;
      choosingRangeEnd = false;
      hoverDate = null;
      searchStartDate.value = '';
      searchEndDate.value = '';
      calendarHelp.textContent = 'Escolha a data inicial.';
      updateDateRangeText();
      renderCalendar();
    }}

    function isDateDisabled(date) {{
      if (choosingRangeEnd && rangeStart) {{
        return Math.abs(dateDiffDays(rangeStart, date)) > MAX_RANGE_DAYS - 1;
      }}
      return false;
    }}

    function previewRangeEnd() {{
      if (choosingRangeEnd && hoverDate && !isDateDisabled(hoverDate)) return hoverDate;
      return rangeEnd;
    }}

    function visibleCalendarRange() {{
      if (!rangeStart) return null;
      const candidateEnd = previewRangeEnd() || rangeStart;
      return candidateEnd < rangeStart
        ? {{ start: candidateEnd, end: rangeStart }}
        : {{ start: rangeStart, end: candidateEnd }};
    }}

    function isDateInRange(date) {{
      const visibleRange = visibleCalendarRange();
      return Boolean(visibleRange && date >= visibleRange.start && date <= visibleRange.end);
    }}

    function updateCalendarRangeStyles() {{
      const visibleRange = visibleCalendarRange();
      calendarGrid.querySelectorAll('.calendar-day').forEach(button => {{
        const date = parseInputDate(button.dataset.date);
        const inRange = Boolean(date && visibleRange && date >= visibleRange.start && date <= visibleRange.end);
        button.classList.toggle('in-range', inRange);
        button.classList.toggle('range-start', Boolean(date && visibleRange && sameDay(date, visibleRange.start)));
        button.classList.toggle('range-end', Boolean(date && visibleRange && sameDay(date, visibleRange.end)));
        button.setAttribute('aria-pressed', inRange ? 'true' : 'false');
      }});
    }}

    function renderCalendar() {{
      const monthTitle = calendarMonth.toLocaleDateString('pt-BR', {{ month: 'long', year: 'numeric' }});
      calendarTitle.textContent = monthTitle.charAt(0).toUpperCase() + monthTitle.slice(1);
      calendarGrid.innerHTML = '';
      const firstDay = startOfMonth(calendarMonth);
      const gridStart = addDays(firstDay, -firstDay.getDay());
      for (let offset = 0; offset < 42; offset += 1) {{
        const date = addDays(gridStart, offset);
        const button = document.createElement('button');
        button.type = 'button';
        button.className = 'calendar-day';
        button.textContent = String(date.getDate());
        button.dataset.date = dateToInputValue(date);
        button.setAttribute('aria-label', formatDateBR(date));
        if (date.getMonth() !== calendarMonth.getMonth()) button.classList.add('outside-month');
        if (isDateDisabled(date)) button.disabled = true;
        button.addEventListener('click', () => selectCalendarDate(date));
        button.addEventListener('mouseenter', () => {{
          if (choosingRangeEnd && !isDateDisabled(date) && !sameDay(hoverDate, date)) {{
            hoverDate = date;
            updateCalendarRangeStyles();
          }}
        }});
        calendarGrid.appendChild(button);
      }}
      updateCalendarRangeStyles();
    }}

    function openDateRangePanel() {{
      if (!choosingRangeEnd && rangeStart) calendarMonth = startOfMonth(rangeStart);
      dateRangePanel.hidden = false;
      dateRangeButton.setAttribute('aria-expanded', 'true');
      calendarHelp.textContent = choosingRangeEnd
        ? 'Escolha a data final. Periodo maximo: 30 dias.'
        : 'Escolha a data inicial.';
      renderCalendar();
    }}

    function closeDateRangePanel() {{
      dateRangePanel.hidden = true;
      dateRangeButton.setAttribute('aria-expanded', 'false');
      hoverDate = null;
    }}

    function selectCalendarDate(date) {{
      const selected = startOfDay(date);
      if (!choosingRangeEnd || !rangeStart) {{
        setRangeStart(selected);
        choosingRangeEnd = true;
        calendarMonth = startOfMonth(selected);
        calendarHelp.textContent = 'Escolha a data final. Periodo maximo: 30 dias.';
        renderCalendar();
        return;
      }}
      if (Math.abs(dateDiffDays(rangeStart, selected)) > MAX_RANGE_DAYS - 1) {{
        calendarHelp.textContent = 'Periodo maximo: 30 dias. Escolha uma data final mais proxima.';
        hoverDate = null;
        renderCalendar();
        return;
      }}
      setRange(rangeStart, selected);
      choosingRangeEnd = false;
      hoverDate = null;
      calendarHelp.textContent = 'Periodo selecionado.';
      renderCalendar();
      closeDateRangePanel();
    }}

    setDefaultDateRange();

    function formatPncpDateKey(value) {{
      const text = String(value || '');
      if (!/^\\d{{8}}$/.test(text)) return '';
      return `${{text.slice(6, 8)}}/${{text.slice(4, 6)}}/${{text.slice(0, 4)}}`;
    }}

    function describeDateRange(start, end) {{
      const startLabel = formatPncpDateKey(start);
      const endLabel = formatPncpDateKey(end);
      if (startLabel && endLabel) return `${{startLabel}} — ${{endLabel}}`;
      if (startLabel) return `a partir de ${{startLabel}}`;
      if (endLabel) return `até ${{endLabel}}`;
      return 'sem filtro de data';
    }}

    function hasValidDateRange() {{
      const start = parseInputDate(searchStartDate.value);
      const end = parseInputDate(searchEndDate.value);
      if (!start && !end) return true;
      if (!start || !end) return false;
      if (end < start) return false;
      return dateDiffDays(start, end) <= MAX_RANGE_DAYS - 1;
    }}

    function escapeHtml(value) {{
      return String(value ?? '').replace(/[&<>"']/g, char => ({{
        '&': '&amp;', '<': '&lt;', '>': '&gt;', '"': '&quot;', "'": '&#39;'
      }}[char]));
    }}

    function setMessage(text, kind = '') {{
      message.className = 'message' + (kind ? ' ' + kind : '');
      message.textContent = text;
    }}

    function setSearchMessage(text, kind = '') {{
      searchResults.innerHTML = `<div class="message${{kind ? ' ' + kind : ''}}">${{escapeHtml(text)}}</div>`;
    }}

    function setIdentificationMessage(text, kind = '') {{
      itemIdentificationStatus.textContent = text;
      itemIdentificationStatus.className = 'identification-status' + (kind ? ' ' + kind : '');
    }}

    function isValidPncpLink(value) {{
      try {{
        const url = new URL(String(value || '').trim());
        return ['http:', 'https:'].includes(url.protocol)
          && url.hostname === 'pncp.gov.br'
          && /^\\/app\\/editais\\/\\d{{14}}\\/\\d{{4}}\\/\\d+\\/?$/.test(url.pathname);
      }} catch {{
        return false;
      }}
    }}

    let lastIdentificationSourceLabel = '';
    let lastIdentificationReviewLabel = '';
    let lastIdentificationHasDivergence = false;

    function updateIdentificationSelectionStatus() {{
      const boxes = Array.from(itemIdentificationPreview.querySelectorAll('[data-identify-select]'));
      if (!boxes.length) return;
      const selectedCount = boxes.filter(input => input.checked).length;
      const suffix = lastIdentificationSourceLabel ? ` pela ${{lastIdentificationSourceLabel}}` : '';
      const review = lastIdentificationReviewLabel ? ` ${{lastIdentificationReviewLabel}}` : '';
      const kind = selectedCount && !lastIdentificationHasDivergence ? 'ok' : 'warn';
      setIdentificationMessage(`${{selectedCount}} de ${{boxes.length}} item(ns) selecionado(s)${{suffix}}.${{review}}`, kind);
    }}

    function syncWantedItemsFromSelection() {{
      const boxes = Array.from(itemIdentificationPreview.querySelectorAll('[data-identify-select]'));
      const selected = boxes.filter(input => input.checked)
        .map(input => input.dataset.itemKey || String(Number(input.value)))
        .filter(value => value && value !== 'NaN')
        .sort((a, b) => a.localeCompare(b, 'pt-BR', {{ numeric: true }}));
      wantedItems.value = selected.join(', ');
      updateIdentificationSelectionStatus();
    }}

    function collectIdentificationUnitValues() {{
      const values = new Map();
      itemIdentificationPreview.querySelectorAll('[data-identify-value]').forEach(input => {{
        const item = input.dataset.itemKey || String(Number(input.dataset.item || ''));
        if (item && item !== 'NaN') values.set(item, input.value.trim());
      }});
      return values;
    }}

    function handleIdentificationValueInput(input) {{
      const row = input.closest('tr');
      const checkbox = row ? row.querySelector('[data-identify-select]') : null;
      if (checkbox && input.value.trim()) {{
        checkbox.checked = true;
        syncWantedItemsFromSelection();
      }}
    }}

    function formatIdentificationValue(input) {{
      const parsed = parseNumber(input.value);
      const valid = Number.isFinite(parsed) && parsed >= 0;
      input.setCustomValidity(input.value.trim() && !valid ? 'Informe um valor monetário válido e não negativo.' : '');
      if (valid) input.value = formatMoney(parsed);
      return valid;
    }}

    function renderItemIdentifications(items) {{
      if (!items.length) {{
        itemIdentificationPreview.innerHTML = '<div class="message">Nenhum item identificado.</div>';
        return;
      }}
      const showLote = items.some(item => String(item.lote || '').trim());
      const rows = items.map(item => `
        <tr>
          <td class="short"><input type="checkbox" data-identify-select data-item-key="${{escapeHtml(itemSelectionKey(item))}}" value="${{escapeHtml(item.item || '')}}"></td>
          <td class="money"><input class="cell-input money-input identify-value" data-identify-value data-item-key="${{escapeHtml(itemSelectionKey(item))}}" data-item="${{escapeHtml(item.item || '')}}" type="text" inputmode="decimal" value="${{escapeHtml(item.valor_unitario || '')}}" placeholder="R$ 0,00"></td>
          ${{showLote ? `<td class="short">${{escapeHtml(item.lote || '')}}</td>` : ''}}
          <td class="short">${{escapeHtml(item.item || '')}}</td>
          <td class="short">${{escapeHtml(item.quantidade || '')}}</td>
          <td class="short">${{escapeHtml(item.unidade || '')}}</td>
          <td>${{escapeHtml(item.categoria || '')}}</td>
          <td class="desc">${{escapeHtml(item.descricao || '')}}</td>
        </tr>
      `).join('');
      itemIdentificationPreview.innerHTML = `
        <table>
          <thead>
            <tr>
              <th class="short">SEL.</th>
              <th class="money">VALOR UNITÁRIO</th>
              ${{showLote ? '<th class="short">LOTE</th>' : ''}}
              <th class="short">ITEM</th>
              <th class="short">QTD</th>
              <th class="short">UND</th>
              <th>IDENTIFICAÇÃO</th>
              <th>DESCRIÇÃO DO TR/EDITAL</th>
            </tr>
          </thead>
          <tbody>${{rows}}</tbody>
        </table>
      `;
      itemIdentificationPreview.querySelectorAll('[data-identify-select]').forEach(input => {{
        input.addEventListener('change', syncWantedItemsFromSelection);
      }});
      itemIdentificationPreview.querySelectorAll('[data-identify-value]').forEach(input => {{
        input.addEventListener('input', () => handleIdentificationValueInput(input));
        input.addEventListener('blur', () => {{
          formatIdentificationValue(input);
          handleIdentificationValueInput(input);
        }});
      }});
      updateIdentificationSelectionStatus();
    }}

    let identificationTimer = null;
    let identificationRequestId = 0;
    let lastIdentifiedLink = '';

    async function identifyItemsFromLink() {{
      const link = pncpLink.value.trim();
      if (!link) {{
        lastIdentifiedLink = '';
        lastIdentificationSourceLabel = '';
        lastIdentificationReviewLabel = '';
        lastIdentificationHasDivergence = false;
        setIdentificationMessage('Aguardando link PNCP.');
        itemIdentificationPreview.innerHTML = '<div class="message">Informe o link para identificar o que cada item representa.</div>';
        return;
      }}
      if (!isValidPncpLink(link)) {{
        setIdentificationMessage('Informe uma URL pública válida do domínio pncp.gov.br.', 'warn');
        if (!itemIdentificationPreview.querySelector('table')) {{
          itemIdentificationPreview.innerHTML = '<div class="message warn">Use o formato https://pncp.gov.br/app/editais/CNPJ/ANO/SEQUENCIAL.</div>';
        }}
        return;
      }}
      if (link === lastIdentifiedLink) return;

      const hasExistingItems = Boolean(itemIdentificationPreview.querySelector('table'));
      lastIdentifiedLink = link;
      const requestId = ++identificationRequestId;
      setIdentificationMessage('Consultando edital...');
      if (!hasExistingItems) {{
        itemIdentificationPreview.innerHTML = '<div class="message">Consultando o edital e transcrevendo os itens disponíveis.</div>';
      }}

      try {{
        const params = new URLSearchParams({{ pncp_link: link }});
        const response = await fetch('/identify-items?' + params.toString());
        const payload = await response.json();
        if (!response.ok) {{
          throw new Error(payload.error || 'Não foi possível identificar os itens.');
        }}
        if (requestId !== identificationRequestId) return;
        const documentType = payload.pncp?.documento_tipo || 'arquivo oficial';
        const documentName = payload.pncp?.documento_usado || '';
        const source = documentName
          ? `descrição do ${{documentType}} (${{documentName}})`
          : `descrição do ${{documentType}}`;
        lastIdentificationSourceLabel = source;
        const review = payload.description_review;
        const check = payload.pncp_items_check;
        lastIdentificationHasDivergence = Boolean(check?.has_divergence);
        const reviewLabel = review
          ? `Descrições: ${{review.complete_count}}/${{review.reviewed_count}} revisadas.`
          : '';
        const divergenceLabel = check?.has_divergence
          ? `Atenção: arquivo com ${{check.file_count}} itens e aba PNCP com ${{check.pncp_count}}.`
          : '';
        lastIdentificationReviewLabel = [reviewLabel, divergenceLabel].filter(Boolean).join(' ');
        renderItemIdentifications(payload.items || []);
      }} catch (error) {{
        if (requestId !== identificationRequestId) return;
        lastIdentifiedLink = '';
        setIdentificationMessage('Não foi possível consultar o edital. Verifique o link informado e tente novamente.', 'danger');
        if (!hasExistingItems) {{
          itemIdentificationPreview.innerHTML = '<div class="message danger">Não foi possível consultar o edital. Verifique o link informado e tente novamente.</div>';
        }}
      }}
    }}

    function scheduleItemIdentification() {{
      window.clearTimeout(identificationTimer);
      identificationTimer = window.setTimeout(identifyItemsFromLink, 300);
    }}

    function dateToPncp(value) {{
      return String(value || '').replace(/-/g, '');
    }}

    function formatDateTime(value) {{
      if (!value) return '';
      const date = new Date(value);
      if (Number.isNaN(date.getTime())) return value;
      return date.toLocaleDateString('pt-BR');
    }}

    function parseNumber(value) {{
      const text = String(value ?? '').trim();
      if (!text) return NaN;
      const cleaned = text.replace(/R\\$/gi, '').replace(/\\s/g, '');
      if (cleaned.includes(',')) {{
        return Number(cleaned.replace(/\\./g, '').replace(',', '.'));
      }}
      return Number(cleaned.replace(/,/g, ''));
    }}

    function formatMoney(value) {{
      if (!Number.isFinite(value)) return '';
      return value.toLocaleString('pt-BR', {{ style: 'currency', currency: 'BRL' }});
    }}

    function calculateTotal(row) {{
      const quantidade = parseNumber(row.quantidade);
      const unitario = parseNumber(row.valor_unitario);
      if (!Number.isFinite(quantidade) || !Number.isFinite(unitario)) return '';
      return formatMoney(quantidade * unitario);
    }}

    function normalizeItemRef(value) {{
      const text = String(value ?? '').trim();
      if (/^\\d+$/.test(text)) return String(Number(text));
      return text;
    }}

    function itemSelectionKey(item) {{
      const itemRef = normalizeItemRef(item.item);
      const loteRef = normalizeItemRef(item.lote);
      if (!itemRef) return '';
      return loteRef ? `${{loteRef}}/${{itemRef}}` : itemRef;
    }}

    function parseWantedItems(value) {{
      const text = String(value ?? '').trim();
      if (!text) return null;
      const selected = new Set();
      const parts = text.split(/[;,\\s]+/).filter(Boolean);
      for (const part of parts) {{
        if (/^[^/]+\\/.+$/.test(part)) {{
          selected.add(part.trim());
          continue;
        }}
        const range = part.match(/^(\\d+)\\s*-\\s*(\\d+)$/);
        if (range) {{
          const start = Number(range[1]);
          const end = Number(range[2]);
          const min = Math.min(start, end);
          const max = Math.max(start, end);
          for (let item = min; item <= max; item += 1) selected.add(String(item));
          continue;
        }}
        if (/^\\d+$/.test(part)) selected.add(String(Number(part)));
      }}
      return selected.size ? selected : null;
    }}

    function filterWantedItems(items) {{
      const selected = parseWantedItems(wantedItems.value);
      if (!selected) return [];
      return items.filter(item => selected.has(itemSelectionKey(item)) || selected.has(normalizeItemRef(item.item)));
    }}

    function refreshTotals() {{
      document.querySelectorAll('tr[data-index]').forEach(tr => {{
        const index = Number(tr.dataset.index);
        currentItems[index].marca = tr.querySelector('[data-field="marca"]').value;
        currentItems[index].valor_unitario = tr.querySelector('[data-field="valor_unitario"]').value;
        currentItems[index].valor_total = calculateTotal(currentItems[index]);
        const total = tr.querySelector('[data-field="valor_total"]');
        total.textContent = currentItems[index].valor_total || 'em branco';
        total.classList.toggle('empty', !currentItems[index].valor_total);
      }});
    }}

    function renderTable(items) {{
      const labels = labelsForItems(items);
      const head = labels.map(([key, label]) => {{
        const cls = ['lote', 'item', 'quantidade', 'unidade'].includes(key) ? 'short' : (key.includes('valor') ? 'money' : '');
        return `<th class="${{cls}}">${{label}}</th>`;
      }}).join('');
      const rows = items.map((row, index) => {{
        return `<tr data-index="${{index}}">` + labels.map(([key]) => {{
          const value = row[key] || '';
          const cls = ['lote', 'item', 'quantidade', 'unidade'].includes(key) ? 'short' : (key.includes('valor') ? 'money' : (key === 'descricao' ? 'desc' : ''));
          if (key === 'marca') {{
            return `<td class="${{cls}}"><input class="cell-input" data-field="marca" value="${{escapeHtml(value)}}" placeholder="Marca"></td>`;
          }}
          if (key === 'valor_unitario') {{
            return `<td class="${{cls}}"><input class="cell-input money-input" data-field="valor_unitario" value="${{escapeHtml(value)}}" placeholder="R$ 0,00"></td>`;
          }}
          if (key === 'valor_total') {{
            const total = value || calculateTotal(row);
            row.valor_total = total;
            return `<td class="${{cls}}"><span data-field="valor_total" class="${{total ? '' : 'empty'}}">${{total ? escapeHtml(total) : 'em branco'}}</span></td>`;
          }}
          return `<td class="${{cls}}">${{value ? escapeHtml(value) : '<span class="empty">em branco</span>'}}</td>`;
        }}).join('') + '</tr>';
      }}).join('');
      preview.innerHTML = `<table><thead><tr>${{head}}</tr></thead><tbody>${{rows}}</tbody></table>`;
      preview.querySelectorAll('.cell-input').forEach(input => {{
        input.addEventListener('input', refreshTotals);
        input.addEventListener('blur', () => {{
          if (input.dataset.field === 'valor_unitario') {{
            const value = parseNumber(input.value);
            if (Number.isFinite(value)) input.value = formatMoney(value);
          }}
          refreshTotals();
        }});
      }});
    }}

    function createGenerateButton() {{
      const button = document.createElement('button');
      button.type = 'button';
      button.textContent = 'Gerar Word';
      button.addEventListener('click', async () => {{
        refreshTotals();
        if (!responsibleId.value) {{
          setMessage('Selecione um responsável pela proposta.', 'warn');
          responsibleId.focus();
          return;
        }}
        if (!currentTemplateRef) {{
          setMessage('Selecione um template disponível para gerar o documento.', 'warn');
          templateChoice.focus();
          return;
        }}
        if (currentItems.some(item => {{
          const value = parseNumber(item.valor_unitario);
          return !Number.isFinite(value) || value < 0;
        }})) {{
          setMessage('Revise os valores unitários antes de gerar o Word.', 'warn');
          return;
        }}
        button.disabled = true;
        button.textContent = 'Gerando Word...';
        setMessage('Gerando Word com marca, valor unitário e valor total...', '');
        try {{
          const response = await fetch('/generate', {{
            method: 'POST',
            headers: {{ 'Content-Type': 'application/json' }},
            body: JSON.stringify({{
              items: currentItems,
              template_ref: currentTemplateRef,
              source_name: currentSourceName,
              responsible_id: responsibleId.value,
              commercial_terms: currentCommercialTerms
            }})
          }});
          const payload = await response.json();
          if (!response.ok) {{
            throw new Error(payload.error || 'Não foi possível gerar o Word.');
          }}
          actions.innerHTML = '';
          const link = document.createElement('a');
          link.className = 'download';
          link.href = payload.download_url;
          link.textContent = 'Baixar Word';
          actions.appendChild(link);
          setMessage('Word gerado. O valor total foi calculado pela quantidade multiplicada pelo valor unitário.', 'ok');
        }} catch (error) {{
          setMessage(error.message || 'Não foi possível gerar o Word.', 'danger');
          button.disabled = false;
          button.textContent = 'Gerar Word';
        }}
      }});
      return button;
    }}

    function renderSearchResults(results) {{
      if (!results.length) {{
        setSearchMessage('Nenhum pregão encontrado com estes filtros.', 'warn');
        return;
      }}
      const rows = results.map((row, index) => `
        <tr>
          <td>${{escapeHtml(row.orgao || '')}}</td>
          <td>${{escapeHtml([row.municipio, row.uf].filter(Boolean).join(' / '))}}</td>
          <td>${{escapeHtml(row.numeroCompra || '')}}</td>
          <td class="desc">${{escapeHtml(row.objeto || '')}}</td>
          <td>${{escapeHtml(formatDateTime(row.encerramento))}}</td>
          <td>
            <button type="button" class="link-button secondary" data-use-link="${{escapeHtml(row.link || '')}}">Usar link</button>
          </td>
          <td>
            <button type="button" class="link-button secondary" data-copy-link="${{escapeHtml(row.link || '')}}">Copiar</button>
          </td>
        </tr>
      `).join('');
      searchResults.innerHTML = `
        <table>
          <thead>
            <tr>
              <th>Órgão</th>
              <th>Local</th>
              <th>Número</th>
              <th>Objeto</th>
              <th>Encerra</th>
              <th>Usar</th>
              <th>Copiar</th>
            </tr>
          </thead>
          <tbody>${{rows}}</tbody>
        </table>
      `;
      searchResults.querySelectorAll('[data-use-link]').forEach(button => {{
        button.addEventListener('click', () => {{
          setSourceMode('pncp');
          pncpLink.value = button.dataset.useLink;
          setProposalVisible(true);
          scheduleItemIdentification();
          pncpLink.focus();
          setMessage('Link PNCP selecionado. Preencha os dados e clique em extrair.', 'ok');
        }});
      }});
      searchResults.querySelectorAll('[data-copy-link]').forEach(button => {{
        button.addEventListener('click', async () => {{
          const link = button.dataset.copyLink;
          try {{
            await navigator.clipboard.writeText(link);
            button.textContent = 'Copiado';
          }} catch {{
            pncpLink.value = link;
            button.textContent = 'No campo';
            scheduleItemIdentification();
          }}
        }});
      }});
    }}

    showProposalBtn.addEventListener('click', () => {{
      setProposalVisible(true);
    }});

    hideProposalBtn.addEventListener('click', () => {{
      setProposalVisible(false);
      showProposalBtn.focus();
    }});

    dateRangeButton.addEventListener('click', () => {{
      if (dateRangePanel.hidden) openDateRangePanel();
      else closeDateRangePanel();
    }});

    dateRangePicker.addEventListener('click', event => {{
      event.stopPropagation();
    }});

    calendarPrev.addEventListener('click', () => {{
      calendarMonth = new Date(calendarMonth.getFullYear(), calendarMonth.getMonth() - 1, 1);
      renderCalendar();
    }});

    calendarNext.addEventListener('click', () => {{
      calendarMonth = new Date(calendarMonth.getFullYear(), calendarMonth.getMonth() + 1, 1);
      renderCalendar();
    }});

    calendarGrid.addEventListener('mouseleave', () => {{
      if (choosingRangeEnd) {{
        hoverDate = null;
        updateCalendarRangeStyles();
      }}
    }});

    document.addEventListener('click', event => {{
      if (!dateRangePicker.contains(event.target)) closeDateRangePanel();
    }});

    document.addEventListener('keydown', event => {{
      if (event.key === 'Escape') closeDateRangePanel();
    }});

    clearSearchBtn.addEventListener('click', () => {{
      clearDateRange();
      searchUf.value = '';
      searchKeyword.value = '';
      searchObjectType.value = '';
      searchModality.value = '';
      searchStatus.textContent = 'Filtros limpos';
      setSearchMessage('Filtros limpos. Informe novos criterios ou busque sem filtro de data.', '');
    }});

    pncpSearchForm.addEventListener('submit', async (event) => {{
      event.preventDefault();
      if (!hasValidDateRange()) {{
        searchStatus.textContent = 'Revise as datas';
        setSearchMessage('Selecione um período de até 30 dias.', 'warn');
        return;
      }}
      searchBtn.disabled = true;
      searchStatus.textContent = 'Consultando PNCP...';
      setSearchMessage('Buscando contratações abertas...', '');
      try {{
        const data = new FormData(pncpSearchForm);
        const params = new URLSearchParams();
        params.set('dataInicial', dateToPncp(data.get('dataInicial')));
        params.set('dataFinal', dateToPncp(data.get('dataFinal')));
        params.set('uf', data.get('uf') || '');
        params.set('palavraChave', data.get('palavraChave') || '');
        params.set('tipoObjeto', data.get('tipoObjeto') || '');
        params.set('codigoModalidadeContratacao', data.get('codigoModalidadeContratacao') || '');
        params.set('tamanhoPagina', '10');
        const response = await fetch('/pncp-search?' + params.toString());
        const payload = await response.json();
        if (!response.ok) {{
          throw new Error(payload.error || 'Não foi possível consultar o PNCP.');
        }}
        renderSearchResults(payload.results || []);
        const source = payload.source === 'api/search' ? 'filtro oficial' : 'consulta pública';
        const cache = payload.cache_hit ? ', cache' : '';
        const dateRange = describeDateRange(payload.dataInicial, payload.dataFinal);
        searchStatus.textContent = `${{(payload.results || []).length}} resultado(s), ${{source}}, ${{payload.pages_checked || 1}} página(s), ${{dateRange}}${{cache}}`;
        if (payload.rate_limited && !(payload.results || []).length) {{
          setSearchMessage('O PNCP limitou temporariamente a consulta. Aguarde alguns segundos e tente novamente com menos filtros.', 'warn');
        }}
        if (payload.timed_out && !(payload.results || []).length) {{
          setSearchMessage('O PNCP não respondeu dentro do tempo limite. A consulta foi feita com o intervalo informado; aguarde alguns instantes e tente novamente.', 'warn');
        }}
      }} catch (error) {{
        searchStatus.textContent = 'Erro na consulta';
        setSearchMessage(error.message || 'Não foi possível consultar o PNCP.', 'danger');
      }} finally {{
        searchBtn.disabled = false;
      }}
    }});

    pncpLink.addEventListener('input', scheduleItemIdentification);
    pncpLink.addEventListener('change', identifyItemsFromLink);
    pncpLink.addEventListener('blur', identifyItemsFromLink);

    form.addEventListener('submit', async (event) => {{
      event.preventDefault();
      const pncpValue = pncpLink.value.trim();
      if (!isValidPncpLink(pncpValue)) {{
        setMessage('Informe uma URL válida do domínio pncp.gov.br.', 'warn');
        pncpLink.focus();
        return;
      }}
      if (!responsibleId.value) {{
        setMessage('Selecione um responsável pela proposta.', 'warn');
        responsibleId.focus();
        return;
      }}
      if (!templateChoice.value && !(templateFile.files && templateFile.files[0])) {{
        setMessage('Selecione um template disponível para a proposta.', 'warn');
        templateChoice.focus();
        return;
      }}
      const selectedBoxes = Array.from(itemIdentificationPreview.querySelectorAll('[data-identify-select]:checked'));
      if (!selectedBoxes.length) {{
        setMessage('Selecione pelo menos um item para a proposta.', 'warn');
        return;
      }}
      for (const checkbox of selectedBoxes) {{
        const valueInput = checkbox.closest('tr')?.querySelector('[data-identify-value]');
        if (!valueInput || !valueInput.value.trim() || !formatIdentificationValue(valueInput)) {{
          setMessage('Informe um valor monetário válido para todos os itens selecionados.', 'warn');
          valueInput?.focus();
          return;
        }}
      }}
      presetBrand.value = presetBrand.value.trim();
      const submitLabel = btn.textContent;
      setMessage('Processando proposta...', '');
      btn.disabled = true;
      btn.textContent = 'Processando proposta...';

      try {{
        const data = new FormData(form);
        const response = await fetch('/process', {{ method: 'POST', body: data }});
        const payload = await response.json();
        if (!response.ok) {{
          throw new Error(payload.error || 'Não foi possível processar o arquivo.');
        }}
        const brand = presetBrand.value.trim();
        const unitValuesByItem = collectIdentificationUnitValues();

        const filteredItems = filterWantedItems(payload.items);
        if (!filteredItems.length) {{
          throw new Error('Selecione pelo menos um item na identificação.');
        }}

        currentItems = filteredItems.map(item => {{
          const row = {{
            ...item,
            marca: brand || item.marca || '',
            valor_unitario: unitValuesByItem.get(itemSelectionKey(item)) || item.valor_unitario || ''
          }};
          row.valor_total = calculateTotal(row);
          return row;
        }});
        currentTemplateRef = payload.template_ref || '';
        currentSourceName = payload.source_name || '';
        currentCommercialTerms = payload.commercial_terms || {{}};
        renderTable(currentItems);
        actions.innerHTML = '';
        actions.appendChild(createGenerateButton());
        const check = payload.pncp_items_check;
        let processMessage = `Proposta processada com sucesso. ${{currentItems.length}} item(ns) selecionado(s) de ${{payload.count}} encontrado(s). Os valores unitários informados por item foram aplicados; o valor total foi calculado automaticamente.`;
        const sourceType = payload.pncp?.documento_tipo || 'arquivo oficial';
        const sourceDocument = payload.pncp?.documento_usado || '';
        processMessage += ` Fonte das descrições: ${{sourceType}}${{sourceDocument ? ` (${{sourceDocument}})` : ''}}.`;
        let processKind = 'ok';
        if (check && check.has_divergence) {{
          processKind = 'warn';
          if (check.file_error) {{
            processMessage += ' Verificação PNCP: houve falha parcial na leitura do arquivo; nenhuma descrição foi substituída pela aba de itens do PNCP.';
          }} else if ((check.added_from_pncp || []).length) {{
            processMessage += ` Verificação PNCP: o arquivo trouxe ${{check.file_count}} item(ns) e a aba oficial trouxe ${{check.pncp_count}}. A divergência foi registrada, sem alterar as descrições do arquivo.`;
          }} else {{
            processMessage += ` Verificação PNCP: o arquivo trouxe ${{check.file_count}} item(ns), a aba oficial trouxe ${{check.pncp_count}}. Confira a divergência antes de gerar a proposta.`;
          }}
        }}
        const descriptionReview = payload.description_review;
        if (descriptionReview) {{
          processMessage += ` ${{descriptionReview.message}}`;
          if (descriptionReview.status === 'warn') processKind = 'warn';
        }}
        setMessage(processMessage, processKind);
      }} catch (error) {{
        setMessage(error.message || 'Não foi possível processar a proposta. Nenhum dado anterior foi alterado.', 'danger');
      }} finally {{
        btn.disabled = false;
        btn.textContent = submitLabel;
      }}
    }});
  </script>
</body>
</html>"""


def render_templates_page():
    page = """<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Templates Word</title>
  <style>
    :root {
      --orange: #ee8627;
      --orange-deep: #d96f16;
      --navy: #102844;
      --navy-hover: #1b3b60;
      --red: #c93434;
      --red-hover: #a92424;
      --ink: #17212c;
      --muted: #68717c;
      --surface: #ffffff;
      --line: #e4e6e9;
      --success: #176b3a;
      --error: #a32323;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      min-height: 100vh;
      background: var(--orange);
      color: var(--ink);
      font-family: Arial, Helvetica, sans-serif;
      font-size: 16px;
    }
    button, input { font: inherit; }
    button { cursor: pointer; }
    button:disabled { cursor: wait; opacity: .62; }
    .page {
      width: min(100% - 32px, 1040px);
      margin: 0 auto;
      padding: 48px 0 64px;
    }
    .topbar {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 20px;
      margin-bottom: 28px;
    }
    h1 {
      margin: 0;
      color: #ffffff;
      font-size: clamp(34px, 6vw, 58px);
      line-height: 1;
      letter-spacing: 0;
    }
    .back-link {
      color: var(--navy);
      font-weight: 700;
      text-decoration: none;
    }
    .back-link:hover { text-decoration: underline; }
    .workspace {
      overflow: hidden;
      border: 1px solid rgba(16, 40, 68, .14);
      border-radius: 8px;
      background: var(--surface);
      box-shadow: 0 18px 45px rgba(73, 37, 7, .2);
    }
    .workspace-head {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 16px;
      padding: 22px 24px;
      border-bottom: 1px solid var(--line);
    }
    .workspace-title {
      margin: 0;
      font-size: 18px;
      color: var(--navy);
    }
    .primary, .replace, .delete {
      min-height: 42px;
      border: 0;
      border-radius: 5px;
      padding: 10px 16px;
      color: #ffffff;
      font-weight: 700;
      white-space: nowrap;
    }
    .primary, .replace { background: var(--navy); }
    .primary:hover, .replace:hover { background: var(--navy-hover); }
    .delete { background: var(--red); }
    .delete:hover { background: var(--red-hover); }
    .status {
      min-height: 24px;
      margin: 0;
      padding: 0 24px;
      font-size: 14px;
      line-height: 24px;
    }
    .status.success { color: var(--success); }
    .status.error { color: var(--error); }
    .status.loading { color: var(--muted); }
    .template-list { border-top: 1px solid var(--line); }
    .template-row {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto;
      align-items: center;
      gap: 20px;
      padding: 20px 24px;
      border-bottom: 1px solid var(--line);
    }
    .template-row:last-child { border-bottom: 0; }
    .template-name {
      display: block;
      overflow-wrap: anywhere;
      color: var(--ink);
      font-size: 16px;
      line-height: 1.35;
    }
    .template-meta {
      display: block;
      margin-top: 5px;
      color: var(--muted);
      font-size: 13px;
    }
    .row-actions {
      display: flex;
      align-items: center;
      gap: 9px;
    }
    .empty, .loading-list {
      padding: 48px 24px;
      color: var(--muted);
      text-align: center;
    }
    .loading-list::before {
      content: "";
      display: inline-block;
      width: 16px;
      height: 16px;
      margin-right: 9px;
      border: 2px solid #c9cdd2;
      border-top-color: var(--navy);
      border-radius: 50%;
      vertical-align: -3px;
      animation: spin .8s linear infinite;
    }
    @keyframes spin { to { transform: rotate(360deg); } }
    @media (max-width: 680px) {
      .page { width: min(100% - 20px, 1040px); padding-top: 28px; }
      .topbar, .workspace-head { align-items: stretch; flex-direction: column; }
      h1 { font-size: 38px; }
      .primary { width: 100%; }
      .template-row { grid-template-columns: 1fr; gap: 14px; padding: 18px; }
      .row-actions { display: grid; grid-template-columns: 1fr 1fr; }
      .replace, .delete { width: 100%; }
      .status { padding: 0 18px; }
    }
    @media (max-width: 380px) {
      .row-actions { grid-template-columns: 1fr; }
    }
  </style>
</head>
<body>
  <main class="page">
    <div class="topbar">
      <h1>TEMPLATES</h1>
      <a class="back-link" href="/">Voltar para propostas</a>
    </div>
    <section class="workspace" aria-labelledby="templateListTitle">
      <div class="workspace-head">
        <h2 id="templateListTitle" class="workspace-title">Arquivos cadastrados</h2>
        <button id="addTemplateButton" class="primary" type="button">Anexar novo template</button>
      </div>
      <p id="statusMessage" class="status" role="status" aria-live="polite"></p>
      <div id="templateList" class="template-list" aria-live="polite">
        <div class="loading-list">Carregando templates...</div>
      </div>
    </section>
    <input id="templateFileInput" type="file" accept=".docx" hidden>
  </main>
  <script>
    const MAX_FILE_SIZE = __MAX_FILE_SIZE__;
    const addButton = document.getElementById('addTemplateButton');
    const fileInput = document.getElementById('templateFileInput');
    const listElement = document.getElementById('templateList');
    const statusElement = document.getElementById('statusMessage');
    let templates = [];
    let pendingAction = null;
    let busy = false;

    function escapeText(value) {
      return String(value || '');
    }

    function formatSize(bytes) {
      if (bytes < 1024) return `${bytes} B`;
      if (bytes < 1024 * 1024) return `${(bytes / 1024).toFixed(1)} KB`;
      return `${(bytes / (1024 * 1024)).toFixed(1)} MB`;
    }

    function formatDate(value) {
      const date = new Date(value);
      return Number.isNaN(date.getTime()) ? '' : date.toLocaleString('pt-BR');
    }

    function setMessage(text, kind = '') {
      statusElement.textContent = text;
      statusElement.className = `status ${kind}`.trim();
    }

    function setBusy(value, message = '') {
      busy = value;
      addButton.disabled = value;
      listElement.querySelectorAll('button').forEach(button => { button.disabled = value; });
      if (value && message) setMessage(message, 'loading');
    }

    function actionButton(label, className, template, action) {
      const button = document.createElement('button');
      button.type = 'button';
      button.className = className;
      button.textContent = label;
      button.disabled = busy;
      button.addEventListener('click', () => action(template));
      return button;
    }

    function renderTemplates() {
      listElement.replaceChildren();
      if (!templates.length) {
        const empty = document.createElement('div');
        empty.className = 'empty';
        empty.textContent = 'Nenhum template cadastrado.';
        listElement.appendChild(empty);
        return;
      }
      templates.forEach(template => {
        const row = document.createElement('article');
        row.className = 'template-row';
        row.dataset.templateId = template.id;

        const info = document.createElement('div');
        const name = document.createElement('strong');
        name.className = 'template-name';
        name.textContent = escapeText(template.name);
        const meta = document.createElement('span');
        meta.className = 'template-meta';
        meta.textContent = `${formatSize(template.size)} - atualizado em ${formatDate(template.updated_at)}`;
        info.append(name, meta);

        const actions = document.createElement('div');
        actions.className = 'row-actions';
        actions.append(
          actionButton('Substituir', 'replace', template, chooseReplacement),
          actionButton('Excluir', 'delete', template, confirmDelete)
        );
        row.append(info, actions);
        listElement.appendChild(row);
      });
    }

    async function loadTemplates({keepMessage = false} = {}) {
      if (!keepMessage) setMessage('Carregando templates...', 'loading');
      try {
        const response = await fetch('/api/templates', {headers: {'Accept': 'application/json'}});
        const payload = await response.json();
        if (!response.ok) throw new Error(payload.error || 'Não foi possível carregar os templates.');
        templates = payload.templates || [];
        renderTemplates();
        if (!keepMessage) setMessage('');
      } catch (error) {
        listElement.innerHTML = '<div class="empty">Não foi possível carregar os templates.</div>';
        setMessage(error.message || 'Não foi possível carregar os templates.', 'error');
      }
    }

    function chooseNewTemplate() {
      if (busy) return;
      pendingAction = {type: 'create'};
      fileInput.value = '';
      fileInput.click();
    }

    function chooseReplacement(template) {
      if (busy) return;
      pendingAction = {type: 'replace', template};
      fileInput.value = '';
      fileInput.click();
    }

    function validateFile(file) {
      if (!file || !file.name.toLowerCase().endsWith('.docx')) {
        throw new Error('Selecione um arquivo no formato .docx.');
      }
      if (file.size === 0) throw new Error('O arquivo selecionado está vazio.');
      if (file.size > MAX_FILE_SIZE) throw new Error('O arquivo excede o limite de 15 MB.');
    }

    async function submitSelectedFile(file) {
      if (busy || !pendingAction) return;
      try {
        validateFile(file);
      } catch (error) {
        setMessage(error.message, 'error');
        pendingAction = null;
        return;
      }
      const action = pendingAction;
      pendingAction = null;
      const replacing = action.type === 'replace';
      setBusy(true, replacing ? `Substituindo ${action.template.name}...` : 'Anexando template...');
      const data = new FormData();
      data.append('template_file', file);
      const url = replacing
        ? `/api/templates/${encodeURIComponent(action.template.id)}/replace`
        : '/api/templates';
      try {
        const response = await fetch(url, {method: 'POST', body: data});
        const payload = await response.json();
        if (!response.ok) throw new Error(payload.error || 'Não foi possível salvar o template.');
        await loadTemplates({keepMessage: true});
        setMessage(replacing ? 'Template substituído com sucesso.' : 'Template anexado com sucesso.', 'success');
      } catch (error) {
        setMessage(error.message || 'Não foi possível salvar o template.', 'error');
      } finally {
        fileInput.value = '';
        setBusy(false);
        renderTemplates();
      }
    }

    async function confirmDelete(template) {
      if (busy) return;
      if (!window.confirm(`Excluir o template "${template.name}"?`)) return;
      setBusy(true, `Excluindo ${template.name}...`);
      try {
        const response = await fetch(`/api/templates/${encodeURIComponent(template.id)}`, {method: 'DELETE'});
        const payload = await response.json();
        if (!response.ok) throw new Error(payload.error || 'Não foi possível excluir o template.');
        await loadTemplates({keepMessage: true});
        setMessage('Template excluído com sucesso.', 'success');
      } catch (error) {
        setMessage(error.message || 'Não foi possível excluir o template.', 'error');
      } finally {
        setBusy(false);
        renderTemplates();
      }
    }

    addButton.addEventListener('click', chooseNewTemplate);
    fileInput.addEventListener('change', () => {
      const file = fileInput.files && fileInput.files[0];
      if (file) submitSelectedFile(file);
      else pendingAction = null;
    });
    loadTemplates();
  </script>
</body>
</html>"""
    return page.replace("__MAX_FILE_SIZE__", str(MAX_TEMPLATE_SIZE))


def parse_template_upload_form(handler):
    content_length = int(handler.headers.get("Content-Length", "0") or 0)
    if content_length <= 0:
        raise ValueError("Selecione um arquivo .docx.")
    if content_length > MAX_TEMPLATE_SIZE + (1024 * 1024):
        raise OverflowError("O arquivo excede o limite de 15 MB.")
    return cgi.FieldStorage(
        fp=handler.rfile,
        headers=handler.headers,
        environ={
            "REQUEST_METHOD": "POST",
            "CONTENT_TYPE": handler.headers.get("Content-Type"),
            "CONTENT_LENGTH": str(content_length),
        },
    )


def template_upload_field(form):
    field = form["template_file"] if "template_file" in form else None
    if field is None or isinstance(field, list) or not getattr(field, "filename", ""):
        raise ValueError("Selecione um arquivo .docx.")
    return field


def template_api_error(handler, exc):
    if isinstance(exc, FileNotFoundError):
        status = 404
    elif isinstance(exc, FileExistsError):
        status = 409
    elif isinstance(exc, OverflowError):
        status = 413
    elif isinstance(exc, ValueError):
        status = 422
    else:
        status = 500
    message = (
        str(exc)
        if status < 500
        else "Não foi possível concluir a operação com o template."
    )
    json_response(handler, status, {"error": message})


def parse_json_body(handler, maximum_size=1024 * 1024):
    content_length = int(handler.headers.get("Content-Length", "0") or 0)
    if content_length <= 0:
        raise ValueError("Envie um corpo JSON válido.")
    if content_length > maximum_size:
        raise OverflowError("Os dados enviados excedem o limite permitido.")
    content_type = handler.headers.get("Content-Type", "").split(";", 1)[0].strip().lower()
    if content_type and content_type != "application/json":
        raise ValueError("O conteúdo deve ser enviado como application/json.")
    try:
        return json.loads(handler.rfile.read(content_length).decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ValueError("Os dados enviados não são um JSON válido.") from exc


def responsible_api_error(handler, exc):
    if isinstance(exc, FileNotFoundError):
        status = 404
    elif isinstance(exc, (FileExistsError, ResponsibleInUseError)):
        status = 409
    elif isinstance(exc, OverflowError):
        status = 413
    elif isinstance(exc, ValueError):
        status = 422
    else:
        status = 500
    message = (
        str(exc)
        if status < 500
        else "Não foi possível concluir a operação com o responsável."
    )
    json_response(handler, status, {"error": message})


def business_api_error(handler, exc):
    if isinstance(exc, FileNotFoundError):
        status = 404
    elif isinstance(exc, OverflowError):
        status = 413
    elif isinstance(exc, ValueError):
        status = 422
    else:
        status = 500
    message = (
        str(exc)
        if status < 500
        else "Não foi possível concluir a operação com o negócio."
    )
    json_response(handler, status, {"error": message})


def etl_repository():
    return ETLRepository(DATABASE_PATH)


def _etl_query_date(value, end_of_day=False):
    text = compact(value)
    if not text:
        return ""
    try:
        parsed = datetime.strptime(text, "%Y%m%d")
    except ValueError:
        try:
            parsed = datetime.fromisoformat(text[:10])
        except ValueError:
            return text
    suffix = "T23:59:59" if end_of_day else "T00:00:00"
    return f"{parsed.date().isoformat()}{suffix}"


def sync_internal_pncp(payload):
    if not isinstance(payload, dict):
        raise ValueError("Dados da sincronizacao invalidos.")
    endpoint = compact(payload.get("source_endpoint") or payload.get("endpoint") or "proposta")
    raw_filters = payload.get("filters")
    raw_filters = raw_filters if isinstance(raw_filters, dict) else {}
    data_final = compact(raw_filters.get("dataFinal"))
    if not data_final:
        data_final = (datetime.now() + timedelta(days=29)).strftime("%Y%m%d")
    base_filters = {"dataFinal": data_final}
    if endpoint in {"publicacao", "atualizacao"}:
        base_filters["dataInicial"] = compact(raw_filters.get("dataInicial")) or datetime.now().strftime("%Y%m%d")
    modality = compact(raw_filters.get("codigoModalidadeContratacao"))
    if modality:
        base_filters["codigoModalidadeContratacao"] = int(modality)
    for name in ("codigoMunicipioIbge", "cnpj", "codigoUnidadeAdministrativa", "idUsuario"):
        value = compact(raw_filters.get(name))
        if value:
            base_filters[name] = value

    ufs = [value.strip().upper() for value in compact(raw_filters.get("uf")).split(",") if value.strip()]
    queries = [{**base_filters, "uf": uf} for uf in ufs] if ufs else [base_filters]
    service = ETLSyncService(
        etl_repository(),
        PNCPConnector(),
        PNCPMapper(),
        OpportunityClassifier(),
    )
    max_pages = min(max(int(payload.get("max_pages") or 1), 1), 20)
    max_records = min(max(int(payload.get("max_records") or 100), 1), 2000)
    results = []
    for filters in queries:
        results.append(service.sync(SyncRequest(
            endpoint=endpoint,
            filters=filters,
            run_type=compact(payload.get("run_type")) or "manual",
            dry_run=bool(payload.get("dry_run", False)),
            max_pages=max_pages,
            max_records=max_records,
            fetch_details=bool(payload.get("fetch_details", True)),
            company_profile=payload.get("company_profile") if isinstance(payload.get("company_profile"), dict) else {},
        )))
    counters = {
        key: sum(int(result.get(key, 0)) for result in results)
        for key in ("fetched", "inserted", "updated", "skipped", "failed")
    }
    return {
        "run_id": results[0]["run_id"] if len(results) == 1 else None,
        "run_ids": [result["run_id"] for result in results],
        "status": "partial" if counters["failed"] else ("dry_run" if payload.get("dry_run") else "success"),
        "total_fetched": counters["fetched"],
        "total_inserted": counters["inserted"],
        "total_updated": counters["updated"],
        "total_skipped": counters["skipped"],
        "total_failed": counters["failed"],
        "dry_run": bool(payload.get("dry_run", False)),
    }


def internal_opportunities_response(query):
    page = max(int(query.get("pagina") or query.get("page") or 1), 1)
    page_size = min(max(int(query.get("tamanhoPagina") or query.get("page_size") or 10), 1), 100)
    filters = {
        "limit": page_size,
        "offset": (page - 1) * page_size,
        "radar_status": query.get("radar_status") or "new,triage,selected,converted_to_proposal",
        "uf": query.get("uf"),
        "keywords": [term.strip() for term in compact(query.get("palavraChave")).split(";") if term.strip()],
        "object_type": query.get("tipoObjeto"),
        "modality_code": query.get("codigoModalidadeContratacao"),
        "purchase_number": query.get("numeroCompra"),
        "uasg": query.get("uasg"),
        "proposal_from": _etl_query_date(query.get("dataInicial")),
        "proposal_to": _etl_query_date(query.get("dataFinal"), end_of_day=True),
        "include_missing_proposal_dates": str(
            query.get("incluirSemDataEncerramento") or ""
        ).strip().lower() in {"1", "true", "yes", "sim"},
        "score_min": query.get("score_min"),
    }
    payload = etl_repository().list_opportunities(filters)
    results = []
    for row in payload["items"]:
        cnpj = compact(row.get("source_cnpj") or row.get("buyer_cnpj"))
        year = row.get("year") or 0
        sequence = row.get("sequence") or 0
        link = safe_public_url(row.get("detail_url") or row.get("source_url"))
        if not link and cnpj and year and sequence:
            link = pncp_app_link(cnpj, year, sequence)
        results.append({
            "id": row["id"],
            "score": row.get("score", 0),
            "radar_status": row.get("radar_status", "new"),
            "orgao": compact(row.get("buyer_name")),
            "cnpj": cnpj,
            "ano": year,
            "sequencial": sequence,
            "numeroCompra": compact(row.get("title")),
            "processo": compact(row.get("process_number")),
            "modalidade": compact(row.get("modality")),
            "objeto": compact(row.get("description") or row.get("title")),
            "uf": compact(row.get("uf")),
            "municipio": compact(row.get("city")),
            "unidade": "",
            "codigoUnidade": compact(row.get("uasg")),
            "valorTotalEstimado": row.get("estimated_value"),
            "modoDisputa": "",
            "situacao": compact(row.get("status")),
            "linkOrigem": safe_public_url(row.get("origin_url")),
            "abertura": compact(row.get("proposal_start_at")),
            "encerramento": compact(row.get("proposal_end_at")),
            "dataEncerramentoInformada": bool(compact(row.get("proposal_end_at"))),
            "link": link,
        })
    total = int(payload["total"])
    total_pages = (total + page_size - 1) // page_size if total else 0
    return {
        "results": results,
        "total": total,
        "source_total": total,
        "pagina": page,
        "tamanhoPagina": page_size,
        "total_pages": total_pages,
        "has_previous": page > 1,
        "has_next": page < total_pages,
        "complete": True,
        "searching": False,
        "source": "sqlite-radar",
    }


def _payload_records(payload):
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []
    for key in ("data", "items", "content", "results", "resultado"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
        if isinstance(value, dict):
            records = _payload_records(value)
            if records:
                return records
    return []


def enrich_detail_documents(opportunity_id, current_detail):
    if not ALLOW_DETAIL_DOCUMENT_ON_DEMAND:
        return current_detail
    row = current_detail["opportunity"]
    cnpj = compact(row.get("source_cnpj") or row.get("buyer_cnpj"))
    year = row.get("year")
    sequence = row.get("sequence")
    if not (cnpj and year and sequence):
        return current_detail

    with DETAIL_DOCUMENT_ENRICHMENT_LOCK:
        fresh = etl_repository().get_opportunity(opportunity_id)
        if fresh and fresh.get("documents"):
            return fresh

        repository = etl_repository()
        run_id = repository.create_run("pncp", "detail_documents", {
            "opportunity_id": opportunity_id,
            "cnpj": cnpj,
            "year": year,
            "sequence": sequence,
        })
        counters = {"fetched": 1, "inserted": 0, "updated": 0, "skipped": 0, "failed": 0}
        try:
            fetched = PNCPConnector().fetch_documents(cnpj, year, sequence)
            raw_documents = _payload_records(fetched.payload)
            documents = PNCPMapper().map_documents(raw_documents, default_source="pncp")
            repository.replace_opportunity_documents(opportunity_id, documents)
            counters["updated"] = 1
            repository.finish_run(run_id, status="success", counters=counters)
        except Exception as exc:
            counters["failed"] = 1
            repository.finish_run(run_id, status="failed", counters=counters, error_message=str(exc))
            raise
    return etl_repository().get_opportunity(opportunity_id) or current_detail


def internal_opportunity_detail(opportunity_id):
    detail = etl_repository().get_opportunity(opportunity_id)
    if detail is None:
        raise FileNotFoundError("Oportunidade nao localizada no radar interno.")
    enrichment_messages = []
    if not detail["documents"]:
        if ALLOW_DETAIL_DOCUMENT_ON_DEMAND:
            try:
                detail = enrich_detail_documents(opportunity_id, detail)
            except Exception as exc:
                enrichment_messages.append(f"Arquivos oficiais pendentes: {exc}")
        else:
            enrichment_messages.append("Arquivos oficiais ainda nao carregados no banco local.")
    if not detail["items"]:
        if ALLOW_DETAIL_ITEMS_ON_DEMAND:
            row = detail["opportunity"]
            cnpj = compact(row.get("source_cnpj") or row.get("buyer_cnpj"))
            year = row.get("year")
            sequence = row.get("sequence")
            if cnpj and year and sequence:
                try:
                    enrich_opportunity_items(cnpj, year, sequence)
                    detail = etl_repository().get_opportunity(opportunity_id) or detail
                except Exception as exc:
                    enrichment_messages.append(f"Itens oficiais pendentes: {exc}")
            else:
                enrichment_messages.append(
                    "Itens ainda nao carregados: a oportunidade nao possui identidade PNCP completa."
                )
        else:
            enrichment_messages.append("Itens oficiais ainda nao carregados no banco local.")
    if enrichment_messages:
        detail["enrichment_error"] = " ".join(enrichment_messages)
    row = detail["opportunity"]
    items = [{
        "numero": compact(item.get("item_number")),
        "lote": compact(item.get("lot_number")),
        "descricao": compact(item.get("description") or item.get("title")),
        "quantidade": format_pncp_quantity(item.get("quantity")),
        "unidade": compact(item.get("unit")) or STANDARD_UNIT,
        "valor_unitario_estimado": item.get("estimated_unit_value"),
        "valor_total_estimado": item.get("estimated_total_value"),
        "criterio_julgamento": compact(item.get("technical_object")),
        "situacao": compact(item.get("status")),
        "tipo": compact(item.get("title")),
        "granularity": compact(item.get("granularity")),
        "confidence": item.get("confidence"),
    } for item in detail["items"]]
    link = safe_public_url(row.get("detail_url") or row.get("source_url"))
    opportunity = {
        "numero_compra": compact(row.get("title")),
        "processo": compact(row.get("process_number")),
        "modalidade": compact(row.get("modality")),
        "objeto": compact(row.get("description") or row.get("title")),
        "orgao": compact(row.get("buyer_name")),
        "orgao_cnpj": compact(row.get("buyer_cnpj") or row.get("source_cnpj")),
        "unidade": "",
        "codigo_unidade": compact(row.get("uasg")),
        "municipio": compact(row.get("city")),
        "uf": compact(row.get("uf")),
        "abertura": compact(row.get("proposal_start_at")),
        "encerramento": compact(row.get("proposal_end_at")),
        "situacao": compact(row.get("status")),
        "valor_total_estimado": row.get("estimated_value"),
        "modo_disputa": "",
        "link_pncp": link,
        "link_origem": safe_public_url(row.get("origin_url")),
        "portal_origem": opportunity_source_portal(row.get("origin_url") or link),
        "categorias": opportunity_categories(row.get("description") or row.get("title") or "", items),
        "numero_controle_pncp": compact(row.get("pncp_control_number")),
    }
    return {
        "oportunidade": opportunity,
        "arquivos": [{
            "titulo": compact(document.get("title") or document.get("filename") or "Arquivo oficial"),
            "tipo": compact(document.get("document_type")),
            "url": compact(document.get("url")),
        } for document in detail["documents"]],
        "itens": items,
        "fontes": {
            "oportunidade": "Banco interno normalizado",
            "arquivos": "Auditoria ETL do PNCP",
            "itens": "Auditoria ETL do PNCP",
        },
        "aviso_enriquecimento": detail.get("enrichment_error", ""),
    }


def enrich_internal_opportunity(opportunity_id, current_detail):
    row = current_detail["opportunity"]
    cnpj = compact(row.get("source_cnpj") or row.get("buyer_cnpj"))
    year = row.get("year")
    sequence = row.get("sequence")
    if not (cnpj and year and sequence):
        return current_detail
    connector = PNCPConnector()
    detail_payload = connector.fetch_detail(cnpj, year, sequence)
    raw_items = []
    item_pages = []
    for page in connector.iter_items(cnpj, year, sequence, 20):
        raw_items.extend(page.records)
        item_pages.append(page.raw_payload)
    documents_payload = connector.fetch_documents(cnpj, year, sequence)
    documents = documents_payload.payload
    if isinstance(documents, dict):
        documents = (
            documents.get("data")
            or documents.get("items")
            or documents.get("content")
            or []
        )
    if not isinstance(documents, list):
        documents = []
    listing = {
        "numeroControlePNCP": row.get("pncp_control_number"),
        "numeroCnpj": cnpj,
        "anoCompra": year,
        "sequencialCompra": sequence,
        "numeroCompra": row.get("title"),
        "processo": row.get("process_number"),
        "objetoCompra": row.get("description"),
    }
    mapper = PNCPMapper()
    normalized = mapper.map(
        listing,
        detail=detail_payload.payload if isinstance(detail_payload.payload, dict) else None,
        items=raw_items,
        documents=documents,
    )
    repository = etl_repository()
    run_id = repository.create_run("pncp", "detail", {
        "opportunity_id": opportunity_id,
        "cnpj": cnpj,
        "year": year,
        "sequence": sequence,
    })
    counters = {"fetched": 1, "inserted": 0, "updated": 0, "skipped": 0, "failed": 0}
    try:
        outcome, _ = repository.persist_record(
            run_id=run_id,
            source_endpoint="detail",
            request_url=detail_payload.request_url,
            raw_payload={
                "listing": listing,
                "detail": detail_payload.payload,
                "items": item_pages,
                "documents": documents_payload.payload,
            },
            opportunity=normalized,
            match=OpportunityClassifier().classify(normalized, {}),
        )
        counters[outcome] += 1
        repository.finish_run(run_id, status="success", counters=counters)
    except Exception as exc:
        counters["failed"] = 1
        repository.finish_run(run_id, status="failed", counters=counters, error_message=str(exc))
        raise
    return repository.get_opportunity(opportunity_id) or current_detail


def convert_internal_opportunity(opportunity_id, payload):
    detail = internal_opportunity_detail(opportunity_id)
    selected_items = payload.get("itens") if isinstance(payload, dict) else None
    result = import_business({
        "pncp_link": detail["oportunidade"]["link_pncp"],
        "empresa": compact(payload.get("empresa")) if isinstance(payload, dict) else "",
        "itens": selected_items,
        "oportunidade": detail["oportunidade"],
    })
    etl_repository().update_radar_status(
        opportunity_id,
        "converted_to_proposal",
        int(result["negocio"]["id"]),
    )
    return result


class App(BaseHTTPRequestHandler):
    def end_headers(self):
        self.send_header("X-Content-Type-Options", "nosniff")
        self.send_header("X-Frame-Options", "SAMEORIGIN")
        self.send_header("Referrer-Policy", "no-referrer")
        self.send_header("Cross-Origin-Resource-Policy", "same-origin")
        self.send_header("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
        self.send_header(
            "Content-Security-Policy",
            "default-src 'self'; base-uri 'self'; object-src 'none'; frame-ancestors 'self'; "
            "form-action 'self'; connect-src 'self'; img-src 'self' data: blob:; "
            "style-src 'self' 'unsafe-inline'; script-src 'self' 'unsafe-inline'; frame-src 'self'",
        )
        super().end_headers()

    def request_allowed(self, *, mutation=False):
        if not local_request_host(self.headers.get("Host", "")):
            json_response(self, 403, {"error": "Host local inválido."})
            return False
        request_path = urlparse(self.path).path
        protected_get = request_path.startswith((
            "/api/", "/internal/", "/identify-items", "/pncp-search",
            "/catalog/draft", "/proposal-preview", "/download/", "/template/",
        ))
        if (
            self.headers.get("Sec-Fetch-Site", "").strip().lower() == "cross-site"
            and (mutation or protected_get)
        ):
            json_response(self, 403, {"error": "Requisição externa não permitida."})
            return False
        if mutation and not local_request_origin(self.headers.get("Origin", "")):
            json_response(self, 403, {"error": "Origem não permitida."})
            return False
        if len(self.path) > 8192:
            json_response(self, 414, {"error": "Endereço da requisição excede o limite permitido."})
            return False
        return True

    def do_GET(self):
        if not self.request_allowed():
            return
        request_path = urlparse(self.path).path
        if request_path in {"/internal/opportunities", "/api/internal/opportunities"}:
            try:
                query = {
                    key: values[0]
                    for key, values in parse_qs(urlparse(self.path).query).items()
                }
                started_at = time.perf_counter()
                payload = internal_opportunities_response(query)
                elapsed_ms = (time.perf_counter() - started_at) * 1000
                json_response(
                    self,
                    200,
                    payload,
                    {"Server-Timing": f"opportunity-search;dur={elapsed_ms:.1f}"},
                )
            except ValueError as exc:
                json_response(self, 422, {"error": str(exc)})
            except Exception:
                json_response(self, 500, {"error": "Nao foi possivel consultar o radar interno."})
            return
        internal_detail_match = re.fullmatch(
            r"/(?:api/)?internal/opportunities/([a-f0-9]{32})",
            request_path,
        )
        if internal_detail_match:
            try:
                json_response(self, 200, internal_opportunity_detail(internal_detail_match.group(1)))
            except FileNotFoundError as exc:
                json_response(self, 404, {"error": str(exc)})
            except Exception:
                json_response(self, 500, {"error": "Nao foi possivel carregar a oportunidade."})
            return
        if request_path in {"/internal/etl/runs", "/api/internal/etl/runs"}:
            query = parse_qs(urlparse(self.path).query)
            limit = int((query.get("limit") or ["50"])[0])
            offset = int((query.get("offset") or ["0"])[0])
            json_response(self, 200, etl_repository().list_runs(limit, offset))
            return
        if request_path == "/api/kanban":
            json_response(self, 200, kanban_store.board(DATABASE_PATH))
            return
        history_match = re.fullmatch(r"/api/kanban/proposals/(\d+)/history", request_path)
        if history_match:
            json_response(self, 200, {"history": kanban_store.history(DATABASE_PATH, history_match.group(1))})
            return
        if self.path == "/" or self.path.startswith("/?"):
            frontend_response(self, self.path)
            return
        if self.path == "/react" or self.path.startswith("/react?"):
            frontend_response(self, self.path)
            return
        if self.path == "/legacy" or self.path.startswith("/legacy?"):
            html_response(self, render_page())
            return
        if urlparse(self.path).path.startswith("/assets/"):
            frontend_response(self, self.path)
            return
        if self.path == "/templates" or self.path.startswith("/templates?"):
            self.send_response(302)
            self.send_header("Location", "/")
            self.send_header("Content-Length", "0")
            self.end_headers()
            return
        if self.path == "/api/templates" or self.path.startswith("/api/templates?"):
            json_response(self, 200, {"templates": list_templates()})
            return
        if self.path == "/api/responsaveis" or self.path.startswith("/api/responsaveis?"):
            json_response(self, 200, {"responsaveis": list_responsibles()})
            return
        request_path = urlparse(self.path).path
        if request_path == "/api/negocios":
            query = parse_qs(urlparse(self.path).query)
            include_archived = (query.get("arquivados") or ["0"])[0] == "1"
            json_response(
                self,
                200,
                {
                    "negocios": list_businesses(include_archived),
                    "etapas": list(BUSINESS_STAGES),
                },
            )
            return
        if request_path == "/api/oportunidades/detalhe":
            try:
                query = parse_qs(urlparse(self.path).query)
                link = (query.get("pncp_link") or [""])[0]
                fallback = {
                    key: (query.get(key) or [""])[0]
                    for key in (
                        "numero_compra", "processo", "modalidade", "objeto",
                        "orgao", "unidade", "municipio", "uf", "abertura",
                        "encerramento", "situacao", "modo_disputa",
                        "codigo_unidade", "link_sistema_origem",
                        "valor_total_estimado",
                    )
                }
                json_response(
                    self,
                    200,
                    opportunity_detail_from_pncp_link(link, fallback),
                )
            except ValueError as exc:
                json_response(self, 400, {"error": str(exc) or "Link PNCP inválido."})
            except Exception:
                json_response(self, 500, {"error": "Não foi possível carregar a oportunidade."})
            return
        business_match = re.fullmatch(r"/api/negocios/(\d+)", request_path)
        if business_match:
            try:
                json_response(
                    self,
                    200,
                    {"negocio": get_business(business_match.group(1), True)},
                )
            except Exception as exc:
                business_api_error(self, exc)
            return
        if self.path.startswith("/pncp-search"):
            try:
                query = {
                    key: values[0]
                    for key, values in parse_qs(urlparse(self.path).query).items()
                }
                if not ENABLE_PNCP_SEARCH:
                    page = max(int(query.get("pagina") or 1), 1)
                    page_size = min(max(int(query.get("tamanhoPagina") or 10), 1), 100)
                    json_response(self, 200, {
                        "results": [],
                        "total": 0,
                        "source_total": 0,
                        "pagina": page,
                        "tamanhoPagina": page_size,
                        "total_pages": 0,
                        "has_previous": False,
                        "has_next": False,
                        "complete": True,
                        "searching": False,
                        "source": "pncp-disabled",
                        "reconciliation": {
                            "status": "disabled",
                            "inserted": 0,
                            "updated": 0,
                            "failed": 0,
                        },
                    })
                    return
                if str(query.get("rapido") or "") == "1":
                    result = search_pncp_open_bids_fast(query)
                else:
                    result = search_pncp_open_bids(query)
                json_response(self, 200, result)
            except ValueError as exc:
                json_response(self, 400, {"error": str(exc) or "Revise os filtros da consulta."})
            except Exception:
                json_response(self, 500, {"error": "Não foi possível consultar o PNCP."})
            return
        if self.path.startswith("/identify-items"):
            try:
                query = {
                    key: values[0]
                    for key, values in parse_qs(urlparse(self.path).query).items()
                }
                link = query.get("pncp_link", "")
                json_response(self, 200, identify_items_from_pncp_link(link))
            except ValueError as exc:
                json_response(self, 400, {"error": str(exc) or "Link PNCP inválido."})
            except Exception:
                json_response(self, 500, {"error": "Não foi possível identificar os itens."})
            return
        if self.path.startswith("/catalog/draft"):
            try:
                query = {
                    key: values[0]
                    for key, values in parse_qs(urlparse(self.path).query).items()
                }
                json_response(
                    self,
                    200,
                    catalog_draft_from_pncp_link(
                        query.get("pncp_link", ""),
                        query.get("item_key", ""),
                    ),
                )
            except ValueError as exc:
                json_response(self, 422, {"error": str(exc)})
            except Exception:
                json_response(
                    self,
                    500,
                    {"error": "Não foi possível estruturar o catálogo."},
                )
            return
        preview_match = re.fullmatch(
            r"/proposal-preview/([a-f0-9]{32})\.pdf",
            urlparse(self.path).path,
        )
        if preview_match:
            proposal_preview_response(self, preview_match.group(1))
            return
        if self.path.startswith("/download/"):
            name = unquote(self.path.split("/download/", 1)[1])
            file_response(self, OUTPUT_DIR / Path(name).name)
            return
        if self.path.startswith("/template/"):
            template_key = unquote(urlparse(self.path).path.split("/template/", 1)[1])
            template_response(self, template_key)
            return
        self.send_error(404)

    def do_POST(self):
        if not self.request_allowed(mutation=True):
            return
        request_path = urlparse(self.path).path
        if request_path in {"/internal/etl/pncp-sync", "/api/internal/etl/pncp-sync"}:
            expected_token = os.environ.get("TOTH_ETL_ADMIN_TOKEN", "")
            if expected_token and self.headers.get("X-ETL-Token", "") != expected_token:
                json_response(self, 403, {"error": "Token administrativo do ETL invalido."})
                return
            try:
                json_response(self, 200, sync_internal_pncp(parse_json_body(self)))
            except ValueError as exc:
                json_response(self, 422, {"error": str(exc)})
            except Exception as exc:
                json_response(self, 502, {"error": str(exc) or "Nao foi possivel sincronizar o PNCP."})
            return
        internal_ignore_match = re.fullmatch(
            r"/(?:api/)?internal/opportunities/([a-f0-9]{32})/ignore",
            request_path,
        )
        if internal_ignore_match:
            if not etl_repository().update_radar_status(internal_ignore_match.group(1), "ignored"):
                json_response(self, 404, {"error": "Oportunidade nao localizada."})
            else:
                json_response(self, 200, {"id": internal_ignore_match.group(1), "radar_status": "ignored"})
            return
        internal_convert_match = re.fullmatch(
            r"/(?:api/)?internal/opportunities/([a-f0-9]{32})/convert-to-proposal",
            request_path,
        )
        if internal_convert_match:
            try:
                result = convert_internal_opportunity(
                    internal_convert_match.group(1),
                    parse_json_body(self),
                )
                json_response(self, 201 if result.get("criado") else 200, result)
            except Exception as exc:
                business_api_error(self, exc)
            return
        if request_path == "/api/oportunidades/conversar":
            try:
                payload = parse_json_body(self)
                result = answer_opportunity_question(
                    payload.get("pncp_link", ""),
                    payload.get("pergunta", ""),
                )
                json_response(self, 200, result)
            except ValueError as exc:
                json_response(self, 400, {"error": str(exc)})
            except Exception as exc:
                json_response(self, 500, {"error": str(exc) or "Não foi possível consultar o edital."})
            return
        if request_path == "/api/kanban/columns":
            try:
                json_response(self, 201, {"column": kanban_store.create_column(DATABASE_PATH, parse_json_body(self))})
            except Exception as exc:
                business_api_error(self, exc)
            return
        if request_path == "/api/kanban/proposals":
            try:
                proposal = kanban_store.save_proposal(DATABASE_PATH, parse_json_body(self))
                sync_business_position_from_proposal(proposal)
                json_response(self, 201, {"proposal": proposal})
            except Exception as exc:
                business_api_error(self, exc)

            return
        if request_path == "/api/negocios/importar":
            try:
                result = import_business(parse_json_body(self))
                json_response(self, 201 if result["criado"] else 200, result)
            except Exception as exc:
                business_api_error(self, exc)
            return
        business_task_match = re.fullmatch(
            r"/api/negocios/(\d+)/tarefas",
            request_path,
        )
        if business_task_match:
            try:
                result = create_business_task(
                    business_task_match.group(1),
                    parse_json_body(self),
                )
                json_response(self, 201, result)
            except Exception as exc:
                business_api_error(self, exc)
            return
        if request_path == "/api/responsaveis":
            try:
                responsible = create_responsible(parse_json_body(self))
                json_response(self, 201, {"responsavel": responsible})
            except Exception as exc:
                responsible_api_error(self, exc)
            return
        if request_path == "/api/templates":
            try:
                form = parse_template_upload_form(self)
                template = store_new_template(template_upload_field(form))
                json_response(self, 201, {"template": template})
            except Exception as exc:
                template_api_error(self, exc)
            return

        if request_path == "/catalog/generate":
            try:
                content_length = int(self.headers.get("Content-Length", "0") or 0)
                if content_length <= 0:
                    raise ValueError("Envie os dados do catálogo.")
                if content_length > MAX_CATALOG_REQUEST_SIZE:
                    raise OverflowError("O catálogo excede o limite total de 80 MB.")
                form = cgi.FieldStorage(
                    fp=self.rfile,
                    headers=self.headers,
                    environ={
                        "REQUEST_METHOD": "POST",
                        "CONTENT_TYPE": self.headers.get("Content-Type"),
                        "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                    },
                )
                data_field = catalog_asset_field(form, "data")
                if data_field is None:
                    raise ValueError("Os dados estruturados do catálogo não foram enviados.")
                payload = json.loads(data_field.value)
                saved_assets = save_catalog_assets(form, payload)
                response = generate_catalog_exports(payload.get("data"), saved_assets)
                json_response(self, 200, response)
            except json.JSONDecodeError:
                json_response(
                    self,
                    422,
                    {"error": "Os dados do catálogo não são um JSON válido."},
                )
            except OverflowError as exc:
                json_response(self, 413, {"error": str(exc)})
            except ValueError as exc:
                response = {"error": str(exc)}
                if hasattr(exc, "catalog_alerts"):
                    response["alerts"] = exc.catalog_alerts
                json_response(self, 422, response)
            except Exception:
                json_response(
                    self,
                    500,
                    {"error": "Não foi possível gerar o catálogo."},
                )
            return

        replace_match = re.fullmatch(r"/api/templates/([^/]+)/replace", request_path)
        if replace_match:
            try:
                template_id = unquote(replace_match.group(1))
                form = parse_template_upload_form(self)
                template = replace_template_file(template_id, template_upload_field(form))
                json_response(self, 200, {"template": template})
            except Exception as exc:
                template_api_error(self, exc)
            return

        if request_path not in {"/process", "/generate", "/proposal-preview"}:
            self.send_error(404)
            return

        try:
            if request_path in {"/generate", "/proposal-preview"}:
                payload = parse_json_body(self, MAX_JSON_REQUEST_SIZE)
                context = proposal_generation_context(payload)
                if request_path == "/proposal-preview":
                    json_response(self, 200, create_proposal_preview(context))
                    return

                out_name = (
                    f"Proposta_Final_{context['source_name']}_{int(time.time())}.docx"
                )
                out_path = OUTPUT_DIR / out_name
                build_docx(
                    context["items"],
                    context["template_path"],
                    out_path,
                    responsible=context["responsible"],
                    commercial_terms=context["commercial_terms"],
                )
                try:
                    record_generated_document(context["responsible_id"], out_path)
                except Exception:
                    out_path.unlink(missing_ok=True)
                    raise
                json_response(
                    self,
                    200,
                    {
                        "download_url": f"/download/{out_name}",
                        "filename": out_name,
                    },
                )
                return

            content_length = int(self.headers.get("Content-Length", "0") or 0)
            if content_length <= 0:
                raise ValueError("Envie os dados da proposta.")
            if content_length > MAX_PROPOSAL_REQUEST_SIZE:
                raise OverflowError("A proposta excede o limite total permitido.")
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type"),
                    "CONTENT_LENGTH": self.headers.get("Content-Length", "0"),
                },
            )
            pncp_link_field = form["pncp_link"] if "pncp_link" in form else None
            pncp_link = pncp_link_field.value.strip() if pncp_link_field is not None else ""
            pncp_info = None
            if not pncp_link:
                json_response(self, 400, {"error": "Informe o link do edital no PNCP para baixar o arquivo oficial."})
                return
            responsible_field = form["responsible_id"] if "responsible_id" in form else None
            responsible_id = compact(responsible_field.value if responsible_field is not None else "")
            if not responsible_id or not resolve_responsible(responsible_id):
                json_response(self, 422, {"error": "Selecione um responsável válido pela proposta."})
                return
            wanted_field = form["wanted_items"] if "wanted_items" in form else None
            wanted_items = compact(wanted_field.value if wanted_field is not None else "")
            if not wanted_items:
                json_response(self, 422, {"error": "Selecione pelo menos um item para a proposta."})
                return
            process_data = proposal_process_from_structured_items(pncp_link, wanted_items)

            template_choice_field = form["template_choice"] if "template_choice" in form else None
            template_choice = compact(template_choice_field.value if template_choice_field is not None else "")
            template = form["template_file"] if "template_file" in form else None
            if template is not None and template.filename:
                template_path, _ = save_upload(template, "modelo")
                template_ref = template_path.name
            else:
                selected_template = template_path_from_name(template_choice)
                if not selected_template:
                    json_response(self, 422, {"error": "Modelo Word inválido."})
                    return
                if not selected_template.exists():
                    json_response(self, 422, {"error": f"O modelo {template_choice} não está disponível."})
                    return
                template_ref = f"managed:{template_choice}"

            process_data["template_ref"] = template_ref
            json_response(self, 200, process_data)
        except json.JSONDecodeError:
            json_response(self, 422, {"error": "Os dados enviados não são um JSON válido."})
        except OverflowError as exc:
            json_response(self, 413, {"error": str(exc)})
        except ValueError as exc:
            json_response(self, 422, {"error": str(exc) or "Revise os dados enviados."})
        except Exception:
            LOGGER.exception("Proposal request failed for %s", request_path)
            json_response(self, 500, {"error": "Não foi possível processar a proposta."})

    def do_PUT(self):
        if not self.request_allowed(mutation=True):
            return
        request_path = urlparse(self.path).path
        column_match = re.fullmatch(r"/api/kanban/columns/(\d+)", request_path)
        if column_match:
            try:
                json_response(self, 200, {"column": kanban_store.update_column(DATABASE_PATH, column_match.group(1), parse_json_body(self))})
            except Exception as exc:
                business_api_error(self, exc)
            return
        move_match = re.fullmatch(r"/api/kanban/proposals/(\d+)/move", request_path)
        if move_match:
            try:
                payload = parse_json_body(self)
                json_response(self, 200, {"proposal": kanban_store.move_proposal(DATABASE_PATH, move_match.group(1), payload.get("column_id"))})
            except Exception as exc:
                business_api_error(self, exc)
            return
        proposal_match = re.fullmatch(r"/api/kanban/proposals/(\d+)", request_path)
        if proposal_match:
            try:
                proposal = kanban_store.save_proposal(DATABASE_PATH, parse_json_body(self), proposal_match.group(1))
                sync_business_position_from_proposal(proposal)
                json_response(self, 200, {"proposal": proposal})
            except Exception as exc:
                business_api_error(self, exc)
            return
        business_task_match = re.fullmatch(
            r"/api/negocios/(\d+)/tarefas/(\d+)",
            request_path,
        )
        if business_task_match:
            try:
                result = update_business_task(
                    business_task_match.group(1),
                    business_task_match.group(2),
                    parse_json_body(self),
                )
                json_response(self, 200, {"negocio": result})
            except Exception as exc:
                business_api_error(self, exc)
            return
        business_match = re.fullmatch(r"/api/negocios/(\d+)", request_path)
        if business_match:
            try:
                result = update_business(
                    business_match.group(1),
                    parse_json_body(self),
                )
                json_response(self, 200, {"negocio": result})
            except Exception as exc:
                business_api_error(self, exc)
            return
        responsible_match = re.fullmatch(r"/api/responsaveis/(\d+)", request_path)
        if not responsible_match:
            self.send_error(404)
            return
        try:
            responsible = update_responsible(
                responsible_match.group(1), parse_json_body(self)
            )
            json_response(self, 200, {"responsavel": responsible})
        except Exception as exc:
            responsible_api_error(self, exc)

    def do_DELETE(self):
        if not self.request_allowed(mutation=True):
            return
        request_path = urlparse(self.path).path
        proposal_match = re.fullmatch(r"/api/kanban/proposals/(\d+)", request_path)
        if proposal_match:
            try:
                proposal = kanban_store.delete_proposal(DATABASE_PATH, proposal_match.group(1))
                clear_business_position_from_deleted_proposal(proposal)
                json_response(self, 200, {"deleted": proposal_match.group(1)})
            except Exception as exc:
                business_api_error(self, exc)
            return
        column_match = re.fullmatch(r"/api/kanban/columns/(\d+)", request_path)
        if column_match:
            try:
                kanban_store.delete_column(DATABASE_PATH, column_match.group(1))
                json_response(self, 200, {"deleted": column_match.group(1)})
            except Exception as exc:
                business_api_error(self, exc)
            return
        responsible_match = re.fullmatch(r"/api/responsaveis/(\d+)", request_path)
        if responsible_match:
            try:
                responsible_id = responsible_match.group(1)
                delete_responsible(responsible_id)
                json_response(self, 200, {"deleted": responsible_id})
            except Exception as exc:
                responsible_api_error(self, exc)
            return
        delete_match = re.fullmatch(r"/api/templates/([^/]+)", request_path)
        if not delete_match:
            self.send_error(404)
            return
        try:
            template_id = unquote(delete_match.group(1))
            delete_template_file(template_id)
            json_response(self, 200, {"deleted": template_id})
        except Exception as exc:
            template_api_error(self, exc)

    def log_message(self, format, *args):
        return


def main():
    init_database()
    cleanup_proposal_previews()
    server = ThreadingHTTPServer(("127.0.0.1", PORT), App)
    print(f"Aplicação aberta em http://127.0.0.1:{PORT}")
    print("Pressione Ctrl+C para encerrar.")
    server.serve_forever()


if __name__ == "__main__":
    main()
