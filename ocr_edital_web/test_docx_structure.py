import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from docx_structure import (
    GENERATED_TABLE_BLOCK_ID,
    inspect_docx_structure,
    rebuild_docx_with_mini_box_order,
    resolve_document_block_layout,
    validate_document_block_order,
    validate_mini_box_alignments,
    validate_mini_box_order,
)
import server


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


class DocxStructureTests(unittest.TestCase):
    def create_marker_template(self, root: Path) -> Path:
        image_path = root / "logo.png"
        Image.new("RGB", (8, 8), (255, 128, 64)).save(image_path)

        path = root / "modelo.docx"
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Prefixo ")
        paragraph.add_run("{Alpha ")
        paragraph.add_run("{interno}")
        paragraph.add_run("}")
        paragraph.add_run(" entre ")
        paragraph.add_run("{Beta}")
        paragraph.add_run(" sufixo")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Antes {Tabela} depois"
        document.sections[0].header.paragraphs[0].text = "Topo {Cabecalho} fixo"
        document.add_picture(str(image_path))
        document.save(path)
        return path

    def test_inspection_builds_ast_for_root_markers_split_across_runs(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = self.create_marker_template(Path(temp_dir))

            structure = inspect_docx_structure(path)

        mini_boxes = [node for node in structure["nodes"] if node["type"] == "MINI_BOX"]
        fixed_text = [node["content"] for node in structure["nodes"] if node["type"] == "FIXED_TEXT"]
        self.assertEqual(
            [node["content"] for node in mini_boxes],
            ["Alpha {interno}", "Beta", "Tabela", "Cabecalho"],
        )
        self.assertEqual([node["order"] for node in mini_boxes], [0, 1, 2, 3])
        self.assertEqual(structure["mini_box_count"], 4)
        self.assertEqual(
            structure["generated_table_block"],
            {
                "id": GENERATED_TABLE_BLOCK_ID,
                "type": "GENERATED_TABLE",
                "content": "Tabela gerada da proposta",
            },
        )
        self.assertIn("Prefixo ", fixed_text)
        self.assertIn(" entre ", fixed_text)
        self.assertIn(" sufixo", fixed_text)
        self.assertIn("Antes ", fixed_text)
        self.assertIn(" depois", fixed_text)
        self.assertEqual(structure["warnings"], [])

    def test_rebuild_reorders_only_marker_contents_and_preserves_assets(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = self.create_marker_template(root)
            output = root / "reordenado.docx"
            structure = inspect_docx_structure(source)
            mini_boxes = [node for node in structure["nodes"] if node["type"] == "MINI_BOX"]

            with zipfile.ZipFile(source, "r") as archive:
                original_image = archive.read("word/media/image1.png")
                original_relationships = archive.read("word/_rels/document.xml.rels")

            rebuild_docx_with_mini_box_order(
                source,
                output,
                [node["id"] for node in reversed(mini_boxes)],
            )

            rebuilt = Document(output)
            with zipfile.ZipFile(output, "r") as archive:
                rebuilt_image = archive.read("word/media/image1.png")
                rebuilt_relationships = archive.read("word/_rels/document.xml.rels")

        self.assertEqual(
            paragraph_text(rebuilt.paragraphs[0]),
            "Prefixo {Cabecalho} entre {Tabela} sufixo",
        )
        self.assertEqual(rebuilt.tables[0].cell(0, 0).text, "Antes {Beta} depois")
        self.assertEqual(
            rebuilt.sections[0].header.paragraphs[0].text,
            "Topo {Alpha {interno}} fixo",
        )
        self.assertEqual(rebuilt_image, original_image)
        self.assertEqual(rebuilt_relationships, original_relationships)

    def test_rebuild_keeps_individual_alignment_with_the_moved_mini_box(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "alinhamentos.docx"
            output = root / "alinhado.docx"
            document = Document()
            document.add_paragraph("{Alpha}")
            document.add_paragraph("{Beta}")
            document.save(source)
            structure = inspect_docx_structure(source)
            mini_boxes = [node for node in structure["nodes"] if node["type"] == "MINI_BOX"]
            alpha_id, beta_id = [node["id"] for node in mini_boxes]

            rebuild_docx_with_mini_box_order(
                source,
                output,
                [beta_id, alpha_id],
                {alpha_id: "center", beta_id: "left"},
            )
            rebuilt = Document(output)

        self.assertEqual(rebuilt.paragraphs[0].text, "{Beta}")
        self.assertEqual(rebuilt.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(rebuilt.paragraphs[1].text, "{Alpha}")
        self.assertEqual(rebuilt.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_alignment_contract_reports_defaults_and_rejects_invalid_values(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "alinhamentos.docx"
            document = Document()
            document.add_paragraph("{Esquerda}")
            centered = document.add_paragraph("{Centro}")
            centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.save(path)
            structure = inspect_docx_structure(path)
            mini_boxes = [node for node in structure["nodes"] if node["type"] == "MINI_BOX"]

            self.assertEqual(
                [node["text_align"] for node in mini_boxes],
                ["left", "center"],
            )
            self.assertEqual(
                validate_mini_box_alignments(path, {mini_boxes[0]["id"]: "center"}),
                {mini_boxes[0]["id"]: "center"},
            )
            with self.assertRaisesRegex(ValueError, "alinhamentos"):
                validate_mini_box_alignments(path, {mini_boxes[0]["id"]: "diagonal"})
            with self.assertRaisesRegex(ValueError, "estrutura do modelo"):
                validate_mini_box_alignments(path, {"stale-id": "center"})

    def test_rebuild_rejects_different_alignments_in_the_same_paragraph(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = self.create_marker_template(Path(temp_dir))
            output = Path(temp_dir) / "invalido.docx"
            structure = inspect_docx_structure(source)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]

            with self.assertRaisesRegex(ValueError, "mesmo parágrafo"):
                rebuild_docx_with_mini_box_order(
                    source,
                    output,
                    identifiers,
                    {identifiers[0]: "center", identifiers[1]: "left"},
                )

    def test_validation_rejects_missing_duplicate_and_stale_identifiers(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = self.create_marker_template(Path(temp_dir))
            structure = inspect_docx_structure(path)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]

            with self.assertRaisesRegex(ValueError, "estrutura do modelo"):
                validate_mini_box_order(path, identifiers[:-1])
            with self.assertRaisesRegex(ValueError, "duplicados"):
                validate_mini_box_order(path, [identifiers[0]] * len(identifiers))
            with self.assertRaisesRegex(ValueError, "estrutura do modelo"):
                validate_mini_box_order(path, [*identifiers[:-1], "stale-id"])

            document_order = [*identifiers, GENERATED_TABLE_BLOCK_ID]
            self.assertEqual(
                validate_document_block_order(path, document_order),
                document_order,
            )
            with self.assertRaisesRegex(ValueError, "estrutura do modelo"):
                validate_document_block_order(path, identifiers)
            with self.assertRaisesRegex(ValueError, "duplicados"):
                validate_document_block_order(
                    path,
                    [*identifiers, identifiers[-1]],
                )

    def test_unmatched_braces_remain_fixed_and_produce_warnings(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "incompleto.docx"
            document = Document()
            document.add_paragraph("Texto {sem fechamento")
            document.add_paragraph("Texto } sem abertura")
            document.save(path)

            structure = inspect_docx_structure(path)

        self.assertEqual(structure["mini_box_count"], 0)
        self.assertEqual(len(structure["warnings"]), 2)
        self.assertEqual(
            [node["content"] for node in structure["nodes"]],
            ["Texto {sem fechamento", "Texto } sem abertura"],
        )


class ProposalDocxIntegrationTests(unittest.TestCase):
    def create_template(self, root: Path) -> Path:
        path = root / "proposta.docx"
        document = Document()
        document.add_paragraph("Fixo {Primeiro} entre {Segundo} final")
        document.save(path)
        return path

    def proposal_item(self) -> dict:
        return {
            "item": "1",
            "quantidade": "2",
            "unidade": "UND",
            "descricao": "Item de teste",
            "marca": "Marca",
            "valor_unitario": "10,00",
            "valor_total": "20,00",
        }

    def create_separate_marker_template(self, root: Path) -> Path:
        path = root / "proposta_blocos.docx"
        document = Document()
        document.add_paragraph("Antes {Primeiro}")
        document.add_paragraph("Depois {Segundo}")
        document.save(path)
        return path

    def test_build_docx_applies_order_before_adding_proposal_content(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = self.create_template(root)
            output = root / "saida.docx"
            structure = inspect_docx_structure(template)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]

            server.build_docx(
                [self.proposal_item()],
                template,
                output,
                mini_box_order=list(reversed(identifiers)),
            )
            generated = Document(output)

        self.assertEqual(
            generated.paragraphs[0].text,
            "Fixo {Segundo} entre {Primeiro} final",
        )
        self.assertEqual(len(generated.tables), 1)
        self.assertEqual(generated.tables[0].rows[1].cells[3].text, "Item de teste")

    def test_build_docx_places_generated_table_at_the_selected_visual_position(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            template = self.create_separate_marker_template(root)
            output = root / "saida_com_tabela_intermediaria.docx"
            structure = inspect_docx_structure(template)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]
            document_order = [
                identifiers[1],
                GENERATED_TABLE_BLOCK_ID,
                identifiers[0],
            ]

            layout = resolve_document_block_layout(template, document_order)
            self.assertEqual(list(layout.mini_box_order), [identifiers[1], identifiers[0]])
            self.assertEqual(layout.table_paragraph_index, 1)

            server.build_docx(
                [self.proposal_item()],
                template,
                output,
                mini_box_order=list(layout.mini_box_order),
                document_block_order=document_order,
                mini_box_alignments={
                    identifiers[0]: "center",
                    identifiers[1]: "left",
                },
            )
            generated = Document(output)
            body_children = list(generated._element.body)
            first_blocks = [child.tag.rsplit("}", 1)[-1] for child in body_children[:3]]

        self.assertEqual(first_blocks, ["p", "tbl", "p"])
        self.assertEqual(generated.paragraphs[0].text, "Antes {Segundo}")
        self.assertEqual(generated.paragraphs[1].text, "Depois {Primeiro}")
        self.assertEqual(generated.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(generated.paragraphs[1].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(generated.tables[0].rows[1].cells[3].text, "Item de teste")

    def test_table_cannot_split_two_markers_from_the_same_paragraph(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = self.create_template(Path(temp_dir))
            structure = inspect_docx_structure(template)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]

            with self.assertRaisesRegex(ValueError, "mesmo parágrafo"):
                resolve_document_block_layout(
                    template,
                    [identifiers[0], GENERATED_TABLE_BLOCK_ID, identifiers[1]],
                )

    def test_generation_context_validates_and_fingerprints_selected_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = self.create_template(Path(temp_dir))
            structure = inspect_docx_structure(template)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]
            payload = {
                "items": [self.proposal_item()],
                "template_ref": "managed:proposta.docx",
                "source_name": "edital",
                "responsible_id": "1",
                "commercial_terms": {},
                "mini_box_order": identifiers,
                "mini_box_alignments": {identifiers[0]: "center"},
            }
            with (
                patch.object(server, "resolve_template", return_value=template),
                patch.object(server, "resolve_responsible", return_value={"id": "1"}),
            ):
                original = server.proposal_generation_context(payload)
                reordered = server.proposal_generation_context(
                    {**payload, "mini_box_order": list(reversed(identifiers))}
                )

        self.assertEqual(original["mini_box_order"], identifiers)
        self.assertEqual(original["mini_box_alignments"], {identifiers[0]: "center"})
        self.assertNotEqual(
            server.proposal_preview_fingerprint(original),
            server.proposal_preview_fingerprint(reordered),
        )

    def test_generation_context_accepts_the_complete_document_order(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = self.create_separate_marker_template(Path(temp_dir))
            structure = inspect_docx_structure(template)
            identifiers = [
                node["id"] for node in structure["nodes"] if node["type"] == "MINI_BOX"
            ]
            document_order = [
                identifiers[1],
                GENERATED_TABLE_BLOCK_ID,
                identifiers[0],
            ]
            payload = {
                "items": [self.proposal_item()],
                "template_ref": "managed:proposta_blocos.docx",
                "source_name": "edital",
                "responsible_id": "1",
                "commercial_terms": {},
                "mini_box_order": [identifiers[1], identifiers[0]],
                "document_block_order": document_order,
            }
            with (
                patch.object(server, "resolve_template", return_value=template),
                patch.object(server, "resolve_responsible", return_value={"id": "1"}),
            ):
                context = server.proposal_generation_context(payload)

        self.assertEqual(context["document_block_order"], document_order)
        self.assertEqual(context["mini_box_order"], [identifiers[1], identifiers[0]])

    def test_structure_endpoint_contract_resolves_the_requested_template(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            template = self.create_template(Path(temp_dir))
            with patch.object(server, "resolve_template", return_value=template):
                response = server.docx_structure_response(
                    {"template_ref": "managed:proposta.docx"}
                )

        self.assertEqual(response["mini_box_count"], 2)
        self.assertEqual(
            response["generated_table_block"]["id"],
            GENERATED_TABLE_BLOCK_ID,
        )
        self.assertEqual(
            [node["content"] for node in response["nodes"] if node["type"] == "MINI_BOX"],
            ["Primeiro", "Segundo"],
        )


if __name__ == "__main__":
    unittest.main()
