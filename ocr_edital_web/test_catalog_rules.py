import json
import tempfile
import unittest
from pathlib import Path

import pdfplumber
from docx import Document

import catalog_generator
from catalog_rules import (
    analyze_catalog_item,
    apply_user_catalog_repertoire,
    catalog_policy_summary,
    catalog_repertoire_key,
    catalog_summary,
    repertoire_summary,
)


def opportunity_item(description, *, number="1"):
    return {
        "numeroItem": number,
        "descricao": description,
        "quantidade": "10",
        "unidade": "UN",
    }


class CatalogRulesTests(unittest.TestCase):
    def test_policy_is_portrait_generic_and_backed_by_structured_repertoire(self):
        policy = catalog_policy_summary()
        repertoire = repertoire_summary()

        self.assertEqual(policy["orientation"], "retrato")
        self.assertEqual(policy["page_size"], "A4")
        self.assertEqual(policy["section_title"], "Características")
        self.assertEqual(policy["architecture"], "catalogo_generico_com_auditoria_separada")
        self.assertEqual(repertoire["structured_models"], 6)
        self.assertEqual(repertoire["source_documents"], 47)
        self.assertEqual(len(repertoire["models"]), 6)

    def test_matching_uses_repertoire_and_never_declares_automatic_compliance(self):
        description = (
            "Cadeira universitária fixa estofada com prancheta, compensado e tecido "
            "100% poliéster. Capacidade mínima de 120 kg. ABNT NBR 16671:2018. "
            "Assento mínimo 39 x 42 x 5 cm."
        )
        item = catalog_generator.normalize_items(
            [opportunity_item(description)],
            "https://pncp.gov.br/app/editais/1/2026/1",
        )[0]

        self.assertEqual(item["modelo_referencia"]["id"], "ufrgs-estofada-prancheta")
        self.assertFalse(item["analise_aderencia"]["declaracao_atendimento_automatica"])
        capacity = next(
            criterion
            for criterion in item["analise_aderencia"]["criterios"]
            if criterion["id"] == "capacidade_carga"
        )
        self.assertEqual(capacity["estado"], "potencialmente_atende")
        self.assertEqual(capacity["diferenca_absoluta_kg"], 10)
        norm = next(
            criterion
            for criterion in item["analise_aderencia"]["criterios"]
            if criterion["criterio"] == "Norma técnica"
        )
        self.assertEqual(norm["estado"], "evidenciado_na_referencia")
        self.assertEqual(item["observacao_repertorio"]["status"], "evidencia_completa")
        self.assertTrue(item["observacao_repertorio"]["evidencias"])
        self.assertEqual(item["observacao_repertorio"]["faltantes"], [])

    def test_dimension_below_express_minimum_blocks_catalog_without_hidden_tolerance(self):
        description = (
            "Cadeira giratória em tela Mesh. Encosto mínimo 56 x 45 x 8 cm, "
            "apoio lombar e pistão classe 4."
        )
        item = catalog_generator.normalize_items(
            [opportunity_item(description)],
            "https://pncp.gov.br/app/editais/1/2026/1",
        )[0]
        dimension = next(
            criterion
            for criterion in item["analise_aderencia"]["criterios"]
            if criterion["criterio"] == "Dimensão"
        )

        self.assertEqual(item["modelo_referencia"]["id"], "pinhais-giratoria-tela")
        self.assertEqual(dimension["estado"], "divergente")
        self.assertEqual(dimension["diferencas_absolutas_mm"][0], -20)
        self.assertIsNone(dimension["tolerancia_percentual"])
        self.assertEqual(item["status_catalogo"], "bloqueado_por_divergencia")
        self.assertEqual(item["observacao_repertorio"]["status"], "evidencia_parcial")
        self.assertTrue(
            any("Dimensão" in entry for entry in item["observacao_repertorio"]["faltantes"])
        )

    def test_unknown_product_does_not_promote_opportunity_text_to_catalog_claim(self):
        analyzed = analyze_catalog_item({
            "produto": "Mesa especial",
            "descricao": "Mesa em material secreto com resistência de 500 kg.",
            "especificacao_tecnica": "",
            "criterios_aceitacao": "",
            "observacoes": "",
            "categoria": "Mobiliário",
            "subcategoria": "Mesas",
        })

        self.assertIsNone(analyzed["modelo_referencia"])
        self.assertEqual(analyzed["caracteristicas_catalogo"], [])
        self.assertEqual(analyzed["status_catalogo"], "bloqueado_sem_modelo")
        self.assertEqual(analyzed["observacao_repertorio"]["status"], "sem_repertorio")
        self.assertEqual(catalog_summary([analyzed])["status_liberacao"], "bloqueado_sem_modelo")

    def test_user_repertoire_applies_numeric_ruler_without_automatic_approval(self):
        item = analyze_catalog_item({
            "produto": "Peça X especial",
            "descricao": "Peça X com tamanho de 5 cm",
            "especificacao_tecnica": "Tamanho exigido: 5 cm",
            "criterios_aceitacao": "",
            "observacoes": "",
            "categoria": "Componentes",
            "subcategoria": "Peças",
        })
        repertoire = {
            "id": "a" * 32,
            "item_key": catalog_repertoire_key(item),
            "produto_nome": "Peça X Goldflex",
            "cobertura_completa": True,
            "parametros": [{
                "id": "b" * 32,
                "componente": "Peça X",
                "atributo": "Tamanho",
                "comparacao": "intervalo",
                "valor_requerido": 5,
                "valor_minimo": 3,
                "valor_maximo": 6,
                "unidade": "cm",
                "evidencia": "Ficha técnica interna FT-001",
            }],
        }

        analyzed = apply_user_catalog_repertoire(item, repertoire)

        self.assertEqual(analyzed["observacao_repertorio"]["status"], "evidencia_completa")
        self.assertEqual(analyzed["observacao_repertorio"]["faltantes"], [])
        self.assertEqual(analyzed["analise_aderencia"]["resultado"], "referencia_identificada")
        self.assertFalse(analyzed["analise_aderencia"]["declaracao_atendimento_automatica"])
        self.assertEqual(analyzed["status_catalogo"], "rascunho_para_revisao")

    def test_user_repertoire_lists_divergence_and_incomplete_coverage(self):
        item = analyze_catalog_item({
            "produto": "Peça X especial",
            "descricao": "Peça X com tamanho de 8 cm",
            "especificacao_tecnica": "Tamanho exigido: 8 cm",
            "criterios_aceitacao": "",
            "observacoes": "",
            "categoria": "Componentes",
            "subcategoria": "Peças",
        })
        repertoire = {
            "id": "c" * 32,
            "produto_nome": "Peça X Goldflex",
            "cobertura_completa": False,
            "parametros": [{
                "id": "d" * 32,
                "componente": "Peça X",
                "atributo": "Tamanho",
                "comparacao": "intervalo",
                "valor_requerido": 8,
                "valor_minimo": 3,
                "valor_maximo": 6,
                "unidade": "cm",
                "evidencia": "Ficha técnica interna FT-001",
            }],
        }

        analyzed = apply_user_catalog_repertoire(item, repertoire)

        self.assertEqual(analyzed["observacao_repertorio"]["status"], "evidencia_parcial")
        self.assertEqual(analyzed["status_catalogo"], "bloqueado_por_divergencia")
        self.assertTrue(any("8 cm" in entry for entry in analyzed["observacao_repertorio"]["faltantes"]))
        self.assertTrue(any("marcado como parcial" in entry for entry in analyzed["observacao_repertorio"]["faltantes"]))

    def test_export_separates_generic_portrait_catalog_from_opportunity_audit(self):
        sentinel = "REQUISITO_SENTINELA_NAO_PUBLICAR"
        items = catalog_generator.normalize_items(
            [opportunity_item(
                "Cadeira giratória em tela Mesh, apoio lombar, pistão classe 4 e "
                f"cinco rodízios. {sentinel}"
            )],
            "https://pncp.gov.br/app/editais/1/2026/1",
        )
        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir)
            exports = catalog_generator.export_catalog(
                output,
                {"objeto": "EDITAL_SENTINELA_NAO_PUBLICAR", "orgao": "Órgão de teste"},
                items,
                "a" * 32,
            )

            self.assertEqual(set(exports), {"docx", "pdf", "xlsx", "csv", "json"})
            pdf_path = output / exports["pdf"]["filename"]
            with pdfplumber.open(pdf_path) as pdf:
                self.assertTrue(all(page.width < page.height for page in pdf.pages))
                pdf_text = "\n".join(page.extract_text() or "" for page in pdf.pages)
            self.assertIn("Catálogo técnico Goldflex", pdf_text)
            self.assertIn("Características", pdf_text)
            self.assertNotIn(sentinel, pdf_text)
            self.assertNotIn("EDITAL_SENTINELA", pdf_text)

            docx_path = output / exports["docx"]["filename"]
            document = Document(docx_path)
            self.assertTrue(all(section.page_width < section.page_height for section in document.sections))
            docx_text = "\n".join(
                [paragraph.text for paragraph in document.paragraphs]
                + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            )
            self.assertIn("Catálogo técnico Goldflex", docx_text)
            self.assertIn("Características", docx_text)
            self.assertNotIn(sentinel, docx_text)
            self.assertNotIn("EDITAL_SENTINELA", docx_text)
            self.assertIn("FFC000", document.part.blob.decode("utf-8"))

            payload = json.loads((output / exports["json"]["filename"]).read_text(encoding="utf-8"))
            self.assertIn("auditoria_oportunidade", payload)
            self.assertIn("catalogos", payload)
            self.assertIn(sentinel, json.dumps(payload["auditoria_oportunidade"], ensure_ascii=False))
            self.assertNotIn(sentinel, json.dumps(payload["catalogos"], ensure_ascii=False))

    def test_export_recomputes_analysis_after_requirement_edit(self):
        original = catalog_generator.normalize_items(
            [opportunity_item("Cadeira giratória em tela Mesh")],
            "https://pncp.gov.br/app/editais/1/2026/1",
        )[0]
        edited = dict(
            original,
            descricao="Mesa de reunião",
            especificacao_tecnica="Mesa de reunião",
            produto="Mesa de reunião",
        )

        prepared = catalog_generator.prepare_catalog_items([edited])[0]

        self.assertIsNone(prepared["modelo_referencia"])
        self.assertEqual(prepared["status_catalogo"], "bloqueado_sem_modelo")


if __name__ == "__main__":
    unittest.main()
